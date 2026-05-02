#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
全自动化图片处理工具 V3.1
==============================
修复问题：
  1. 图片归档：处理完成后自动移动图片到"已处理图片"
  2. 动态分类：根据内容动态创建新分类，而非固定分类
  3. 进度条：添加实时进度条显示
  4. 文档命名优化：根据内容提取更有意义的标题
  5. 内容整理增强：改进结构化逻辑，自动识别章节

使用方法：
  python auto_process_all_v3.py
"""

import os
import sys
import io
import json
import time
import re
import hashlib
import shutil
import logging
from pathlib import Path
from datetime import datetime
from collections import Counter, defaultdict
from typing import Optional, List, Dict

# 修复 Windows 控制台编码
if sys.platform == 'win32':
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
    except Exception:
        pass

# 添加项目根目录到路径
sys.path.insert(0, '.')

# 加载 .env 环境变量配置
from dotenv import load_dotenv
env_path = Path('.') / '.env'
if env_path.exists():
    load_dotenv(dotenv_path=env_path)
else:
    print("[WARN] .env 文件不存在，使用系统环境变量")

# 导入各模块
from local_ocr import LocalOCR
from scripts.classifier_engine import ClassifierEngine

# 尝试导入云端OCR
try:
    from tencent_ocr import TencentOCR
    TENINCENT_AVAILABLE = True
except ImportError:
    TENINCENT_AVAILABLE = False

try:
    from baidu_ocr import BaiduOCR
    BAIDU_AVAILABLE = True
except ImportError:
    BAIDU_AVAILABLE = False

# 日志配置
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(message)s',
    handlers=[
        logging.FileHandler('处理结果/process.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# ============================================================
# 进度条显示
# ============================================================
class ProgressBar:
    """终端进度条"""
    
    def __init__(self, total: int, width: int = 40, prefix: str = "进度"):
        self.total = total
        self.width = width
        self.prefix = prefix
        self.current = 0
        
    def update(self, current: int, info: str = ""):
        self.current = current
        percent = current / self.total if self.total > 0 else 0
        filled = int(self.width * percent)
        bar = '█' * filled + '░' * (self.width - filled)
        
        # 计算ETA
        if current > 0:
            elapsed = time.time() - self.start_time
            eta = elapsed / current * (self.total - current)
            if eta > 60:
                eta_str = f"{int(eta/60)}分{int(eta%60)}秒"
            else:
                eta_str = f"{int(eta)}秒"
        else:
            eta_str = "--"
            
        print(f"\r{self.prefix}: |{bar}| {current}/{self.total} ({percent:.0%}) ETA:{eta_str} {info}", 
              end='', flush=True)
        
    def start(self):
        self.start_time = time.time()
        self.update(0)
        
    def finish(self):
        self.update(self.total, "完成!")
        print()


# ============================================================
# 主题分类配置 - 基础关键词 + 动态扩展
# ============================================================
BASE_THEME_KEYWORDS = {
    "中医养生": ["中医", "养生", "药膳", "食疗", "中药", "经络", "穴位", "调理", "体质", "气血", "肝", "肾", "脾", "胃", "湿气", "虚", "寒", "热", "上火", "补", "滋阴", "温阳", "祛湿", "健脾", "润肺", "养心", "补肾"],
    "健康饮食": ["饮食", "营养", "食物", "蔬菜", "水果", "蛋白质", "碳水", "健康", "抗炎", "抗氧化", "免疫力", "肠道", "益生元", "膳食纤维", "维生素", "矿物质", "有机", "纯天然", "少油", "少盐", "清淡"],
    "疾病防治": ["疾病", "预防", "治疗", "症状", "指标", "血压", "血糖", "血脂", "胆固醇", "尿酸", "脂肪肝", "结节", "囊肿", "肿瘤", "癌症", "慢性病", "并发症", "吃药", "服药", "手术", "复查"],
    "生活方式": ["运动", "睡眠", "压力", "情绪", "心理健康", "作息", "习惯", "减肥", "增重", "美容", "护肤", "跑步", "走路", "瑜伽", "冥想", "放松"],
    "营养科普": ["科普", "知识", "研究", "发现", "实验", "数据", "结论", "专家", "建议", "指南", "推荐", "科学", "原理", "机制", "分析", "解读"],
    "旅游攻略": ["旅游", "旅行", "自驾", "景点", "打卡", "必去", "攻略", "路线", "酒店", "民宿", "拍照"],
    "历史文化": ["历史", "文化", "典故", "人物", "朝代", "皇帝", "诗词", "文物", "博物馆", "遗址", "传承"],
    "育儿教育": ["育儿", "宝宝", "孩子", "教育", "辅食", "喂养", "早教", "亲子", "成长", "发育"],
    "美食烹饪": ["美食", "烹饪", "做法", "菜谱", "食谱", "食材", "配料", "做饭", "煮", "炒", "炖"],
    "植物养护": ["植物", "养护", "种植", "花草", "绿植", "养花", "盆栽", "浇水", "施肥"]
}

# 主题特征词（用于识别内容类型）
TOPIC_SIGNATURES = {
    "养生方": ["方子", "秘方", "配方", "做法", "煮", "熬", "炖"],
    "食疗食谱": ["食谱", "菜谱", "做法", "食材", "配料", "烹饪"],
    "禁忌注意": ["禁忌", "注意", "不宜", "少吃", "避免", "千万", "不能", "不可"],
    "健康功效": ["功效", "作用", "好处", "益处", "价值", "可以", "有助于"],
    "疾病症状": ["症状", "表现", "特征", "体征", "不舒服", "疼痛", "难受"]
}


class ImageArchiver:
    """图片归档器 - 处理完成后移动图片到已处理文件夹"""
    
    def __init__(self, processed_dir: str = '已处理图片'):
        self.processed_dir = Path(processed_dir)
        self.processed_dir.mkdir(parents=True, exist_ok=True)
        self.archived_count = 0
        
    def archive_image(self, image_path: str) -> Optional[str]:
        """移动图片到已处理文件夹"""
        src = Path(image_path)
        if not src.exists():
            return None
            
        # 生成唯一文件名（避免覆盖）
        dest_name = src.name
        dest = self.processed_dir / dest_name
        
        if dest.exists():
            # 文件已存在，添加时间戳
            timestamp = datetime.now().strftime('%H%M%S')
            name_parts = src.stem, src.suffix
            dest = self.processed_dir / f"{name_parts[0]}_{timestamp}{name_parts[1]}"
        
        try:
            # 移动文件
            shutil.move(str(src), str(dest))
            self.archived_count += 1
            return str(dest)
        except Exception as e:
            logger.warning(f"[归档] 移动失败 {src.name}: {e}")
            # 尝试复制
            try:
                shutil.copy2(str(src), str(dest))
                src.unlink()  # 删除原文件
                self.archived_count += 1
                return str(dest)
            except:
                return None
                
    def get_archive_stats(self) -> Dict:
        """获取归档统计"""
        total = len(list(self.processed_dir.glob('*.*')))
        return {
            'archived_count': self.archived_count,
            'total_archived': total
        }


class DynamicClassifier:
    """动态分类器 - 根据内容自动识别和创建新分类"""
    
    def __init__(self):
        # 基础主题
        self.themes = dict(BASE_THEME_KEYWORDS)
        # 已创建的自定义主题
        self.custom_themes: Dict[str, List[str]] = {}
        # 主题创建历史
        self.theme_history: List[str] = []
        
    def add_custom_theme(self, theme_name: str, keywords: List[str]):
        """添加自定义主题"""
        if theme_name not in self.themes:
            self.themes[theme_name] = keywords
            self.custom_themes[theme_name] = keywords
            self.theme_history.append(theme_name)
            logger.info(f"[分类] 创建新主题: {theme_name}")
            
    def extract_topic_keywords(self, text: str) -> List[str]:
        """从内容中提取主题关键词，用于创建新分类"""
        # 提取名词短语
        patterns = [
            r'[\u4e00-\u9fa5]{2,4}(?:方法|指南|攻略|知识|技巧|教程|原理)',
            r'关于[\u4e00-\u9fa5]{2,6}',
            r'[\u4e00-\u9fa5]{2,6}(?:是什么|怎么做|如何做)',
        ]
        
        keywords = []
        for pattern in patterns:
            matches = re.findall(pattern, text)
            keywords.extend(matches)
            
        return list(set(keywords))[:10]  # 去重，最多10个
        
    def detect_or_create_theme(self, text: str, min_score: float = 2.0) -> str:
        """检测主题，如果没有匹配则动态创建新分类"""
        max_score = 0
        best_theme = "综合知识"
        
        # 计算每个主题的得分
        for theme, keywords in self.themes.items():
            score = sum(3 if kw in text else 0 for kw in keywords)
            if score > max_score:
                max_score = score
                best_theme = theme
                
        # 如果得分低于阈值，创建新主题
        if max_score < min_score:
            # 尝试从内容中提取新主题名
            new_theme = self._extract_theme_name(text)
            if new_theme:
                extracted_keywords = self.extract_topic_keywords(text)
                if extracted_keywords:
                    self.add_custom_theme(new_theme, extracted_keywords[:5])
                    return new_theme
                    
        return best_theme
        
    def _extract_theme_name(self, text: str) -> Optional[str]:
        """从内容中提取新主题名称"""
        # 查找模式：关于XXX、XXX指南、XXX知识
        patterns = [
            r'关于([\u4e00-\u9fa5]{2,6})',
            r'([\u4e00-\u9fa5]{2,4})知识',
            r'([\u4e00-\u9fa5]{2,4})指南',
            r'([\u4e00-\u9fa5]{2,4})攻略',
            r'([\u4e00-\u9fa5]{2,4})技巧',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                theme_name = match.group(1)
                # 排除已存在的主题
                if theme_name not in self.themes and len(theme_name) >= 2:
                    return theme_name
                    
        return None
        
    def get_all_themes(self) -> List[str]:
        """获取所有主题"""
        return list(self.themes.keys())


class ContentOrganizer:
    """内容整理器 - 对OCR识别的原始内容进行智能整理"""

    def __init__(self):
        self.source_keywords = ['来源', '出处', '作者', '小红书', '微信', '抖音', '微博', 'B站', '公众号']
        
        # 增强的标题关键词
        self.heading_keywords = [
            # 功效类
            '功效', '作用', '好处', '益处', '价值', '有助于',
            # 材料类
            '原料', '食材', '配料', '材料', '需要',
            # 做法类
            '做法', '制作', '步骤', '方法', '教程', '如何', '怎么',
            # 注意类
            '禁忌', '注意', '提醒', '警告', '不宜', '避免',
            # 介绍类
            '简介', '介绍', '概述', '说明', '是什么', '什么是',
            # 推荐类
            '推荐', '建议', '指南', '攻略'
        ]
        
        # 段落分隔关键词
        self.section_keywords = [
            '一、', '二、', '三、', '四、', '五、', '六、',
            '1.', '2.', '3.', '4.', '5.', '6.',
            '（1）', '（2）', '（3）', '（4）',
            '第一', '第二', '第三', '第四', '第五',
            '首先', '其次', '然后', '最后', '另外', '此外'
        ]

    def extract_source(self, text: str) -> Dict:
        """从内容中提取来源信息"""
        lines = text.split('\n')
        source_info = {}

        for line in lines[:20]:
            line = line.strip()
            
            for kw in ['来源', '出处', '作者', '摘录自']:
                if kw in line:
                    parts = line.split(kw)
                    if len(parts) > 1:
                        source = parts[1].strip().lstrip('：:').strip()
                        if source:
                            source_info['platform'] = source
                            break

            for platform in ['小红书', '微信读书', '微信视频号', '微信', '抖音', 'B站', 'BTV']:
                if platform in line:
                    source_info['platform'] = platform
                    break

            if '作者' in line:
                match = re.search(r'作者[：:]\s*([^\n]+)', line)
                if match:
                    source_info['author'] = match.group(1).strip()

        return source_info

    def clean_text(self, text: str) -> str:
        """清理OCR识别的原始文本"""
        lines = text.split('\n')
        cleaned_lines = []
        prev_line_empty = False

        for line in lines:
            line = line.strip()

            if len(line) < 2 and line:
                continue

            if re.match(r'^[_\-=]{3,}$', line):
                continue

            if re.match(r'^[\u4e00-\u9fa5]*[a-zA-Z0-9]{5,}[\u4e00-\u9fa5]*$', line):
                if not re.search(r'[\u4e00-\u9fa5]', line):
                    continue

            if not line:
                if not prev_line_empty:
                    cleaned_lines.append('')
                    prev_line_empty = True
                continue

            prev_line_empty = False
            line = self._fix_common_ocr_errors(line)
            cleaned_lines.append(line)

        result = '\n'.join(cleaned_lines)
        result = re.sub(r'\n{3,}', '\n\n', result)
        return result.strip()

    def _fix_common_ocr_errors(self, line: str) -> str:
        """修复常见的OCR错误"""
        line = line.replace(chr(8220), '"')
        line = line.replace(chr(8221), '"')
        line = line.replace(chr(8216), "'")
        line = line.replace(chr(8217), "'")
        line = line.replace(chr(8218), "'")
        line = line.replace(chr(8222), '"')
        line = line.replace(chr(8223), '"')
        line = line.replace(chr(8242), "'")
        line = line.replace(chr(8243), "'")
        line = line.replace(chr(180), "")
        line = line.replace(chr(9032), "")
        line = line.replace(chr(8212), "-")
        line = line.replace(chr(183), "-")
        line = line.replace("`", "")
        line = line.replace('V', '✓').replace('X', '✗')
        line = re.sub(r' {2,}', ' ', line)
        return line

    def structure_content(self, text: str, source_info: Dict = None) -> str:
        """对内容进行智能结构化处理"""
        lines = text.split('\n')
        structured = []
        current_h1 = None
        current_h2 = None
        h2_content = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            heading_type = self._detect_heading_type(line)
            
            if heading_type == 1:  # 一级标题
                # 保存之前的H2内容
                if h2_content and current_h2:
                    structured.extend(h2_content)
                    structured.append('')
                if current_h1:
                    structured.append('')
                structured.append(f"## {line}")
                current_h1 = line
                current_h2 = None
                h2_content = []
                
            elif heading_type == 2:  # 二级标题
                if h2_content and current_h2:
                    structured.extend(h2_content)
                    structured.append('')
                if current_h1 is None:
                    structured.append(f"## {line}")
                    current_h1 = line
                else:
                    structured.append(f"### {line}")
                current_h2 = line
                h2_content = []
                
            else:  # 普通内容
                if current_h2:
                    h2_content.append(line)
                else:
                    structured.append(line)

        # 添加最后的H2内容
        if h2_content:
            structured.extend(h2_content)

        return '\n'.join(structured)

    def _detect_heading_type(self, line: str) -> int:
        """检测标题类型：0=普通内容, 1=一级标题, 2=二级标题"""
        if len(line) < 4:
            return 0

        # 检查是否包含标题关键词
        heading_count = sum(1 for kw in self.heading_keywords if kw in line)
        if heading_count > 0 and len(line) <= 35:
            if line.startswith('#'):
                return 1
            # 特定格式开头
            if re.match(r'^[◆★▼▪▸●○]', line) and len(line) <= 25:
                return 2
            if re.match(r'^[\u4e00-\u9fa5]{4,15}$', line) and len(line) >= 4:
                return 1
            return 2

        # 检查序号格式
        if re.match(r'^[一二三四五六七八九十\d]+[、.．:：]', line):
            if len(line) <= 20:
                return 2
                
        # 检查分隔关键词
        for sep in self.section_keywords:
            if line.startswith(sep) and len(line) <= 40:
                return 2
                
        return 0

    def organize(self, raw_text: str, theme: str = None) -> Dict:
        """综合整理内容"""
        source_info = self.extract_source(raw_text)
        cleaned_text = self.clean_text(raw_text)
        structured_text = self.structure_content(cleaned_text, source_info)

        return {
            'content': structured_text,
            'source': source_info,
            'theme': theme
        }


class ContentAnalyzer:
    """AI内容分析器 - 自动识别主题并生成文档名"""

    def __init__(self, dynamic_classifier: DynamicClassifier = None):
        self.cache = {}
        self.organizer = ContentOrganizer()
        self.dynamic_classifier = dynamic_classifier or DynamicClassifier()

    def extract_title(self, text: str) -> Optional[str]:
        """从内容中提取标题 - 增强版"""
        lines = text.split('\n')
        
        # 第一轮：查找明确的标题格式
        for line in lines[:15]:
            line = line.strip()
            
            # # 标题
            if line.startswith('#'):
                title = line.lstrip('#').strip()
                if 2 <= len(title) <= 25:
                    return title
                    
            # 特殊符号标题
            if re.match(r'^[◆★▼▪▸]+', line):
                title = re.sub(r'^[◆★▼▪▸]+\s*', '', line)
                if 2 <= len(title) <= 20:
                    return title
                    
            # 居中大标题（纯中文，4-15字）
            if re.match(r'^[\u4e00-\u9fa5]{4,15}$', line):
                return line
                
            # 关于XXX模式
            match = re.match(r'^关于([\u4e00-\u9fa5]{2,8})', line)
            if match:
                return match.group(1)
                
        # 第二轮：查找关键词组合
        for line in lines[:20]:
            line = line.strip()
            if len(line) >= 4 and len(line) <= 20:
                # 检查是否包含多个标题关键词
                kw_count = sum(1 for kw in self.organizer.heading_keywords if kw in line)
                if kw_count >= 2:
                    return line
                # 检查是否是"XXX知识/指南/攻略"格式
                match = re.match(r'([\u4e00-\u9fa5]{2,8})(?:知识|指南|攻略|技巧|方法)', line)
                if match:
                    return line
                    
        # 第三轮：第一个有意义的短句
        for line in lines[:15]:
            line = line.strip()
            if 4 <= len(line) <= 20 and re.match(r'^[\u4e00-\u9fa5]', line):
                if 'http' not in line.lower() and not line.startswith('#'):
                    return line
                    
        return None

    def generate_doc_name(self, text: str, theme: str, index: int) -> str:
        """生成有意义的文档名 - 增强版"""
        # 1. 尝试从内容中提取标题
        title = self.extract_title(text)
        if title:
            # 清理标题
            title = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '', title)
            if len(title) >= 3:
                # 截取有意义的长度
                if len(title) > 20:
                    title = title[:20]
                return title

        # 2. 根据主题和内容类型生成
        topic_type = self._detect_topic_type(text)
        if topic_type != "知识整理":
            return f"{topic_type}"

        # 3. 根据主题生成
        if theme and theme != "综合知识":
            return f"{theme}_整理"

        # 4. 默认命名
        timestamp = datetime.now().strftime('%m%d%H%M')
        return f"知识_{timestamp}_{index:02d}"

    def _detect_topic_type(self, text: str) -> str:
        """检测内容类型"""
        for topic, signatures in TOPIC_SIGNATURES.items():
            score = sum(2 if sig in text else 0 for sig in signatures)
            if score >= 2:
                return topic
        return "知识整理"

    def compute_content_hash(self, text: str) -> str:
        """计算内容哈希"""
        cleaned = re.sub(r'\s+', '', text)
        cleaned = cleaned[:500]
        return hashlib.md5(cleaned.encode('utf-8')).hexdigest()[:8]

    def analyze(self, text: str, theme: str, index: int) -> Dict:
        """综合分析内容"""
        cache_key = self.compute_content_hash(text)

        if cache_key in self.cache:
            return self.cache[cache_key]

        # 使用动态分类器
        detected_theme = self.dynamic_classifier.detect_or_create_theme(text)
        topic_type = self._detect_topic_type(text)
        doc_name = self.generate_doc_name(text, detected_theme, index)

        # 整理内容
        organized = self.organizer.organize(text, detected_theme)

        result = {
            'theme': detected_theme,
            'topic_type': topic_type,
            'doc_name': doc_name,
            'content_hash': cache_key,
            'organized_content': organized['content'],
            'source': organized['source']
        }

        self.cache[cache_key] = result
        return result


class MultiEngineOCR:
    """多引擎OCR管理器"""

    def __init__(self):
        self.current_engine = None
        self.engine_status = []

        secret_id = os.getenv('TENCENT_SECRET_ID', '')
        secret_key = os.getenv('TENCENT_SECRET_KEY', '')
        if secret_id and secret_key and '替换' not in secret_id and TENINCENT_AVAILABLE:
            self.engine_status.append(('腾讯云', True, '已配置'))
        else:
            self.engine_status.append(('腾讯云', False, '未配置'))

        app_id = os.getenv('BAIDU_APP_ID', '')
        api_key = os.getenv('BAIDU_API_KEY', '')
        secret_key = os.getenv('BAIDU_SECRET_KEY', '')
        if app_id and api_key and secret_key and '替换' not in app_id and BAIDU_AVAILABLE:
            self.engine_status.append(('百度云', True, '已配置'))
        else:
            self.engine_status.append(('百度云', False, '未配置'))

        local = LocalOCR()
        if local.tesseract_available:
            self.engine_status.append(('本地Tesseract', True, '可用'))
        else:
            self.engine_status.append(('本地Tesseract', False, getattr(local, 'error_message', '未知错误')))

        self._select_best_engine()

    def _select_best_engine(self):
        for name, available, msg in self.engine_status:
            if available:
                self.current_engine = name
                logger.info(f"[OCR] 选用引擎: {name} ({msg})")
                return True
        logger.warning("[OCR] 没有可用的OCR引擎!")
        return False

    def recognize(self, image_path: str) -> Dict:
        last_error = None
        for name, available, _ in self.engine_status:
            if not available:
                continue
            try:
                logger.info(f"[OCR] 尝试 {name}...")
                if name == '腾讯云' and TENINCENT_AVAILABLE:
                    engine = TencentOCR()
                elif name == '百度云' and BAIDU_AVAILABLE:
                    engine = BaiduOCR()
                else:
                    engine = LocalOCR()

                result = engine.recognize(image_path)
                if result and result.get('success'):
                    logger.info(f"[OCR] {name} 识别成功!")
                    return result
                else:
                    last_error = result.get('error', '未知错误') if result else '返回为空'
            except Exception as e:
                last_error = str(e)
        return {'success': False, 'error': last_error or '无可用OCR引擎', 'text': ''}


class DocumentMerger:
    """文档合并器"""

    def __init__(self, output_dir: str = '处理结果'):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.theme_docs = defaultdict(list)
        self._scan_existing_docs()

    def _scan_existing_docs(self):
        for theme_folder in self.output_dir.iterdir():
            if theme_folder.is_dir() and not theme_folder.name.startswith('.'):
                for md_file in theme_folder.glob('*.md'):
                    topic = self._extract_topic(md_file)
                    self.theme_docs[theme_folder.name].append({
                        'file': md_file,
                        'topic': topic
                    })

    def _extract_topic(self, md_file: Path) -> str:
        name = md_file.stem
        name = re.sub(r'^\d+_', '', name)
        name = re.sub(r'_\d+$', '', name)
        name = re.sub(r'_[a-f0-9]{8}$', '', name)
        return name

    def find_similar_doc(self, theme: str, doc_name: str, content: str) -> Optional[Path]:
        doc_topic = re.sub(r'^\d+_', '', doc_name)
        doc_topic = re.sub(r'_\d+$', '', doc_topic)

        for existing_doc in self.theme_docs.get(theme, []):
            existing_topic = existing_doc['topic']
            if self._calc_similarity(doc_topic, existing_topic) > 0.6:
                return existing_doc['file']
        return None

    def _calc_similarity(self, s1: str, s2: str) -> float:
        s1_set = set(s1)
        s2_set = set(s2)
        if not s1_set or not s2_set:
            return 0
        return len(s1_set & s2_set) / len(s1_set | s2_set)

    def merge_content(self, existing_file: Path, new_content: str, new_image_name: str, content_hash: str = None) -> bool:
        try:
            with open(existing_file, 'r', encoding='utf-8') as f:
                existing = f.read()

            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
            append_section = f"""

---

## 📌 补充内容

> 来源图片: {new_image_name}
> 追加时间: {timestamp}
> 内容标识: {content_hash}

{new_content}

"""

            new_content_merged = existing.replace(
                '\n\n---\n\n*本文档由图片知识库整理工具自动生成*',
                append_section + '\n\n---\n\n*本文档由图片知识库整理工具自动生成*'
            )

            with open(existing_file, 'w', encoding='utf-8') as f:
                f.write(new_content_merged)

            return True
        except Exception as e:
            logger.warning(f"[合并] 失败: {e}")
            return False


class SmartDocumentGenerator:
    """智能文档生成器"""

    def __init__(self, output_dir: str = '处理结果'):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.merger = DocumentMerger(output_dir)
        self.existing_hashes = set()
        self._scan_hashes()
        self.generated_docs = set()

    def _scan_hashes(self):
        for md_file in self.output_dir.rglob('*.md'):
            try:
                file_content = md_file.read_text(encoding='utf-8')
                hash_match = re.search(r'content_hash:\s*([a-f0-9]{8})', file_content)
                if hash_match:
                    self.existing_hashes.add(hash_match.group(1))
            except:
                pass

    def _get_doc_key(self, theme: str, doc_name: str) -> str:
        safe_name = re.sub(r'[<>:"/\|?*]', '', doc_name)
        if len(safe_name) > 25:
            safe_name = safe_name[:25]
        return f"{theme}/{safe_name}"

    def is_duplicate(self, content_hash: str) -> bool:
        return content_hash in self.existing_hashes

    def is_doc_generated(self, theme: str, doc_name: str) -> bool:
        return self._get_doc_key(theme, doc_name) in self.generated_docs

    def mark_doc_generated(self, theme: str, doc_name: str):
        self.generated_docs.add(self._get_doc_key(theme, doc_name))

    def generate_markdown(self, text: str, theme: str, doc_name: str, category: str, 
                         image_name: str, keywords: List = None, content_hash: str = None, 
                         source_info: Dict = None) -> str:
        """生成Markdown文档"""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')

        # 查找是否有相似文档可以合并
        similar_doc = self.merger.find_similar_doc(theme, doc_name, text)
        if similar_doc:
            print(f"  → 发现相似文档 {similar_doc.name}，合并内容...")
            if self.merger.merge_content(similar_doc, text, image_name, content_hash):
                self.mark_doc_generated(theme, doc_name)
                return str(similar_doc)

        keywords_str = ', '.join(keywords) if keywords else '无'

        source_section = ""
        if source_info:
            if source_info.get('platform'):
                source_section += f"**来源**：{source_info['platform']}\n\n"
            if source_info.get('author'):
                source_section += f"**作者**：{source_info['author']}\n\n"

        md_content = f"""# {doc_name}

> 来源图片: {image_name}
> 识别时间: {timestamp}
> 主题分类: {theme}
> 原分类: {category}
> 关键词: {keywords_str}
> content_hash: {content_hash or 'N/A'}

---



{source_section}{text}

---

*本文档由图片知识库整理工具自动生成*
"""

        # 创建主题文件夹
        theme_folder = self.output_dir / theme
        theme_folder.mkdir(exist_ok=True)

        safe_name = re.sub(r'[<>:"/\|?*]', '', doc_name)
        if len(safe_name) > 25:
            safe_name = safe_name[:25]

        md_file = theme_folder / f"{safe_name}.md"
        counter = 1
        while md_file.exists():
            md_file = theme_folder / f"{safe_name}_{counter}.md"
            counter += 1

        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(md_content)

        if content_hash:
            self.existing_hashes.add(content_hash)
            
        self.merger.theme_docs[theme].append({
            'file': md_file,
            'topic': doc_name
        })
        self.mark_doc_generated(theme, doc_name)

        logger.info(f"[文档] Markdown已生成: {theme}/{md_file.name}")
        return str(md_file)

    def generate_word(self, text: str, theme: str, doc_name: str, image_name: str,
                     keywords: List = None, content_hash: str = None, existing_md_file: str = None) -> Optional[str]:
        """生成Word文档"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.oxml.ns import qn
        except ImportError:
            logger.warning("[文档] python-docx未安装，跳过Word生成")
            return None

        if self.is_doc_generated(theme, doc_name):
            safe_name = re.sub(r'[<>:"/\|?*]', '', doc_name)
            if len(safe_name) > 25:
                safe_name = safe_name[:25]
            existing_docx = self.output_dir / theme / f"{safe_name}.docx"
            if existing_docx.exists():
                print(f"  → Word文档已存在，跳过")
                return str(existing_docx)
            for i in range(1, 100):
                existing_docx = self.output_dir / theme / f"{safe_name}_{i}.docx"
                if existing_docx.exists():
                    print(f"  → Word文档已存在，跳过")
                    return str(existing_docx)

        theme_folder = self.output_dir / theme
        theme_folder.mkdir(exist_ok=True)

        safe_name = re.sub(r'[<>:"/\|?*]', '', doc_name)
        if len(safe_name) > 25:
            safe_name = safe_name[:25]

        docx_file = theme_folder / f"{safe_name}.docx"
        counter = 1
        while docx_file.exists():
            docx_file = theme_folder / f"{safe_name}_{counter}.docx"
            counter += 1

        if existing_md_file and Path(existing_md_file).exists():
            try:
                md_content = Path(existing_md_file).read_text(encoding='utf-8')
                main_content = md_content.split('## 内容\n\n')[1] if '## 内容\n\n' in md_content else ''
                main_content = main_content.split('## 📌 补充内容')[0] if '## 📌 补充内容' in main_content else main_content
                main_content = main_content.replace('\n\n---\n\n*本文档由图片知识库整理工具自动生成*', '')
                text = main_content.strip()
            except Exception as e:
                logger.warning(f"[Word] 读取合并内容失败: {e}")

        doc = Document()
        doc.styles['Normal'].font.name = 'Microsoft YaHei'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        doc.add_heading(doc_name, 0)
        doc.add_paragraph(f"来源图片: {image_name}")
        doc.add_paragraph(f"识别时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"主题分类: {theme}")
        keywords_str = ', '.join(keywords) if keywords else '无'
        doc.add_paragraph(f"关键词: {keywords_str}")
        doc.add_paragraph("─" * 30)
        doc.add_paragraph(text)

        doc.save(str(docx_file))
        self.mark_doc_generated(theme, doc_name)
        logger.info(f"[文档] Word已生成: {theme}/{docx_file.name}")
        return str(docx_file)


class IMASyncer:
    """IMA笔记同步器"""

    def __init__(self):
        self.client_id = os.getenv('IMA_OPENAPI_CLIENTID', '')
        self.api_key = os.getenv('IMA_OPENAPI_APIKEY', '')
        self.base_url = 'https://ima.qq.com/openapi/note/v1'
        self.enabled = bool(self.client_id and self.api_key and '填入' not in self.client_id)
        self.sync_log_file = Path('处理结果/ima_sync_log.json')
        self.sync_log = self._load_sync_log()
        self.rate_limited = False

        if self.enabled:
            logger.info("[IMA] 已配置")
        else:
            logger.warning("[IMA] 未配置")

    def _load_sync_log(self) -> Dict:
        if self.sync_log_file.exists():
            try:
                return json.loads(self.sync_log_file.read_text(encoding='utf-8'))
            except:
                pass
        return {}

    def _save_sync_log(self):
        self.sync_log_file.write_text(
            json.dumps(self.sync_log, ensure_ascii=False, indent=2),
            encoding='utf-8'
        )

    def _api_call(self, endpoint: str, payload: Dict, retries: int = 3) -> Optional[Dict]:
        if not self.enabled or self.rate_limited:
            return None

        for attempt in range(retries):
            try:
                import requests
                headers = {
                    'ima-openapi-clientid': self.client_id,
                    'ima-openapi-apikey': self.api_key,
                    'Content-Type': 'application/json'
                }
                url = f"{self.base_url}/{endpoint}"
                response = requests.post(url, json=payload, headers=headers, timeout=30)

                if response.status_code == 200:
                    return response.json()
                elif response.status_code == 403:
                    result = response.json()
                    if '请求超量' in result.get('msg', ''):
                        logger.warning("[IMA] API请求超量，已触发限流，明天再试")
                        self.rate_limited = True
                        return None
                return None
            except Exception as e:
                logger.warning(f"[IMA] API调用失败 ({attempt+1}/{retries}): {e}")
                if attempt < retries - 1:
                    time.sleep(2)
        return None

    def sync_note(self, title: str, content: str, theme: str = None, 
                  content_hash: str = None, doc_path: str = None) -> Optional[str]:
        if not self.enabled or self.rate_limited:
            return None

        doc_key = content_hash or title
        existing = self.sync_log.get(doc_key, {})

        full_content = f"# {title}\n\n"
        if theme:
            full_content += f"> 主题: {theme}\n"
        if content_hash:
            full_content += f"> 内容标识: {content_hash}\n"
        full_content += f"\n{content}\n\n---\n*自动同步自图片知识库*\n"

        if doc_path and existing.get('doc_id'):
            result = self._api_call('append_doc', {
                'doc_id': existing['doc_id'],
                'content_format': 1,
                'content': f"\n\n---\n*更新于 {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n" + content
            })
            if result:
                self.sync_log[doc_key] = {
                    'doc_id': existing['doc_id'],
                    'last_sync': datetime.now().isoformat(),
                    'update_count': existing.get('update_count', 0) + 1
                }
                self._save_sync_log()
                return existing['doc_id']

        result = self._api_call('import_doc', {
            'content_format': 1,
            'content': full_content
        })

        if result and result.get('code') == 0:
            doc_id = result.get('data', {}).get('note_id', '')
            if doc_id:
                self.sync_log[doc_key] = {
                    'doc_id': doc_id,
                    'first_sync': datetime.now().isoformat(),
                    'last_sync': datetime.now().isoformat(),
                    'theme': theme,
                    'update_count': 1
                }
                self._save_sync_log()
                logger.info(f"[IMA] 笔记已同步: {title[:20]}")
                return doc_id
            else:
                return "imported"

        return None


def process_single_image(ocr, classifier, content_analyzer, doc_gen, ima_syncer, 
                         archiver, image_path: str, index: int, total: int, 
                         progress: ProgressBar = None):
    """处理单张图片"""
    image_name = Path(image_path).name
    
    # 更新进度条
    if progress:
        progress.update(index, image_name[:20])
    
    print(f"\n{'='*50}")
    print(f"[{index}/{total}] 处理: {image_name}")
    print(f"{'='*50}")

    # 1. OCR识别
    print("\n📷 OCR识别中...")
    result = ocr.recognize(image_path)

    if not result.get('success'):
        print(f"  ⚠️ OCR失败: {result.get('error')}")
        archiver.archive_image(image_path)  # 归档失败图片
        return None

    text = result.get('text', '').strip()
    if not text or len(text) < 10:
        print("  ⚠️ 无文字内容或内容过少")
        archiver.archive_image(image_path)  # 归档空内容图片
        return None

    print(f"  ✓ 识别到 {len(text)} 个字符")

    # 2. AI分类
    print("\n🏷️ 内容分类中...")
    categories = classifier.classify(text, threshold=0.25)

    if categories:
        primary = categories[0]
        category_name = primary.category_name
        keywords = primary.matched_keywords if primary.matched_keywords else []
        confidence = primary.confidence
        print(f"  ✓ 分类: {category_name} (置信度 {confidence:.0%})")
        if keywords:
            print(f"     关键词: {', '.join(keywords[:5])}")
    else:
        category_name = "综合知识"
        keywords = []
        print("  ⚠️ 未匹配到分类，使用默认分类")

    # 3. AI内容分析（包括内容整理）
    print("\n🧠 AI内容分析中...")
    analysis = content_analyzer.analyze(text, category_name, index)
    theme = analysis['theme']
    topic_type = analysis['topic_type']
    doc_name = analysis['doc_name']
    content_hash = analysis['content_hash']
    organized_content = analysis.get('organized_content', text)
    source_info = analysis.get('source', {})
    
    print(f"  ✓ 主题分类: {theme}")
    print(f"  ✓ 内容类型: {topic_type}")
    print(f"  ✓ 文档命名: {doc_name}")
    if source_info.get('platform'):
        print(f"  ✓ 来源: {source_info['platform']}")
        
    # 打印新创建的主题
    if theme != category_name and theme not in BASE_THEME_KEYWORDS:
        print(f"  ★ 新建主题: {theme}")

    # 4. 检测重复
    print("\n🔍 检测重复内容...")
    if doc_gen.is_duplicate(content_hash):
        print("  ⚠️ 检测到重复内容，跳过存储")
        archiver.archive_image(image_path)  # 归档重复图片
        return {
            'image': image_name,
            'text_length': len(text),
            'category': category_name,
            'theme': theme,
            'doc_name': doc_name,
            'is_duplicate': True,
            'content_hash': content_hash
        }
    print("  ✓ 新内容")

    # 5. 生成Markdown
    print("\n📝 生成Markdown...")
    md_file = doc_gen.generate_markdown(
        organized_content, theme, doc_name, category_name,
        image_name, keywords, content_hash, source_info
    )

    # 6. 生成Word
    print("\n📄 生成Word...")
    docx_file = doc_gen.generate_word(
        organized_content, theme, doc_name, image_name, keywords, content_hash, existing_md_file=md_file
    )
    if docx_file:
        rel_path = Path(docx_file).relative_to(doc_gen.output_dir)
        print(f"  ✓ 已保存到: {rel_path}")
    else:
        print("  ⚠️ Word文档生成失败或未安装python-docx")

    # 7. 同步到IMA
    print("\n☁️ 同步到IMA...")
    if ima_syncer.rate_limited:
        print("  ⚠️ IMA被限流，请明天再试")
    else:
        ima_id = ima_syncer.sync_note(doc_name, organized_content, theme, content_hash, md_file)
        if ima_id:
            print(f"  ✓ IMA同步成功")
        else:
            print("  ⚠️ IMA同步失败")

    # 8. 图片归档（处理成功后才归档）
    print("\n📦 归档图片...")
    archived_path = archiver.archive_image(image_path)
    if archived_path:
        archived_name = Path(archived_path).name
        print(f"  ✓ 已归档到: 已处理图片/{archived_name}")

    return {
        'image': image_name,
        'text_length': len(text),
        'category': category_name,
        'theme': theme,
        'topic_type': topic_type,
        'doc_name': doc_name,
        'keywords': keywords,
        'is_duplicate': False,
        'md_file': md_file,
        'docx_file': docx_file,
        'content_hash': content_hash,
        'source': source_info
    }


def main():
    print("=" * 60)
    print("全自动化图片处理工具 V3.1")
    print("OCR → AI分类 → 动态主题 → 文档合并 → MD+Word → 归档")
    print("=" * 60)

    # 初始化组件
    print("\n[1/7] 初始化OCR引擎...")
    ocr = MultiEngineOCR()
    if not ocr.current_engine:
        print("❌ 没有可用的OCR引擎!")
        return

    print("\n[2/7] 初始化分类器...")
    classifier = ClassifierEngine()

    print("\n[3/7] 初始化动态分类器...")
    dynamic_classifier = DynamicClassifier()

    print("\n[4/7] 初始化内容分析器...")
    content_analyzer = ContentAnalyzer(dynamic_classifier)

    print("\n[5/7] 初始化文档生成器...")
    doc_gen = SmartDocumentGenerator()

    print("\n[6/7] 初始化IMA同步器...")
    ima_syncer = IMASyncer()
    if ima_syncer.rate_limited:
        print("  ⚠️ IMA被限流，本次将跳过IMA同步")

    print("\n[7/7] 初始化图片归档器...")
    archiver = ImageArchiver()

    print("\n[8/8] 扫描待处理图片...")

    # 扫描待处理图片
    source_dir = Path('待处理图片')
    if not source_dir.exists():
        source_dir = Path('待处理图片/历史')

    images = []
    for ext in ['*.jpg', '*.png', '*.jpeg', '*.webp', '*.bmp']:
        images.extend(source_dir.rglob(ext))

    if not images:
        print("❌ 未找到待处理图片")
        return

    total = len(images)
    print(f"✓ 找到 {total} 张图片\n")

    # 初始化进度条
    progress = ProgressBar(total, prefix="📊 处理进度")

    # 批量处理
    results = []
    start_time = time.time()
    progress.start()

    for i, img_path in enumerate(images, 1):
        try:
            result = process_single_image(
                ocr, classifier, content_analyzer,
                doc_gen, ima_syncer, archiver,
                str(img_path), i, total, progress
            )
            if result:
                results.append(result)
        except Exception as e:
            print(f"\n  ✗ 处理失败: {e}")
            logger.error(f"处理失败 {img_path}: {e}")

        if i < total:
            time.sleep(0.3)

    progress.finish()
    
    # 统计
    elapsed = time.time() - start_time

    print("\n" + "=" * 60)
    print("处理完成!")
    print("=" * 60)
    print(f"总计处理: {len(results)}/{total} 张")
    print(f"总耗时: {elapsed:.1f} 秒")

    # 主题分布
    if results:
        print("\n📊 主题分布:")
        themes = Counter(r['theme'] for r in results if not r.get('is_duplicate'))
        for theme, count in themes.most_common():
            print(f"  - {theme}: {count} 张")

    # 重复检测
    dup_count = sum(1 for r in results if r.get('is_duplicate'))
    if dup_count > 0:
        print(f"\n🔍 重复检测: 跳过 {dup_count} 条重复内容")

    # 输出统计
    md_count = len([r for r in results if r.get('md_file')])
    docx_count = len([r for r in results if r.get('docx_file')])
    print(f"\n📁 输出统计:")
    print(f"  - Markdown: {md_count} 个")
    print(f"  - Word: {docx_count} 个")
    
    # 归档统计
    archive_stats = archiver.get_archive_stats()
    print(f"\n📦 归档统计:")
    print(f"  - 本次归档: {archive_stats['archived_count']} 张")
    print(f"  - 已归档总数: {archive_stats['total_archived']} 张")

    # 新建主题
    if dynamic_classifier.custom_themes:
        print(f"\n🆕 新建主题:")
        for theme in dynamic_classifier.custom_themes:
            print(f"  - {theme}")

    if ima_syncer.rate_limited:
        print(f"\n☁️ IMA: 被限流，请明天再试同步")

    print("\n" + "=" * 60)

    # 保存报告
    report_file = Path('处理结果/处理报告.json')
    report = {
        'timestamp': datetime.now().isoformat(),
        'total_images': total,
        'processed': len(results),
        'new_documents': md_count,
        'word_documents': docx_count,
        'duplicates': dup_count,
        'archived': archive_stats['archived_count'],
        'elapsed_seconds': elapsed,
        'new_themes': list(dynamic_classifier.custom_themes.keys()),
        'results': [{k: v for k, v in r.items() if k not in ['md_file', 'docx_file']} for r in results],
        'theme_stats': dict(Counter(r['theme'] for r in results if not r.get('is_duplicate')))
    }
    with open(report_file, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    print(f"\n📊 报告已保存: {report_file}")


if __name__ == '__main__':
    main()
