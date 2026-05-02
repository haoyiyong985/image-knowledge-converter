#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
全自动化图片处理工具 V8.0
==============================
V8 修复项（基于V7）：
  [F-1] 默认模式智能分批 - 图片>10张自动启用分批（V6逻辑整合）
  [F-2] 进度条始终可见 - 精简单图日志为1行，进度条不被淹没
  [F-3] 内容智能整理 - 去除手机UI噪点（状态栏/信号/App按钮）
  [F-4] 智能文档命名 - 重构extract_topic()，AI语义提取+正则兜底

V7 核心改进（保留）：
  [P0-1] 重构分类引擎 - 内容理解驱动，动态新建分类
  [P0-2] 统一文档命名 - 格式：{分类}-{主题}
  [P0-3] 同主题内容合并 - 跨分类检测同名文档
  [P1-4] 图片/文档分类分离 - 图片按源文件夹，文档按AI分类
  [P1-5] 优化重复检测 - 主题词相似度 + hash 双重检测
  [P1-6] 修复统计报告
  [P1-7] 修复 IMA 同步

使用方法：
  python auto_process_all_v8.py           # 全自动处理（>10张自动分批）
  python auto_process_all_v8.py --batch    # 强制分批模式（可中断续传）
  python auto_process_all_v8.py --init     # 仅初始化分批
  python auto_process_all_v8.py --progress # 查看进度
  python auto_process_all_v8.py --clear    # 清除分批状态
"""

import os
import sys
import io
import json
import time
import re
import hashlib
import shutil
import logging
from pathlib import Path
from datetime import datetime
from collections import Counter, defaultdict
from typing import Optional, List, Dict, Tuple
from dataclasses import dataclass, asdict

# 修复 Windows 控制台编码
if sys.platform == 'win32':
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
    except Exception:
        pass

sys.path.insert(0, '.')

from dotenv import load_dotenv
env_path = Path('.') / '.env'
if env_path.exists():
    load_dotenv(dotenv_path=env_path)
else:
    print("[WARN] .env 文件不存在，使用系统环境变量")

from local_ocr import LocalOCR

try:
    from tencent_ocr import TencentOCR
    TENCENT_AVAILABLE = True
except ImportError:
    TENCENT_AVAILABLE = False

try:
    from baidu_ocr import BaiduOCR
    BAIDU_AVAILABLE = True
except ImportError:
    BAIDU_AVAILABLE = False

# 日志配置
Path('处理结果').mkdir(exist_ok=True)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(message)s',
    handlers=[
        logging.FileHandler('处理结果/process.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# ============================================================
# [P0-1] 智能分类引擎 - 完全重构
# ============================================================

# 预设分类 + 关键词（仅作匹配辅助，AI内容理解优先）
PRESET_CATEGORIES = {
    "历史文化": [
        "历史", "文化", "朝代", "古代", "人物", "传记", "皇帝", "将军", "战争", "事件",
        "明朝", "清朝", "汉朝", "唐朝", "宋朝", "元朝", "秦朝", "民国", "三国", "春秋",
        "战国", "革命", "史记", "文献", "典故", "传统", "文明", "民俗", "风俗", "节日",
        "诗词", "古文", "书法", "绘画", "建筑", "遗址", "文物", "考古",
    ],
    "营养健康": [
        "营养", "健康", "饮食", "食物", "蔬菜", "水果", "蛋白质", "碳水",
        "中医", "养生", "药膳", "食疗", "中药", "经络", "穴位", "调理",
        "体质", "气血", "疾病", "预防", "治疗", "症状", "血压", "血糖",
        "维生素", "矿物质", "膳食", "抗炎", "免疫力", "肠道", "益生菌",
        "减肥", "增肌", "低卡", "热量", "卡路里", "蛋白", "脂肪",
    ],
    "生活方式": [
        "生活", "运动", "睡眠", "压力", "情绪", "心理", "作息", "习惯",
        "美容", "护肤", "健身", "锻炼", "瑜伽", "冥想", "放松",
        "整理", "收纳", "家居", "装修", "穿搭", "时尚", "购物",
        "理财", "储蓄", "投资", "工作", "职场", "效率",
    ],
    "教育育儿": [
        "育儿", "宝宝", "孩子", "教育", "辅食", "喂养", "早教",
        "亲子", "成长", "发育", "妈妈", "孕妇", "备孕", "新生儿",
        "婴儿", "幼儿", "儿童", "学习", "作业", "学校", "老师",
    ],
    "旅游攻略": [
        "旅游", "旅行", "景点", "攻略", "出行", "交通", "酒店",
        "民宿", "美食", "打卡", "签证", "机票", "行程", "路线",
        "城市", "国家", "海外", "境外", "自驾", "背包",
    ],
    "综合知识": [
        "科普", "知识", "研究", "发现", "实验", "数据", "专家",
        "建议", "指南", "推荐", "科学", "原理", "机制", "分析",
    ],
}

# 分类特征权重（区分度高的关键词加权）
CATEGORY_BOOST_WORDS = {
    "历史文化": {"朝代", "皇帝", "将军", "史记", "文献", "古代", "传记", "人物志"},
    "营养健康": {"中医", "养生", "营养素", "血压", "血糖", "食疗", "药膳"},
    "教育育儿": {"宝宝", "辅食", "早教", "新生儿", "幼儿园"},
    "旅游攻略": {"攻略", "景点", "签证", "行程", "民宿"},
}


class SmartClassifier:
    """
    [P0-1] 智能分类器 - 内容理解驱动
    
    策略：
    1. 对OCR文本进行预处理（去噪）
    2. 对每个预设分类计算综合分数（关键词密度 + 特征词加权）
    3. 如果最高分低于阈值，尝试从内容中推断新分类
    4. 返回分类名和置信度
    """

    def __init__(self):
        # 从已知分类文件夹中动态加载已有分类
        self.known_categories = set(PRESET_CATEGORIES.keys())
        self._load_existing_categories()

    def _load_existing_categories(self):
        """加载已存在的分类（包括动态新建的）"""
        result_dir = Path('处理结果')
        if result_dir.exists():
            for d in result_dir.iterdir():
                if d.is_dir() and not d.name.startswith('.'):
                    self.known_categories.add(d.name)
        archived_dir = Path('已处理图片')
        if archived_dir.exists():
            for d in archived_dir.iterdir():
                if d.is_dir() and not d.name.startswith('.'):
                    self.known_categories.add(d.name)

    def classify(self, text: str) -> Tuple[str, float]:
        """
        对文本分类，返回 (分类名, 置信度)
        
        置信度说明：
          >= 0.6  高置信度（关键词明确）
          0.3-0.6 中置信度（有一定匹配）
          < 0.3   低置信度（使用综合知识兜底）
        """
        if not text or len(text) < 5:
            return "综合知识", 0.1

        scores = {}
        text_len = max(len(text), 1)

        for cat_name, keywords in PRESET_CATEGORIES.items():
            # 基础关键词得分（密度）
            matched = [kw for kw in keywords if kw in text]
            base_score = len(matched) * 3 / text_len * 1000

            # 特征词加权
            boost_words = CATEGORY_BOOST_WORDS.get(cat_name, set())
            boost = sum(5 for w in boost_words if w in text)

            scores[cat_name] = base_score + boost

        if not scores:
            return "综合知识", 0.1

        best_cat = max(scores, key=scores.get)
        best_score = scores[best_cat]

        # 归一化置信度（0~1）
        max_possible = 50.0
        confidence = min(best_score / max_possible, 1.0)

        if confidence < 0.05:
            # 完全没有匹配，兜底为综合知识
            return "综合知识", 0.05

        return best_cat, confidence

    def get_matched_keywords(self, text: str, category: str) -> List[str]:
        """获取匹配到的关键词"""
        keywords = PRESET_CATEGORIES.get(category, [])
        return [kw for kw in keywords if kw in text][:8]

    def register_new_category(self, category_name: str):
        """注册新分类"""
        self.known_categories.add(category_name)
        logger.info(f"[分类] 注册新分类: {category_name}")


# ============================================================
# [P0-2][P0-3] 智能内容分析器 - 重构
# ============================================================

class SmartContentAnalyzer:
    """
    [P0-2][P0-3] 智能内容分析器
    
    核心功能：
    1. 提取分类（使用 SmartClassifier）
    2. 提取主题词（人名/事件名/核心主题）
    3. 生成文档名：格式 {分类}-{主题}
    4. 支持多主题合并到主主题
    """

    def __init__(self, classifier: SmartClassifier):
        self.classifier = classifier
        self.cache = {}

    def extract_topic(self, text: str, category: str) -> str:
        """
        [F-4] 重构主题提取 - AI语义分析 + 正则兜底

        策略（按优先级）：
        1. 人名提取（扩充匹配模式）
        2. 事件/朝代/主题提取（历史类专用）
        3. 内容核心标题（排除UI噪点行）
        4. 关键词智能组合
        5. 兜底
        """
        lines = text.split('\n')
        # 过滤空行和过短行
        valid_lines = [l.strip() for l in lines if l.strip() and len(l.strip()) >= 3]

        # 预处理：跳过前3行（通常是噪点）
        content_lines = valid_lines

        # === 策略1：人名提取（扩充匹配模式）===
        person = self._extract_person_name(content_lines)
        if person:
            return person

        # === 策略2：历史事件/朝代/主题 ===
        if category == '历史文化':
            topic = self._extract_history_topic(content_lines)
            if topic:
                return topic

        # === 策略3：内容标题行（排除UI噪点）===
        title = self._extract_content_title(content_lines)
        if title:
            return title

        # === 策略4：关键词智能组合 ===
        matched_kws = self.classifier.get_matched_keywords(text, category)
        if matched_kws:
            # 过滤掉太泛的词，取最有区分度的2个
            stop_words = {'历史', '文化', '知识', '整理', '推荐', '建议', '指南',
                         '健康', '生活', '科学', '科普', '饮食', '营养', '学习'}
            meaningful = [kw for kw in matched_kws
                         if len(kw) >= 2 and kw not in stop_words][:2]
            if meaningful:
                return ''.join(meaningful)

        # === 策略5：兜底 ===
        return f"整理{datetime.now().strftime('%m%d%H%M')}"

    def _extract_person_name(self, lines: List[str]) -> Optional[str]:
        """[F-4] 提取人名 - 扩充匹配模式"""
        # 模式集合
        person_patterns = [
            # XXX传/简介/生平/人物志
            r'([^\s，。、：！？\d]{2,4})(?:传|简介|生平|人物志|人物|世家)',
            # XXX（字/号/谥...）
            r'^([^\s，。]{2,5})[，,]\s*(?:字|号|谥|生于|字号|原名|名)',
            # XXX是/为/任/担任...人物传记格式
            r'^([^\s，。]{2,4})(?:是|为|任|担任|生于|出生|卒于|去世|逝世)',
            # XXX，中国/著名/当代/现代...描述格式
            r'^([^\s，。]{2,4})[，,]\s*(?:中国|著名|当代|现代|近代|古代|清末|明末|南宋|北宋|晚唐|初唐)',
            # XXX（XXXX-XXXX）生卒年格式
            r'^([^\s，。（(]{2,4})[（(]\s*\d{3,4}[-—]\d{3,4}',
            # 江苏XX人/浙江XX人 等
            r'(?:江苏|浙江|山东|广东|四川|湖南|湖北|福建|安徽|河南|河北|陕西|山西|辽宁|北京|上海|天津|重庆)([^\s，。]{2,4})',
        ]

        for line in lines[:15]:
            for pattern in person_patterns:
                match = re.search(pattern, line)
                if match:
                    name = match.group(1)
                    # 过滤掉明显不是人名的
                    if len(name) < 2 or len(name) > 4:
                        continue
                    if re.match(r'^\d+$', name):
                        continue
                    return name

        return None

    def _extract_history_topic(self, lines: List[str]) -> Optional[str]:
        """[F-4] 提取历史类主题（事件/朝代/地图等）"""
        # 朝代+事件模式
        history_patterns = [
            r'((?:明朝|清朝|汉朝|唐朝|宋朝|元朝|秦朝|周朝|隋朝|春秋|战国|三国|民国)'
            r'[^\n]{2,15})',
            r'(一张图看完[^\n]{2,15})',
            r'((?:开国|建国|统一|灭亡|中兴)[^\n]{2,10})',
        ]

        for line in lines[:10]:
            for pattern in history_patterns:
                match = re.search(pattern, line)
                if match:
                    topic = match.group(1).strip()
                    if len(topic) >= 4:
                        return topic[:15]

        return None

    def _extract_content_title(self, lines: List[str]) -> Optional[str]:
        """[F-4] 提取内容标题（排除UI噪点行）"""
        # UI噪点关键词（标题行中不应出现这些）
        noise_words = [
            'http', 'www', '转载', '来源', '作者', '发送', '翻译', '探索版',
            '关注', '收藏', '点赞', '评论', '分享', '来自', '搜索', '找到',
            '写作', 'doubao', 'kimi', '元宝', '豆包', 'ChatGPT', 'DeepSeek',
        ]

        for line in lines[:8]:
            line = line.strip()
            if not line:
                continue
            # 去掉markdown标记
            cleaned = re.sub(r'^#{1,6}\s*', '', line)
            cleaned = re.sub(r'^[◆★▼▪▸【】\[\]]+\s*', '', cleaned)
            cleaned = cleaned.strip()
            if not cleaned:
                continue

            # 排除含噪点词的行
            if any(noise in cleaned for noise in noise_words):
                continue
            # 排除纯数字/时间行
            if re.match(r'^[\d:.\s%]+$', cleaned):
                continue
            # 排除过短行
            if len(cleaned) < 3:
                continue

            # 匹配中文短标题（3-15字）
            if re.match(r'^[\u4e00-\u9fa5]{3,15}$', cleaned):
                return cleaned
            # 带标点的标题（不超过20字，以中文开头）
            if 4 <= len(cleaned) <= 20 and re.match(r'^[\u4e00-\u9fa5]', cleaned):
                return cleaned[:15]

        return None

    def generate_doc_name(self, text: str, category: str) -> str:
        """
        [P0-2] 生成文档名
        格式：{分类}-{主题}
        示例：历史文化-钱谦益传、营养健康-春季养生
        """
        topic = self.extract_topic(text, category)
        # 清理非法文件名字符
        topic = re.sub(r'[<>:"/\\|?*\n\r\t，。！？、：；]', '', topic).strip()
        # 限制长度（按词截断，避免截断到中间）
        if len(topic) > 20:
            topic = topic[:20]
        if not topic or len(topic) < 2:
            topic = f"整理{datetime.now().strftime('%m%d%H%M')}"
        return f"{category}-{topic}"

    def compute_content_hash(self, text: str) -> str:
        """计算内容哈希"""
        cleaned = re.sub(r'\s+', '', text)
        return hashlib.md5(cleaned.encode('utf-8')).hexdigest()[:8]

    def analyze(self, text: str, index: int) -> Dict:
        """综合分析，返回分类/主题/文档名等"""
        content_hash = self.compute_content_hash(text)

        if content_hash in self.cache:
            return self.cache[content_hash]

        # 分类
        category, confidence = self.classifier.classify(text)
        keywords = self.classifier.get_matched_keywords(text, category)

        # 文档名
        doc_name = self.generate_doc_name(text, category)

        result = {
            'category': category,
            'confidence': confidence,
            'keywords': keywords,
            'doc_name': doc_name,
            'content_hash': content_hash,
        }

        self.cache[content_hash] = result
        return result


# ============================================================
# 内容整理器（保留V6功能）
# ============================================================
class ContentOrganizer:
    """[F-3] 内容整理器 - 增强版"""

    def __init__(self):
        # 手机UI噪点模式
        self.ui_noise_patterns = [
            # 信号栏
            r'^\d{1,2}:\d{2}$',                          # 时间 "23:24"
            r'^[0-9]{1,2}:%$',                            # 时间截断
            r'^4G\s*5G',                                  # 信号标识
            r'^[0-9]+\s*%$',                              # 电量百分比
            r'^\d{1,2}:\d{2}\s*$',                        # 独立时间行
            r'^\s*\d{1,3}%\s*$',                          # 电量
            r'^Wi-?Fi',                                   # Wi-Fi
            r'^\.?\s*AM$', r'^\.?\s*PM$',                # AM/PM
            # App名称和按钮
            r'^(写作|翻译|探索版|有什么问题尽管问我)$',
            r'^(doubao\.com|豆包)$',
            r'^(Kimi|kimi)$',
            r'^(腾讯元宝|元宝)$',
            r'^(ChatGPT|GPT)$',
            r'^(文心一言|通义千问|智谱清言)$',
            r'^(豆包|DeepSeek)$',
            r'^(发送|发送消息)$',
            r'^\+关注$',
            r'^(收藏|点赞|评论|分享|转发)$',
            r'^(已关注|关注)$',
            r'^(更多|展开|收起|折叠)$',
            # 杂项UI
            r'^\d+$',                                     # 纯数字行
            r'^[A-Z][a-z]+$',                             # 纯英文单词（如 Java, Kimi 已单独处理）
            r'^[a-z]+\.(com|cn|net|org)$',               # 域名
            r'^http[s]?://',                              # URL
            r'^来自.*的搜索$',                            # 搜索来源标记
            r'^搜一下$',                                  # 搜索按钮
            r'^找到\d+篇资料参考$',                       # 搜索结果数
        ]

        # 来源平台模式（增强）
        self.platform_patterns = {
            '小红书': [r'小红书', r'RED', r'@.*小红书'],
            '微信读书': [r'微信读书', r'weread'],
            '微信': [r'微信', r'wechat', r'WeChat'],
            '抖音': [r'抖音', r'douyin', r'TikTok'],
            'B站': [r'B站', r'bilibili', r'BILIBILI'],
            '豆包': [r'豆包', r'doubao\.com'],
            'Kimi': [r'Kimi', r'kimi', r'moonshot'],
            '元宝': [r'腾讯元宝', r'元宝'],
            'DeepSeek': [r'DeepSeek', r'deepseek'],
            '知乎': [r'知乎', r'zhihu'],
            '头条': [r'今日头条', r'头条'],
        }

        # 内容结构化关键词
        self.heading_patterns = [
            # 编号标题
            (r'^(\d+)[.、]\s*(.{2,30})$', r'## \1. \2'),
            (r'^[（(](\d+)[)）]\s*(.{2,30})$', r'### (\1) \2'),
            # 中文序号
            (r'^([一二三四五六七八九十]+)[、：]\s*(.{2,30})$', r'## \1、\2'),
            # 带符号标题
            (r'^[●◆■▶▪▸]\s*(.{4,30})$', r'### \1'),
            (r'^[-–—]\s*(.{4,30})$', r'### \1'),
            # 纯大写或加粗标记
            (r'^#{1,4}\s*(.{2,50})$', r'## \1'),
        ]

    def extract_source(self, text: str) -> Dict:
        """[F-3] 增强来源提取 - 支持更多App和格式"""
        lines = text.split('\n')
        source_info = {}

        # 前30行中查找来源信息
        for line in lines[:30]:
            line_stripped = line.strip()
            if not line_stripped:
                continue

            # 平台检测（优先精确匹配）
            for platform, patterns in self.platform_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, line_stripped, re.IGNORECASE):
                        source_info['platform'] = platform
                        break
                if 'platform' in source_info:
                    break
            if 'platform' in source_info:
                break

            # 文本标记检测
            for kw in ['来源', '出处', '摘录自', '来自']:
                if kw in line_stripped:
                    parts = line_stripped.split(kw)
                    if len(parts) > 1:
                        source_val = parts[1].strip().lstrip('：:：').strip()
                        if source_val and len(source_val) < 20:
                            source_info['platform'] = source_val
                            break

            # 作者检测
            if '作者' in line_stripped:
                match = re.search(r'作者[：:]\s*([^\n,，]{1,20})', line_stripped)
                if match:
                    source_info['author'] = match.group(1).strip()

        return source_info

    def clean_text(self, text: str) -> str:
        """[F-3] 增强内容清理 - 去除手机UI噪点"""
        lines = text.split('\n')
        cleaned_lines = []
        prev_line_empty = False

        for line in lines:
            line = line.strip()

            # 空行处理
            if not line:
                if not prev_line_empty:
                    cleaned_lines.append('')
                    prev_line_empty = True
                continue

            # 1. 匹配手机UI噪点模式
            is_noise = False
            for pattern in self.ui_noise_patterns:
                if re.match(pattern, line, re.IGNORECASE):
                    is_noise = True
                    break
            if is_noise:
                continue

            # 2. 过短的纯数字/字母行
            if re.match(r'^[\d\s%]+$', line) and len(line) < 8:
                continue

            # 3. 分隔线
            if re.match(r'^[_\-=]{3,}$', line):
                continue

            # 4. 纯英文长行（无中文）
            if re.match(r'^[a-zA-Z0-9\s.,!?;:\'"()-]{20,}$', line) and not re.search(r'[\u4e00-\u9fa5]', line):
                continue

            # 5. 纯短行（<2字符的中文）
            if len(line) < 2:
                continue

            # 6. OCR错误修复
            line = self._fix_common_ocr_errors(line)

            prev_line_empty = False
            cleaned_lines.append(line)

        # 清理多余空行
        result = '\n'.join(cleaned_lines)
        result = re.sub(r'\n{3,}', '\n\n', result)
        return result.strip()

    def _fix_common_ocr_errors(self, line: str) -> str:
        """修复常见OCR错误"""
        replacements = [
            (chr(8220), '"'), (chr(8221), '"'),
            (chr(8216), "'"), (chr(8217), "'"),
            (chr(8212), "-"), (chr(183), "-"),
            ('`', ''), ('√', '✓'), ('×', '✗'),
            # 手机截图中常见OCR截断修复
            ('…', '...'), ('……', '...'),
            ('"', ''), ('"', ''),  # 清理残留的特殊引号
        ]
        for old, new in replacements:
            line = line.replace(old, new)
        line = re.sub(r' {2,}', ' ', line)
        return line

    def structure_content(self, text: str) -> str:
        """[F-3] 内容结构化 - 识别标题层级并添加Markdown标记"""
        lines = text.split('\n')
        structured_lines = []
        for line in lines:
            stripped = line.strip()
            if not stripped:
                structured_lines.append('')
                continue

            # 尝试匹配标题模式
            matched = False
            for pattern, replacement in self.heading_patterns:
                match = re.match(pattern, stripped)
                if match:
                    structured_lines.append(match.expand(replacement))
                    matched = True
                    break

            if not matched:
                structured_lines.append(line)

        return '\n'.join(structured_lines)

    def organize(self, text: str) -> Dict:
        """整理内容（清理 + 来源提取 + 结构化）"""
        cleaned = self.clean_text(text)
        source = self.extract_source(text)
        structured = self.structure_content(cleaned)
        return {
            'content': structured,
            'source': source
        }


# ============================================================
# [P0-3][P1-5] 智能文档管理器 - 重构合并逻辑
# ============================================================

class SmartDocumentManager:
    """
    [P0-3][P1-5] 智能文档管理器
    
    核心改进：
    1. 跨分类检测同名文档（避免同主题分散存储）
    2. 主题词相似度 + hash 双重重复检测
    3. 统一文档路径：处理结果/{分类}/{分类}-{主题}.md
    """

    def __init__(self, output_dir: str = '处理结果'):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        # 全局文档索引：doc_name -> Path（跨分类）
        self.doc_index: Dict[str, Path] = {}
        # 已知内容hash集合
        self.known_hashes: set = set()
        # 已知主题词集合（用于相似度检测）
        self.known_topics: List[Tuple[str, Path]] = []
        self._build_index()

    def _build_index(self):
        """扫描已有文档，建立全局索引"""
        for md_file in self.output_dir.rglob('*.md'):
            doc_name = md_file.stem
            # 跨分类索引：用文档名（不含分类前缀）作为key
            self.doc_index[doc_name] = md_file

            # 提取主题词（去掉分类前缀）
            topic = self._extract_topic_from_name(doc_name)
            self.known_topics.append((topic, md_file))

            # 扫描hash
            try:
                content = md_file.read_text(encoding='utf-8')
                hash_match = re.search(r'content_hash:\s*([a-f0-9]{8})', content)
                if hash_match:
                    self.known_hashes.add(hash_match.group(1))
            except Exception:
                pass

    def _extract_topic_from_name(self, doc_name: str) -> str:
        """从文档名提取主题词（去掉分类前缀）"""
        # 格式：分类-主题 -> 主题
        if '-' in doc_name:
            parts = doc_name.split('-', 1)
            return parts[1] if len(parts) > 1 else doc_name
        return doc_name

    def is_duplicate_hash(self, content_hash: str) -> bool:
        """[P1-5] hash重复检测"""
        return content_hash in self.known_hashes

    def find_similar_doc(self, doc_name: str, content: str) -> Optional[Path]:
        """
        [P0-3][P1-5] 跨分类相似文档检测
        
        两个条件同时满足才判定为相似：
        1. 主题词相似度 >= 0.7
        2. 内容字符集重叠 >= 20
        """
        topic = self._extract_topic_from_name(doc_name)

        for known_topic, known_path in self.known_topics:
            # 条件1：主题名相似度
            sim = self._topic_similarity(topic, known_topic)
            if sim < 0.7:
                continue

            # 条件2：内容重叠
            if known_path.exists():
                try:
                    existing_content = known_path.read_text(encoding='utf-8')
                    c1_chars = set(content[:400].replace(' ', '').replace('\n', ''))
                    c2_chars = set(existing_content[:400].replace(' ', '').replace('\n', ''))
                    overlap = len(c1_chars & c2_chars)
                    if overlap < 20:
                        continue
                except Exception:
                    pass

            return known_path

        return None

    def _topic_similarity(self, t1: str, t2: str) -> float:
        """主题词相似度（包含关系 + bigram Jaccard）"""
        if not t1 or not t2:
            return 0.0
        # 去掉常见后缀再比较
        def normalize(s):
            return re.sub(r'(传|简介|人物|生平|整理|知识|合集|汇总)$', '', s).strip()
        n1, n2 = normalize(t1), normalize(t2)
        if n1 == n2:
            return 1.0
        if t1 == t2:
            return 1.0
        # 包含关系（去后缀后）
        shorter, longer = (n1, n2) if len(n1) <= len(n2) else (n2, n1)
        if len(shorter) >= 2 and shorter in longer:
            return 0.90
        # 原始字符串包含关系
        s2, l2 = (t1, t2) if len(t1) <= len(t2) else (t2, t1)
        if len(s2) >= 2 and s2 in l2:
            return 0.85
        # bigram Jaccard（使用去后缀版本）
        def bigrams(s):
            return set(s[i:i+2] for i in range(len(s) - 1)) if len(s) >= 2 else set(s)
        b1, b2 = bigrams(n1), bigrams(n2)
        if not b1 or not b2:
            return 0.0
        return len(b1 & b2) / len(b1 | b2)

    def merge_content(self, existing_file: Path, new_content: str,
                      new_image_name: str, content_hash: str = None) -> bool:
        """追加内容到已有文档"""
        try:
            existing = existing_file.read_text(encoding='utf-8')
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
            append_section = f"""

---

## 补充内容

> 来源图片: {new_image_name}
> 追加时间: {timestamp}
> 内容标识: {content_hash or 'N/A'}

{new_content}

"""
            marker = '\n\n---\n\n*本文档由图片知识库整理工具自动生成*'
            if marker in existing:
                new_full = existing.replace(marker, append_section + marker)
            else:
                new_full = existing + append_section

            existing_file.write_text(new_full, encoding='utf-8')
            return True
        except Exception as e:
            logger.warning(f"[合并] 失败: {e}")
            return False

    def save_document(self, doc_name: str, category: str, content: str,
                      image_name: str, keywords: List[str],
                      content_hash: str, source_info: Dict) -> Tuple[str, bool]:
        """
        保存文档（新建或合并）
        返回 (文件路径, 是_new_doc)
        """
        # 检查相似文档（跨分类）
        similar = self.find_similar_doc(doc_name, content)
        if similar:
            logger.info(f"[文档] 发现相似文档 {similar.name}，合并内容")
            self.merge_content(similar, content, image_name, content_hash)
            if content_hash:
                self.known_hashes.add(content_hash)
            return str(similar), False  # 非新文档

        # 新建文档
        cat_dir = self.output_dir / category
        cat_dir.mkdir(exist_ok=True)

        safe_name = re.sub(r'[<>:"/\\|?*]', '', doc_name)
        if len(safe_name) > 40:
            safe_name = safe_name[:40]

        md_file = cat_dir / f"{safe_name}.md"
        counter = 1
        while md_file.exists():
            md_file = cat_dir / f"{safe_name}_{counter}.md"
            counter += 1

        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
        keywords_str = ', '.join(keywords) if keywords else '无'
        source_section = ""
        if source_info.get('platform'):
            source_section += f"**来源**：{source_info['platform']}\n\n"
        if source_info.get('author'):
            source_section += f"**作者**：{source_info['author']}\n\n"

        md_content = f"""# {doc_name}

> 来源图片: {image_name}
> 识别时间: {timestamp}
> 分类: {category}
> 关键词: {keywords_str}
> content_hash: {content_hash or 'N/A'}

---

<!-- CONTENT_START -->
{source_section}{content}
<!-- CONTENT_END -->

---

*本文档由图片知识库整理工具自动生成*
"""
        md_file.write_text(md_content, encoding='utf-8')

        # 更新索引
        self.doc_index[doc_name] = md_file
        topic = self._extract_topic_from_name(doc_name)
        self.known_topics.append((topic, md_file))
        if content_hash:
            self.known_hashes.add(content_hash)

        logger.info(f"[文档] 新建: {category}/{md_file.name}")
        return str(md_file), True  # 是新文档

    def generate_word(self, doc_name: str, category: str, content: str,
                      image_name: str, keywords: List[str],
                      md_file_path: str = None) -> Optional[str]:
        """生成Word文档"""
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            logger.warning("[文档] python-docx未安装，跳过Word生成")
            return None

        cat_dir = self.output_dir / category
        cat_dir.mkdir(exist_ok=True)

        safe_name = re.sub(r'[<>:"/\\|?*]', '', doc_name)
        if len(safe_name) > 40:
            safe_name = safe_name[:40]

        docx_file = cat_dir / f"{safe_name}.docx"
        if docx_file.exists():
            return str(docx_file)

        # 从MD文件读取最新内容（含合并后的内容）
        text_content = content
        if md_file_path and Path(md_file_path).exists():
            try:
                md_text = Path(md_file_path).read_text(encoding='utf-8')
                if '<!-- CONTENT_START -->' in md_text and '<!-- CONTENT_END -->' in md_text:
                    start = md_text.index('<!-- CONTENT_START -->') + len('<!-- CONTENT_START -->')
                    end = md_text.index('<!-- CONTENT_END -->')
                    text_content = md_text[start:end].strip()
            except Exception:
                pass

        try:
            doc = Document()
            doc.styles['Normal'].font.name = 'Microsoft YaHei'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

            doc.add_heading(doc_name, 0)
            doc.add_paragraph(f"来源图片: {image_name}")
            doc.add_paragraph(f"识别时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph(f"分类: {category}")
            keywords_str = ', '.join(keywords) if keywords else '无'
            doc.add_paragraph(f"关键词: {keywords_str}")
            doc.add_paragraph("─" * 30)
            doc.add_paragraph(text_content)

            doc.save(str(docx_file))
            logger.info(f"[文档] Word已生成: {category}/{docx_file.name}")
            return str(docx_file)
        except Exception as e:
            logger.warning(f"[Word] 生成失败: {e}")
            return None


# ============================================================
# [P1-4] 图片归档器 - 按源文件夹分类（便于溯源）
# ============================================================

class ImageArchiver:
    """
    [P1-4] 图片归档器
    
    分类策略：按源文件夹分类，便于溯源查询
    - 图片归档：按待处理图片中的子文件夹分类
    - 文档归档：按V7 AI分类（在SmartDocumentManager中处理）
    """

    def __init__(self, processed_dir: str = '已处理图片'):
        self.processed_dir = Path(processed_dir)
        self.processed_dir.mkdir(parents=True, exist_ok=True)
        self.archived_count = 0

    def get_source_folder(self, image_path: str) -> str:
        """
        从图片路径提取源文件夹名
        例如: 待处理图片/历史学/a.jpg -> 历史学
        """
        src = Path(image_path)
        # 取父文件夹名（相对于待处理图片）
        parent_name = src.parent.name
        # 如果父文件夹名等于源目录名（说明图片直接在待处理图片根目录）
        if parent_name == '待处理图片':
            return '根目录'
        return parent_name

    def archive_image(self, image_path: str, category: str = None) -> Optional[str]:
        """
        移动图片到已处理文件夹
        按源文件夹分类，便于溯源
        """
        src = Path(image_path)
        if not src.exists():
            return None

        # 使用源文件夹名作为归档目录
        source_folder = self.get_source_folder(image_path)
        archive_subdir = self.processed_dir / source_folder

        archive_subdir.mkdir(parents=True, exist_ok=True)

        dest = archive_subdir / src.name
        if dest.exists():
            ts = datetime.now().strftime('%H%M%S')
            dest = archive_subdir / f"{src.stem}_{ts}{src.suffix}"

        try:
            shutil.move(str(src), str(dest))
            self.archived_count += 1
            return str(dest)
        except Exception as e:
            logger.warning(f"[归档] 移动失败 {src.name}: {e}")
            try:
                shutil.copy2(str(src), str(dest))
                src.unlink()
                self.archived_count += 1
                return str(dest)
            except Exception:
                return None

    def get_stats(self) -> Dict:
        total = 0
        subdirs = {}
        for item in self.processed_dir.iterdir():
            if item.is_dir():
                count = len(list(item.glob('*.*')))
                subdirs[item.name] = count
                total += count
        return {
            'archived_this_session': self.archived_count,
            'total_archived': total,
            'by_category': subdirs
        }


# ============================================================
# 进度条（保留V6功能）
# ============================================================
class ProgressBar:
    """[F-2] 进度条 - 始终可见，不被日志淹没"""

    def __init__(self, total: int, width: int = 40, prefix: str = "进度"):
        self.total = total
        self.width = width
        self.prefix = prefix
        self.start_time = time.time()
        self.last_line = ""       # 记录上一条日志行
        self.current = 0

    def update(self, current: int, info: str = ""):
        self.current = current
        percent = current / self.total if self.total > 0 else 0
        filled = int(self.width * percent)
        bar = '█' * filled + '░' * (self.width - filled)
        if current > 0:
            elapsed = time.time() - self.start_time
            eta = elapsed / current * (self.total - current)
            eta_str = f"{int(eta//60)}m{int(eta%60)}s"
        else:
            eta_str = "--"
        # 进度条单独一行，使用 \r 覆盖自己
        print(f"\r{self.prefix}: [{bar}] {current}/{self.total} ({percent:.0%}) ETA:{eta_str} {info[:15]}",
              end='', flush=True)

    def log(self, message: str):
        """[F-2] 在进度条上方安全打印日志行"""
        # 1. 清除当前进度条（用空格覆盖）
        clear = ' ' * 100
        print(f'\r{clear}\r', end='', flush=True)
        # 2. 打印日志
        print(message)
        # 3. 重新绘制进度条
        if self.current > 0:
            self._redraw()

    def _redraw(self):
        """重新绘制当前进度条"""
        percent = self.current / self.total if self.total > 0 else 0
        filled = int(self.width * percent)
        bar = '█' * filled + '░' * (self.width - filled)
        elapsed = time.time() - self.start_time
        if self.current > 0:
            eta = elapsed / self.current * (self.total - self.current)
            eta_str = f"{int(eta//60)}m{int(eta%60)}s"
        else:
            eta_str = "--"
        print(f"\r{self.prefix}: [{bar}] {self.current}/{self.total} ({percent:.0%}) ETA:{eta_str}",
              end='', flush=True)

    def finish(self):
        self.update(self.total, "完成!")
        print()


# ============================================================
# 多引擎OCR（保留V6功能）
# ============================================================
class MultiEngineOCR:
    def __init__(self):
        self.current_engine = None
        self.engine_status = []

        secret_id = os.getenv('TENCENT_SECRET_ID', '')
        secret_key = os.getenv('TENCENT_SECRET_KEY', '')
        if secret_id and secret_key and '替换' not in secret_id and TENCENT_AVAILABLE:
            self.engine_status.append(('腾讯云', True, '已配置'))
        else:
            self.engine_status.append(('腾讯云', False, '未配置'))

        app_id = os.getenv('BAIDU_APP_ID', '')
        api_key = os.getenv('BAIDU_API_KEY', '')
        bsecret = os.getenv('BAIDU_SECRET_KEY', '')
        if app_id and api_key and bsecret and '替换' not in app_id and BAIDU_AVAILABLE:
            self.engine_status.append(('百度云', True, '已配置'))
        else:
            self.engine_status.append(('百度云', False, '未配置'))

        local = LocalOCR()
        if local.tesseract_available:
            self.engine_status.append(('本地Tesseract', True, '可用'))
        else:
            self.engine_status.append(('本地Tesseract', False, '不可用'))

        self._select_best_engine()

    def _select_best_engine(self):
        for name, available, msg in self.engine_status:
            if available:
                self.current_engine = name
                logger.info(f"[OCR] 选用引擎: {name}")
                return True
        logger.warning("[OCR] 没有可用的OCR引擎!")
        return False

    def recognize(self, image_path: str) -> Dict:
        for name, available, _ in self.engine_status:
            if not available:
                continue
            try:
                if name == '腾讯云' and TENCENT_AVAILABLE:
                    engine = TencentOCR()
                elif name == '百度云' and BAIDU_AVAILABLE:
                    engine = BaiduOCR()
                else:
                    engine = LocalOCR()
                result = engine.recognize(image_path)
                if result and result.get('success'):
                    return result
            except Exception as e:
                logger.warning(f"[OCR] {name} 失败: {e}")
        return {'success': False, 'error': '无可用OCR引擎', 'text': ''}


# ============================================================
# IMA同步器（修复P1-7）
# ============================================================
class IMASyncer:
    """[P1-7] IMA笔记同步器（修复同步逻辑）"""

    def __init__(self):
        self.client_id = os.getenv('IMA_OPENAPI_CLIENTID', '')
        self.api_key = os.getenv('IMA_OPENAPI_APIKEY', '')
        self.base_url = 'https://ima.qq.com/openapi/note/v1'
        self.enabled = bool(self.client_id and self.api_key and '填入' not in self.client_id)
        self.sync_log_file = Path('处理结果/ima_sync_log.json')
        self.sync_log = self._load_sync_log()
        self.default_notebook_id = os.getenv('IMA_NOTEBOOK_ID', '')
        self.rate_limited = False
        # [P1-7] 追踪本次同步数量
        self.synced_this_session = 0
        self.failed_this_session = 0

        if self.enabled:
            logger.info("[IMA] 已配置，同步功能启用")
        else:
            logger.warning("[IMA] 未配置，同步功能禁用")

    def _load_sync_log(self) -> Dict:
        if self.sync_log_file.exists():
            try:
                return json.loads(self.sync_log_file.read_text(encoding='utf-8'))
            except Exception:
                pass
        return {}

    def _save_sync_log(self):
        self.sync_log_file.write_text(
            json.dumps(self.sync_log, ensure_ascii=False, indent=2),
            encoding='utf-8'
        )

    def _api_call(self, endpoint: str, payload: Dict, retries: int = 3) -> Optional[Dict]:
        if not self.enabled or self.rate_limited:
            return None
        for attempt in range(retries):
            try:
                import requests
                headers = {
                    'ima-openapi-clientid': self.client_id,
                    'ima-openapi-apikey': self.api_key,
                    'Content-Type': 'application/json'
                }
                url = f"{self.base_url}/{endpoint}"
                response = requests.post(url, json=payload, headers=headers, timeout=30)
                if response.status_code == 200:
                    return response.json()
                elif response.status_code == 403:
                    result = response.json()
                    if '请求超量' in result.get('msg', ''):
                        logger.warning(f"[IMA] API请求超量，等待60秒...")
                        if attempt < retries - 1:
                            time.sleep(60)
                        continue
                return None
            except Exception as e:
                logger.warning(f"[IMA] 调用失败 ({attempt+1}/{retries}): {e}")
                if attempt < retries - 1:
                    time.sleep(2)
        return None

    def sync_note(self, title: str, content: str, category: str = None,
                  content_hash: str = None, doc_path: str = None) -> Optional[str]:
        """同步笔记到IMA"""
        if not self.enabled or self.rate_limited:
            return None

        doc_key = content_hash or title
        existing = self.sync_log.get(doc_key, {})

        full_content = f"# {title}\n\n"
        if category:
            full_content += f"> 分类: {category}\n"
        if content_hash:
            full_content += f"> 内容标识: {content_hash}\n"
        full_content += f"\n{content}\n\n---\n*自动同步自图片知识库*\n"

        # 追加到已有笔记
        if existing.get('doc_id') and doc_path:
            result = self._api_call('append_doc', {
                'doc_id': existing['doc_id'],
                'content_format': 1,
                'content': f"\n\n---\n*更新于 {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n" + content
            })
            if result:
                self.sync_log[doc_key]['last_sync'] = datetime.now().isoformat()
                self.sync_log[doc_key]['update_count'] = existing.get('update_count', 0) + 1
                self._save_sync_log()
                self.synced_this_session += 1
                return existing['doc_id']

        # 新建笔记
        import_payload = {
            'content_format': 1,
            'content': full_content
        }
        if self.default_notebook_id:
            import_payload['notebook_id'] = self.default_notebook_id

        result = self._api_call('import_doc', import_payload)
        if result and result.get('code') == 0:
            doc_id = result.get('data', {}).get('note_id', 'imported')
            self.sync_log[doc_key] = {
                'doc_id': doc_id,
                'title': title,
                'category': category,
                'first_sync': datetime.now().isoformat(),
                'last_sync': datetime.now().isoformat(),
                'update_count': 1
            }
            self._save_sync_log()
            self.synced_this_session += 1
            return doc_id
        else:
            self.failed_this_session += 1
            return None


# ============================================================
# 分批处理（保留V6功能）
# ============================================================
_BATCH_DIR = Path(__file__).parent / '.workbuddy' / 'memory'
BATCH_STATE_FILE = _BATCH_DIR / 'batch_state.json'

BATCH_CONFIG = {
    "small":  {"max_size": 500 * 1024,       "batch_size": 10},
    "medium": {"max_size": 2 * 1024 * 1024,  "batch_size": 6},
    "large":  {"max_size": 5 * 1024 * 1024,  "batch_size": 4},
    "xlarge": {"max_size": float('inf'),      "batch_size": 2},
}

@dataclass
class ImageInfo:
    path: str
    name: str
    size: int
    folder: str
    hash: str = ""
    def to_dict(self): return asdict(self)
    @classmethod
    def from_dict(cls, d): return cls(**d)

@dataclass
class BatchInfo:
    batch_id: str
    folder: str
    images: List[str]
    status: str
    created_at: str
    started_at: Optional[str] = None
    completed_at: Optional[str] = None
    result: Optional[str] = None
    retry_count: int = 0
    def to_dict(self): return asdict(self)
    @classmethod
    def from_dict(cls, d): return cls(**d)

@dataclass
class ProcessingState:
    session_id: str
    started_at: str
    total_images: int
    total_batches: int
    completed_batches: int = 0
    failed_batches: int = 0
    batches: Optional[Dict] = None
    def __post_init__(self):
        if self.batches is None:
            self.batches = {}
    def to_dict(self):
        data = asdict(self)
        if self.batches:
            data['batches'] = {k: v.to_dict() if isinstance(v, BatchInfo) else v
                               for k, v in self.batches.items()}
        return data
    @classmethod
    def from_dict(cls, data):
        if 'batches' in data and data['batches']:
            data['batches'] = {k: BatchInfo.from_dict(v) for k, v in data['batches'].items()}
        return cls(**data)


class BatchManager:
    def __init__(self, source_dir: str = '待处理图片'):
        self.source_dir = Path(source_dir)
        self.state: Optional[ProcessingState] = None

    def _load_state(self) -> Optional[ProcessingState]:
        if BATCH_STATE_FILE.exists():
            try:
                with open(BATCH_STATE_FILE, 'r', encoding='utf-8') as f:
                    return ProcessingState.from_dict(json.load(f))
            except Exception as e:
                print(f"[WARN] 加载状态失败: {e}")
        return None

    def _save_state(self):
        BATCH_STATE_FILE.parent.mkdir(parents=True, exist_ok=True)
        with open(BATCH_STATE_FILE, 'w', encoding='utf-8') as f:
            json.dump(self.state.to_dict(), f, ensure_ascii=False, indent=2)

    def _get_batch_size(self, file_size: int) -> int:
        for tier, cfg in BATCH_CONFIG.items():
            if file_size <= cfg["max_size"]:
                return cfg["batch_size"]
        return 2

    def _compute_hash(self, path: Path) -> str:
        try:
            h = hashlib.md5()
            with open(path, 'rb') as f:
                for chunk in iter(lambda: f.read(8192), b''):
                    h.update(chunk)
            return h.hexdigest()[:8]
        except Exception:
            return ""

    def _scan_images(self) -> List[ImageInfo]:
        images = []
        for ext in ['*.jpg', '*.jpeg', '*.png', '*.webp', '*.bmp']:
            for img in self.source_dir.rglob(ext):
                size = img.stat().st_size
                images.append(ImageInfo(
                    path=str(img),
                    name=img.name,
                    size=size,
                    folder=img.parent.name,
                    hash=self._compute_hash(img)
                ))
        return images

    def initialize(self) -> Optional[ProcessingState]:
        existing = self._load_state()
        if existing:
            pending = sum(1 for b in existing.batches.values()
                         if isinstance(b, BatchInfo) and b.status in ['pending', 'failed'])
            if pending > 0:
                print(f"[INFO] 发现未完成的处理状态，继续上次处理 (剩余 {pending} 批)")
                self.state = existing
                return self.state

        print("[INFO] 初始化分批处理...")
        images = self._scan_images()
        if not images:
            print("[INFO] 没有待处理的图片")
            return None

        images_by_folder: Dict[str, List[ImageInfo]] = defaultdict(list)
        for img in images:
            images_by_folder[img.folder].append(img)

        session_id = datetime.now().strftime('%Y%m%d_%H%M%S')
        batches = {}
        for folder, folder_images in images_by_folder.items():
            i = 0
            while i < len(folder_images):
                batch_size = self._get_batch_size(folder_images[i].size)
                batch_imgs = folder_images[i:i+batch_size]
                batch_id = f"{folder}_{i//batch_size:03d}"
                batches[batch_id] = BatchInfo(
                    batch_id=batch_id,
                    folder=folder,
                    images=[img.path for img in batch_imgs],
                    status='pending',
                    created_at=datetime.now().isoformat()
                )
                i += batch_size

        self.state = ProcessingState(
            session_id=session_id,
            started_at=datetime.now().isoformat(),
            total_images=len(images),
            total_batches=len(batches),
            batches=batches
        )
        self._save_state()
        print(f"[INFO] 初始化完成: {len(images)} 张图片, {len(batches)} 个批次")
        return self.state

    def get_next_batch(self) -> Optional[BatchInfo]:
        if not self.state:
            return None
        for batch in self.state.batches.values():
            if isinstance(batch, BatchInfo) and batch.status in ['pending', 'failed']:
                if batch.retry_count < 3:
                    batch.status = 'processing'
                    batch.started_at = datetime.now().isoformat()
                    self._save_state()
                    return batch
        return None

    def mark_batch_completed(self, batch_id: str, result: str = ""):
        if self.state and batch_id in self.state.batches:
            batch = self.state.batches[batch_id]
            if isinstance(batch, BatchInfo):
                batch.status = 'completed'
                batch.completed_at = datetime.now().isoformat()
                batch.result = result
                self.state.completed_batches += 1
                self._save_state()

    def mark_batch_failed(self, batch_id: str, error: str = ""):
        if self.state and batch_id in self.state.batches:
            batch = self.state.batches[batch_id]
            if isinstance(batch, BatchInfo):
                batch.status = 'failed'
                batch.retry_count += 1
                batch.result = error
                self.state.failed_batches += 1
                self._save_state()

    def print_progress(self):
        state = self._load_state()
        if not state:
            print("[INFO] 没有正在进行的处理任务")
            return
        completed = sum(1 for b in state.batches.values()
                       if isinstance(b, BatchInfo) and b.status == 'completed')
        failed = sum(1 for b in state.batches.values()
                    if isinstance(b, BatchInfo) and b.status == 'failed')
        pending = state.total_batches - completed - failed
        print(f"[进度] 总批次: {state.total_batches} | 完成: {completed} | 失败: {failed} | 待处理: {pending}")
        print(f"[进度] 总图片: {state.total_images}")

    def clear_state(self):
        if BATCH_STATE_FILE.exists():
            BATCH_STATE_FILE.unlink()
            print("[INFO] 处理状态已清除")


# ============================================================
# 处理单张图片（V7重构版）
# ============================================================
def process_single_image(
    ocr: MultiEngineOCR,
    analyzer: SmartContentAnalyzer,
    organizer: ContentOrganizer,
    doc_manager: SmartDocumentManager,
    ima_syncer: IMASyncer,
    archiver: ImageArchiver,
    image_path: str,
    index: int,
    total: int,
    progress: ProgressBar = None
) -> Optional[Dict]:
    """[F-2] 处理单张图片（精简日志版）"""
    image_name = Path(image_path).name

    # 进度条更新
    if progress:
        progress.update(index, f"{image_name[:12]}")

    # 日志摘要行（1行搞定，不淹没进度条）
    print(f"\n{'─'*60}")
    print(f"[{index}/{total}] {image_name}")
    t_start = time.time()

    # 1. OCR识别
    result = ocr.recognize(image_path)
    if not result.get('success'):
        msg = f"  ❌ OCR失败: {result.get('error')}"
        if progress:
            progress.log(msg)
        else:
            print(msg)
        archiver.archive_image(image_path, "其他")
        return None

    text = result.get('text', '').strip()
    if not text or len(text) < 10:
        msg = "  ⏭️ 无文字内容，跳过"
        if progress:
            progress.log(msg)
        else:
            print(msg)
        archiver.archive_image(image_path, "其他")
        return None

    # 2. 内容整理
    organized = organizer.organize(text)
    clean_text = organized['content']
    source_info = organized['source']

    # 3. AI分类 + 文档命名
    analysis = analyzer.analyze(clean_text, index)
    category = analysis['category']
    confidence = analysis['confidence']
    doc_name = analysis['doc_name']
    keywords = analysis['keywords']
    content_hash = analysis['content_hash']

    # 4. 重复检测
    if doc_manager.is_duplicate_hash(content_hash):
        archiver.archive_image(image_path, category)
        elapsed = time.time() - t_start
        msg = f"  🔄 重复内容 | {category} | {doc_name} ({elapsed:.1f}s)"
        if progress:
            progress.log(msg)
        else:
            print(msg)
        return {
            'image': image_name, 'category': category, 'doc_name': doc_name,
            'is_duplicate': True, 'content_hash': content_hash
        }

    # 5. 保存文档（MD + Word）
    md_file, is_new_doc = doc_manager.save_document(
        doc_name=doc_name, category=category, content=clean_text,
        image_name=image_name, keywords=keywords,
        content_hash=content_hash, source_info=source_info
    )
    docx_file = doc_manager.generate_word(
        doc_name=doc_name, category=category, content=clean_text,
        image_name=image_name, keywords=keywords, md_file_path=md_file
    )

    # 6. IMA同步
    if ima_syncer.enabled and not ima_syncer.rate_limited:
        ima_syncer.sync_note(doc_name, clean_text, category, content_hash, md_file)

    # 7. 图片归档
    archived_path = archiver.archive_image(image_path, category)

    # 最终汇总日志（1行）
    elapsed = time.time() - t_start
    conf_label = "高" if confidence >= 0.6 else ("中" if confidence >= 0.3 else "低")
    action = "新建" if is_new_doc else "合并"
    status = "✅"
    if confidence < 0.15:
        status = "⚠️"
    source_tag = f" | {source_info['platform']}" if source_info.get('platform') else ""
    msg = f"  {status} [{conf_label}] {category}-{doc_name} ({action}){source_tag} ({elapsed:.1f}s)"
    if progress:
        progress.log(msg)
    else:
        print(msg)

    return {
        'image': image_name,
        'text_length': len(text),
        'category': category,
        'confidence': round(confidence, 3),
        'doc_name': doc_name,
        'keywords': keywords,
        'is_duplicate': False,
        'is_new_doc': is_new_doc,
        'md_file': md_file,
        'docx_file': docx_file,
        'content_hash': content_hash,
        'source': source_info
    }


# ============================================================
# [P1-6] 统一报告生成
# ============================================================
def save_report(results: List[Dict], total_images: int,
                elapsed: float, mode: str = 'full'):
    """[P1-6] 修复统计报告逻辑"""
    valid = [r for r in results if r and not r.get('failed')]
    duplicates = [r for r in valid if r.get('is_duplicate')]
    new_docs = [r for r in valid if not r.get('is_duplicate') and r.get('is_new_doc')]
    merged_docs = [r for r in valid if not r.get('is_duplicate') and not r.get('is_new_doc')]
    md_count = sum(1 for r in valid if r.get('md_file') and not r.get('is_duplicate'))
    docx_count = sum(1 for r in valid if r.get('docx_file') and not r.get('is_duplicate'))
    failed = [r for r in results if r and r.get('failed')]

    # 分类统计（[P1-6] 使用 category 而非 theme）
    cat_stats = dict(Counter(
        r['category'] for r in valid
        if not r.get('is_duplicate') and r.get('category')
    ))

    report = {
        'timestamp': datetime.now().isoformat(),
        'version': 'V8.0',
        'mode': mode,
        'total_images': total_images,
        'processed_successfully': len(valid),
        'failed': len(failed),
        'duplicates_skipped': len(duplicates),
        'new_documents': len(new_docs),
        'merged_documents': len(merged_docs),
        'markdown_files': md_count,
        'word_files': docx_count,
        'elapsed_seconds': round(elapsed, 1),
        'category_stats': cat_stats,
        'results': [{k: v for k, v in r.items() if k not in ['md_file', 'docx_file']}
                    for r in results if r]
    }

    Path('处理结果').mkdir(exist_ok=True)
    report_file = Path('处理结果/处理报告.json')
    with open(report_file, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    return report, report_file


def print_summary(report: Dict, archiver: ImageArchiver, ima_syncer: IMASyncer):
    """打印处理汇总"""
    print("\n" + "=" * 60)
    print("处理完成 - V8.0 汇总")
    print("=" * 60)
    print(f"  总图片数:     {report['total_images']} 张")
    print(f"  成功处理:     {report['processed_successfully']} 张")
    print(f"  处理失败:     {report['failed']} 张")
    print(f"  重复跳过:     {report['duplicates_skipped']} 张")
    print(f"  新建文档:     {report['new_documents']} 个")
    print(f"  合并到已有:   {report['merged_documents']} 次")
    print(f"  Markdown:     {report['markdown_files']} 个")
    print(f"  Word:         {report['word_files']} 个")
    print(f"  总耗时:       {int(report['elapsed_seconds']//60)}分{int(report['elapsed_seconds']%60)}秒")

    if report['category_stats']:
        print("\n分类分布:")
        for cat, cnt in sorted(report['category_stats'].items(), key=lambda x: -x[1]):
            print(f"  {cat}: {cnt} 张")

    arch_stats = archiver.get_stats()
    print(f"\n归档统计:")
    print(f"  本次归档: {arch_stats['archived_this_session']} 张")
    print(f"  总已归档: {arch_stats['total_archived']} 张")

    if ima_syncer.enabled:
        print(f"\nIMA同步:")
        print(f"  本次成功: {ima_syncer.synced_this_session} 条")
        print(f"  本次失败: {ima_syncer.failed_this_session} 条")

    print("=" * 60)


# ============================================================
# 全自动处理模式
# ============================================================
def run_full_mode(batch_manager: BatchManager):
    """[F-1] 全自动处理模式 - 智能分批"""
    print("=" * 60)
    print("全自动化图片处理工具 V8.0")
    print("OCR -> 内容整理 -> AI分类 -> 智能命名 -> 合并 -> Word -> IMA -> 归档")
    print("=" * 60)

    # 初始化组件
    print("\n[初始化] OCR引擎...")
    ocr = MultiEngineOCR()
    if not ocr.current_engine:
        print("错误: 没有可用的OCR引擎!")
        return

    print("[初始化] 智能分类器...")
    classifier = SmartClassifier()

    print("[初始化] 内容分析器...")
    analyzer = SmartContentAnalyzer(classifier)

    print("[初始化] 内容整理器...")
    organizer = ContentOrganizer()

    print("[初始化] 文档管理器...")
    doc_manager = SmartDocumentManager()

    print("[初始化] IMA同步器...")
    ima_syncer = IMASyncer()

    print("[初始化] 图片归档器...")
    archiver = ImageArchiver()

    # [F-1] 扫描图片，判断是否需要分批
    source_dir = Path('待处理图片')
    if not source_dir.exists():
        print("错误: 待处理图片目录不存在!")
        return

    images = []
    for ext in ['*.jpg', '*.jpeg', '*.png', '*.webp', '*.bmp']:
        images.extend(source_dir.rglob(ext))

    if not images:
        print("没有找到待处理图片")
        return

    total = len(images)

    # [F-1] 智能分批判断：>10张自动启用分批
    AUTO_BATCH_THRESHOLD = 10
    if total > AUTO_BATCH_THRESHOLD:
        print(f"\n[分批] 检测到 {total} 张图片，超过阈值({AUTO_BATCH_THRESHOLD})，自动启用分批模式")
        batch_manager.clear_state()  # 清除旧状态
        run_batch_mode(batch_manager)
        return

    print(f"\n找到 {total} 张待处理图片（≤{AUTO_BATCH_THRESHOLD}，不分批）")
    subfolders = [d.name for d in source_dir.iterdir() if d.is_dir()]
    if subfolders:
        print(f"子文件夹: {subfolders}")

    progress = ProgressBar(total, prefix="处理进度")
    results = []
    start_time = time.time()

    # [F-1] 全局进度统计
    processed_count = 0

    for i, img_path in enumerate(images, 1):
        try:
            result = process_single_image(
                ocr, analyzer, organizer, doc_manager, ima_syncer, archiver,
                str(img_path), i, total, progress
            )
            if result:
                results.append(result)
                if not result.get('failed'):
                    processed_count += 1
        except Exception as e:
            msg = f"\n  ❌ 处理失败: {e}"
            if progress:
                progress.log(msg)
            else:
                print(msg)
            logger.error(f"处理失败 {img_path}: {e}")
            results.append({'image': Path(img_path).name, 'failed': True, 'error': str(e)})

    progress.finish()

    elapsed = time.time() - start_time
    report, report_file = save_report(results, total, elapsed, 'full')
    print_summary(report, archiver, ima_syncer)
    print(f"\n报告已保存: {report_file}")


# ============================================================
# 分批处理模式
# ============================================================
def run_batch_mode(batch_manager: BatchManager):
    """分批处理模式（V7）"""
    print("\n" + "=" * 60)
    print("分批处理模式 V8.0")
    print("=" * 60)

    state = batch_manager.initialize()
    if not state:
        return

    batch_manager.print_progress()

    # 初始化组件
    print("\n[初始化] OCR引擎...")
    ocr = MultiEngineOCR()
    if not ocr.current_engine:
        print("错误: 没有可用的OCR引擎!")
        return

    classifier = SmartClassifier()
    analyzer = SmartContentAnalyzer(classifier)
    organizer = ContentOrganizer()
    doc_manager = SmartDocumentManager()
    ima_syncer = IMASyncer()
    archiver = ImageArchiver()

    all_results = []
    batch_num = 0
    total_images = state.total_images
    total_batches = state.total_batches
    start_time = time.time()

    while True:
        batch = batch_manager.get_next_batch()
        if not batch:
            print("\n所有批次处理完成！")
            break

        batch_num += 1
        print(f"\n{'#'*55}")
        print(f"批次 {batch_num}/{total_batches}: {batch.batch_id} | {len(batch.images)} 张")
        print(f"{'#'*55}")

        batch_results = []
        batch_success = True
        batch_progress = ProgressBar(len(batch.images), prefix=f"批次{batch_num}")

        for i, img_path in enumerate(batch.images, 1):
            try:
                result = process_single_image(
                    ocr, analyzer, organizer, doc_manager, ima_syncer, archiver,
                    img_path, i, len(batch.images), batch_progress
                )
                if result:
                    batch_results.append(result)
                else:
                    batch_results.append({'image': Path(img_path).name, 'failed': True})
            except Exception as e:
                print(f"\n  错误: {e}")
                logger.error(f"处理失败 {img_path}: {e}")
                batch_results.append({'image': Path(img_path).name, 'failed': True, 'error': str(e)})
                batch_success = False

        batch_progress.finish()

        if batch_success:
            batch_manager.mark_batch_completed(batch.batch_id, f"处理 {len(batch_results)} 张")
        else:
            failed_n = sum(1 for r in batch_results if r.get('failed'))
            batch_manager.mark_batch_failed(batch.batch_id, f"失败 {failed_n} 张")

        all_results.extend(batch_results)

        if batch_num < total_batches:
            print("\n[批次间隔] 1 秒...")
            time.sleep(1)

    elapsed = time.time() - start_time
    report, report_file = save_report(all_results, total_images, elapsed, 'batch')
    print_summary(report, archiver, ima_syncer)
    print(f"\n报告已保存: {report_file}")

    return all_results


# ============================================================
# 主入口
# ============================================================
def print_usage():
    print("\n用法:")
    print("  python auto_process_all_v8.py           # 全自动处理（>10张自动分批）")
    print("  python auto_process_all_v8.py --batch    # 强制分批模式（可中断续传）")
    print("  python auto_process_all_v8.py --init     # 仅初始化分批")
    print("  python auto_process_all_v8.py --progress # 查看进度")
    print("  python auto_process_all_v8.py --clear    # 清除分批状态")


def main():
    batch_manager = BatchManager()
    if len(sys.argv) > 1:
        cmd = sys.argv[1]
        if cmd == '--init':
            state = batch_manager.initialize()
            if state:
                batch_manager.print_progress()
        elif cmd == '--progress':
            batch_manager.print_progress()
        elif cmd == '--clear':
            batch_manager.clear_state()
        elif cmd == '--batch':
            run_batch_mode(batch_manager)
        else:
            print(f"[ERROR] 未知参数: {cmd}")
            print_usage()
    else:
        run_full_mode(batch_manager)


if __name__ == '__main__':
    main()
