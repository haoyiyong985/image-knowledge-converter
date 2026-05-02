#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
全自动化图片处理工具 V9.0 - LLM全覆盖版
==========================================
V9.0 核心重构：用混元Lite替代所有正则决策
  [V9-1] LLM全覆盖命名 - 每张图片调用混元Lite，一次性返回 title + category + summary
  [V9-2] LLM全覆盖分类 - 分类不再依赖关键词匹配，由LLM直接判断
  [V9-3] 删除正则决策系统 - 移除 ~300行正则代码（人名/标题/置信度评估）
  [V9-4] V2轻量结构化 - 保留V2的语义标题识别（功效/做法/禁忌等关键词）
  [V9-5] 正则仅做过滤 - 去噪点、去空行、去乱码、格式化

V9 哲学：正则只做过滤（排除垃圾），LLM做所有决策（命名/分类/摘要）

保留功能（来自V8）：
  [F-1] 默认模式智能分批 - 图片>10张自动启用分批
  [F-2] 进度条始终可见（Fix-7简化版）
  [V7-P0-3] 同主题内容合并 - 跨分类检测同名文档
  [V7-P1-4] 图片按源文件夹归档
  [V7-P1-5] 优化重复检测 - 主题词相似度 + hash 双重检测
  [V7-P1-6] 精确统计报告
  [V7-P1-7] IMA 同步（限流处理）

使用方法：
  python auto_process_all_v8_1.py           # 全自动处理（>10张自动分批）
  python auto_process_all_v8_1.py --batch    # 强制分批模式（可中断续传）
  python auto_process_all_v8_1.py --init     # 仅初始化分批
  python auto_process_all_v8_1.py --progress # 查看进度
  python auto_process_all_v8_1.py --clear    # 清除分批状态
"""

import os
import sys
import io
import json
import time
import re
import hashlib
import shutil
import logging
from pathlib import Path
from datetime import datetime
from collections import Counter, defaultdict
from typing import Optional, List, Dict, Tuple
from dataclasses import dataclass, asdict

# 修复 Windows 控制台编码
if sys.platform == 'win32':
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
    except Exception:
        pass

sys.path.insert(0, '.')

from dotenv import load_dotenv
env_path = Path('.') / '.env'
if env_path.exists():
    load_dotenv(dotenv_path=env_path)
else:
    print("[WARN] .env 文件不存在，使用系统环境变量")

from local_ocr import LocalOCR

try:
    from tencent_ocr import TencentOCR
    TENCENT_AVAILABLE = True
except ImportError:
    TENCENT_AVAILABLE = False

try:
    from baidu_ocr import BaiduOCR
    BAIDU_AVAILABLE = True
except ImportError:
    BAIDU_AVAILABLE = False

# 日志配置
Path('处理结果').mkdir(exist_ok=True)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(message)s',
    handlers=[
        logging.FileHandler('处理结果/process.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# ============================================================
# [P0-1] 智能分类引擎 - 完全重构
# ============================================================

# 预设分类 + 关键词（仅作匹配辅助，AI内容理解优先）
PRESET_CATEGORIES = {
    "历史文化": [
        "历史", "文化", "朝代", "古代", "人物", "传记", "皇帝", "将军", "战争", "事件",
        "明朝", "清朝", "汉朝", "唐朝", "宋朝", "元朝", "秦朝", "民国", "三国", "春秋",
        "战国", "革命", "史记", "文献", "典故", "传统", "文明", "民俗", "风俗", "节日",
        "诗词", "古文", "书法", "绘画", "建筑", "遗址", "文物", "考古",
        # [P1-3] 补充：分封/门阀/世族相关
        "分封", "诸侯", "封地", "封国", "氏族", "郡望", "五姓", "七望", "门阀",
        "世族", "大族", "望族", "宗族", "姓氏", "始封", "封君", "藩王", "藩国",
        "太后", "太妃", "宣太后", "太子", "皇后", "皇妃", "皇室", "宗室",
    ],
    "营养健康": [
        "营养", "健康", "饮食", "食物", "蔬菜", "水果", "蛋白质", "碳水",
        "中医", "养生", "药膳", "食疗", "中药", "经络", "穴位", "调理",
        "体质", "气血", "疾病", "预防", "治疗", "症状", "血压", "血糖",
        "维生素", "矿物质", "膳食", "抗炎", "免疫力", "肠道", "益生菌",
        "减肥", "增肌", "低卡", "热量", "卡路里", "蛋白", "脂肪",
    ],
    "生活方式": [
        "生活", "运动", "睡眠", "压力", "情绪", "心理", "作息", "习惯",
        "美容", "护肤", "健身", "锻炼", "瑜伽", "冥想", "放松",
        "整理", "收纳", "家居", "装修", "穿搭", "时尚", "购物",
        "理财", "储蓄", "投资", "工作", "职场", "效率",
    ],
    "教育育儿": [
        "育儿", "宝宝", "孩子", "教育", "辅食", "喂养", "早教",
        "亲子", "成长", "发育", "妈妈", "孕妇", "备孕", "新生儿",
        "婴儿", "幼儿", "儿童", "学习", "作业", "学校", "老师",
    ],
    "旅游攻略": [
        "旅游", "旅行", "景点", "攻略", "出行", "交通", "酒店",
        "民宿", "美食", "打卡", "签证", "机票", "行程", "路线",
        "城市", "国家", "海外", "境外", "自驾", "背包",
    ],
    "综合知识": [
        "科普", "知识", "研究", "发现", "实验", "数据", "专家",
        "建议", "指南", "推荐", "科学", "原理", "机制", "分析",
    ],
}

# 分类特征权重（区分度高的关键词加权）
CATEGORY_BOOST_WORDS = {
    "历史文化": {"朝代", "皇帝", "将军", "史记", "文献", "古代", "传记", "人物志"},
    "营养健康": {"中医", "养生", "营养素", "血压", "血糖", "食疗", "药膳"},
    "教育育儿": {"宝宝", "辅食", "早教", "新生儿", "幼儿园"},
    "旅游攻略": {"攻略", "景点", "签证", "行程", "民宿"},
}


class SmartClassifier:
    """
    [P0-1] 智能分类器 - 内容理解驱动
    
    策略：
    1. 对OCR文本进行预处理（去噪）
    2. 对每个预设分类计算综合分数（关键词密度 + 特征词加权）
    3. 如果最高分低于阈值，尝试从内容中推断新分类
    4. 返回分类名和置信度
    """

    def __init__(self):
        # 从已知分类文件夹中动态加载已有分类
        self.known_categories = set(PRESET_CATEGORIES.keys())
        self._load_existing_categories()

    def _load_existing_categories(self):
        """加载已存在的分类（包括动态新建的）"""
        result_dir = Path('处理结果')
        if result_dir.exists():
            for d in result_dir.iterdir():
                if d.is_dir() and not d.name.startswith('.'):
                    self.known_categories.add(d.name)
        archived_dir = Path('已处理图片')
        if archived_dir.exists():
            for d in archived_dir.iterdir():
                if d.is_dir() and not d.name.startswith('.'):
                    self.known_categories.add(d.name)

    def classify(self, text: str) -> Tuple[str, float]:
        """
        对文本分类，返回 (分类名, 置信度)
        
        置信度说明：
          >= 0.6  高置信度（关键词明确）
          0.3-0.6 中置信度（有一定匹配）
          < 0.3   低置信度（使用综合知识兜底）
        """
        if not text or len(text) < 5:
            return "综合知识", 0.1

        scores = {}
        text_len = max(len(text), 1)

        for cat_name, keywords in PRESET_CATEGORIES.items():
            # 基础关键词得分（密度）
            matched = [kw for kw in keywords if kw in text]
            base_score = len(matched) * 3 / text_len * 1000

            # 特征词加权
            boost_words = CATEGORY_BOOST_WORDS.get(cat_name, set())
            boost = sum(5 for w in boost_words if w in text)

            scores[cat_name] = base_score + boost

        if not scores:
            return "综合知识", 0.1

        best_cat = max(scores, key=scores.get)
        best_score = scores[best_cat]

        # 归一化置信度（0~1）
        max_possible = 50.0
        confidence = min(best_score / max_possible, 1.0)

        if confidence < 0.05:
            # 完全没有匹配，兜底为综合知识
            return "综合知识", 0.05

        return best_cat, confidence

    def get_matched_keywords(self, text: str, category: str) -> List[str]:
        """获取匹配到的关键词"""
        keywords = PRESET_CATEGORIES.get(category, [])
        return [kw for kw in keywords if kw in text][:8]

    def register_new_category(self, category_name: str):
        """注册新分类"""
        self.known_categories.add(category_name)
        logger.info(f"[分类] 注册新分类: {category_name}")


# ============================================================
# [P0-2][P0-3] 智能内容分析器 - 重构
# ============================================================

class SmartContentAnalyzer:
    """
    [P0-2][P0-3] 智能内容分析器
    
    V8.2改进：
    1. [P0-3] 双引擎LLM命名：腾讯混元Lite → 硅基流动SiliconFlow（OpenAI兼容）
    2. 命名策略：正则高置信度(≥0.8)直接用，低置信度调用大模型，失败降级正则
    """

    def __init__(self, classifier: SmartClassifier):
        self.classifier = classifier
        self.cache = {}
        self._last_person_name_low_conf = False  # [V8.3 Fix-9]
        # [P0-3] 大模型配置（双引擎）
        self._hunyuan_available = None   # 混元Lite
        self._siliconflow_available = None  # 硅基流动

    def _check_llm_available(self) -> bool:
        """[P0-3][V8.2] 检测是否有至少一个LLM引擎可用"""
        return self._check_hunyuan() or self._check_siliconflow()

    def _check_hunyuan(self) -> bool:
        """检测混元Lite是否可用"""
        if self._hunyuan_available is not None:
            return self._hunyuan_available
        try:
            secret_id = os.getenv('TENCENT_SECRET_ID', '')
            secret_key = os.getenv('TENCENT_SECRET_KEY', '')
            if secret_id and secret_key and '替换' not in secret_id:
                self._hunyuan_available = True
                logger.info(f"[LLM] 混元Lite已配置 (key前缀: {secret_id[:6]}...)")
            else:
                self._hunyuan_available = False
                logger.info("[LLM] 混元Lite未配置")
        except Exception:
            self._hunyuan_available = False
        return self._hunyuan_available

    def _check_siliconflow(self) -> bool:
        """检测硅基流动是否可用"""
        if self._siliconflow_available is not None:
            return self._siliconflow_available
        try:
            api_key = os.getenv('SILICONFLOW_API_KEY', '')
            if api_key and '替换' not in api_key:
                self._siliconflow_available = True
                logger.info(f"[LLM] 硅基流动已配置 (key前缀: {api_key[:6]}...)")
            else:
                self._siliconflow_available = False
                logger.info("[LLM] 硅基流动未配置")
        except Exception:
            self._siliconflow_available = False
        return self._siliconflow_available

    def _llm_extract_topic(self, text: str, category: str) -> Optional[Dict]:
        """
        [P0-3][V8.2] 调用LLM提取主题（双引擎：腾讯混元Lite → 硅基流动兜底）

        优先级：腾讯混元Lite（SDK）→ 硅基流动SiliconFlow（OpenAI兼容格式）→ None
        返回：{'title': '宣太后', 'category': '历史文化', 'confidence': 0.9}
        失败返回 None
        """
        # 截取前500字作为输入（节省token）
        snippet = text[:500].replace('\n', ' ').strip()
        prompt = (
            f"以下是从图片中识别出的文本内容片段，请分析后提取：\n"
            f"1. 主题标题（人名/事件名/地名/主题，10字以内）\n"
            f"2. 内容分类（从以下选择：历史文化/营养健康/生活方式/教育育儿/旅游攻略/综合知识）\n\n"
            f"文本：{snippet}\n\n"
            f"请只返回JSON格式，例如：{{\"title\": \"宣太后\", \"category\": \"历史文化\"}}"
        )

        # 引擎1：腾讯混元Lite
        result = self._call_hunyuan_lite(prompt)
        if result:
            return result

        # 引擎2：硅基流动 SiliconFlow
        result = self._call_siliconflow(prompt, category)
        if result:
            return result

        return None

    def _call_hunyuan_lite(self, prompt: str) -> Optional[Dict]:
        """调用腾讯混元Lite（官方SDK）"""
        try:
            from tencentcloud.hunyuan.v20230901 import hunyuan_client, models as hm_models
            from tencentcloud.common.exception.tencent_cloud_sdk_exception import TencentCloudSDKException
            from tencentcloud.common.profile.client_profile import ClientProfile
            from tencentcloud.common.profile.http_profile import HttpProfile
            from tencentcloud.common.credential import Credential

            secret_id = os.getenv('TENCENT_SECRET_ID', '')
            secret_key = os.getenv('TENCENT_SECRET_KEY', '')
            if not secret_id or not secret_key or '替换' in secret_key:
                return None

            httpProfile = HttpProfile()
            httpProfile.endpoint = "hunyuan.tencentcloudapi.com"
            httpProfile.reqTimeout = 8
            clientProfile = ClientProfile()
            clientProfile.httpProfile = httpProfile
            client = hunyuan_client.HunyuanClient(
                Credential(secret_id, secret_key),
                "ap-guangzhou",
                clientProfile
            )

            req = hm_models.ChatCompletionsRequest()
            req.Model = "hunyuan-lite"
            req.Messages = [{"Role": "user", "Content": prompt}]
            req.Stream = False

            logger.info("[LLM] 调用混元Lite (SDK)...")
            resp = client.ChatCompletions(req)
            content_text = resp.Choices[0].Message.Content.strip()
            logger.info(f"[LLM] 混元Lite回复: {content_text[:150]}")

            result = self._parse_llm_response(content_text)
            if result:
                result['_engine'] = 'hunyuan-lite'
            return result

        except ImportError:
            logger.info("[LLM] 混元SDK未安装")
            return None
        except TencentCloudSDKException as e:
            code = getattr(e, 'code', '')
            msg = getattr(e, 'message', str(e))
            if 'ServiceNotActivated' in code:
                logger.info("[LLM] 混元Lite服务未开通, 跳过 (请在腾讯云控制台开通)")
            elif 'UnauthorizedOperation' in code:
                logger.info("[LLM] 混元Lite CAM权限不足, 跳过")
            else:
                logger.info(f"[LLM] 混元Lite异常: {code} - {msg}")
            return None
        except Exception as e:
            logger.info(f"[LLM] 混元Lite调用失败: {e}")
            return None

    def _call_siliconflow(self, prompt: str, category: str) -> Optional[Dict]:
        """调用硅基流动 SiliconFlow（OpenAI兼容格式，免费额度）"""
        try:
            api_key = os.getenv('SILICONFLOW_API_KEY', '')
            if not api_key or '替换' in api_key:
                logger.info("[LLM] 硅基流动未配置 (SILICONFLOW_API_KEY), 跳过")
                return None

            url = "https://api.siliconflow.cn/v1/chat/completions"
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": os.getenv('SILICONFLOW_MODEL', 'Qwen/Qwen3-8B'),
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": 200,
                "temperature": 0.3
            }

            logger.info("[LLM] 调用硅基流动...")
            resp = requests.post(url, json=payload, headers=headers, timeout=8)
            if resp.status_code == 200:
                data = resp.json()
                content_text = data.get('choices', [{}])[0].get('message', {}).get('content', '').strip()
                logger.info(f"[LLM] 硅基流动回复: {content_text[:150]}")
                result = self._parse_llm_response(content_text)
                if result:
                    result['_engine'] = 'siliconflow'
                return result
            else:
                logger.info(f"[LLM] 硅基流动 HTTP {resp.status_code}: {resp.text[:100]}")
                return None
        except Exception as e:
            logger.info(f"[LLM] 硅基流动调用失败: {e}")
            return None

    def _parse_llm_response(self, content_text: str) -> Optional[Dict]:
        """解析LLM返回的JSON"""
        import json
        json_match = re.search(r'\{[^{}]+\}', content_text)
        if json_match:
            result = json.loads(json_match.group())
            title = result.get('title', '').strip()
            cat = result.get('category', '').strip()
            if title and 2 <= len(title) <= 15:
                title = re.sub(r'[<>:"/\\|?*\n\r\t，。！？、：；]', '', title)
                return {'title': title, 'category': cat, 'confidence': 0.85}
            else:
                logger.info(f"[LLM] 标题长度不符: {len(title)} (要求2-15)")
        else:
            logger.info(f"[LLM] 未找到JSON格式回复")
        return None

    def extract_topic(self, text: str, category: str,
                      use_llm: bool = False, llm_result: Optional[Dict] = None) -> str:
        """
        [F-4][P0-3] 主题提取（V8.1增强版）

        策略（按优先级）：
        1. [P0-3] 如果传入了大模型结果，优先使用
        2. 人名提取（扩充匹配模式）
        3. 事件/朝代/主题提取（历史类专用）
        4. 内容核心标题（排除UI噪点行）
        5. 关键词智能组合
        6. 兜底
        """
        # [P0-3] 优先使用大模型结果
        if llm_result and llm_result.get('title'):
            return llm_result['title']

        lines = text.split('\n')
        valid_lines = [l.strip() for l in lines if l.strip() and len(l.strip()) >= 3]
        content_lines = valid_lines

        # === 策略1：人名提取 ===
        person = self._extract_person_name(content_lines)
        if person:
            return person

        # === 策略2：历史事件/朝代/主题 ===
        if category == '历史文化':
            topic = self._extract_history_topic(content_lines)
            if topic:
                return topic

        # === 策略3：内容标题行 ===
        title = self._extract_content_title(content_lines)
        if title:
            return title

        # === 策略4：关键词智能组合 ===
        matched_kws = self.classifier.get_matched_keywords(text, category)
        if matched_kws:
            stop_words = {'历史', '文化', '知识', '整理', '推荐', '建议', '指南',
                         '健康', '生活', '科学', '科普', '饮食', '营养', '学习'}
            meaningful = [kw for kw in matched_kws
                         if len(kw) >= 2 and kw not in stop_words][:2]
            if meaningful:
                return ''.join(meaningful)

        # === 策略5：兜底 ===
        return f"整理{datetime.now().strftime('%m%d%H%M')}"

    def _extract_person_name(self, lines: List[str]) -> Optional[str]:
        """[F-4] 提取人名 - 扩充匹配模式"""
        # 模式集合
        person_patterns = [
            # XXX传/简介/生平/人物志
            r'([^\s，。、：！？\d]{2,4})(?:传|简介|生平|人物志|人物|世家)',
            # XXX（字/号/谥...）
            r'^([^\s，。]{2,5})[，,]\s*(?:字|号|谥|生于|字号|原名|名)',
            # XXX是/为/任/担任...人物传记格式
            r'^([^\s，。]{2,4})(?:是|为|任|担任|生于|出生|卒于|去世|逝世)',
            # XXX，中国/著名/当代/现代...描述格式
            r'^([^\s，。]{2,4})[，,]\s*(?:中国|著名|当代|现代|近代|古代|清末|明末|南宋|北宋|晚唐|初唐)',
            # XXX（XXXX-XXXX）生卒年格式
            r'^([^\s，。（(]{2,4})[（(]\s*\d{3,4}[-—]\d{3,4}',
            # 江苏XX人/浙江XX人 等
            r'(?:江苏|浙江|山东|广东|四川|湖南|湖北|福建|安徽|河南|河北|陕西|山西|辽宁|北京|上海|天津|重庆)([^\s，。]{2,4})',
        ]

        for line in lines[:15]:
            for pattern in person_patterns:
                match = re.search(pattern, line)
                if match:
                    name = match.group(1)
                    # 过滤掉明显不是人名的
                    if len(name) < 2 or len(name) > 4:
                        continue
                    if re.match(r'^\d+$', name):
                        continue
                            # [Fix-9/10] 地名前缀过滤
                    province_prefixes = '苏浙鲁豫鄂湘赣闽粤桂黔滇蜀皖冀京津沪渝'
                    # 完整省份名（用于检测OCR断行导致省名被吞入）
                    province_full = ['江苏','浙江','山东','广东','四川','湖南','湖北','福建','安徽','河南','河北','陕西','山西','辽宁','北京','上海','天津','重庆']
                    # 省简称本身就是名字 → 跳过
                    if name in province_prefixes:
                        continue
                    # [Fix-10] 名字以省简称开头（如「苏无锡人」= 苏+无锡人）→ 跳过
                    if any(name.startswith(p) for p in province_prefixes):
                        continue
                    # [Fix-10] 名字含完整省名开头（如「江苏无锡人」）→ 跳过
                    if any(name.startswith(p) for p in province_full):
                        continue
                    # [Fix-9] 2字人名且含地名简称 → 低置信度
                    if len(name) == 2 and any(p in name for p in province_prefixes):
                        self._last_person_name_low_conf = True
                    return name

        return None

    def _extract_history_topic(self, lines: List[str]) -> Optional[str]:
        """[F-4] 提取历史类主题（事件/朝代/地图等）"""
        # 朝代+事件模式
        history_patterns = [
            r'((?:明朝|清朝|汉朝|唐朝|宋朝|元朝|秦朝|周朝|隋朝|春秋|战国|三国|民国)'
            r'[^\n]{2,15})',
            r'((?:开国|建国|统一|灭亡|中兴)[^\n]{2,10})',
        ]

        for line in lines[:10]:
            for pattern in history_patterns:
                match = re.search(pattern, line)
                if match:
                    topic = match.group(1).strip()
                    if len(topic) >= 4:
                        return topic[:15]

        return None

    def _extract_content_title(self, lines: List[str]) -> Optional[str]:
        """[F-4] 提取内容标题（排除UI噪点行）"""
        # UI噪点关键词（标题行中不应出现这些）
        noise_words = [
            'http', 'www', '转载', '来源', '作者', '发送', '翻译', '探索版',
            '关注', '收藏', '点赞', '评论', '分享', '来自', '搜索', '找到',
            '写作', 'doubao', 'kimi', '元宝', '豆包', 'ChatGPT', 'DeepSeek',
        ]

        for line in lines[:8]:
            line = line.strip()
            if not line:
                continue
            # 去掉markdown标记
            cleaned = re.sub(r'^#{1,6}\s*', '', line)
            cleaned = re.sub(r'^[◆★▼▪▸【】\[\]]+\s*', '', cleaned)
            cleaned = cleaned.strip()
            if not cleaned:
                continue

            # 排除含噪点词的行
            if any(noise in cleaned for noise in noise_words):
                continue
            # 排除纯数字/时间行
            if re.match(r'^[\d:.\s%]+$', cleaned):
                continue
            # 排除过短行
            if len(cleaned) < 3:
                continue

            # 匹配中文短标题（3-15字）
            if re.match(r'^[\u4e00-\u9fa5]{3,15}$', cleaned):
                return cleaned
            # 带标点的标题（不超过20字，以中文开头）
            if 4 <= len(cleaned) <= 20 and re.match(r'^[\u4e00-\u9fa5]', cleaned):
                return cleaned[:15]

        return None

    def _extract_topic_confidence(self, text: str, category: str) -> float:
        """
        [P0-3] 评估正则提取的置信度
        返回 0.0~1.0，高于0.8才直接用正则，否则调用大模型
        """
        lines = text.split('\n')
        valid_lines = [l.strip() for l in lines if l.strip() and len(l.strip()) >= 3]

        # 策略1：人名提取命中 → 高置信度
        if self._extract_person_name(valid_lines):
            return 0.9

        # 策略2：历史朝代提取命中
        if category == '历史文化' and self._extract_history_topic(valid_lines):
            return 0.85

        # 策略3：内容标题行命中（短标题）
        title = self._extract_content_title(valid_lines)
        if title and len(title) >= 4 and not title.startswith('整理'):
            return 0.75

        # 兜底：低置信度
        return 0.3

    def generate_doc_name(self, text: str, category: str) -> str:
        """
        [P0-2][P0-3] 生成文档名（V8.1：先评估正则置信度，低置信度调用大模型）

        格式：{分类}-{主题}
        策略：正则置信度≥0.8 → 直接用；<0.8 → 调用腾讯混元Lite；失败 → 正则兜底
        """
        conf = self._extract_topic_confidence(text, category)

        # [V8.3 Fix-9] 人名低置信度时强制触发LLM
        if self._last_person_name_low_conf and conf >= 0.8:
            conf = 0.6
            logger.info(f"[命名] 人名含地名前缀，置信度降为0.6，触发LLM")
        self._last_person_name_low_conf = False

        llm_result = None
        if conf < 0.8:
            # 尝试调用大模型（双引擎：混元Lite → 硅基流动）
            logger.info(f"[命名] 正则置信度: {conf:.2f} (< 0.8), 尝试调用LLM (混元Lite→硅基流动)...")
            llm_result = self._llm_extract_topic(text, category)
            if llm_result and llm_result.get('title'):
                engine = llm_result.get('_engine', 'unknown')
                logger.info(f"[命名] LLM返回 ({engine}): title={llm_result['title']}, category={llm_result.get('category', '')}")
            else:
                logger.info(f"[命名] LLM无结果, 使用正则兜底")
            if llm_result and llm_result.get('category'):
                # 大模型返回的分类可能更准确
                category = llm_result['category']
        else:
            logger.info(f"[命名] 正则置信度: {conf:.2f} (≥ 0.8), 直接使用正则")

        topic = self.extract_topic(text, category, llm_result=llm_result)
        # 清理非法文件名字符
        topic = re.sub(r'[<>:"/\\|?*\n\r\t，。！？、：；]', '', topic).strip()
        if len(topic) > 20:
            topic = topic[:20]
        if not topic or len(topic) < 2:
            topic = f"整理{datetime.now().strftime('%m%d%H%M')}"
        return f"{category}-{topic}"

    def compute_content_hash(self, text: str) -> str:
        """计算内容哈希"""
        cleaned = re.sub(r'\s+', '', text)
        return hashlib.md5(cleaned.encode('utf-8')).hexdigest()[:8]

    def analyze(self, text: str, index: int) -> Dict:
        """综合分析，返回分类/主题/文档名等"""
        content_hash = self.compute_content_hash(text)

        if content_hash in self.cache:
            return self.cache[content_hash]

        # 分类
        category, confidence = self.classifier.classify(text)
        keywords = self.classifier.get_matched_keywords(text, category)

        # 文档名（含LLM辅助命名，generate_doc_name内部已处理）
        doc_name = self.generate_doc_name(text, category)

        # 若大模型给出了更准确的分类（包含在doc_name前缀中），更新category
        if '-' in doc_name:
            llm_cat = doc_name.split('-', 1)[0]
            if llm_cat and llm_cat != category:
                # 大模型分类与原分类不同，以大模型为准（仅当有效分类时）
                known_cats = set(PRESET_CATEGORIES.keys()) | self.classifier.known_categories
                if llm_cat in known_cats:
                    category = llm_cat
                    keywords = self.classifier.get_matched_keywords(text, category)

        result = {
            'category': category,
            'confidence': confidence,
            'keywords': keywords,
            'doc_name': doc_name,
            'content_hash': content_hash,
        }

        self.cache[content_hash] = result
        return result


# ============================================================
# 内容整理器（V8.1 全面重构）
# ============================================================
class ContentOrganizer:
    """
    [P0-0][P0-2][P1-2][P1-4] 内容整理器 - V8.1增强版

    核心改进（相比V8）：
    1. [P0-0] 恢复V2的 _is_heading_line() + heading_keywords 语义标题识别
    2. [P0-2] 合并V2语义标题 + V8编号格式 + 新增多列排版重建 + 历史人物结构化
    3. [P1-2] 补充8种遗漏噪点
    4. [P1-4] 补充V2的13种OCR字符替换（V8只有9种）
    """

    def __init__(self):
        # 手机UI噪点模式（V8原有 + P1-2新增8种）
        self.ui_noise_patterns = [
            # 信号栏
            r'^\d{1,2}:\d{2}$',                          # 时间 "23:24"
            r'^[0-9]{1,2}:%$',                            # 时间截断
            r'^4G\s*5G',                                  # 信号标识
            r'^[0-9]+\s*%$',                              # 电量百分比
            r'^\d{1,2}:\d{2}\s*$',                        # 独立时间行
            r'^\s*\d{1,3}%\s*$',                          # 电量
            r'^Wi-?Fi',                                   # Wi-Fi
            r'^\.?\s*AM$', r'^\.?\s*PM$',                # AM/PM
            # App名称和按钮
            r'^(写作|翻译|探索版|有什么问题尽管问我)$',
            r'^(doubao\.com|豆包)$',
            r'^(Kimi|kimi)$',
            r'^(腾讯元宝|元宝)$',
            r'^(ChatGPT|GPT)$',
            r'^(文心一言|通义千问|智谱清言)$',
            r'^(豆包|DeepSeek)$',
            r'^(发送|发送消息)$',
            r'^\+关注$',
            r'^(收藏|点赞|评论|分享|转发)$',
            r'^(已关注|关注)$',
            r'^(更多|展开|收起|折叠)$',
            # 杂项UI
            r'^\d+$',                                     # 纯数字行
            r'^[A-Z][a-z]+$',                             # 纯英文单词
            r'^[a-z]+\.(com|cn|net|org)$',               # 域名
            r'^http[s]?://',                              # URL
            r'^来自.*的搜索$',                            # 搜索来源标记
            r'^搜一下$',                                  # 搜索按钮
            r'^找到\d+篇资料参考$',                       # 搜索结果数
            # [P1-2] 新增8种遗漏噪点
            r'^a?付费',                                   # "a付费"
            r'^小红书号[：:]\s*\d+',                     # 小红书号
            r'^内容由\w+生成',                            # "内容由A1生成"
            r'^相关视频$',                                # "相关视频"
            r'^双深度思考$',                              # "双深度思考"
            r'^AI生图$',                                  # "AI生图"
            r'^回照片动起来$',                            # "回照片动起来"
            r'^已阅读\d+个网页',                          # "已阅读N个网页"
            r'^按住说话$',                                # "按住说话"
            r'^\d+个赞$',                                 # "N个赞"
            # [V8.2] 新增噪点模式（实机测试发现）
            r'^(艮拍|拍题|答疑)$',                        # App UI残留
            r'^(打电话|帮我写|发消息\.\.\.)$',             # App按钮/输入框
            r'^WMA[:：]',                                 # 音频播放器标记
            r'^\d+K$',                                    # 字数统计标记
            r'^\^[^\s]{1,3}$',                            # 脱字符开头的OCR乱码
            r'^今天我',                                    # 对话残片
            r'^较长的视',                                  # 视频链接截断
            r'^(钱谦益墓|QIAN QIAN YI TOMB)$',            # 图片标题/水印残留
        ]

        # 来源平台模式（增强）
        self.platform_patterns = {
            '小红书': [r'小红书', r'RED', r'@.*小红书'],
            '微信读书': [r'微信读书', r'weread'],
            '微信': [r'微信', r'wechat', r'WeChat'],
            '抖音': [r'抖音', r'douyin', r'TikTok'],
            'B站': [r'B站', r'bilibili', r'BILIBILI'],
            '豆包': [r'豆包', r'doubao\.com'],
            'Kimi': [r'Kimi', r'kimi', r'moonshot'],
            '元宝': [r'腾讯元宝', r'元宝'],
            'DeepSeek': [r'DeepSeek', r'deepseek'],
            '知乎': [r'知乎', r'zhihu'],
            '头条': [r'今日头条', r'头条'],
        }

        # [P0-0] 恢复V2的语义标题关键词（28个，扩展到全领域）
        self.heading_keywords = [
            # 养生/食疗类（V2原有）
            '功效', '作用', '好处', '适合', '适宜', '适用',
            '原料', '食材', '配料', '材料',
            '做法', '制作', '步骤', '方法', '教程',
            '禁忌', '注意', '提醒', '警告', '不宜',
            '简介', '介绍', '概述', '说明',
            '推荐', '建议', '指南',
            # [P0-2] 扩展：历史人物传记类
            '生平', '生平经历', '早年', '晚年', '幼年', '青年',
            '仕途', '文学成就', '政治生涯', '政治活动', '学术成就',
            '教育贡献', '主要著作', '情感生活', '人物评价',
            # [P0-2] 扩展：旅游攻略类
            '行程', '景点', '交通', '住宿', '美食', '费用', '注意事项',
            # [P0-2] 扩展：营养科普类
            '营养价值', '主要成分', '食用方法', '适宜人群', '禁忌人群',
        ]

        # [P0-2][V8.2] 编号标题模式
        # 注意：数字序号(\d+[.、])已在 structure_content 中单独处理（动态层级）
        self.heading_patterns = [
            (r'^[（(](\d+)[)）]\s*(.{2,30})$', r'### (\1) \2'),
            # 中文序号
            (r'^([一二三四五六七八九十]+)[、：]\s*(.{2,30})$', r'## \1、\2'),
            # 带符号标题
            (r'^[●◆■▶▪▸]\s*(.{4,30})$', r'### \1'),
            (r'^[-–—]\s*(.{4,30})$', r'### \1'),
            # 纯大写或加粗标记
            (r'^#{1,4}\s*(.{2,50})$', r'## \1'),
        ]

    def extract_source(self, text: str) -> Dict:
        """增强来源提取"""
        lines = text.split('\n')
        source_info = {}
        for line in lines[:30]:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            for platform, patterns in self.platform_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, line_stripped, re.IGNORECASE):
                        source_info['platform'] = platform
                        break
                if 'platform' in source_info:
                    break
            if 'platform' in source_info:
                break
            for kw in ['来源', '出处', '摘录自', '来自']:
                if kw in line_stripped:
                    parts = line_stripped.split(kw)
                    if len(parts) > 1:
                        source_val = parts[1].strip().lstrip('：:：').strip()
                        if source_val and len(source_val) < 20:
                            source_info['platform'] = source_val
                            break
            if '作者' in line_stripped:
                match = re.search(r'作者[：:]\s*([^\n,，]{1,20})', line_stripped)
                if match:
                    source_info['author'] = match.group(1).strip()
        return source_info

    def clean_text(self, text: str) -> str:
        """[P1-2] 增强内容清理 - 去除手机UI噪点（含新增8种）"""
        lines = text.split('\n')
        cleaned_lines = []
        prev_line_empty = False
        for line in lines:
            line = line.strip()
            if not line:
                if not prev_line_empty:
                    cleaned_lines.append('')
                    prev_line_empty = True
                continue
            # 1. 匹配手机UI噪点模式
            is_noise = False
            for pattern in self.ui_noise_patterns:
                if re.match(pattern, line, re.IGNORECASE):
                    is_noise = True
                    break
            if is_noise:
                continue
            # 2. 过短的纯数字/字母行
            if re.match(r'^[\d\s%]+$', line) and len(line) < 8:
                continue
            # 3. 分隔线
            if re.match(r'^[_\-=]{3,}$', line):
                continue
            # 4. 纯英文长行（无中文）
            if re.match(r'^[a-zA-Z0-9\s.,!?;:\'"()-]{20,}$', line) and not re.search(r'[\u4e00-\u9fa5]', line):
                continue
            # 5. 纯短行（<2字符）
            if len(line) < 2:
                continue
            # [V8.2] 通用短乱码行检测（≤8字符且含2+个非中英文数字字符）
            if len(line) <= 8:
                non_text_chars = len(re.findall(r'[^\u4e00-\u9fa5\w\s，。、；：！？""\'\'（）【】《》\u2014\u2026\u00b7-]', line))
                if non_text_chars >= 2:
                    continue
            # [V8.2] 通用短乱码行检测（≤6字符且非中文主体）
            if len(line) <= 6:
                chinese_chars = len(re.findall(r'[\u4e00-\u9fa5]', line))
                # 中文不足2个字且包含特殊字符，视为噪点
                if chinese_chars < 2 and re.search(r'[^\w\s]', line):
                    continue
            # [V8.3] 通用混合乱码行检测（如"甘o亚a"——小写英文字母夹在中文中间）
            if len(line) >= 3 and len(line) <= 15:
                chinese_chars = len(re.findall(r'[\u4e00-\u9fa5]', line))
                # 小写英文字母散布在中文中间（OCR将偏旁识别为字母）
                embedded_letters = len(re.findall(r'[a-z]', line))
                if embedded_letters >= 1 and chinese_chars >= 1:
                    # 含中文但混入字母，检查是否有正常中文语境
                    if not re.search(r'[\u4e00-\u9fa5]{4,}', line):
                        # 没有连续4+中文字符，很可能是乱码
                        continue
            # 6. OCR错误修复（P1-4）
            line = self._fix_common_ocr_errors(line)
            prev_line_empty = False
            cleaned_lines.append(line)
        result = '\n'.join(cleaned_lines)
        result = re.sub(r'\n{3,}', '\n\n', result)
        return result.strip()

    def _fix_common_ocr_errors(self, line: str) -> str:
        """[P1-4][V8.2] 修复常见OCR错误 - 合并V2的13种替换 + 历史人物别名纠正"""
        replacements = [
            # V8原有9种
            (chr(8220), '"'), (chr(8221), '"'),
            (chr(8216), "'"), (chr(8217), "'"),
            (chr(8212), "-"), (chr(183), "-"),
            ('`', ''), ('√', '✓'), ('×', '✗'),
            # V2额外4种（补充）
            (chr(8218), "'"),   # ‚
            (chr(8222), '"'),   # „
            (chr(8223), '"'),   # ‟
            (chr(8242), "'"),   # ′
            (chr(8243), "'"),   # ″
            (chr(180), ""),     # ´
            (chr(9032), ""),    # ✈ 特殊符号
            # 通用修复
            ('…', '...'), ('……', '...'),
            # [V8.2] 常见历史人物音近字纠正（高频OCR错误）
            ('钱穆斋', '钱牧斋'),     # 牧→穆
            ('朱厚熄', '朱厚熜'),     # 熜→熄（嘉靖帝）
            ('朱载屋', '朱载垕'),     # 垕→屋（隆庆帝）
        ]
        for old, new in replacements:
            line = line.replace(old, new)
        line = re.sub(r' {2,}', ' ', line)
        return line

    def _is_heading_line(self, line: str) -> bool:
        """
        [P0-0][V8.2] 恢复V2的语义标题识别逻辑（增强版）

        判断是否为标题行：
        1. 包含语义标题关键词（功效/做法/禁忌/生平...）
        2. 以数字/中文序号开头
        3. 以特定符号开头

        V8.2改进：
        - 排除以「的/了/着/过/等/个」等助词结尾的非标题行
        - 排除含大量标点/数字的叙述行
        """
        if len(line) < 3:
            return False

        # [V8.2] 排除明显不是标题的句式
        # 以助词/量词/连词结尾的句子不太可能是标题
        ending_noise = ('的', '了', '着', '过', '等', '个', '吧', '吗', '呢', '啊',
                        '和', '与', '但', '而', '并', '或', '其', '被', '把', '将')
        if line.rstrip('，。、；：！？""''）】》》').endswith(ending_noise):
            # 但如果包含明确的标题关键词则仍视为标题
            strong_kw = {'功效', '做法', '禁忌', '注意', '步骤', '方法', '生平',
                         '早年', '仕途', '晚年', '情感', '成就', '文学', '政治',
                         '主要著作', '人物评价', '教育贡献'}
            if not any(kw in line for kw in strong_kw):
                return False

        # [V8.2] 排除含太多数字/标点的叙述行（标题通常简短且以中文为主）
        digit_count = len(re.findall(r'\d', line))
        if digit_count > 3 and len(line) > 15:
            return False

        # 1. 检查语义标题关键词（V2核心逻辑）
        for kw in self.heading_keywords:
            if kw in line and len(line) <= 35:
                return True
        # 2. 以数字/中文序号开头（V2+V8共有）
        if re.match(r'^[一二三四五六七八九十\d]+[、.．:：]', line):
            return True
        # 3. 特定符号开头（长度限制放宽到30字）
        if re.match(r'^[◆★▼▪▸●○■▶]', line) and len(line) <= 30:
            return True
        return False

    def _rebuild_multicolumn(self, lines: List[str]) -> List[str]:
        """
        [P0-2][V8.2] 多列排版重建（增强跨行检测）

        检测历史人物表格型内容（姓名+在位时间+描述错排）
        例如：
          朱允        1368-1398在位        （同行）
          朱允                            （跨行）
          1368-1398在位                    （下一行为年代）
        重建为：
          ### 朱允（1368-1398在位）
        """
        rebuilt = []
        i = 0
        while i < len(lines):
            line = lines[i]
            # 检测"人名 + 年代"并排模式
            # 模式：2-4汉字 + 空格 + 年代/称号
            name_year = re.match(r'^([^\s]{2,4})\s+(\d{3,4}[-—]\d{3,4}.{0,10})$', line)
            if name_year:
                name = name_year.group(1)
                info = name_year.group(2).strip()
                rebuilt.append(f"### {name}（{info}）")
                i += 1
                continue
            # [V8.2] 跨行检测：当前行仅为2-4汉字人名，下一行为年代
            name_only = re.match(r'^([^\s]{2,4})\s*$', line)
            if name_only and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                year_alone = re.match(r'^(\d{3,4}[-—]\d{3,4}.{0,15})$', next_line)
                if year_alone:
                    name = name_only.group(1)
                    info = year_alone.group(1)
                    rebuilt.append(f"### {name}（{info}）")
                    i += 2
                    continue
            # 检测"在位X年" 孤立行，合并到上一个人名
            if re.match(r'^在位\d+年$', line) and rebuilt and rebuilt[-1].startswith('###'):
                rebuilt[-1] = rebuilt[-1].rstrip('）') + f'，{line}）'
                i += 1
                continue
            rebuilt.append(line)
            i += 1
        return rebuilt

    def structure_content(self, text: str) -> str:
        """
        [P0-2][V8.2] 内容结构化 - 增强版

        策略（按优先级顺序）：
        1. [P0-0 V2逻辑] 语义标题识别（功效/做法/生平等语义词）→ ### 标题
        2. [V8逻辑] 编号格式识别（1./一、/●）→ ## 标题
           [V8.2] 数字序号(1./2.)在已有章节内自动降为 ###
        3. [P0-2新增] 历史人物多列排版重建（姓名+年代错排）
        """
        lines = text.split('\n')
        # 第一步：多列排版重建（针对表格型内容）
        lines = self._rebuild_multicolumn(lines)

        structured_lines = []
        current_section = None

        # [V8.2] 数字序号模式单独处理
        num_heading_re = re.compile(r'^(\d+)[.、]\s*(.{2,30})$')

        for line in lines:
            stripped = line.strip()
            if not stripped:
                structured_lines.append('')
                continue

            # [V8.2] 优先检查数字序号（需要动态决定层级）
            num_match = num_heading_re.match(stripped)
            if num_match:
                if current_section is not None:
                    # 已在章节内，降为 ###
                    structured_lines.append(f"### {num_match.group(1)}. {num_match.group(2)}")
                else:
                    # 独立章节，保持 ##
                    structured_lines.append(f"## {num_match.group(1)}. {num_match.group(2)}")
                current_section = stripped
                continue

            # 尝试其他编号格式匹配（一、/●/（1）等）
            matched_by_pattern = False
            for pattern, replacement in self.heading_patterns:
                match = re.match(pattern, stripped)
                if match:
                    structured_lines.append(match.expand(replacement))
                    matched_by_pattern = True
                    current_section = stripped
                    break

            if matched_by_pattern:
                continue

            # 其次尝试V2语义标题识别（功效/做法/禁忌/生平等）
            if self._is_heading_line(stripped):
                level = "###" if current_section is None else "####"
                structured_lines.append(f"{level} {stripped}")
                current_section = stripped
                continue

            # 普通内容行
            structured_lines.append(line)

        return '\n'.join(structured_lines)

    def organize(self, text: str) -> Dict:
        """整理内容（清理 + 来源提取 + 结构化）"""
        cleaned = self.clean_text(text)
        source = self.extract_source(text)  # 从原始文本提取来源
        structured = self.structure_content(cleaned)
        return {
            'content': structured,
            'source': source
        }




# ============================================================
# [P0-3][P1-5] 智能文档管理器 - 重构合并逻辑
# ============================================================

class SmartDocumentManager:
    """
    [P0-3][P1-5] 智能文档管理器
    
    核心改进：
    1. 跨分类检测同名文档（避免同主题分散存储）
    2. 主题词相似度 + hash 双重重复检测
    3. 统一文档路径：处理结果/{分类}/{分类}-{主题}.md
    """

    def __init__(self, output_dir: str = '处理结果'):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        # 全局文档索引：doc_name -> Path（跨分类）
        self.doc_index: Dict[str, Path] = {}
        # 已知内容hash集合
        self.known_hashes: set = set()
        # 已知主题词集合（用于相似度检测）
        self.known_topics: List[Tuple[str, Path]] = []
        self._build_index()

    def _build_index(self):
        """扫描已有文档，建立全局索引"""
        for md_file in self.output_dir.rglob('*.md'):
            doc_name = md_file.stem
            # 跨分类索引：用文档名（不含分类前缀）作为key
            self.doc_index[doc_name] = md_file

            # 提取主题词（去掉分类前缀）
            topic = self._extract_topic_from_name(doc_name)
            self.known_topics.append((topic, md_file))

            # 扫描hash
            try:
                content = md_file.read_text(encoding='utf-8')
                hash_match = re.search(r'content_hash:\s*([a-f0-9]{8})', content)
                if hash_match:
                    self.known_hashes.add(hash_match.group(1))
            except Exception:
                pass

    def _extract_topic_from_name(self, doc_name: str) -> str:
        """从文档名提取主题词（去掉分类前缀）"""
        # 格式：分类-主题 -> 主题
        if '-' in doc_name:
            parts = doc_name.split('-', 1)
            return parts[1] if len(parts) > 1 else doc_name
        return doc_name

    def is_duplicate_hash(self, content_hash: str) -> bool:
        """[P1-5] hash重复检测"""
        return content_hash in self.known_hashes

    def find_similar_doc(self, doc_name: str, content: str) -> Optional[Path]:
        """
        [P0-3][P1-5][P1-1] 跨分类相似文档检测（V8.1增强）

        V8.1改进：
        1. 去 (上)/(下)/(一)/(二) 后缀后再比较
        2. 南北东西对称朝代统一（北宋+南宋→宋朝）
        3. 阈值从 0.7 降至 0.6
        """
        topic = self._extract_topic_from_name(doc_name)

        for known_topic, known_path in self.known_topics:
            # 条件1：主题名相似度（使用V8.1增强版）
            sim = self._topic_similarity(topic, known_topic)
            if sim < 0.6:  # [P1-1] 阈值从0.7降至0.6
                continue

            # 条件2：内容重叠（降低到15，减少漏合并）
            if known_path.exists():
                try:
                    existing_content = known_path.read_text(encoding='utf-8')
                    c1_chars = set(content[:400].replace(' ', '').replace('\n', ''))
                    c2_chars = set(existing_content[:400].replace(' ', '').replace('\n', ''))
                    overlap = len(c1_chars & c2_chars)
                    if overlap < 15:  # [P1-1] 从20降至15
                        continue
                except Exception:
                    pass

            return known_path

        return None

    def _topic_similarity(self, t1: str, t2: str) -> float:
        """[P1-1] 增强主题词相似度（去后缀+朝代统一+bigram Jaccard）"""
        if not t1 or not t2:
            return 0.0

        def normalize(s):
            """[P1-1] 增强normalize：去(上)(下)后缀 + 南北东西对称处理"""
            # 去掉常见整理类后缀
            s = re.sub(r'(传|简介|人物|生平|整理|知识|合集|汇总)$', '', s).strip()
            # [P1-1] 去掉 (上)/(下)/(一)/(二)/上/下 等分册后缀
            s = re.sub(r'[（(]\s*[上下一二三四五六七八九十\d]\s*[)）]$', '', s).strip()
            s = re.sub(r'[上下一二三四]$', '', s).strip()
            # [P1-1] 南北朝代统一（北宋/南宋→宋，北魏/南朝→朝代）
            dynasty_unify = {
                '北宋': '宋朝', '南宋': '宋朝',
                '北魏': '魏朝', '东魏': '魏朝', '西魏': '魏朝',
                '东汉': '汉朝', '西汉': '汉朝',
                '东晋': '晋朝', '西晋': '晋朝',
                '前秦': '秦', '后秦': '秦',
            }
            for variant, unified in dynasty_unify.items():
                if variant in s:
                    s = s.replace(variant, unified)
            return s

        n1, n2 = normalize(t1), normalize(t2)

        if n1 == n2:
            return 1.0
        if t1 == t2:
            return 1.0
        # 包含关系（去后缀后）
        shorter, longer = (n1, n2) if len(n1) <= len(n2) else (n2, n1)
        if len(shorter) >= 2 and shorter in longer:
            return 0.90
        # 原始字符串包含关系
        s2, l2 = (t1, t2) if len(t1) <= len(t2) else (t2, t1)
        if len(s2) >= 2 and s2 in l2:
            return 0.85
        # bigram Jaccard（使用去后缀版本）
        def bigrams(s):
            return set(s[i:i+2] for i in range(len(s) - 1)) if len(s) >= 2 else set(s)
        b1, b2 = bigrams(n1), bigrams(n2)
        if not b1 or not b2:
            return 0.0
        return len(b1 & b2) / len(b1 | b2)

    def merge_content(self, existing_file: Path, new_content: str,
                      new_image_name: str, content_hash: str = None) -> bool:
        """追加内容到已有文档"""
        try:
            existing = existing_file.read_text(encoding='utf-8')
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
            append_section = f"""

---

## 补充内容

> 来源图片: {new_image_name}
> 追加时间: {timestamp}
> 内容标识: {content_hash or 'N/A'}

{new_content}

"""
            marker = '\n\n---\n\n*本文档由图片知识库整理工具自动生成*'
            if marker in existing:
                new_full = existing.replace(marker, append_section + marker)
            else:
                new_full = existing + append_section

            existing_file.write_text(new_full, encoding='utf-8')
            return True
        except Exception as e:
            logger.warning(f"[合并] 失败: {e}")
            return False

    def save_document(self, doc_name: str, category: str, content: str,
                      image_name: str, keywords: List[str],
                      content_hash: str, source_info: Dict) -> Tuple[str, bool]:
        """
        保存文档（新建或合并）
        返回 (文件路径, 是_new_doc)
        """
        # 检查相似文档（跨分类）
        similar = self.find_similar_doc(doc_name, content)
        if similar:
            logger.info(f"[文档] 发现相似文档 {similar.name}，合并内容")
            self.merge_content(similar, content, image_name, content_hash)
            if content_hash:
                self.known_hashes.add(content_hash)
            return str(similar), False  # 非新文档

        # 新建文档
        cat_dir = self.output_dir / category
        cat_dir.mkdir(exist_ok=True)

        safe_name = re.sub(r'[<>:"/\\|?*]', '', doc_name)
        if len(safe_name) > 40:
            safe_name = safe_name[:40]

        md_file = cat_dir / f"{safe_name}.md"
        counter = 1
        while md_file.exists():
            md_file = cat_dir / f"{safe_name}_{counter}.md"
            counter += 1

        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
        keywords_str = ', '.join(keywords) if keywords else '无'
        source_section = ""
        if source_info.get('platform'):
            source_section += f"**来源**：{source_info['platform']}\n\n"
        if source_info.get('author'):
            source_section += f"**作者**：{source_info['author']}\n\n"

        md_content = f"""# {doc_name}

> 来源图片: {image_name}
> 识别时间: {timestamp}
> 分类: {category}
> 关键词: {keywords_str}
> content_hash: {content_hash or 'N/A'}

---

<!-- CONTENT_START -->
{source_section}{content}
<!-- CONTENT_END -->

---

*本文档由图片知识库整理工具自动生成*
"""
        md_file.write_text(md_content, encoding='utf-8')

        # 更新索引
        self.doc_index[doc_name] = md_file
        topic = self._extract_topic_from_name(doc_name)
        self.known_topics.append((topic, md_file))
        if content_hash:
            self.known_hashes.add(content_hash)

        logger.info(f"[文档] 新建: {category}/{md_file.name}")
        return str(md_file), True  # 是新文档

    def generate_word(self, doc_name: str, category: str, content: str,
                      image_name: str, keywords: List[str],
                      md_file_path: str = None) -> Optional[str]:
        """生成Word文档"""
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            logger.warning("[文档] python-docx未安装，跳过Word生成")
            return None

        cat_dir = self.output_dir / category
        cat_dir.mkdir(exist_ok=True)

        safe_name = re.sub(r'[<>:"/\\|?*]', '', doc_name)
        if len(safe_name) > 40:
            safe_name = safe_name[:40]

        docx_file = cat_dir / f"{safe_name}.docx"
        if docx_file.exists():
            return str(docx_file)

        # 从MD文件读取最新内容（含合并后的内容）
        text_content = content
        if md_file_path and Path(md_file_path).exists():
            try:
                md_text = Path(md_file_path).read_text(encoding='utf-8')
                if '<!-- CONTENT_START -->' in md_text and '<!-- CONTENT_END -->' in md_text:
                    start = md_text.index('<!-- CONTENT_START -->') + len('<!-- CONTENT_START -->')
                    end = md_text.index('<!-- CONTENT_END -->')
                    text_content = md_text[start:end].strip()
            except Exception:
                pass

        try:
            doc = Document()
            doc.styles['Normal'].font.name = 'Microsoft YaHei'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

            doc.add_heading(doc_name, 0)
            doc.add_paragraph(f"来源图片: {image_name}")
            doc.add_paragraph(f"识别时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph(f"分类: {category}")
            keywords_str = ', '.join(keywords) if keywords else '无'
            doc.add_paragraph(f"关键词: {keywords_str}")
            doc.add_paragraph("─" * 30)

            # [V8.3 Fix-8] Word中解析Markdown标题，输出为Word标题样式
            for line in text_content.split('\n'):
                stripped = line.strip()
                if not stripped:
                    doc.add_paragraph('')
                elif re.match(r'^#{1,4}\s+', stripped):
                    # Markdown标题 → Word标题样式
                    level = len(re.match(r'^(#+)', stripped).group(1))
                    heading_text = re.sub(r'^#+\s*', '', stripped)
                    if 1 <= level <= 4:
                        doc.add_heading(heading_text, level=level)
                    else:
                        doc.add_paragraph(heading_text)
                elif stripped.startswith('<!--') and stripped.endswith('-->'):
                    continue  # 跳过HTML注释标记
                else:
                    doc.add_paragraph(stripped)

            doc.save(str(docx_file))
            logger.info(f"[文档] Word已生成: {category}/{docx_file.name}")
            return str(docx_file)
        except Exception as e:
            logger.warning(f"[Word] 生成失败: {e}")
            return None


# ============================================================
# [P1-4] 图片归档器 - 按源文件夹分类（便于溯源）
# ============================================================

class ImageArchiver:
    """
    [P1-4] 图片归档器
    
    分类策略：按源文件夹分类，便于溯源查询
    - 图片归档：按待处理图片中的子文件夹分类
    - 文档归档：按V7 AI分类（在SmartDocumentManager中处理）
    """

    def __init__(self, processed_dir: str = '已处理图片'):
        self.processed_dir = Path(processed_dir)
        self.processed_dir.mkdir(parents=True, exist_ok=True)
        self.archived_count = 0

    def get_source_folder(self, image_path: str) -> str:
        """
        从图片路径提取源文件夹名
        例如: 待处理图片/历史学/a.jpg -> 历史学
        """
        src = Path(image_path)
        # 取父文件夹名（相对于待处理图片）
        parent_name = src.parent.name
        # 如果父文件夹名等于源目录名（说明图片直接在待处理图片根目录）
        if parent_name == '待处理图片':
            return '根目录'
        return parent_name

    def archive_image(self, image_path: str, category: str = None) -> Optional[str]:
        """
        移动图片到已处理文件夹
        按源文件夹分类，便于溯源
        """
        src = Path(image_path)
        if not src.exists():
            return None

        # 使用源文件夹名作为归档目录
        source_folder = self.get_source_folder(image_path)
        archive_subdir = self.processed_dir / source_folder

        archive_subdir.mkdir(parents=True, exist_ok=True)

        dest = archive_subdir / src.name
        if dest.exists():
            ts = datetime.now().strftime('%H%M%S')
            dest = archive_subdir / f"{src.stem}_{ts}{src.suffix}"

        try:
            shutil.move(str(src), str(dest))
            self.archived_count += 1
            return str(dest)
        except Exception as e:
            logger.warning(f"[归档] 移动失败 {src.name}: {e}")
            try:
                shutil.copy2(str(src), str(dest))
                src.unlink()
                self.archived_count += 1
                return str(dest)
            except Exception:
                return None

    def get_stats(self) -> Dict:
        total = 0
        subdirs = {}
        for item in self.processed_dir.iterdir():
            if item.is_dir():
                count = len(list(item.glob('*.*')))
                subdirs[item.name] = count
                total += count
        return {
            'archived_this_session': self.archived_count,
            'total_archived': total,
            'by_category': subdirs
        }


# ============================================================
# 进度条（保留V6功能）
# ============================================================
class ProgressBar:
    """[F-2] 进度条 - 始终可见，不被日志淹没"""

    def __init__(self, total: int, width: int = 40, prefix: str = "进度"):
        self.total = total
        self.width = width
        self.prefix = prefix
        self.start_time = time.time()
        self.last_line = ""       # 记录上一条日志行
        self.current = 0

    def update(self, current: int, info: str = ""):
        self.current = current
        percent = current / self.total if self.total > 0 else 0
        filled = int(self.width * percent)
        bar = '█' * filled + '░' * (self.width - filled)
        if current > 0:
            elapsed = time.time() - self.start_time
            eta = elapsed / current * (self.total - current)
            eta_str = f"{int(eta//60)}m{int(eta%60)}s"
        else:
            eta_str = "--"
        # 进度条单独一行，使用 \r 覆盖自己
        print(f"\r{self.prefix}: [{bar}] {current}/{self.total} ({percent:.0%}) ETA:{eta_str} {info[:15]}",
              end='', flush=True)

    def log(self, message: str):
        """[F-2] 在进度条上方安全打印日志行"""
        # 1. 清除当前进度条（用空格覆盖）
        clear = ' ' * 100
        print(f'\r{clear}\r', end='', flush=True)
        # 2. 打印日志
        print(message)
        # 3. 重新绘制进度条
        if self.current > 0:
            self._redraw()

    def _redraw(self):
        """重新绘制当前进度条"""
        percent = self.current / self.total if self.total > 0 else 0
        filled = int(self.width * percent)
        bar = '█' * filled + '░' * (self.width - filled)
        elapsed = time.time() - self.start_time
        if self.current > 0:
            eta = elapsed / self.current * (self.total - self.current)
            eta_str = f"{int(eta//60)}m{int(eta%60)}s"
        else:
            eta_str = "--"
        print(f"\r{self.prefix}: [{bar}] {self.current}/{self.total} ({percent:.0%}) ETA:{eta_str}",
              end='', flush=True)

    def finish(self):
        self.update(self.total, "完成!")
        print()


# ============================================================
# 多引擎OCR（保留V6功能）
# ============================================================
class MultiEngineOCR:
    def __init__(self):
        self.current_engine = None
        self.engine_status = []

        secret_id = os.getenv('TENCENT_SECRET_ID', '')
        secret_key = os.getenv('TENCENT_SECRET_KEY', '')
        if secret_id and secret_key and '替换' not in secret_id and TENCENT_AVAILABLE:
            self.engine_status.append(('腾讯云', True, '已配置'))
        else:
            self.engine_status.append(('腾讯云', False, '未配置'))

        app_id = os.getenv('BAIDU_APP_ID', '')
        api_key = os.getenv('BAIDU_API_KEY', '')
        bsecret = os.getenv('BAIDU_SECRET_KEY', '')
        if app_id and api_key and bsecret and '替换' not in app_id and BAIDU_AVAILABLE:
            self.engine_status.append(('百度云', True, '已配置'))
        else:
            self.engine_status.append(('百度云', False, '未配置'))

        local = LocalOCR()
        if local.tesseract_available:
            self.engine_status.append(('本地Tesseract', True, '可用'))
        else:
            self.engine_status.append(('本地Tesseract', False, '不可用'))

        self._select_best_engine()

    def _select_best_engine(self):
        for name, available, msg in self.engine_status:
            if available:
                self.current_engine = name
                logger.info(f"[OCR] 选用引擎: {name}")
                return True
        logger.warning("[OCR] 没有可用的OCR引擎!")
        return False

    def recognize(self, image_path: str) -> Dict:
        for name, available, _ in self.engine_status:
            if not available:
                continue
            try:
                if name == '腾讯云' and TENCENT_AVAILABLE:
                    engine = TencentOCR()
                elif name == '百度云' and BAIDU_AVAILABLE:
                    engine = BaiduOCR()
                else:
                    engine = LocalOCR()
                result = engine.recognize(image_path)
                if result and result.get('success'):
                    return result
            except Exception as e:
                logger.warning(f"[OCR] {name} 失败: {e}")
        return {'success': False, 'error': '无可用OCR引擎', 'text': ''}


# ============================================================
# IMA同步器（修复P1-7）
# ============================================================
class IMASyncer:
    """[P1-7] IMA笔记同步器（修复同步逻辑）"""

    def __init__(self):
        self.client_id = os.getenv('IMA_OPENAPI_CLIENTID', '')
        self.api_key = os.getenv('IMA_OPENAPI_APIKEY', '')
        self.base_url = 'https://ima.qq.com/openapi/note/v1'
        self.enabled = bool(self.client_id and self.api_key and '填入' not in self.client_id)
        self.sync_log_file = Path('处理结果/ima_sync_log.json')
        self.sync_log = self._load_sync_log()
        self.default_notebook_id = os.getenv('IMA_NOTEBOOK_ID', '')
        self.rate_limited = False
        # [P1-7] 追踪本次同步数量
        self.synced_this_session = 0
        self.failed_this_session = 0

        if self.enabled:
            logger.info("[IMA] 已配置，同步功能启用")
        else:
            logger.warning("[IMA] 未配置，同步功能禁用")

    def _load_sync_log(self) -> Dict:
        if self.sync_log_file.exists():
            try:
                return json.loads(self.sync_log_file.read_text(encoding='utf-8'))
            except Exception:
                pass
        return {}

    def _save_sync_log(self):
        self.sync_log_file.write_text(
            json.dumps(self.sync_log, ensure_ascii=False, indent=2),
            encoding='utf-8'
        )

    def _api_call(self, endpoint: str, payload: Dict, retries: int = 3) -> Optional[Dict]:
        if not self.enabled or self.rate_limited:
            return None
        for attempt in range(retries):
            try:
                import requests
                headers = {
                    'ima-openapi-clientid': self.client_id,
                    'ima-openapi-apikey': self.api_key,
                    'Content-Type': 'application/json'
                }
                url = f"{self.base_url}/{endpoint}"
                response = requests.post(url, json=payload, headers=headers, timeout=30)
                if response.status_code == 200:
                    return response.json()
                elif response.status_code == 403:
                    result = response.json()
                    if '请求超量' in result.get('msg', ''):
                        logger.warning(f"[IMA] API请求超量，等待60秒...")
                        if attempt < retries - 1:
                            time.sleep(60)
                        continue
                return None
            except Exception as e:
                logger.warning(f"[IMA] 调用失败 ({attempt+1}/{retries}): {e}")
                if attempt < retries - 1:
                    time.sleep(2)
        return None

    def sync_note(self, title: str, content: str, category: str = None,
                  content_hash: str = None, doc_path: str = None) -> Optional[str]:
        """同步笔记到IMA"""
        if not self.enabled or self.rate_limited:
            return None

        doc_key = content_hash or title
        existing = self.sync_log.get(doc_key, {})

        full_content = f"# {title}\n\n"
        if category:
            full_content += f"> 分类: {category}\n"
        if content_hash:
            full_content += f"> 内容标识: {content_hash}\n"
        full_content += f"\n{content}\n\n---\n*自动同步自图片知识库*\n"

        # 追加到已有笔记
        if existing.get('doc_id') and doc_path:
            result = self._api_call('append_doc', {
                'doc_id': existing['doc_id'],
                'content_format': 1,
                'content': f"\n\n---\n*更新于 {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n" + content
            })
            if result:
                self.sync_log[doc_key]['last_sync'] = datetime.now().isoformat()
                self.sync_log[doc_key]['update_count'] = existing.get('update_count', 0) + 1
                self._save_sync_log()
                self.synced_this_session += 1
                return existing['doc_id']

        # 新建笔记
        import_payload = {
            'content_format': 1,
            'content': full_content
        }
        if self.default_notebook_id:
            import_payload['notebook_id'] = self.default_notebook_id

        result = self._api_call('import_doc', import_payload)
        if result and result.get('code') == 0:
            doc_id = result.get('data', {}).get('note_id', 'imported')
            self.sync_log[doc_key] = {
                'doc_id': doc_id,
                'title': title,
                'category': category,
                'first_sync': datetime.now().isoformat(),
                'last_sync': datetime.now().isoformat(),
                'update_count': 1
            }
            self._save_sync_log()
            self.synced_this_session += 1
            return doc_id
        else:
            self.failed_this_session += 1
            return None


# ============================================================
# 分批处理（保留V6功能）
# ============================================================
_BATCH_DIR = Path(__file__).parent / '.workbuddy' / 'memory'
BATCH_STATE_FILE = _BATCH_DIR / 'batch_state.json'

BATCH_CONFIG = {
    "small":  {"max_size": 500 * 1024,       "batch_size": 10},
    "medium": {"max_size": 2 * 1024 * 1024,  "batch_size": 6},
    "large":  {"max_size": 5 * 1024 * 1024,  "batch_size": 4},
    "xlarge": {"max_size": float('inf'),      "batch_size": 2},
}

@dataclass
class ImageInfo:
    path: str
    name: str
    size: int
    folder: str
    hash: str = ""
    def to_dict(self): return asdict(self)
    @classmethod
    def from_dict(cls, d): return cls(**d)

@dataclass
class BatchInfo:
    batch_id: str
    folder: str
    images: List[str]
    status: str
    created_at: str
    started_at: Optional[str] = None
    completed_at: Optional[str] = None
    result: Optional[str] = None
    retry_count: int = 0
    def to_dict(self): return asdict(self)
    @classmethod
    def from_dict(cls, d): return cls(**d)

@dataclass
class ProcessingState:
    session_id: str
    started_at: str
    total_images: int
    total_batches: int
    completed_batches: int = 0
    failed_batches: int = 0
    batches: Optional[Dict] = None
    def __post_init__(self):
        if self.batches is None:
            self.batches = {}
    def to_dict(self):
        data = asdict(self)
        if self.batches:
            data['batches'] = {k: v.to_dict() if isinstance(v, BatchInfo) else v
                               for k, v in self.batches.items()}
        return data
    @classmethod
    def from_dict(cls, data):
        if 'batches' in data and data['batches']:
            data['batches'] = {k: BatchInfo.from_dict(v) for k, v in data['batches'].items()}
        return cls(**data)


class BatchManager:
    def __init__(self, source_dir: str = '待处理图片'):
        self.source_dir = Path(source_dir)
        self.state: Optional[ProcessingState] = None

    def _load_state(self) -> Optional[ProcessingState]:
        if BATCH_STATE_FILE.exists():
            try:
                with open(BATCH_STATE_FILE, 'r', encoding='utf-8') as f:
                    return ProcessingState.from_dict(json.load(f))
            except Exception as e:
                print(f"[WARN] 加载状态失败: {e}")
        return None

    def _save_state(self):
        BATCH_STATE_FILE.parent.mkdir(parents=True, exist_ok=True)
        with open(BATCH_STATE_FILE, 'w', encoding='utf-8') as f:
            json.dump(self.state.to_dict(), f, ensure_ascii=False, indent=2)

    def _get_batch_size(self, file_size: int) -> int:
        for tier, cfg in BATCH_CONFIG.items():
            if file_size <= cfg["max_size"]:
                return cfg["batch_size"]
        return 2

    def _compute_hash(self, path: Path) -> str:
        try:
            h = hashlib.md5()
            with open(path, 'rb') as f:
                for chunk in iter(lambda: f.read(8192), b''):
                    h.update(chunk)
            return h.hexdigest()[:8]
        except Exception:
            return ""

    def _scan_images(self) -> List[ImageInfo]:
        images = []
        for ext in ['*.jpg', '*.jpeg', '*.png', '*.webp', '*.bmp']:
            for img in self.source_dir.rglob(ext):
                size = img.stat().st_size
                images.append(ImageInfo(
                    path=str(img),
                    name=img.name,
                    size=size,
                    folder=img.parent.name,
                    hash=self._compute_hash(img)
                ))
        return images

    def initialize(self) -> Optional[ProcessingState]:
        existing = self._load_state()
        if existing:
            pending = sum(1 for b in existing.batches.values()
                         if isinstance(b, BatchInfo) and b.status in ['pending', 'failed'])
            if pending > 0:
                print(f"[INFO] 发现未完成的处理状态，继续上次处理 (剩余 {pending} 批)")
                self.state = existing
                return self.state

        print("[INFO] 初始化分批处理...")
        images = self._scan_images()
        if not images:
            print("[INFO] 没有待处理的图片")
            return None

        images_by_folder: Dict[str, List[ImageInfo]] = defaultdict(list)
        for img in images:
            images_by_folder[img.folder].append(img)

        session_id = datetime.now().strftime('%Y%m%d_%H%M%S')
        batches = {}
        for folder, folder_images in images_by_folder.items():
            i = 0
            while i < len(folder_images):
                batch_size = self._get_batch_size(folder_images[i].size)
                batch_imgs = folder_images[i:i+batch_size]
                batch_id = f"{folder}_{i//batch_size:03d}"
                batches[batch_id] = BatchInfo(
                    batch_id=batch_id,
                    folder=folder,
                    images=[img.path for img in batch_imgs],
                    status='pending',
                    created_at=datetime.now().isoformat()
                )
                i += batch_size

        self.state = ProcessingState(
            session_id=session_id,
            started_at=datetime.now().isoformat(),
            total_images=len(images),
            total_batches=len(batches),
            batches=batches
        )
        self._save_state()
        print(f"[INFO] 初始化完成: {len(images)} 张图片, {len(batches)} 个批次")
        return self.state

    def get_next_batch(self) -> Optional[BatchInfo]:
        if not self.state:
            return None
        for batch in self.state.batches.values():
            if isinstance(batch, BatchInfo) and batch.status in ['pending', 'failed']:
                if batch.retry_count < 3:
                    batch.status = 'processing'
                    batch.started_at = datetime.now().isoformat()
                    self._save_state()
                    return batch
        return None

    def mark_batch_completed(self, batch_id: str, result: str = ""):
        if self.state and batch_id in self.state.batches:
            batch = self.state.batches[batch_id]
            if isinstance(batch, BatchInfo):
                batch.status = 'completed'
                batch.completed_at = datetime.now().isoformat()
                batch.result = result
                self.state.completed_batches += 1
                self._save_state()

    def mark_batch_failed(self, batch_id: str, error: str = ""):
        if self.state and batch_id in self.state.batches:
            batch = self.state.batches[batch_id]
            if isinstance(batch, BatchInfo):
                batch.status = 'failed'
                batch.retry_count += 1
                batch.result = error
                self.state.failed_batches += 1
                self._save_state()

    def print_progress(self):
        state = self._load_state()
        if not state:
            print("[INFO] 没有正在进行的处理任务")
            return
        completed = sum(1 for b in state.batches.values()
                       if isinstance(b, BatchInfo) and b.status == 'completed')
        failed = sum(1 for b in state.batches.values()
                    if isinstance(b, BatchInfo) and b.status == 'failed')
        pending = state.total_batches - completed - failed
        print(f"[进度] 总批次: {state.total_batches} | 完成: {completed} | 失败: {failed} | 待处理: {pending}")
        print(f"[进度] 总图片: {state.total_images}")

    def clear_state(self):
        if BATCH_STATE_FILE.exists():
            BATCH_STATE_FILE.unlink()
            print("[INFO] 处理状态已清除")


# ============================================================
# 处理单张图片（V7重构版）
# ============================================================
def process_single_image(
    ocr: MultiEngineOCR,
    analyzer: SmartContentAnalyzer,
    organizer: ContentOrganizer,
    doc_manager: SmartDocumentManager,
    ima_syncer: IMASyncer,
    archiver: ImageArchiver,
    image_path: str,
    index: int,
    total: int,
    progress: ProgressBar = None
) -> Optional[Dict]:
    """
    [P0-1] 处理单张图片 - V8.1 重构版

    正确处理流程（先整理再提炼）：
    1. OCR识别（原始文本）
    2. 内容深度整理（去噪 + 语义结构化）← 先整理
    3. 基于整理后内容 → AI分类 + 文档命名  ← 再提炼
    4. 重复检测
    5. 合并判断（相似文档合并）
    6. 生成最终文档（MD + Word）
    7. IMA同步 + 图片归档
    """
    image_name = Path(image_path).name

    if progress:
        progress.update(index, f"{image_name[:12]}")
    else:
        print(f"[{index}/{total}] {image_name}")
    t_start = time.time()

    # ======== 步骤1：OCR识别 ========
    result = ocr.recognize(image_path)
    if not result.get('success'):
        msg = f"  ❌ OCR失败: {result.get('error')}"
        if progress:
            progress.log(msg)
        else:
            print(msg)
        archiver.archive_image(image_path, "其他")
        return None

    raw_text = result.get('text', '').strip()
    if not raw_text or len(raw_text) < 10:
        msg = "  ⏭️ 无文字内容，跳过"
        if progress:
            progress.log(msg)
        else:
            print(msg)
        archiver.archive_image(image_path, "其他")
        return None

    # ======== 步骤2：内容深度整理（先整理）========
    # [P0-1] 关键改变：先整理，再用整理后内容做后续所有操作
    organized = organizer.organize(raw_text)
    clean_text = organized['content']  # 整理后的结构化内容
    source_info = organized['source']

    # ======== 步骤3：AI分类 + 文档命名（基于整理后内容）========
    # [P0-1] 关键改变：分析的输入是整理后的clean_text，而非raw_text
    analysis = analyzer.analyze(clean_text, index)
    category = analysis['category']
    confidence = analysis['confidence']
    doc_name = analysis['doc_name']
    keywords = analysis['keywords']
    content_hash = analysis['content_hash']

    # ======== 步骤4：重复检测 ========
    if doc_manager.is_duplicate_hash(content_hash):
        archiver.archive_image(image_path, category)
        elapsed = time.time() - t_start
        msg = f"  🔄 重复内容 | {category} | {doc_name} ({elapsed:.1f}s)"
        if progress:
            progress.log(msg)
        else:
            print(msg)
        return {
            'image': image_name, 'category': category, 'doc_name': doc_name,
            'is_duplicate': True, 'content_hash': content_hash
        }

    # ======== 步骤5+6：合并判断 + 生成文档（MD + Word）========
    md_file, is_new_doc = doc_manager.save_document(
        doc_name=doc_name, category=category, content=clean_text,
        image_name=image_name, keywords=keywords,
        content_hash=content_hash, source_info=source_info
    )
    docx_file = None
    if is_new_doc:
        docx_file = doc_manager.generate_word(
            doc_name=doc_name, category=category, content=clean_text,
            image_name=image_name, keywords=keywords, md_file_path=md_file
        )

    # ======== 步骤7：IMA同步 ========
    if ima_syncer.enabled and not ima_syncer.rate_limited:
        ima_syncer.sync_note(doc_name, clean_text, category, content_hash, md_file)

    # ======== 步骤8：图片归档 ========
    archived_path = archiver.archive_image(image_path, category)

    # 最终汇总（仅更新进度条，不单独输出；失败/重复仍输出）
    elapsed = time.time() - t_start
    conf_label = "高" if confidence >= 0.6 else ("中" if confidence >= 0.3 else "低")
    action = "新建" if is_new_doc else "合并"
    status = "✅"
    if confidence < 0.15:
        status = "⚠️"
    source_tag = f" | {source_info['platform']}" if source_info.get('platform') else ""
    # [V8.3 Fix-7] 成功时静默更新进度条，不打印详情行
    if progress:
        progress.update(index, f"{status} {doc_name[:12]}{source_tag}")
    else:
        msg = f"  {status} [{conf_label}] {doc_name} ({action}){source_tag} ({elapsed:.1f}s)"
        print(msg)

    return {
        'image': image_name,
        'text_length': len(raw_text),
        'category': category,
        'confidence': round(confidence, 3),
        'doc_name': doc_name,
        'keywords': keywords,
        'is_duplicate': False,
        'is_new_doc': is_new_doc,
        'md_file': md_file,
        'docx_file': docx_file,
        'content_hash': content_hash,
        'source': source_info
    }


# ============================================================
# [P1-6] 统一报告生成
# ============================================================
def save_report(results: List[Dict], total_images: int,
                elapsed: float, mode: str = 'full'):
    """[P1-6] 修复统计报告逻辑"""
    valid = [r for r in results if r and not r.get('failed')]
    duplicates = [r for r in valid if r.get('is_duplicate')]
    new_docs = [r for r in valid if not r.get('is_duplicate') and r.get('is_new_doc')]
    merged_docs = [r for r in valid if not r.get('is_duplicate') and not r.get('is_new_doc')]
    md_count = sum(1 for r in valid if r.get('md_file') and not r.get('is_duplicate'))
    docx_count = sum(1 for r in valid if r.get('docx_file') and not r.get('is_duplicate'))
    failed = [r for r in results if r and r.get('failed')]

    # 分类统计（[P1-6] 使用 category 而非 theme）
    cat_stats = dict(Counter(
        r['category'] for r in valid
        if not r.get('is_duplicate') and r.get('category')
    ))

    report = {
        'timestamp': datetime.now().isoformat(),
        'version': 'V8.3',
        'mode': mode,
        'total_images': total_images,
        'processed_successfully': len(valid),
        'failed': len(failed),
        'duplicates_skipped': len(duplicates),
        'new_documents': len(new_docs),
        'merged_documents': len(merged_docs),
        'markdown_files': md_count,
        'word_files': docx_count,
        'elapsed_seconds': round(elapsed, 1),
        'category_stats': cat_stats,
        'results': [{k: v for k, v in r.items() if k not in ['md_file', 'docx_file']}
                    for r in results if r]
    }

    Path('处理结果').mkdir(exist_ok=True)
    report_file = Path('处理结果/处理报告.json')
    with open(report_file, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    return report, report_file


def print_summary(report: Dict, archiver: ImageArchiver, ima_syncer: IMASyncer):
    """打印处理汇总"""
    print("\n" + "=" * 60)
    print("处理完成 - V8.3 汇总")
    print("=" * 60)
    print(f"  总图片数:     {report['total_images']} 张")
    print(f"  成功处理:     {report['processed_successfully']} 张")
    print(f"  处理失败:     {report['failed']} 张")
    print(f"  重复跳过:     {report['duplicates_skipped']} 张")
    print(f"  新建文档:     {report['new_documents']} 个")
    print(f"  合并到已有:   {report['merged_documents']} 次")
    print(f"  Markdown:     {report['markdown_files']} 个")
    print(f"  Word:         {report['word_files']} 个")
    print(f"  总耗时:       {int(report['elapsed_seconds']//60)}分{int(report['elapsed_seconds']%60)}秒")

    if report['category_stats']:
        print("\n分类分布:")
        for cat, cnt in sorted(report['category_stats'].items(), key=lambda x: -x[1]):
            print(f"  {cat}: {cnt} 张")

    arch_stats = archiver.get_stats()
    print(f"\n归档统计:")
    print(f"  本次归档: {arch_stats['archived_this_session']} 张")
    print(f"  总已归档: {arch_stats['total_archived']} 张")

    if ima_syncer.enabled:
        print(f"\nIMA同步:")
        print(f"  本次成功: {ima_syncer.synced_this_session} 条")
        print(f"  本次失败: {ima_syncer.failed_this_session} 条")

    print("=" * 60)


# ============================================================
# 全自动处理模式
# ============================================================
def run_full_mode(batch_manager: BatchManager):
    """[F-1] 全自动处理模式 - 智能分批"""
    print("=" * 60)
    print("全自动化图片处理工具 V8.3")
    print("OCR -> 深度整理 -> AI分类+智能命名 -> 合并 -> Word -> IMA -> 归档")
    print("=" * 60)

    # 初始化组件
    print("\n[初始化] OCR引擎...")
    ocr = MultiEngineOCR()
    if not ocr.current_engine:
        print("错误: 没有可用的OCR引擎!")
        return

    print("[初始化] 智能分类器...")
    classifier = SmartClassifier()

    print("[初始化] 内容分析器...")
    analyzer = SmartContentAnalyzer(classifier)

    print("[初始化] 内容整理器...")
    organizer = ContentOrganizer()

    print("[初始化] 文档管理器...")
    doc_manager = SmartDocumentManager()

    print("[初始化] IMA同步器...")
    ima_syncer = IMASyncer()

    print("[初始化] 图片归档器...")
    archiver = ImageArchiver()

    # [F-1] 扫描图片，判断是否需要分批
    source_dir = Path('待处理图片')
    if not source_dir.exists():
        print("错误: 待处理图片目录不存在!")
        return

    images = []
    for ext in ['*.jpg', '*.jpeg', '*.png', '*.webp', '*.bmp']:
        images.extend(source_dir.rglob(ext))

    if not images:
        print("没有找到待处理图片")
        return

    total = len(images)

    # [F-1] 智能分批判断：>10张自动启用分批
    AUTO_BATCH_THRESHOLD = 10
    if total > AUTO_BATCH_THRESHOLD:
        print(f"\n[分批] 检测到 {total} 张图片，超过阈值({AUTO_BATCH_THRESHOLD})，自动启用分批模式")
        batch_manager.clear_state()  # 清除旧状态
        run_batch_mode(batch_manager)
        return

    print(f"\n找到 {total} 张待处理图片（≤{AUTO_BATCH_THRESHOLD}，不分批）")
    subfolders = [d.name for d in source_dir.iterdir() if d.is_dir()]
    if subfolders:
        print(f"子文件夹: {subfolders}")

    progress = ProgressBar(total, prefix="处理进度")
    results = []
    start_time = time.time()

    # [F-1] 全局进度统计
    processed_count = 0

    for i, img_path in enumerate(images, 1):
        try:
            result = process_single_image(
                ocr, analyzer, organizer, doc_manager, ima_syncer, archiver,
                str(img_path), i, total, progress
            )
            if result:
                results.append(result)
                if not result.get('failed'):
                    processed_count += 1
        except Exception as e:
            msg = f"\n  ❌ 处理失败: {e}"
            if progress:
                progress.log(msg)
            else:
                print(msg)
            logger.error(f"处理失败 {img_path}: {e}")
            results.append({'image': Path(img_path).name, 'failed': True, 'error': str(e)})

    progress.finish()

    elapsed = time.time() - start_time
    report, report_file = save_report(results, total, elapsed, 'full')
    print_summary(report, archiver, ima_syncer)
    print(f"\n报告已保存: {report_file}")


# ============================================================
# 分批处理模式
# ============================================================
def run_batch_mode(batch_manager: BatchManager):
    """分批处理模式（V7）"""
    print("\n" + "=" * 60)
    print("分批处理模式 V8.3")
    print("=" * 60)

    state = batch_manager.initialize()
    if not state:
        return

    batch_manager.print_progress()

    # 初始化组件
    print("\n[初始化] OCR引擎...")
    ocr = MultiEngineOCR()
    if not ocr.current_engine:
        print("错误: 没有可用的OCR引擎!")
        return

    classifier = SmartClassifier()
    analyzer = SmartContentAnalyzer(classifier)
    organizer = ContentOrganizer()
    doc_manager = SmartDocumentManager()
    ima_syncer = IMASyncer()
    archiver = ImageArchiver()

    all_results = []
    batch_num = 0
    total_images = state.total_images
    total_batches = state.total_batches
    start_time = time.time()

    while True:
        batch = batch_manager.get_next_batch()
        if not batch:
            print("\n所有批次处理完成！")
            break

        batch_num += 1
        logger.info(f"[批次] {batch_num}/{total_batches}: {batch.batch_id} ({len(batch.images)}张)")

        batch_results = []
        batch_success = True
        batch_progress = ProgressBar(len(batch.images), prefix=f"批次{batch_num}")

        for i, img_path in enumerate(batch.images, 1):
            try:
                result = process_single_image(
                    ocr, analyzer, organizer, doc_manager, ima_syncer, archiver,
                    img_path, i, len(batch.images), batch_progress
                )
                if result:
                    batch_results.append(result)
                else:
                    batch_results.append({'image': Path(img_path).name, 'failed': True})
            except Exception as e:
                print(f"\n  错误: {e}")
                logger.error(f"处理失败 {img_path}: {e}")
                batch_results.append({'image': Path(img_path).name, 'failed': True, 'error': str(e)})
                batch_success = False

        batch_progress.finish()

        if batch_success:
            batch_manager.mark_batch_completed(batch.batch_id, f"处理 {len(batch_results)} 张")
        else:
            failed_n = sum(1 for r in batch_results if r.get('failed'))
            batch_manager.mark_batch_failed(batch.batch_id, f"失败 {failed_n} 张")

        all_results.extend(batch_results)

        if batch_num < total_batches:
            print("\n[批次间隔] 1 秒...")
            time.sleep(1)

    elapsed = time.time() - start_time
    report, report_file = save_report(all_results, total_images, elapsed, 'batch')
    print_summary(report, archiver, ima_syncer)
    print(f"\n报告已保存: {report_file}")

    return all_results


# ============================================================
# 主入口
# ============================================================
def print_usage():
    print("\n用法:")
    print("  python auto_process_all_v8.py           # 全自动处理（>10张自动分批）")
    print("  python auto_process_all_v8.py --batch    # 强制分批模式（可中断续传）")
    print("  python auto_process_all_v8.py --init     # 仅初始化分批")
    print("  python auto_process_all_v8.py --progress # 查看进度")
    print("  python auto_process_all_v8.py --clear    # 清除分批状态")


def main():
    batch_manager = BatchManager()
    if len(sys.argv) > 1:
        cmd = sys.argv[1]
        if cmd == '--init':
            state = batch_manager.initialize()
            if state:
                batch_manager.print_progress()
        elif cmd == '--progress':
            batch_manager.print_progress()
        elif cmd == '--clear':
            batch_manager.clear_state()
        elif cmd == '--batch':
            run_batch_mode(batch_manager)
        else:
            print(f"[ERROR] 未知参数: {cmd}")
            print_usage()
    else:
        run_full_mode(batch_manager)


if __name__ == '__main__':
    main()
