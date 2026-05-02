#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
全自动化图片处理工具 V9.3 - 合并逻辑重构版
=====================================================
V9.3 核心修复：三处系统性问题根治

  [V9.3-1] 合并逻辑简化（改动一，普适性修复）
    - 彻底删除 find_similar_doc 中的60行模糊匹配逻辑
      （包含：normalize朝代映射、包含关系0.90、核心实体匹配0.80、bigram Jaccard）
    - 改为：只有完全同名（doc_name精确相同）才合并
    - 原则：宁可多建文档，不误合并不同内容（V2哲学）
    - 适用所有分类：历史/中医/养生/旅游均不再有跨主题误合并

  [V9.3-2] 乱码过滤收紧（改动二）
    - 原规则：ltr_chars>=1 and zh_chars>=1 and not has_long_chinese(4+字)
      → 误删短行（"始封君"、"洪武年号"等）
    - 新规则：字母占比>30% 且 无连续2字以上汉字词 才视为乱码
    - 效果：保留结构化短行，只过滤真正乱码（"甘o亚a"类）

  [V9.3-3] LLM失败改为跳过（改动三）
    - 原逻辑：LLM失败 → 正则兜底 → 生成错误命名文档（如"始封君"）
    - 新逻辑：LLM失败 → 返回 _skip 标记 → 图片留在原地待人工处理
    - 原则：宁可不生成文档，也不生成错误文档

V9.2 历史升级：
  [V9.2-1] 分次调用 - 第1次LLM返回title/category/summary（短JSON不易坏）
  [V9.2-2] 第2次LLM返回整理后content（纯Markdown输出，无JSON包裹）

V9 哲学：正则只做过滤（排除垃圾），LLM做所有决策（命名/分类/摘要/内容整理）

使用方法：
  python auto_process_all_v9_3.py           # 全自动处理
  python auto_process_all_v9_3.py --batch    # 强制分批模式
  python auto_process_all_v9_3.py --init     # 仅初始化分批
  python auto_process_all_v9_3.py --progress # 查看进度
  python auto_process_all_v9_3.py --clear    # 清除分批状态
"""

import os
import sys
import io
import json
import time
import re
import hashlib
import shutil
import logging
import requests
from pathlib import Path
from datetime import datetime
from collections import Counter
from typing import Optional, List, Dict, Tuple
from dataclasses import dataclass

# 修复 Windows 控制台编码
if sys.platform == 'win32':
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
    except Exception:
        pass

sys.path.insert(0, '.')

from dotenv import load_dotenv
env_path = Path('.') / '.env'
if env_path.exists():
    load_dotenv(dotenv_path=env_path)
else:
    print("[WARN] .env 文件不存在，使用系统环境变量")

from local_ocr import LocalOCR

try:
    from tencent_ocr import TencentOCR
    TENCENT_AVAILABLE = True
except ImportError:
    TENCENT_AVAILABLE = False

try:
    from baidu_ocr import BaiduOCR
    BAIDU_AVAILABLE = True
except ImportError:
    BAIDU_AVAILABLE = False

# 日志配置
Path('处理结果').mkdir(exist_ok=True)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(message)s',
    handlers=[
        logging.FileHandler('处理结果/process.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# 预设分类列表（仅用于文档文件夹归档和LLM提示）
PRESET_CATEGORIES = [
    "历史文化", "营养健康", "生活方式", "教育育儿", "旅游攻略", "综合知识", "其他"
]


# ============================================================
# V9 核心：LLM全覆盖分析器
# ============================================================
class LLMAnalyzer:
    """
    V9.0 核心：用混元Lite替代所有正则决策

    每张图片调用一次LLM，一次性返回：
    - title: 文档标题
    - category: 分类
    - summary: 内容摘要（用于文档开头）

    兜底策略：LLM失败时使用简单提取（V2风格）
    """

    def __init__(self):
        self._hunyuan_available = None
        self._siliconflow_available = None
        self.cache = {}

    # ---------- LLM引擎 ----------

    def _check_hunyuan(self) -> bool:
        if self._hunyuan_available is not None:
            return self._hunyuan_available
        try:
            sid = os.getenv('TENCENT_SECRET_ID', '')
            skey = os.getenv('TENCENT_SECRET_KEY', '')
            self._hunyuan_available = bool(sid and skey and '替换' not in sid)
            logger.info(f"[LLM] 混元Lite: {'已配置' if self._hunyuan_available else '未配置'}")
        except Exception:
            self._hunyuan_available = False
        return self._hunyuan_available

    def _check_siliconflow(self) -> bool:
        if self._siliconflow_available is not None:
            return self._siliconflow_available
        try:
            key = os.getenv('SILICONFLOW_API_KEY', '')
            self._siliconflow_available = bool(key and '替换' not in key)
            logger.info(f"[LLM] 硅基流动: {'已配置' if self._siliconflow_available else '未配置'}")
        except Exception:
            self._siliconflow_available = False
        return self._siliconflow_available

    def _call_hunyuan_lite(self, prompt: str) -> Optional[str]:
        """调用腾讯混元Lite"""
        try:
            from tencentcloud.hunyuan.v20230901 import hunyuan_client, models as hm_models
            from tencentcloud.common.exception.tencent_cloud_sdk_exception import TencentCloudSDKException
            from tencentcloud.common.profile.client_profile import ClientProfile
            from tencentcloud.common.profile.http_profile import HttpProfile
            from tencentcloud.common.credential import Credential

            sid = os.getenv('TENCENT_SECRET_ID', '')
            skey = os.getenv('TENCENT_SECRET_KEY', '')
            if not sid or not skey or '替换' in skey:
                return None

            hp = HttpProfile()
            hp.endpoint = "hunyuan.tencentcloudapi.com"
            hp.reqTimeout = 10
            cp = ClientProfile()
            cp.httpProfile = hp
            client = hunyuan_client.HunyuanClient(Credential(sid, skey), "ap-guangzhou", cp)

            req = hm_models.ChatCompletionsRequest()
            req.Model = "hunyuan-lite"
            req.Messages = [{"Role": "user", "Content": prompt}]
            req.Stream = False

            resp = client.ChatCompletions(req)
            return resp.Choices[0].Message.Content.strip()

        except ImportError:
            logger.info("[LLM] 混元SDK未安装")
            return None
        except TencentCloudSDKException as e:
            code = getattr(e, 'code', '')
            if 'ServiceNotActivated' in code:
                logger.info("[LLM] 混元Lite服务未开通")
            elif 'UnauthorizedOperation' in code:
                logger.info("[LLM] 混元Lite CAM权限不足")
            else:
                logger.info(f"[LLM] 混元Lite异常: {code}")
            return None
        except Exception as e:
            logger.info(f"[LLM] 混元Lite调用失败: {e}")
            return None

    def _call_siliconflow(self, prompt: str) -> Optional[str]:
        """调用硅基流动"""
        try:
            api_key = os.getenv('SILICONFLOW_API_KEY', '')
            if not api_key or '替换' in api_key:
                return None
            url = "https://api.siliconflow.cn/v1/chat/completions"
            resp = requests.post(url, json={
                "model": os.getenv('SILICONFLOW_MODEL', 'Qwen/Qwen3-8B'),
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": 300, "temperature": 0.3
            }, headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }, timeout=10)
            if resp.status_code == 200:
                return resp.json().get('choices', [{}])[0].get('message', {}).get('content', '').strip()
            return None
        except Exception as e:
            logger.info(f"[LLM] 硅基流动失败: {e}")
            return None

    # ---------- 核心分析 ----------

    def analyze(self, text: str, index: int) -> Dict:
        """
        V9.2核心：分两次调用LLM
          第1次：获取 title + category + summary（短JSON，不会截断）
          第2次：获取整理后的 content（纯Markdown输出，无JSON包裹）
        兜底：LLM全部失败时使用简单提取（仅去噪，不重写）
        """
        content_hash = hashlib.md5(re.sub(r'\s+', '', text).encode('utf-8')).hexdigest()[:8]

        if content_hash in self.cache:
            return self.cache[content_hash]

        # 传给LLM的全文（最多1200字，保留更多上下文以便内容整理）
        full_text = text[:1200].strip()
        categories_str = '/'.join(PRESET_CATEGORIES)

        # ====== 第1次调用：获取 title/category/summary（短JSON，不易坏） ======
        prompt_meta = (
            f"你是一个知识整理助手。以下是从手机截图中OCR识别的原始文本，可能包含界面元素和识别错误。\n\n"
            f"请返回JSON格式（不要用代码块包裹）：\n"
            f'{{"title": "简洁标题2-15字", "category": "从以下选一个：{categories_str}", "summary": "一句话概括20-50字"}}\n\n'
            f"原始OCR文本：\n{full_text}\n\n"
            f"只返回JSON，不加其他文字。"
        )

        result = None
        engine = None

        for check, call, name in [
            (self._check_hunyuan, self._call_hunyuan_lite, "混元Lite"),
            (self._check_siliconflow, self._call_siliconflow, "硅基流动"),
        ]:
            if check():
                logger.info(f"[LLM-{index}] 第1次调用{name}（元数据）...")
                raw = call(prompt_meta)
                if raw:
                    result = self._parse_llm_json(raw)
                    if result:
                        engine = name
                        logger.info(f"[LLM-{index}] 元数据成功: title={result.get('title')}, category={result.get('category')}")
                        break
                    else:
                        logger.info(f"[LLM-{index}] {name}元数据解析失败，尝试下一个引擎")
                else:
                    logger.info(f"[LLM-{index}] {name}无响应，尝试下一个引擎")

        # V9.3：LLM全部失败 → 跳过，不走正则兜底，不生成错误文档
        if not result:
            logger.warning(f"[LLM-{index}] 所有LLM调用失败，返回跳过标记（图片留原地待人工处理）")
            return {
                'category': '', 'confidence': 0.0,
                'doc_name': '', 'keywords': [],
                'content_hash': content_hash,
                'summary': '', 'llm_content': '',
                '_engine': None, '_skip': True,
            }

        # ====== 第2次调用：获取整理后的content（纯Markdown，无JSON包裹） ======
        llm_content = ''
        if engine:
            for check, call, name in [
                (self._check_hunyuan, self._call_hunyuan_lite, "混元Lite"),
                (self._check_siliconflow, self._call_siliconflow, "硅基流动"),
            ]:
                if check():
                    prompt_content = (
                        f"你是一个知识整理助手。以下是从手机截图中OCR识别的原始文本，"
                        f"可能包含界面元素（如按钮文字、时间栏、APP名称）和识别错误。\n\n"
                        f"请将以下OCR文本整理为Markdown格式正文，要求：\n"
                        f"1. 去掉所有APP界面残片（如'豆包>'、时间栏、按钮等）\n"
                        f"2. 修正明显的OCR错误\n"
                        f"3. 用##/###添加合适的章节标题\n"
                        f"4. 保留所有实质性知识内容，不要删减\n"
                        f"5. 如有列表用-或数字编号格式化\n\n"
                        f"直接返回整理后的Markdown正文，不要加任何前缀、说明或代码块包裹。\n\n"
                        f"原始OCR文本：\n{full_text}"
                    )
                    logger.info(f"[LLM-{index}] 第2次调用{name}（内容整理）...")
                    raw_content = call(prompt_content)
                    if raw_content and len(raw_content.strip()) > 20:
                        # 去掉可能的代码块包裹
                        cleaned = raw_content.strip()
                        if cleaned.startswith('```'):
                            cleaned = re.sub(r'^```[a-z]*\n?', '', cleaned)
                            cleaned = re.sub(r'\n?```$', '', cleaned).strip()
                        llm_content = cleaned
                        logger.info(f"[LLM-{index}] 内容整理成功，长度={len(llm_content)}")
                        break
                    else:
                        logger.info(f"[LLM-{index}] 内容整理返回为空或过短，跳过")

        # 构建最终结果
        title = result.get('title', '').strip()
        category = result.get('category', '综合知识').strip()
        summary = result.get('summary', '').strip()
        # V9.2：第2次LLM调用整理后的内容（若失败则为空，后续用原始清理内容）
        # 注意：第1次调用的降级提取不再含content字段

        # 清理标题（去掉非法文件名字符）
        title = re.sub(r'[<>:"/\\|?*\n\r\t，。！？、：；]', '', title)
        if len(title) > 20:
            title = title[:20]
        if not title or len(title) < 2:
            title = f"整理{datetime.now().strftime('%m%d%H%M')}"

        # 校验分类
        if category not in PRESET_CATEGORIES:
            category = '综合知识'

        # 提取关键词（用于文档元信息）
        keywords = self._extract_keywords(text, category)

        doc_name = f"{category}-{title}"

        final = {
            'category': category,
            'confidence': 0.85 if engine else 0.5,
            'doc_name': doc_name,
            'keywords': keywords,
            'content_hash': content_hash,
            'summary': summary,
            'llm_content': llm_content,   # V9.1新增：LLM整理后的正文
            '_engine': engine,
        }

        self.cache[content_hash] = final
        return final

    def _parse_llm_json(self, raw: str) -> Optional[Dict]:
        """解析LLM返回的JSON（V9.2：只解析title/category/summary，content由第2次调用单独获取）"""
        raw_stripped = raw.strip()
        # 去掉可能的代码块包裹
        if raw_stripped.startswith('```'):
            raw_stripped = re.sub(r'^```[a-z]*\n?', '', raw_stripped)
            raw_stripped = re.sub(r'\n?```$', '', raw_stripped).strip()

        # 第1轮：尝试整体解析（完整JSON）
        try:
            data = json.loads(raw_stripped)
            return self._validate_and_extract(data)
        except json.JSONDecodeError:
            pass

        # 第2轮：正则找最外层JSON块（DOTALL）
        json_match = re.search(r'\{.*\}', raw_stripped, re.DOTALL)
        if json_match:
            try:
                data = json.loads(json_match.group())
                return self._validate_and_extract(data)
            except json.JSONDecodeError:
                pass  # 继续降级尝试

        # 第3轮：JSON被截断时，单独用正则提取 title/category/summary
        # V9.2：content由第2次LLM调用单独获取，这里不再尝试提取content
        extracted = {}
        for field in ('title', 'category', 'summary'):
            # 匹配 "field": "value" 形式（value不含换行的单行字符串）
            m = re.search(
                r'"' + field + r'"\s*:\s*"([^"\\]*(?:\\.[^"\\]*)*)"',
                raw_stripped
            )
            if m:
                extracted[field] = m.group(1).replace('\\"', '"').strip()

        if extracted.get('title') and len(extracted['title']) >= 2:
            logger.info(f"[LLM] JSON截断，降级提取: title={extracted.get('title')}, category={extracted.get('category')}")
            extracted.setdefault('category', '')
            extracted.setdefault('summary', '')
            return self._validate_and_extract(extracted)

        logger.info(f"[LLM] JSON解析彻底失败: {raw_stripped[:120]}")
        return None

    def _validate_and_extract(self, data: Dict) -> Optional[Dict]:
        """验证并提取LLM JSON字段（V9.2：只验证title/category/summary）"""
        title = str(data.get('title', '')).strip()
        category = str(data.get('category', '')).strip()
        summary = str(data.get('summary', '')).strip()

        if not title or len(title) < 2:
            logger.info(f"[LLM] 标题为空或过短: '{title}'")
            return None
        if len(title) > 25:
            # 截断过长标题
            title = title[:20]
        return {'title': title, 'category': category, 'summary': summary}

    def _fallback_extract(self, text: str) -> Dict:
        """V2风格兜底：提取标题 + 关键词推断分类"""
        lines = [l.strip() for l in text.split('\n') if l.strip() and len(l.strip()) >= 3]

        # 取前8行中第一个有意义的短句
        noise_words = {'http', 'www', '关注', '收藏', '点赞', '评论', '分享', '发送',
                       '翻译', '探索版', '写作', 'doubao', 'kimi', '元宝', '豆包',
                       '一张图看完', 'AI生成', '内容由'}

        title = ''
        for line in lines[:8]:
            cleaned = re.sub(r'^#{1,6}\s*', '', line)
            cleaned = re.sub(r'^[◆★▼▪▸【】\[\]]+\s*', '', cleaned).strip()
            if not cleaned:
                continue
            if any(noise in cleaned for noise in noise_words):
                continue
            if re.match(r'^[\d:.\s%]+$', cleaned):
                continue
            if 3 <= len(cleaned) <= 15 and re.match(r'^[\u4e00-\u9fa5]+$', cleaned):
                title = cleaned
                break

        if not title:
            title = f"整理{datetime.now().strftime('%m%d%H%M')}"

        # 关键词推断分类（而非写死综合知识）
        category = self._infer_category_by_keywords(text)
        return {'title': title, 'category': category, 'summary': ''}

    def _infer_category_by_keywords(self, text: str) -> str:
        """根据内容关键词推断分类"""
        cat_keywords = {
            "历史文化": ["历史", "朝代", "皇帝", "皇后", "太后", "太子", "王朝", "诸侯",
                         "古代", "传记", "人物", "将领", "战役", "文化", "典故", "世家",
                         "门阀", "宰相", "藩王", "大臣", "年号", "庙号", "谥号",
                         "明朝", "清朝", "宋朝", "唐朝", "汉朝", "元朝", "周朝", "秦朝",
                         "北宋", "南宋", "东汉", "西汉", "五代", "春秋", "战国",
                         "钱谦益", "钱穆", "明末", "淮盐", "诸侯国", "始封"],
            "营养健康": ["营养", "健康", "饮食", "中医", "养生", "药膳", "食疗", "免疫",
                         "调理", "滋补", "中药", "穴位", "经络", "茶饮", "功效"],
            "生活方式": ["生活", "运动", "睡眠", "心理", "习惯", "健身", "护肤",
                         "整理", "收纳", "时间管理", "效率"],
            "教育育儿": ["育儿", "宝宝", "教育", "孩子", "亲子", "学习", "学校",
                         "辅食", "早教", "发育", "成长"],
            "旅游攻略": ["旅游", "旅行", "景点", "攻略", "路线", "酒店", "美食",
                         "打卡", "民宿", "自驾", "行程"],
            "综合知识": ["科普", "知识", "研究", "科学", "原理", "数据", "统计"],
        }
        text_lower = text[:800]  # 只看前800字
        scores = {}
        for cat, kws in cat_keywords.items():
            score = sum(1 for kw in kws if kw in text_lower)
            if score > 0:
                scores[cat] = score
        if scores:
            best = max(scores, key=scores.get)
            logger.info(f"[分类] 关键词推断: {best}（得分={scores}）")
            return best
        return '综合知识'

    def _extract_keywords(self, text: str, category: str) -> List[str]:
        """简单关键词提取（从分类关键词中匹配）"""
        cat_keywords = {
            "历史文化": ["历史", "朝代", "人物", "皇帝", "战争", "古代", "传统", "文献"],
            "营养健康": ["营养", "健康", "饮食", "中医", "养生", "药膳", "食疗", "免疫"],
            "生活方式": ["生活", "运动", "睡眠", "心理", "习惯", "健身", "护肤"],
            "教育育儿": ["育儿", "宝宝", "教育", "孩子", "亲子", "学习", "学校"],
            "旅游攻略": ["旅游", "旅行", "景点", "攻略", "路线", "酒店", "美食"],
            "综合知识": ["科普", "知识", "研究", "科学", "原理", "数据"],
            "其他": [],
        }
        kws = cat_keywords.get(category, [])
        return [kw for kw in kws if kw in text][:6]


# ============================================================
# V9 内容整理器（V2轻量版）
# ============================================================
class ContentOrganizer:
    """
    V9 内容整理器 - V2轻量版

    只做三件事：
    1. 提取来源信息（小红书/微信等）
    2. 清理OCR错误和噪点
    3. 语义标题识别（V2风格，只识别明确的语义关键词）
    """

    def __init__(self):
        # 手机UI噪点模式
        self.ui_noise = [
            r'^\d{1,2}:\d{2}$', r'^[0-9]{1,2}:%$', r'^4G\s*5G',
            r'^[0-9]+\s*%$', r'^Wi-?Fi', r'^\.?\s*AM$', r'^\.?\s*PM$',
            r'^(写作|翻译|探索版|有什么问题尽管问我)$',
            r'^(doubao\.com|豆包|Kimi|kimi|腾讯元宝|元宝|ChatGPT|GPT)$',
            r'^(文心一言|通义千问|智谱清言|DeepSeek)$',
            r'^(发送|发送消息|\+关注|收藏|点赞|评论|分享|转发)$',
            r'^(已关注|更多|展开|收起|折叠)$',
            r'^\d+$', r'^http[s]?://', r'^来自.*的搜索$', r'^搜一下$',
            r'^a?付费', r'^小红书号[：:]\s*\d+', r'^内容由\w+生成',
            r'^相关视频$', r'^双深度思考$', r'^AI生图$', r'^回照片动起来$',
            r'^已阅读\d+个网页', r'^按住说话$', r'^\d+个赞$',
            r'^(艮拍|拍题|答疑)$', r'^(打电话|帮我写|发消息\.\.\.)$',
            r'^WMA:', r'^较长的视', r'^今天我',
        ]

        # V2 OCR错误替换
        self.ocr_fixes = {
            '√': '✓', '✖': '✗', '╳': '✗',
            '—': '—', '～': '~',
            '．': '.', '，': '，', '．．': '..',
            '＇': "'", '＇＇': '""',
            '＊': '*', '＋': '+', '＜': '<', '＞': '>',
        }

        # V2语义标题关键词（明确的功能性关键词才识别为标题）
        self.heading_keywords = [
            '功效', '作用', '做法', '做法步骤', '材料', '食材', '原料',
            '禁忌', '注意事项', '副作用', '不宜',
            '症状', '表现', '诊断', '治疗', '预防', '调理',
            '营养价值', '营养成分', '适用人群', '食用方法', '食用禁忌',
            '生长习性', '种植方法', '栽培技术',
            '人物简介', '生平', '人物志', '历史背景', '主要成就',
        ]

    def organize(self, raw_text: str) -> Dict:
        """整理内容：提取来源 → 清理 → 结构化"""
        source = self._extract_source(raw_text)
        cleaned = self._clean_text(raw_text)
        structured = self._structure_content(cleaned)

        return {
            'content': structured,
            'source': source,
        }

    def _extract_source(self, text: str) -> Dict:
        """提取来源信息"""
        platform = ''
        author = ''

        if '@' in text and ('小红书' in text or 'red' in text.lower()):
            platform = '小红书'
        elif '微信读书' in text or 'Weread' in text:
            platform = '微信读书'
        elif '抖音' in text or 'TikTok' in text:
            platform = '抖音'
        elif 'B站' in text or 'bilibili' in text:
            platform = 'B站'
        elif '知乎' in text:
            platform = '知乎'
        elif '得到' in text or 'iget' in text:
            platform = '得到'

        # 提取@作者
        author_match = re.search(r'@([\w\u4e00-\u9fa5]{2,15})', text)
        if author_match:
            author = author_match.group(1)

        return {'platform': platform, 'author': author}

    def _clean_text(self, text: str) -> str:
        """清理OCR错误和噪点"""
        lines = text.split('\n')
        cleaned = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # 过滤UI噪点
            is_noise = False
            for pattern in self.ui_noise:
                if re.match(pattern, stripped, re.IGNORECASE):
                    is_noise = True
                    break
            if is_noise:
                continue

            # [V9.3] 过滤真正乱码行（字母占比>30% 且 无连续汉字词）
            # 原规则(V8.3)太宽松：ltr>=1 且 zh>=1 且 无4连汉字 → 误删短行如"洪武年号"、"始封君"
            # 新规则收紧：只有字母占比明显偏高 且 没有任何连续汉字词 才视为乱码
            ltr_chars = len(re.findall(r'[a-z]', stripped))
            total_chars = max(len(stripped.replace(' ', '')), 1)
            has_chinese_words = bool(re.search(r'[\u4e00-\u9fa5]{2,}', stripped))
            if (ltr_chars / total_chars) > 0.3 and not has_chinese_words:
                continue

            # OCR字符替换
            for old, new in self.ocr_fixes.items():
                stripped = stripped.replace(old, new)

            # 清理行首行尾空白
            stripped = stripped.strip()

            if stripped:
                cleaned.append(stripped)

        return '\n\n'.join(cleaned)

    def _structure_content(self, text: str) -> str:
        """V2轻量结构化：只识别语义标题关键词"""
        lines = text.split('\n')
        result = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # 检查是否包含语义标题关键词（在行首或独立成行时）
            for keyword in self.heading_keywords:
                if stripped == keyword or stripped.startswith(keyword + '：') or stripped.startswith(keyword + ':'):
                    result.append(f'\n### {stripped}\n')
                    break
            else:
                result.append(stripped)

        return '\n'.join(result)


# ============================================================
# 文档管理器（保留V8.3合并逻辑）
# ============================================================
class SmartDocumentManager:
    """文档管理器：合并检测 + MD/Word生成"""

    def __init__(self, output_dir: str = '处理结果'):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.doc_index = {}
        self.known_topics = []
        self.known_hashes = set()
        self._load_existing_docs()

    def _load_existing_docs(self):
        """加载已有文档索引"""
        for cat_dir in self.output_dir.iterdir():
            if cat_dir.is_dir() and not cat_dir.name.startswith('.'):
                for md_file in cat_dir.glob('*.md'):
                    base = md_file.stem
                    self.doc_index[base] = md_file
                    # 从文档名提取主题
                    topic = base.split('-', 1)[1] if '-' in base else base
                    self.known_topics.append((topic, md_file))
                    # 读取已有content_hash
                    try:
                        content = md_file.read_text(encoding='utf-8')
                        hash_match = re.search(r'content_hash:\s*(\w+)', content)
                        if hash_match:
                            self.known_hashes.add(hash_match.group(1))
                    except Exception:
                        pass

    def _extract_topic_from_name(self, doc_name: str) -> str:
        """从文档名提取主题（去掉分类前缀）"""
        if '-' in doc_name:
            return doc_name.split('-', 1)[1]
        return doc_name

    def is_duplicate_hash(self, content_hash: str) -> bool:
        return content_hash in self.known_hashes

    def find_similar_doc(self, doc_name: str, content: str) -> Optional[Path]:
        """
        V9.3：精确同名匹配（替换原60行模糊匹配逻辑）

        原逻辑删除：
          - normalize 朝代映射表（北宋→宋朝等，导致跨朝代误合并）
          - 包含关系匹配（0.90，导致"补气食材"合并进"补气养血"）
          - 核心实体匹配（0.80，导致南宋合并进北宋）
          - bigram Jaccard 相似度

        V9.3 原则：宁可多建文档，不误合并不同内容（V2哲学）
        只有 doc_name 完全相同（LLM命名一致）才合并。
        """
        if doc_name in self.doc_index:
            return self.doc_index[doc_name]
        return None



    def merge_content(self, existing_file: Path, new_content: str,
                      new_image_name: str, content_hash: str = None) -> bool:
        """追加内容到已有文档"""
        try:
            existing = existing_file.read_text(encoding='utf-8')
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
            append = f"\n\n---\n\n## 补充内容\n\n> 来源图片: {new_image_name}\n> 追加时间: {timestamp}\n> 内容标识: {content_hash or 'N/A'}\n\n{new_content}\n"
            marker = '\n\n---\n\n*本文档由图片知识库整理工具自动生成*'
            if marker in existing:
                new_full = existing.replace(marker, append + marker)
            else:
                new_full = existing + append
            existing_file.write_text(new_full, encoding='utf-8')
            return True
        except Exception as e:
            logger.warning(f"[合并] 失败: {e}")
            return False

    def save_document(self, doc_name: str, category: str, content: str,
                      image_name: str, keywords: List[str],
                      content_hash: str, source_info: Dict,
                      summary: str = '') -> Tuple[str, bool]:
        """保存文档（新建或合并）"""
        similar = self.find_similar_doc(doc_name, content)
        if similar:
            logger.info(f"[文档] 发现相似文档 {similar.name}，合并内容")
            self.merge_content(similar, content, image_name, content_hash)
            if content_hash:
                self.known_hashes.add(content_hash)
            return str(similar), False

        cat_dir = self.output_dir / category
        cat_dir.mkdir(exist_ok=True)

        safe_name = re.sub(r'[<>:"/\\|?*]', '', doc_name)
        if len(safe_name) > 40:
            safe_name = safe_name[:40]

        md_file = cat_dir / f"{safe_name}.md"
        counter = 1
        while md_file.exists():
            md_file = cat_dir / f"{safe_name}_{counter}.md"
            counter += 1

        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
        keywords_str = ', '.join(keywords) if keywords else '无'
        source_section = ""
        if source_info.get('platform'):
            source_section += f"**来源**：{source_info['platform']}\n\n"
        if source_info.get('author'):
            source_section += f"**作者**：{source_info['author']}\n\n"
        summary_section = f"**摘要**：{summary}\n\n" if summary else ""

        md_content = f"""# {doc_name}

> 来源图片: {image_name}
> 识别时间: {timestamp}
> 分类: {category}
> 关键词: {keywords_str}
> content_hash: {content_hash or 'N/A'}

---

<!-- CONTENT_START -->
{source_section}{summary_section}{content}
<!-- CONTENT_END -->

---

*本文档由图片知识库整理工具V9自动生成*
"""
        md_file.write_text(md_content, encoding='utf-8')

        self.doc_index[doc_name] = md_file
        topic = self._extract_topic_from_name(doc_name)
        self.known_topics.append((topic, md_file))
        if content_hash:
            self.known_hashes.add(content_hash)

        logger.info(f"[文档] 新建: {category}/{md_file.name}")
        return str(md_file), True

    def generate_word(self, doc_name: str, category: str, content: str,
                      image_name: str, keywords: List[str],
                      md_file_path: str = None) -> Optional[str]:
        """生成Word文档"""
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            logger.warning("[文档] python-docx未安装，跳过Word生成")
            return None

        cat_dir = self.output_dir / category
        cat_dir.mkdir(exist_ok=True)

        safe_name = re.sub(r'[<>:"/\\|?*]', '', doc_name)
        if len(safe_name) > 40:
            safe_name = safe_name[:40]

        docx_file = cat_dir / f"{safe_name}.docx"
        if docx_file.exists():
            return str(docx_file)

        # 从MD文件读取最新内容
        text_content = content
        if md_file_path and Path(md_file_path).exists():
            try:
                md_text = Path(md_file_path).read_text(encoding='utf-8')
                if '<!-- CONTENT_START -->' in md_text and '<!-- CONTENT_END -->' in md_text:
                    start = md_text.index('<!-- CONTENT_START -->') + len('<!-- CONTENT_START -->')
                    end = md_text.index('<!-- CONTENT_END -->')
                    text_content = md_text[start:end].strip()
            except Exception:
                pass

        try:
            doc = Document()
            doc.styles['Normal'].font.name = 'Microsoft YaHei'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

            doc.add_heading(doc_name, 0)
            doc.add_paragraph(f"来源图片: {image_name}")
            doc.add_paragraph(f"识别时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph(f"分类: {category}")
            keywords_str = ', '.join(keywords) if keywords else '无'
            doc.add_paragraph(f"关键词: {keywords_str}")
            doc.add_paragraph("─" * 30)

            for line in text_content.split('\n'):
                stripped = line.strip()
                if not stripped:
                    doc.add_paragraph('')
                elif re.match(r'^#{1,4}\s+', stripped):
                    level = len(re.match(r'^(#+)', stripped).group(1))
                    heading_text = re.sub(r'^#+\s*', '', stripped)
                    if 1 <= level <= 4:
                        doc.add_heading(heading_text, level=level)
                    else:
                        doc.add_paragraph(heading_text)
                elif stripped.startswith('<!--') and stripped.endswith('-->'):
                    continue
                elif stripped.startswith('**摘要') and stripped.endswith('**') and '：' in stripped:
                    # 摘要行 → 普通段落（加粗处理）
                    text = stripped.strip('*').split('：', 1)[1] if '：' in stripped else stripped.strip('*')
                    doc.add_paragraph(f"摘要：{text}")
                elif stripped.startswith('**来源') and stripped.endswith('**') and '：' in stripped:
                    text = stripped.strip('*').split('：', 1)[1] if '：' in stripped else stripped.strip('*')
                    doc.add_paragraph(f"来源：{text}")
                else:
                    doc.add_paragraph(stripped)

            doc.save(str(docx_file))
            logger.info(f"[文档] Word已生成: {category}/{docx_file.name}")
            return str(docx_file)
        except Exception as e:
            logger.warning(f"[Word] 生成失败: {e}")
            return None


# ============================================================
# 图片归档器
# ============================================================
class ImageArchiver:
    def __init__(self, processed_dir: str = '已处理图片'):
        self.processed_dir = Path(processed_dir)
        self.processed_dir.mkdir(parents=True, exist_ok=True)
        self.archived_count = 0

    def get_source_folder(self, image_path: str) -> str:
        src = Path(image_path)
        parent_name = src.parent.name
        return '根目录' if parent_name == '待处理图片' else parent_name

    def archive_image(self, image_path: str, category: str = None) -> Optional[str]:
        src = Path(image_path)
        if not src.exists():
            return None
        source_folder = self.get_source_folder(image_path)
        archive_subdir = self.processed_dir / source_folder
        archive_subdir.mkdir(parents=True, exist_ok=True)
        dest = archive_subdir / src.name
        if dest.exists():
            ts = datetime.now().strftime('%H%M%S')
            dest = archive_subdir / f"{src.stem}_{ts}{src.suffix}"
        try:
            shutil.move(str(src), str(dest))
            self.archived_count += 1
            return str(dest)
        except Exception as e:
            logger.warning(f"[归档] 移动失败 {src.name}: {e}")
            try:
                shutil.copy2(str(src), str(dest))
                src.unlink()
                self.archived_count += 1
                return str(dest)
            except Exception:
                return None

    def get_stats(self) -> Dict:
        total = 0
        for item in self.processed_dir.iterdir():
            if item.is_dir():
                total += len(list(item.glob('*.*')))
        return {'archived_this_session': self.archived_count, 'total_archived': total}


# ============================================================
# 进度条
# ============================================================
class ProgressBar:
    def __init__(self, total: int, width: int = 40, prefix: str = "进度"):
        self.total = total
        self.width = width
        self.prefix = prefix
        self.start_time = time.time()
        self.current = 0

    def update(self, current: int, info: str = ""):
        self.current = current
        percent = current / self.total if self.total > 0 else 0
        filled = int(self.width * percent)
        bar = '█' * filled + '░' * (self.width - filled)
        if current > 0:
            elapsed = time.time() - self.start_time
            eta = elapsed / current * (self.total - current)
            eta_str = f"{int(eta//60)}m{int(eta%60)}s"
        else:
            eta_str = "--"
        print(f"\r{self.prefix}: [{bar}] {current}/{self.total} ({percent:.0%}) ETA:{eta_str} {info[:15]}",
              end='', flush=True)

    def log(self, message: str):
        clear = ' ' * 100
        print(f'\r{clear}\r', end='', flush=True)
        print(message)
        if self.current > 0:
            self._redraw()

    def _redraw(self):
        percent = self.current / self.total if self.total > 0 else 0
        filled = int(self.width * percent)
        bar = '█' * filled + '░' * (self.width - filled)
        elapsed = time.time() - self.start_time
        if self.current > 0:
            eta = elapsed / self.current * (self.total - self.current)
            eta_str = f"{int(eta//60)}m{int(eta%60)}s"
        else:
            eta_str = "--"
        print(f"\r{self.prefix}: [{bar}] {self.current}/{self.total} ({percent:.0%}) ETA:{eta_str}",
              end='', flush=True)

    def finish(self):
        self.update(self.total, "完成!")
        print()


# ============================================================
# 多引擎OCR
# ============================================================
class MultiEngineOCR:
    def __init__(self):
        self.current_engine = None
        self.engine_status = []

        sid = os.getenv('TENCENT_SECRET_ID', '')
        skey = os.getenv('TENCENT_SECRET_KEY', '')
        self.engine_status.append(('腾讯云', bool(sid and skey and '替换' not in sid and TENCENT_AVAILABLE), '已配置' if sid else '未配置'))

        aid = os.getenv('BAIDU_APP_ID', '')
        ak = os.getenv('BAIDU_API_KEY', '')
        bs = os.getenv('BAIDU_SECRET_KEY', '')
        self.engine_status.append(('百度云', bool(aid and ak and bs and '替换' not in aid and BAIDU_AVAILABLE), '已配置' if aid else '未配置'))

        local = LocalOCR()
        self.engine_status.append(('本地Tesseract', local.tesseract_available, '可用' if local.tesseract_available else '不可用'))
        self._select_best_engine()

    def _select_best_engine(self):
        for name, available, _ in self.engine_status:
            if available:
                self.current_engine = name
                logger.info(f"[OCR] 选用引擎: {name}")
                return True
        return False

    def recognize(self, image_path: str) -> Dict:
        for name, available, _ in self.engine_status:
            if not available:
                continue
            try:
                if name == '腾讯云' and TENCENT_AVAILABLE:
                    engine = TencentOCR()
                elif name == '百度云' and BAIDU_AVAILABLE:
                    engine = BaiduOCR()
                else:
                    engine = LocalOCR()
                result = engine.recognize(image_path)
                if result and result.get('success'):
                    return result
            except Exception as e:
                logger.warning(f"[OCR] {name} 失败: {e}")
        return {'success': False, 'error': '无可用OCR引擎', 'text': ''}


# ============================================================
# IMA同步器
# ============================================================
class IMASyncer:
    def __init__(self):
        self.client_id = os.getenv('IMA_OPENAPI_CLIENTID', '')
        self.api_key = os.getenv('IMA_OPENAPI_APIKEY', '')
        self.base_url = 'https://ima.qq.com/openapi/note/v1'
        self.enabled = bool(self.client_id and self.api_key and '填入' not in self.client_id)
        self.sync_log_file = Path('处理结果/ima_sync_log.json')
        self.sync_log = self._load_sync_log()
        self.default_notebook_id = os.getenv('IMA_NOTEBOOK_ID', '')
        self.rate_limited = False
        self.synced_this_session = 0
        self.failed_this_session = 0

    def _load_sync_log(self) -> Dict:
        if self.sync_log_file.exists():
            try:
                return json.loads(self.sync_log_file.read_text(encoding='utf-8'))
            except Exception:
                pass
        return {}

    def _save_sync_log(self):
        self.sync_log_file.write_text(json.dumps(self.sync_log, ensure_ascii=False, indent=2), encoding='utf-8')

    def _api_call(self, endpoint: str, payload: Dict, retries: int = 3) -> Optional[Dict]:
        if not self.enabled or self.rate_limited:
            return None
        for attempt in range(retries):
            try:
                headers = {
                    'ima-openapi-clientid': self.client_id,
                    'ima-openapi-apikey': self.api_key,
                    'Content-Type': 'application/json'
                }
                response = requests.post(f"{self.base_url}/{endpoint}", json=payload, headers=headers, timeout=30)
                if response.status_code == 200:
                    return response.json()
                elif response.status_code == 403:
                    result = response.json()
                    if '请求超量' in result.get('msg', '') and attempt < retries - 1:
                        time.sleep(60)
                        continue
                return None
            except Exception as e:
                if attempt < retries - 1:
                    time.sleep(2)
        return None

    def sync_note(self, title: str, content: str, category: str = None,
                  content_hash: str = None, doc_path: str = None) -> Optional[str]:
        if not self.enabled or self.rate_limited:
            return None

        doc_key = content_hash or title
        existing = self.sync_log.get(doc_key, {})

        full_content = f"# {title}\n\n"
        if category:
            full_content += f"> 分类: {category}\n"
        if content_hash:
            full_content += f"> 内容标识: {content_hash}\n"
        full_content += f"\n{content}\n\n---\n*自动同步自图片知识库*\n"

        if existing.get('doc_id') and doc_path:
            result = self._api_call('append_doc', {
                'doc_id': existing['doc_id'],
                'content_format': 1,
                'content': f"\n\n---\n*更新于 {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n" + content
            })
            if result:
                self.sync_log[doc_key]['last_sync'] = datetime.now().isoformat()
                self.sync_log[doc_key]['update_count'] = existing.get('update_count', 0) + 1
                self._save_sync_log()
                self.synced_this_session += 1
                return existing['doc_id']

        import_payload = {'content_format': 1, 'content': full_content}
        if self.default_notebook_id:
            import_payload['notebook_id'] = self.default_notebook_id

        result = self._api_call('import_doc', import_payload)
        if result and result.get('code') == 0:
            doc_id = result.get('data', {}).get('note_id', 'imported')
            self.sync_log[doc_key] = {
                'doc_id': doc_id, 'title': title, 'category': category,
                'first_sync': datetime.now().isoformat(), 'last_sync': datetime.now().isoformat(),
                'update_count': 1
            }
            self._save_sync_log()
            self.synced_this_session += 1
            return doc_id
        else:
            self.failed_this_session += 1
            return None


# ============================================================
# 分批处理
# ============================================================
_BATCH_DIR = Path(__file__).parent / 'progress'
BATCH_STATE_FILE = _BATCH_DIR / 'auto_batch_state.json'

BATCH_CONFIG = {
    "small":  {"max_size": 500 * 1024,       "batch_size": 10},
    "medium": {"max_size": 2 * 1024 * 1024,  "batch_size": 6},
    "large":  {"max_size": 5 * 1024 * 1024,  "batch_size": 4},
    "xlarge": {"max_size": float('inf'),      "batch_size": 2},
}

@dataclass
class ImageInfo:
    path: str
    name: str
    size: int

@dataclass
class Batch:
    batch_id: str
    images: List[str]
    size: int
    count: int

class BatchManager:
    def __init__(self):
        _BATCH_DIR.mkdir(parents=True, exist_ok=True)
        self.state_file = BATCH_STATE_FILE
        self.state = self._load_state()

    def _load_state(self) -> Dict:
        if self.state_file.exists():
            try:
                return json.loads(self.state_file.read_text(encoding='utf-8'))
            except Exception:
                pass
        return {'total_images': 0, 'total_batches': 0, 'completed_batches': 0,
                'current_batch': 0, 'processed_images': [], 'failed_images': [],
                'start_time': None, 'end_time': None}

    def _save_state(self):
        self.state_file.write_text(json.dumps(self.state, ensure_ascii=False, indent=2), encoding='utf-8')

    def initialize(self) -> Optional[Dict]:
        source_dir = Path('待处理图片')
        if not source_dir.exists():
            print("待处理图片目录不存在")
            return None

        images = []
        for ext in ['*.jpg', '*.jpeg', '*.png', '*.webp', '*.bmp']:
            images.extend(source_dir.rglob(ext))

        if not images:
            print("没有找到待处理图片")
            return None

        all_images = []
        for img in images:
            all_images.append(ImageInfo(str(img), img.name, img.stat().st_size))

        batches = self._create_batches(all_images)
        self.state = {
            'total_images': len(all_images),
            'total_batches': len(batches),
            'completed_batches': 0,
            'current_batch': 0,
            'processed_images': [],
            'failed_images': [],
            'batches': [{'id': f"batch_{i+1}", 'images': b.images, 'size': b.size, 'count': b.count}
                        for i, b in enumerate(batches)],
            'start_time': datetime.now().isoformat(),
            'end_time': None,
        }
        self._save_state()
        print(f"已初始化: {len(all_images)}张图片, {len(batches)}个批次")
        return self.state

    def _create_batches(self, images: List[ImageInfo]) -> List[Batch]:
        # 按大小排序（大的先处理，容易早发现问题）
        images.sort(key=lambda x: x.size, reverse=True)
        batches = []
        current = []
        current_size = 0
        current_count = 0

        def get_batch_size(img_size):
            for cfg_name, cfg in BATCH_CONFIG.items():
                if img_size <= cfg['max_size']:
                    return cfg['batch_size']
            return BATCH_CONFIG['xlarge']['batch_size']

        for img in images:
            batch_size = get_batch_size(img.size)
            if current_count >= batch_size:
                batches.append(Batch(
                    batch_id=f"batch_{len(batches)+1}",
                    images=[i.path for i in current],
                    size=sum(i.size for i in current),
                    count=len(current)
                ))
                current = []
                current_size = 0
                current_count = 0
            current.append(img)
            current_size += img.size
            current_count += 1

        if current:
            batches.append(Batch(
                batch_id=f"batch_{len(batches)+1}",
                images=[i.path for i in current],
                size=sum(i.size for i in current),
                count=len(current)
            ))

        return batches

    def get_next_batch(self) -> Optional[Batch]:
        batches = self.state.get('batches', [])
        idx = self.state['completed_batches']
        if idx >= len(batches):
            return None
        batch_data = batches[idx]
        return Batch(batch_data['id'], batch_data['images'], batch_data['size'], batch_data['count'])

    def mark_batch_completed(self, batch_id: str, note: str = ""):
        self.state['completed_batches'] += 1
        self._save_state()

    def mark_batch_failed(self, batch_id: str, note: str = ""):
        self.state['completed_batches'] += 1
        self._save_state()

    def clear_state(self):
        self.state = self._load_state()
        self._save_state()
        print("分批状态已清除")

    def print_progress(self):
        s = self.state
        print(f"总图片: {s.get('total_images', 0)}")
        print(f"总批次: {s.get('total_batches', 0)}")
        print(f"已完成: {s.get('completed_batches', 0)}")
        print(f"剩余: {s.get('total_batches', 0) - s.get('completed_batches', 0)}")


# ============================================================
# 单张图片处理
# ============================================================
def process_single_image(ocr, analyzer, organizer, doc_manager, ima_syncer, archiver,
                         image_path: str, index: int, total: int, progress=None):
    """处理单张图片：OCR → LLM分析 → 文档生成 → IMA → 归档"""
    t_start = time.time()
    image_name = Path(image_path).name

    # 步骤1：OCR识别
    result = ocr.recognize(image_path)
    if not result or not result.get('success'):
        msg = f"  ❌ OCR失败: {result.get('error', '未知错误') if result else '无结果'}"
        if progress:
            progress.log(msg)
        else:
            print(msg)
        archiver.archive_image(image_path, "其他")
        return None

    raw_text = result.get('text', '').strip()
    if not raw_text or len(raw_text) < 10:
        msg = "  ⏭️ 无文字内容，跳过"
        if progress:
            progress.log(msg)
        else:
            print(msg)
        archiver.archive_image(image_path, "其他")
        return None

    # 步骤2：内容整理
    organized = organizer.organize(raw_text)
    clean_text = organized['content']
    source_info = organized['source']

    # 步骤3：LLM分析（命名+分类+摘要+内容重写）
    analysis = analyzer.analyze(clean_text, index)

    # V9.3：LLM失败时跳过，图片留在原地待人工处理，不生成错误文档
    if analysis.get('_skip'):
        msg = f"  ⏭️ LLM失败，跳过（图片留待人工处理）: {image_name}"
        if progress:
            progress.log(msg)
        else:
            print(msg)
        logger.warning(f"[跳过] LLM失败，未归档，待人工处理: {image_name}")
        return {'image': image_name, 'failed': True, 'error': 'LLM调用失败，待人工处理',
                'content_hash': analysis.get('content_hash', '')}

    category = analysis['category']
    confidence = analysis['confidence']
    doc_name = analysis['doc_name']
    keywords = analysis['keywords']
    content_hash = analysis['content_hash']
    summary = analysis.get('summary', '')
    # V9.1：优先使用LLM整理后的内容，若LLM未返回则退回OCR清理文本
    llm_content = analysis.get('llm_content', '').strip()
    final_content = llm_content if llm_content else clean_text

    # 步骤4：重复检测
    if doc_manager.is_duplicate_hash(content_hash):
        archiver.archive_image(image_path, category)
        elapsed = time.time() - t_start
        msg = f"  🔄 重复内容 | {category} | {doc_name} ({elapsed:.1f}s)"
        if progress:
            progress.log(msg)
        else:
            print(msg)
        return {'image': image_name, 'category': category, 'doc_name': doc_name, 'is_duplicate': True, 'content_hash': content_hash}

    # 步骤5+6：合并判断 + 生成文档（使用LLM整理后内容）
    md_file, is_new_doc = doc_manager.save_document(
        doc_name=doc_name, category=category, content=final_content,
        image_name=image_name, keywords=keywords,
        content_hash=content_hash, source_info=source_info,
        summary=summary
    )
    docx_file = None
    if is_new_doc:
        docx_file = doc_manager.generate_word(
            doc_name=doc_name, category=category, content=final_content,
            image_name=image_name, keywords=keywords, md_file_path=md_file
        )

    # 步骤7：IMA同步
    if ima_syncer.enabled and not ima_syncer.rate_limited:
        ima_syncer.sync_note(doc_name, final_content, category, content_hash, md_file)

    # 步骤8：图片归档
    archiver.archive_image(image_path, category)

    # 汇总（静默更新进度条）
    elapsed = time.time() - t_start
    action = "新建" if is_new_doc else "合并"
    source_tag = f" | {source_info['platform']}" if source_info.get('platform') else ""
    engine_tag = f" | {analysis.get('_engine', '正则')}" if analysis.get('_engine') else ""
    if progress:
        progress.update(index, f"✅ {doc_name[:12]}{source_tag}{engine_tag}")
    else:
        print(f"  ✅ [{action}] {doc_name}{source_tag}{engine_tag} ({elapsed:.1f}s)")

    return {
        'image': image_name, 'text_length': len(final_content),
        'category': category, 'confidence': round(confidence, 3),
        'doc_name': doc_name, 'keywords': keywords,
        'is_duplicate': False, 'is_new_doc': is_new_doc,
        'md_file': md_file, 'docx_file': docx_file,
        'content_hash': content_hash, 'source': source_info,
    }


# ============================================================
# 报告生成
# ============================================================
def save_report(results: List[Dict], total_images: int, elapsed: float, mode: str = 'full'):
    valid = [r for r in results if r and not r.get('failed')]
    duplicates = [r for r in valid if r.get('is_duplicate')]
    new_docs = [r for r in valid if not r.get('is_duplicate') and r.get('is_new_doc')]
    merged_docs = [r for r in valid if not r.get('is_duplicate') and not r.get('is_new_doc')]
    md_count = sum(1 for r in valid if r.get('md_file') and not r.get('is_duplicate'))
    docx_count = sum(1 for r in valid if r.get('docx_file') and not r.get('is_duplicate'))
    failed = [r for r in results if r and r.get('failed')]
    cat_stats = dict(Counter(r['category'] for r in valid if not r.get('is_duplicate') and r.get('category')))

    report = {
        'timestamp': datetime.now().isoformat(), 'version': 'V9.0', 'mode': mode,
        'total_images': total_images, 'processed_successfully': len(valid),
        'failed': len(failed), 'duplicates_skipped': len(duplicates),
        'new_documents': len(new_docs), 'merged_documents': len(merged_docs),
        'markdown_files': md_count, 'word_files': docx_count,
        'elapsed_seconds': round(elapsed, 1), 'category_stats': cat_stats,
        'results': [{k: v for k, v in r.items() if k not in ['md_file', 'docx_file']} for r in results if r]
    }

    Path('处理结果').mkdir(exist_ok=True)
    report_file = Path('处理结果/处理报告.json')
    with open(report_file, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    return report, report_file


def print_summary(report: Dict, archiver: ImageArchiver, ima_syncer: IMASyncer):
    print("\n" + "=" * 60)
    print("处理完成 - V9.0 汇总")
    print("=" * 60)
    print(f"  总图片数:     {report['total_images']} 张")
    print(f"  成功处理:     {report['processed_successfully']} 张")
    print(f"  处理失败:     {report['failed']} 张")
    print(f"  重复跳过:     {report['duplicates_skipped']} 张")
    print(f"  新建文档:     {report['new_documents']} 个")
    print(f"  合并到已有:   {report['merged_documents']} 次")
    print(f"  Markdown:     {report['markdown_files']} 个")
    print(f"  Word:         {report['word_files']} 个")
    print(f"  总耗时:       {int(report['elapsed_seconds']//60)}分{int(report['elapsed_seconds']%60)}秒")

    if report['category_stats']:
        print("\n分类分布:")
        for cat, cnt in sorted(report['category_stats'].items(), key=lambda x: -x[1]):
            print(f"  {cat}: {cnt} 张")

    arch_stats = archiver.get_stats()
    print(f"\n归档统计:")
    print(f"  本次归档: {arch_stats['archived_this_session']} 张")
    print(f"  总已归档: {arch_stats['total_archived']} 张")

    if ima_syncer.enabled:
        print(f"\nIMA同步:")
        print(f"  本次成功: {ima_syncer.synced_this_session} 条")
        print(f"  本次失败: {ima_syncer.failed_this_session} 条")

    print("=" * 60)


# ============================================================
# 全自动处理模式
# ============================================================
def run_full_mode(batch_manager: BatchManager):
    print("=" * 60)
    print("全自动化图片处理工具 V9.3 (合并逻辑重构)")
    print("OCR → LLM分析(命名+分类+摘要) → 文档 → Word → IMA → 归档")
    print("合并策略：精确同名匹配（V9.3新）| LLM失败跳过不兜底（V9.3新）")
    print("=" * 60)

    print("\n[初始化] OCR引擎...")
    ocr = MultiEngineOCR()
    if not ocr.current_engine:
        print("错误: 没有可用的OCR引擎!")
        return

    print("[初始化] LLM分析器(混元Lite)...")
    analyzer = LLMAnalyzer()

    print("[初始化] 内容整理器...")
    organizer = ContentOrganizer()

    print("[初始化] 文档管理器...")
    doc_manager = SmartDocumentManager()

    print("[初始化] IMA同步器...")
    ima_syncer = IMASyncer()

    print("[初始化] 图片归档器...")
    archiver = ImageArchiver()

    source_dir = Path('待处理图片')
    if not source_dir.exists():
        print("错误: 待处理图片目录不存在!")
        return

    images = []
    for ext in ['*.jpg', '*.jpeg', '*.png', '*.webp', '*.bmp']:
        images.extend(source_dir.rglob(ext))

    if not images:
        print("没有找到待处理图片")
        return

    total = len(images)
    print(f"\n找到 {total} 张待处理图片")

    progress = ProgressBar(total, prefix="处理进度")
    results = []
    start_time = time.time()

    for i, img_path in enumerate(images, 1):
        try:
            result = process_single_image(
                ocr, analyzer, organizer, doc_manager, ima_syncer, archiver,
                str(img_path), i, total, progress
            )
            if result:
                results.append(result)
        except Exception as e:
            msg = f"\n  ❌ 处理失败: {e}"
            if progress:
                progress.log(msg)
            else:
                print(msg)
            logger.error(f"处理失败 {img_path}: {e}")
            results.append({'image': Path(img_path).name, 'failed': True, 'error': str(e)})

    progress.finish()

    elapsed = time.time() - start_time
    report, report_file = save_report(results, total, elapsed, 'full')
    print_summary(report, archiver, ima_syncer)
    print(f"\n报告已保存: {report_file}")


# ============================================================
# 分批处理模式
# ============================================================
def run_batch_mode(batch_manager: BatchManager):
    print("\n" + "=" * 60)
    print("分批处理模式 V9.3")
    print("=" * 60)

    state = batch_manager.initialize()
    if not state:
        return

    batch_manager.print_progress()

    print("\n[初始化] OCR引擎...")
    ocr = MultiEngineOCR()
    if not ocr.current_engine:
        print("错误: 没有可用的OCR引擎!")
        return

    analyzer = LLMAnalyzer()
    organizer = ContentOrganizer()
    doc_manager = SmartDocumentManager()
    ima_syncer = IMASyncer()
    archiver = ImageArchiver()

    all_results = []
    batch_num = 0
    total_images = state.total_images
    total_batches = state.total_batches
    start_time = time.time()

    while True:
        batch = batch_manager.get_next_batch()
        if not batch:
            print("\n所有批次处理完成！")
            break

        batch_num += 1
        logger.info(f"[批次] {batch_num}/{total_batches}: {batch.batch_id} ({len(batch.images)}张)")

        batch_results = []
        batch_success = True
        batch_progress = ProgressBar(len(batch.images), prefix=f"批次{batch_num}")

        for i, img_path in enumerate(batch.images, 1):
            try:
                result = process_single_image(
                    ocr, analyzer, organizer, doc_manager, ima_syncer, archiver,
                    img_path, i, len(batch.images), batch_progress
                )
                if result:
                    batch_results.append(result)
                else:
                    batch_results.append({'image': Path(img_path).name, 'failed': True})
            except Exception as e:
                print(f"\n  错误: {e}")
                logger.error(f"处理失败 {img_path}: {e}")
                batch_results.append({'image': Path(img_path).name, 'failed': True, 'error': str(e)})
                batch_success = False

        batch_progress.finish()

        if batch_success:
            batch_manager.mark_batch_completed(batch.batch_id, f"处理 {len(batch_results)} 张")
        else:
            failed_n = sum(1 for r in batch_results if r.get('failed'))
            batch_manager.mark_batch_failed(batch.batch_id, f"失败 {failed_n} 张")

        all_results.extend(batch_results)

        if batch_num < total_batches:
            print("\n[批次间隔] 1 秒...")
            time.sleep(1)

    elapsed = time.time() - start_time
    report, report_file = save_report(all_results, total_images, elapsed, 'batch')
    print_summary(report, archiver, ima_syncer)
    print(f"\n报告已保存: {report_file}")
    return all_results


# ============================================================
# 主入口
# ============================================================
def print_usage():
    print("\n用法:")
    print("  python auto_process_all_v9_3.py           # 全自动处理")
    print("  python auto_process_all_v9_3.py --batch    # 强制分批模式")
    print("  python auto_process_all_v9_3.py --init     # 仅初始化分批")
    print("  python auto_process_all_v9_3.py --progress # 查看进度")
    print("  python auto_process_all_v9_3.py --clear    # 清除分批状态")


def main():
    batch_manager = BatchManager()
    if len(sys.argv) > 1:
        cmd = sys.argv[1]
        if cmd == '--init':
            state = batch_manager.initialize()
            if state:
                batch_manager.print_progress()
        elif cmd == '--progress':
            batch_manager.print_progress()
        elif cmd == '--clear':
            batch_manager.clear_state()
        elif cmd == '--batch':
            run_batch_mode(batch_manager)
        else:
            print(f"[ERROR] 未知参数: {cmd}")
            print_usage()
    else:
        run_full_mode(batch_manager)


if __name__ == '__main__':
    main()
