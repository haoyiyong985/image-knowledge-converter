#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成小白版 Word 配置说明文档
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = Path("D:/新建文件夹/分享包/图片知识库_配置说明（小白版）.docx")
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 工具函数 ──
def set_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    size = {1: 18, 2: 14, 3: 12}[level]
    color = {1: (0x1a, 0x56, 0x76), 2: (0x2e, 0x86, 0xab), 3: (0x33, 0x33, 0x33)}[level]
    set_font(run, size=size, bold=True, color=color)
    p.paragraph_format.space_before = Pt(16 if level==1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_para(doc, text, indent=0, bold_parts=None, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(indent)
    if bold_parts:
        # bold_parts: [(文字, is_bold), ...]
        for txt, is_bold in bold_parts:
            run = p.add_run(txt)
            set_font(run, bold=is_bold, color=color)
    else:
        run = p.add_run(text)
        set_font(run, color=color)
    return p

def add_step(doc, num, title, body_lines):
    """添加带编号的步骤块"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"第 {num} 步  {title}")
    set_font(run, size=12, bold=True, color=(0x2e, 0x86, 0xab))

    for line in body_lines:
        sub = doc.add_paragraph()
        sub.paragraph_format.left_indent = Cm(0.8)
        sub.paragraph_format.space_after  = Pt(2)
        if line.startswith(">>"):
            # 提示框样式
            run = sub.add_run("  提示：" + line[2:].strip())
            set_font(run, size=10, color=(0x80, 0x80, 0x80))
        elif line.startswith("**") and line.endswith("**"):
            run = sub.add_run(line[2:-2])
            set_font(run, bold=True, size=11)
        else:
            run = sub.add_run(line)
            set_font(run, size=11)

def add_tip_box(doc, text):
    """添加提示框段落（灰色背景模拟）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("💡  " + text)
    set_font(run, size=10, color=(0x1a, 0x56, 0x76))
    # 设置段落底纹
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E8F4F8')
    pPr.append(shd)
    return p

def add_table_row(table, col1, col2, header=False):
    row = table.add_row()
    for i, text in enumerate([col1, col2]):
        cell = row.cells[i]
        cell.text = text
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(text)
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.text = text
        set_font(run, bold=header, size=10,
                 color=(0x1a, 0x56, 0x76) if header else (0x33, 0x33, 0x33))
    return row

# ════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("图片知识库转化工具")
set_font(run, size=26, bold=True, color=(0x1a, 0x56, 0x76))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("v1.0 — 配置说明（小白版）")
set_font(run2, size=16, color=(0x2e, 0x86, 0xab))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("发布日期：2026年3月17日")
set_font(run3, size=11, color=(0x99, 0x99, 0x99))


p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(20)
run3 = p3.add_run("把手机截图自动整理成 Word + Markdown 知识文档，并同步到 ima 个人笔记")
set_font(run3, size=12, color=(0x66, 0x66, 0x66))

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(80)
run4 = p4.add_run("阅读本文档约需 5 分钟，配置完成约需 10 分钟")
set_font(run4, size=10, color=(0x99, 0x99, 0x99))

doc.add_page_break()

# ════════════════════════════════════════════════
# 第一章：这个工具是做什么的
# ════════════════════════════════════════════════
add_heading(doc, "一、这个工具是做什么的", 1)

add_para(doc, "你是否有这样的烦恼：")
for item in [
    "手机里存了大量小红书、微信读书的截图，想回头看却找不到",
    "截图里有很多有用的文字，但一张张手动整理太累了",
    "想把同类内容整理到一起，但不知道从哪里开始",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(item)
    set_font(run, size=11)

add_para(doc, "")
add_para(doc, "这个工具可以帮你：")
for item in [
    "自动识别截图里的文字（无需手动输入）",
    "自动判断内容属于哪个分类（养生、饮食、中医……）",
    "自动整理成 Word 文档和 Markdown 文件",
    "同类内容自动合并，不会重复",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(item)
    set_font(run, size=11, color=(0x1a, 0x56, 0x76))

add_tip_box(doc, "你只需要把截图放进指定文件夹，然后在 WorkBuddy 里说一句话，剩下的全交给 AI。")

# ════════════════════════════════════════════════
# 第二章：你需要准备什么
# ════════════════════════════════════════════════
add_heading(doc, "二、你需要准备什么", 1)

add_heading(doc, "2.1  软件要求", 2)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.columns[0].width = Cm(5)
table.columns[1].width = Cm(10)
hdr = table.rows[0]
for i, txt in enumerate(["软件", "说明"]):
    cell = hdr.cells[i]
    run = cell.paragraphs[0].add_run(txt)
    set_font(run, bold=True, size=10, color=(0x1a, 0x56, 0x76))

for sw, desc in [
    ("Python 3.8 以上", "运行处理脚本的环境（免费）"),
    ("python-docx 库",  "生成 Word 文档用（一条命令安装）"),
    ("WorkBuddy",       "AI 对话工具，用来识别图片和智能分类"),
    ("Typora（推荐）",  "查看 MD 文件用，打开效果和 Word 一样整洁（可选，免费版搜索「Typora 1.0」）"),
]:
    row = table.add_row()
    for i, txt in enumerate([sw, desc]):
        run = row.cells[i].paragraphs[0].add_run(txt)
        set_font(run, size=10)

doc.add_paragraph()

add_heading(doc, "2.2  分享包里有什么", 2)

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
table2.columns[0].width = Cm(6)
table2.columns[1].width = Cm(3)
table2.columns[2].width = Cm(6)
for i, txt in enumerate(["文件名", "类型", "作用"]):
    run = table2.rows[0].cells[i].paragraphs[0].add_run(txt)
    set_font(run, bold=True, size=10, color=(0x1a, 0x56, 0x76))

for fname, ftype, desc in [
    ("auto_process.py",       "脚本", "核心处理程序，每次使用前运行它"),
    ("ima_sync.py",           "脚本", "自动同步到 ima 个人笔记"),
    ("ima_config.txt",        "配置", "填入 ima API 凭证（首次配置时填写）"),
    ("test_flow.py",          "脚本", "验证配置是否正确"),
    ("使用说明.md",            "文档", "日常操作手册（中文）"),
    ("本配置说明.docx",        "文档", "你正在看的这份文档"),
    ("【启动提示词】…txt",     "文本", "新对话时粘贴这里面的内容"),
]:
    row = table2.add_row()
    for i, txt in enumerate([fname, ftype, desc]):
        run = row.cells[i].paragraphs[0].add_run(txt)
        set_font(run, size=10)

doc.add_paragraph()

# ════════════════════════════════════════════════
# 第三章：第一次配置（只需做一次）
# ════════════════════════════════════════════════
add_heading(doc, "三、第一次配置（只需做一次）", 1)
add_tip_box(doc, "以下步骤只需要做一次。配置完成后，以后每次用只需要「放图片 → 打开WorkBuddy → 说一句话」。")
doc.add_paragraph()

# 步骤1
add_step(doc, 1, "安装 Python", [
    "打开浏览器，搜索「Python 官网」，或直接访问：https://www.python.org",
    "点击「Downloads」，下载最新版本（3.10 或以上）",
    "运行下载的安装包，安装时务必勾选「Add Python to PATH」选项",
    ">> 如果已经安装过 Python，可以跳过此步骤",
])

# 步骤2
add_step(doc, 2, "安装 python-docx 库", [
    "按下键盘 Windows键 + R，输入 cmd，按回车",
    "在黑色窗口（命令提示符）里输入以下内容，按回车：",
    "**pip install python-docx**",
    "等待安装完成（会显示 Successfully installed），然后关闭窗口",
    ">> 需要连接互联网才能安装",
])

# 步骤3
add_step(doc, 3, "把分享包文件放到你电脑上", [
    "在你的电脑上选一个方便的位置，新建一个文件夹",
    "例如：D:\\我的知识库  或  桌面\\图片整理工具",
    "把分享包里的所有文件复制到这个文件夹里",
    ">> 文件夹路径中最好不要有特殊符号，中文名称没有问题",
])

# 步骤4
add_step(doc, 4, "修改脚本中的路径", [
    "找到文件夹里的 auto_process.py，右键 → 用记事本打开",
    "找到第 34 行，内容如下：",
    "**BASE_DIR = Path(\"D:/新建文件夹\")**",
    "把引号里面的路径改成你自己的文件夹路径，例如：",
    "**BASE_DIR = Path(\"D:/我的知识库\")**",
    "改完后 Ctrl+S 保存，关闭记事本",
    ">> 路径中的斜杠用 / 或 \\\\ 都可以，但要保持一致",
])

# 步骤5
add_step(doc, 5, "建立所需文件夹", [
    "在你的文件夹里，手动新建以下三个子文件夹：",
    "**待处理图片**（以后新图片放这里）",
    "**已处理图片**（程序自动归档用，建好就行）",
    "**处理结果**（生成的 Word 和 MD 文件放这里）",
    ">> 文件夹名称必须和上面完全一致，包括中文",
])

# 步骤6
add_step(doc, 6, "验证配置是否正确", [
    "按下 Windows键 + R，输入 cmd，按回车",
    "在命令窗口输入（把路径换成你自己的）：",
    "**python D:/我的知识库/test_flow.py**",
    "如果看到「全部测试通过」，说明配置成功！",
    "如果看到 FAIL 开头的行，请截图发给分享给你工具的朋友",
])

add_tip_box(doc, "完成以上6步，配置就全部完成了！以后不需要再做这些操作。")

# ════════════════════════════════════════════════
# 第四章：日常使用方法（配置完后每次这样用）
# ════════════════════════════════════════════════
add_heading(doc, "四、日常使用方法", 1)
add_para(doc, "配置完成后，每次有新截图要整理，按以下 3 步操作：")
doc.add_paragraph()

add_step(doc, 1, "把截图放进对应文件夹", [
    "在「待处理图片」文件夹里，按主题新建子文件夹",
    "例如：待处理图片\\健康养生\\  或  待处理图片\\旅行记录\\",
    "把手机截图复制进对应的子文件夹",
    ">> 子文件夹的名字随意，中英文都行，这个名字就是这批图片的主题标签",
])

add_step(doc, 2, "打开 WorkBuddy，粘贴启动提示词", [
    "打开电脑上的 WorkBuddy，点击「新建对话」",
    "打开文件夹里的「【启动提示词】新对话粘贴这段.txt」",
    "全选内容（Ctrl+A）→ 复制（Ctrl+C）",
    "粘贴到 WorkBuddy 对话框（Ctrl+V）→ 按发送",
    "等 AI 回复「已准备好」",
    ">> WorkBuddy 每次新对话都是空白的，必须先粘贴这段内容，AI 才知道你的项目",
])

add_step(doc, 3, "发送指令，开始处理", [
    "AI 准备好后，在对话框里输入：处理新图片",
    "AI 会自动完成识别→分类→整理→归档全部工作",
    "完成后去「处理结果」文件夹查看生成的 Word 和 MD 文档",
    ">> 处理完的截图会自动移动到「已处理图片」文件夹，不用手动操作",
])

# ════════════════════════════════════════════════
# 第五章：常见问题
# ════════════════════════════════════════════════
add_heading(doc, "五、常见问题", 1)

faqs = [
    ("问：图片识别不准确怎么办？",
     "答：图片模糊或字体很小时识别效果会差一些。建议截图时分辨率高一点。AI 会跳过无法识别的图片并告知你。"),
    ("问：内容被归错类了怎么办？",
     "答：告诉 WorkBuddy「把刚才的内容从 XX 文档移到 YY 文档」，AI 会帮你调整。"),
    ("问：重启电脑后怎么继续用？",
     "答：打开 WorkBuddy，粘贴启动提示词（第四章第2步），然后正常发指令即可。"),
    ("问：朋友发来的文件要怎么更新？",
     "答：处理完一批图片后，运行 auto_process.py，它会自动更新启动提示词文件，不需要找朋友重发。"),
    ("问：我的截图内容和工具默认分类不同怎么办？",
     "答：直接告诉 WorkBuddy，比如「这批图片是旅行相关的，帮我新建旅行记录文档」，AI 会自动创建新分类。"),
    ("问：运行脚本时出现英文报错怎么办？",
     "答：截图发给分享你工具的朋友，或者在 WorkBuddy 里粘贴启动提示词后，把报错内容发给 AI，它会帮你解决。"),
    ("问：MD 文件用什么软件打开？",
     "答：推荐用 Typora，打开效果和 Word 一样整洁，不会显示 # 和 ** 这些符号。"
     "可搜索「Typora 1.0 免费版」下载。也可以用 VS Code（免费，微软出品）打开后点右上角预览按钮。"
     "实在不想装软件，用记事本也能打开，只是格式符号会原样显示。"),
]

for q, a in faqs:
    p_q = doc.add_paragraph()
    p_q.paragraph_format.space_before = Pt(8)
    run_q = p_q.add_run(q)
    set_font(run_q, bold=True, size=11, color=(0x2e, 0x86, 0xab))

    p_a = doc.add_paragraph()
    p_a.paragraph_format.left_indent = Cm(0.5)
    p_a.paragraph_format.space_after = Pt(4)
    run_a = p_a.add_run(a)
    set_font(run_a, size=11)

# ════════════════════════════════════════════════
# 第六章：ima 个人笔记自动同步（进阶功能）
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "六、ima 个人笔记自动同步（可选进阶功能）", 1)

add_para(doc, "配置完成后，每次处理图片时，生成的文档会自动同步到你的 ima 个人笔记，无需任何手动操作。")
add_para(doc, "ima 是腾讯出品的 AI 知识库工具，支持智能检索和问答，非常适合作为长期知识积累的平台。")
doc.add_paragraph()

add_tip_box(doc, "ima 同步是可选功能。如果你不使用 ima，跳过本章即可，工具其他功能完全不受影响。")
doc.add_paragraph()

add_heading(doc, "6.1  获取 ima API 凭证（只需做一次）", 2)

add_step(doc, 1, "获取 API 凭证", [
    "用浏览器打开：https://ima.qq.com/agent-interface",
    "用你的 QQ 或微信账号登录",
    "找到「创建应用」或「获取 API Key」按钮，点击生成",
    "页面会显示两个值：Client ID 和 API Key",
    ">> 重要：API Key 只显示一次，请立即复制保存到记事本！不要关闭页面",
    ">> 凭证有效期通常为 30 天，到期后重新获取即可",
])

add_step(doc, 2, "填写凭证配置文件", [
    "找到文件夹里的 ima_config.txt，右键 → 用记事本打开",
    "找到以下两行，把占位文字替换成你刚才复制的真实值：",
    "**IMA_CLIENT_ID=填入你的ClientID**",
    "**IMA_API_KEY=填入你的APIKey**",
    "改完后 Ctrl+S 保存，关闭记事本",
    ">> 等号两边不要加空格，不要加引号",
])

add_step(doc, 3, "验证凭证是否有效", [
    "按 Windows键 + R，输入 cmd，按回车",
    "在命令窗口输入（路径换成你自己的）：",
    "**python D:/我的知识库/ima_sync.py --check**",
    "如果看到「凭证验证通过」，说明配置成功！",
    ">> 如果显示失败，检查凭证是否完整复制，有没有多余空格",
])

add_tip_box(doc, "凭证配置完成后，以后每次处理图片，文档会自动同步到 ima，完全不需要你再操作任何东西。")
doc.add_paragraph()

add_heading(doc, "6.2  ima 同步规则说明", 2)

table_ima = doc.add_table(rows=1, cols=2)
table_ima.style = 'Table Grid'
table_ima.columns[0].width = Cm(5)
table_ima.columns[1].width = Cm(10)
for i, txt in enumerate(["情况", "处理方式"]):
    run = table_ima.rows[0].cells[i].paragraphs[0].add_run(txt)
    set_font(run, bold=True, size=10, color=(0x1a, 0x56, 0x76))

for situation, handling in [
    ("文档首次同步",         "在 ima 「笔记」里新建一篇笔记"),
    ("文档有新内容（追加）", "自动在 ima 原笔记末尾追加新内容"),
    ("文档未变更",           "跳过，不重复同步"),
    ("凭证未填写",           "跳过 ima 同步，其他功能正常"),
]:
    row = table_ima.add_row()
    for i, txt in enumerate([situation, handling]):
        run = row.cells[i].paragraphs[0].add_run(txt)
        set_font(run, size=10)

doc.add_paragraph()

add_heading(doc, "6.3  凭证到期后如何更新", 2)
add_para(doc, "API 凭证到期后，ima 同步会停止，但图片处理功能完全不受影响。更新步骤：")
for item in [
    "重新访问 https://ima.qq.com/agent-interface 获取新凭证",
    "打开 ima_config.txt，替换 IMA_CLIENT_ID 和 IMA_API_KEY 两行的值",
    "保存文件，同步自动恢复",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(item)
    set_font(run, size=11)

doc.add_paragraph()

# ════════════════════════════════════════════════
# 第七章：快速参考卡
# ════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "七、快速参考卡（可剪下来放桌上）", 1)
add_tip_box(doc, "配置完成后，每次使用只需记住下面这张卡片的内容。")
doc.add_paragraph()

table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Table Grid'
table3.columns[0].width = Cm(7)
table3.columns[1].width = Cm(8)
for i, txt in enumerate(["操作", "怎么做"]):
    run = table3.rows[0].cells[i].paragraphs[0].add_run(txt)
    set_font(run, bold=True, size=11, color=(0x1a, 0x56, 0x76))

for op, how in [
    ("有新截图要整理",    "放入「待处理图片\\主题名\\」文件夹"),
    ("开始处理",          "WorkBuddy → 粘贴启动提示词 → 发送「处理新图片」"),
    ("自动同步到 ima",    "处理完成后自动同步，无需操作（需配置第六章凭证）"),
    ("只处理某个主题",    "发送「处理待处理图片/主题名」"),
    ("查看有哪些文档",    "发送「查看现有文档」"),
    ("工具优化",          "发送「工具优化」提升工具处理能力（进阶功能）"),
    ("出问题了",          "截图 + 在 WorkBuddy 里描述问题，AI 帮你排查"),
    ("重启电脑后",        "重新粘贴启动提示词，其他不变"),
    ("ima 凭证到期",      "重新获取凭证，更新 ima_config.txt 两行内容"),
]:
    row = table3.add_row()
    for i, txt in enumerate([op, how]):
        run = row.cells[i].paragraphs[0].add_run(txt)
        set_font(run, size=11)

doc.add_paragraph()
doc.add_paragraph()

p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_end = p_end.add_run("如有疑问，请联系分享给你此工具的朋友  |  ima 同步配置详见第六章")
set_font(run_end, size=10, color=(0x99, 0x99, 0x99))

# ── 保存 ──
doc.save(str(OUTPUT))
print(f"[完成] Word 文档已生成：{OUTPUT}")
