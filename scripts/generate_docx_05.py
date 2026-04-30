# -*- coding: utf-8 -*-
"""
生成 05_石家庄美食地图.docx
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    print("python-docx loaded")
except ImportError:
    print("ERROR: python-docx not installed")
    exit(1)

import os

OUTPUT_DIR = r"D:\新建文件夹\处理结果"

def create_doc_05():
    """生成 05_石家庄美食地图.docx"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('石家庄美食地图', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('整理来源：微信、百度地图、今日头条等平台的美食店铺信息')
    doc.add_paragraph('整理时间：2026年3月18日')
    doc.add_paragraph()
    
    # 一、早餐店
    doc.add_heading('一、早餐店', level=1)
    doc.add_heading('1. 美味苑', level=2)
    p = doc.add_paragraph()
    p.add_run('地址：').bold = True
    p.add_run('翟营大街与建明北路交口南行路西')
    p = doc.add_paragraph()
    p.add_run('电话：').bold = True
    p.add_run('0311-85085004')
    p = doc.add_paragraph()
    p.add_run('营业时间：').bold = True
    p.add_run('周一至周日 早06:00-09:30')
    p = doc.add_paragraph()
    p.add_run('特色：').bold = True
    p.add_run('早餐供应')
    
    # 二、河北菜/本地特色
    doc.add_heading('二、河北菜/本地特色', level=1)
    doc.add_heading('2. 白家罩饼', level=2)
    p = doc.add_paragraph()
    p.add_run('地址：').bold = True
    p.add_run('和平路平安大街交口东行二百米路南')
    p = doc.add_paragraph()
    p.add_run('电话：').bold = True
    p.add_run('0311-87833338')
    p = doc.add_paragraph()
    p.add_run('营业时间：').bold = True
    p.add_run('11:00-14:00, 17:30-22:00')
    p = doc.add_paragraph()
    p.add_run('特色：').bold = True
    p.add_run('清真，百年老店，省级非物质文化遗产')
    
    # 三、火锅
    doc.add_heading('三、火锅', level=1)
    doc.add_heading('3. 聚宝园大铜锅', level=2)
    p = doc.add_paragraph()
    p.add_run('地址：').bold = True
    p.add_run('槐安路与谈固大街交口西南角')
    p = doc.add_paragraph()
    p.add_run('类型：').bold = True
    p.add_run('铜锅火锅')
    
    doc.add_heading('10. 盐池滩羊(翠墨西三庄店)', level=2)
    p = doc.add_paragraph()
    p.add_run('地址：').bold = True
    p.add_run('河北省石家庄市新华区西三庄街道天翼路15号2楼(芳泽园对面)')
    p = doc.add_paragraph()
    p.add_run('电话：').bold = True
    p.add_run('18033832698')
    p = doc.add_paragraph()
    p.add_run('人均：').bold = True
    p.add_run('144元/人')
    p = doc.add_paragraph()
    p.add_run('营业时间：').bold = True
    p.add_run('10:00-14:00; 17:00-21:30')
    p = doc.add_paragraph()
    p.add_run('特色：').bold = True
    p.add_run('火锅，品味餐厅，有包厢')
    p = doc.add_paragraph()
    p.add_run('收藏数：').bold = True
    p.add_run('871人收藏')
    
    # 四、自助餐
    doc.add_heading('四、自助餐', level=1)
    doc.add_heading('4. 凤凰酒楼', level=2)
    p = doc.add_paragraph()
    p.add_run('地址：').bold = True
    p.add_run('石家庄西美花街')
    p = doc.add_paragraph()
    p.add_run('特色：').bold = True
    p.add_run('港式自助，69元/人')
    p = doc.add_paragraph()
    p.add_run('来源：').bold = True
    p.add_run('今日头条推荐')
    
    # 五、烧烤/烤羊
    doc.add_heading('五、烧烤/烤羊', level=1)
    doc.add_heading('5. 香羊部落烤羊腿羊排', level=2)
    p = doc.add_paragraph()
    p.add_run('地址：').bold = True
    p.add_run('石家庄市新华区柏林北区附近（华林国际）')
    p = doc.add_paragraph()
    p.add_run('特色：').bold = True
    p.add_run('烤羊腿、羊排')
    
    # 六、海鲜
    doc.add_heading('六、海鲜', level=1)
    doc.add_heading('7. 沧六海鲜', level=2)
    p = doc.add_paragraph()
    p.add_run('地址：').bold = True
    p.add_run('河北省石家庄市新华区友谊北大街222号')
    p = doc.add_paragraph()
    p.add_run('营业时间：').bold = True
    p.add_run('08:00-19:00')
    p = doc.add_paragraph()
    p.add_run('类型：').bold = True
    p.add_run('海鲜')
    p = doc.add_paragraph()
    p.add_run('收藏数：').bold = True
    p.add_run('446人收藏')
    
    # 七、中餐/小馆
    doc.add_heading('七、中餐/小馆', level=1)
    doc.add_heading('9. 南巷小馆·手艺菜', level=2)
    p = doc.add_paragraph()
    p.add_run('地址：').bold = True
    p.add_run('河北省石家庄市新华区翔翼路与西三庄街交叉口正西方向153米左右')
    p = doc.add_paragraph()
    p.add_run('人均：').bold = True
    p.add_run('58元/人')
    p = doc.add_paragraph()
    p.add_run('营业时间：').bold = True
    p.add_run('11:30-14:00; 17:30-22:30')
    p = doc.add_paragraph()
    p.add_run('特色：').bold = True
    p.add_run('其它中餐厅，有包厢，菜品健康，肉类好')
    p = doc.add_paragraph()
    p.add_run('收藏数：').bold = True
    p.add_run('1696人收藏')
    
    # 八、熟食/卤味
    doc.add_heading('八、熟食/卤味', level=1)
    doc.add_heading('6. 鑫源商店(平安大街店)', level=2)
    p = doc.add_paragraph()
    p.add_run('地址：').bold = True
    p.add_run('河北省石家庄市鹿泉区山尹村镇山尹村平安大街37号')
    p = doc.add_paragraph()
    p.add_run('特色：').bold = True
    p.add_run('石家庄特产，传统美食')
    
    doc.add_heading('8. （压猪头肉店）', level=2)
    p = doc.add_paragraph()
    p.add_run('地址：').bold = True
    p.add_run('河北省石家庄市鹿泉区寺家庄镇沿河南路城市傲居北侧约90米')
    p = doc.add_paragraph()
    p.add_run('营业时间：').bold = True
    p.add_run('06:30-23:15')
    p = doc.add_paragraph()
    p.add_run('特色：').bold = True
    p.add_run('传统压猪头肉')
    
    doc.add_heading('13. 大妞猪头肉', level=2)
    p = doc.add_paragraph()
    p.add_run('地址：').bold = True
    p.add_run('河北省石家庄市鹿泉区铜冶镇北铜冶小学南边一百米路东')
    p = doc.add_paragraph()
    p.add_run('人均：').bold = True
    p.add_run('38元/人')
    p = doc.add_paragraph()
    p.add_run('营业时间：').bold = True
    p.add_run('08:00-21:30')
    p = doc.add_paragraph()
    p.add_run('类型：').bold = True
    p.add_run('小吃快餐，餐馆早餐')
    p = doc.add_paragraph()
    p.add_run('收藏数：').bold = True
    p.add_run('175人收藏')
    
    # 九、农贸市场/生鲜
    doc.add_heading('九、农贸市场/生鲜', level=1)
    doc.add_heading('11. 修红西红柿专卖', level=2)
    p = doc.add_paragraph()
    p.add_run('地址：').bold = True
    p.add_run('河北省石家庄市长安区西兆通镇凌透西街与金星路交口北行50米路西第二间')
    p = doc.add_paragraph()
    p.add_run('电话：').bold = True
    p.add_run('13315999488')
    p = doc.add_paragraph()
    p.add_run('类型：').bold = True
    p.add_run('农贸市场')
    p = doc.add_paragraph()
    p.add_run('距离：').bold = True
    p.add_run('距你8公里，21分钟车程')
    p = doc.add_paragraph()
    p.add_run('收藏数：').bold = True
    p.add_run('17人收藏')
    
    doc.add_heading('12. 鑫德胜牛肉(晶都花苑店)', level=2)
    p = doc.add_paragraph()
    p.add_run('地址：').bold = True
    p.add_run('河北省石家庄市正定县正定镇恒州南街63号晶都花苑')
    p = doc.add_paragraph()
    p.add_run('类型：').bold = True
    p.add_run('农贸市场')
    p = doc.add_paragraph()
    p.add_run('收藏数：').bold = True
    p.add_run('16人收藏')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('本文档由AI整理生成，店铺信息仅供参考').italic = True
    p.runs[0].font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run('建议前往前电话确认营业状态').italic = True
    p.runs[0].font.size = Pt(10)
    
    output_path = os.path.join(OUTPUT_DIR, "05_石家庄美食地图.docx")
    doc.save(output_path)
    print("[OK] 05_石家庄美食地图.docx generated")
    return True

if __name__ == "__main__":
    print("=" * 50)
    print("Generating Word document...")
    print("=" * 50)
    
    try:
        create_doc_05()
        print("\n" + "=" * 50)
        print("[DONE] Word document generated!")
        print("=" * 50)
    except Exception as e:
        print(f"\n[ERROR] Error: {e}")
        import traceback
        traceback.print_exc()
