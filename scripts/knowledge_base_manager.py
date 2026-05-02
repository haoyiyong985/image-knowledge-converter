# -*- coding: utf-8 -*-
"""
智能图片知识库管理系统
功能：
1. 批量处理图片OCR
2. 智能分类（支持增量更新）
3. 自动合并同类内容
4. 生成Word和Markdown文档
5. 支持ima导入（多种方式）
"""

import os
import json
import re
import hashlib
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import shutil

# 尝试导入必要的库
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx未安装，将无法生成Word文档")
    print("请运行: pip install python-docx")


class KnowledgeBaseManager:
    """知识库管理器"""
    
    # 分类定义
    CATEGORIES = {
        "01_抗炎饮食与营养科普": {
            "keywords": ["抗炎", "ORAC", "ω-3", "ω-6", "坚果", "营养", "食物", "饮食", 
                        "氧化", "抗氧化", "炎症", "营养师", "比例", "水果", "蔬菜"],
            "description": "抗炎饮食、营养科普、食物营养价值相关内容"
        },
        "02_肠道健康与饮食分类": {
            "keywords": ["肠道", "益生菌", "益生元", "微生物", "绿灯", "黄灯", "红灯",
                        "发酵", "麸质", "纤维", "消化", "菌群", "有益菌"],
            "description": "肠道健康、饮食红绿灯分类、微生物组相关内容"
        },
        "03_中医养生与食疗": {
            "keywords": ["中医", "养生", "食疗", "穴位", "经络", "三伏", "仲冬",
                        "中药", "五味", "体质", "气血", "阴阳", "养生茶", "汤方"],
            "description": "中医养生、食疗方、经络穴位、中药知识"
        },
        "04_日常饮食建议": {
            "keywords": ["早餐", "食谱", "膳食", "营养搭配", "一日三餐", "饮食建议",
                        "健康餐", "食材", "烹饪", "做法"],
            "description": "日常饮食建议、食谱、膳食搭配"
        }
    }
    
    def __init__(self, base_path: str = "D:/新建文件夹"):
        self.base_path = Path(base_path)
        self.input_dir = self.base_path / "待处理图片" / "示范"
        self.output_dir = self.base_path / "处理结果"
        self.config_file = self.base_path / "knowledge_config.json"
        
        # 确保输出目录存在
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 加载配置
        self.config = self._load_config()
        
    def _load_config(self) -> Dict:
        """加载配置文件"""
        if self.config_file.exists():
            with open(self.config_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {
            "processed_images": {},  # 已处理图片记录 {hash: {filename, category, date}}
            "categories_content": {cat: [] for cat in self.CATEGORIES.keys()},  # 各类别内容
            "last_update": None
        }
    
    def _save_config(self):
        """保存配置文件"""
        self.config["last_update"] = datetime.now().isoformat()
        with open(self.config_file, 'w', encoding='utf-8') as f:
            json.dump(self.config, f, ensure_ascii=False, indent=2)
    
    def _get_image_hash(self, image_path: Path) -> str:
        """计算图片哈希（用于去重）"""
        with open(image_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    
    def _is_new_image(self, image_path: Path) -> bool:
        """检查是否为新增图片"""
        image_hash = self._get_image_hash(image_path)
        return image_hash not in self.config["processed_images"]
    
    def _classify_content(self, content: str) -> str:
        """
        智能分类内容
        返回最适合的分类名称
        """
        content_lower = content.lower()
        scores = {}
        
        for category, info in self.CATEGORIES.items():
            score = 0
            for keyword in info["keywords"]:
                # 计算关键词匹配次数
                matches = len(re.findall(keyword, content))
                score += matches
            scores[category] = score
        
        # 返回得分最高的分类
        if max(scores.values()) > 0:
            return max(scores, key=scores.get)
        
        # 默认分类
        return "04_日常饮食建议"
    
    def _extract_text_from_image(self, image_path: Path) -> str:
        """
        从图片提取文字
        注意：实际使用时需要接入OCR API或本地OCR
        这里预留接口
        """
        # 这里应该调用OCR服务
        # 为了演示，返回一个标记
        return f"[OCR_TEXT_FROM_{image_path.name}]"
    
    def process_new_images(self, image_dir: Optional[str] = None) -> Dict:
        """
        处理新增图片
        返回处理结果统计
        """
        if image_dir is None:
            image_dir = self.input_dir
        else:
            image_dir = Path(image_dir)
        
        if not image_dir.exists():
            print(f"目录不存在: {image_dir}")
            return {"processed": 0, "skipped": 0, "errors": 0}
        
        # 获取所有图片文件
        image_files = list(image_dir.glob("*.jpg")) + list(image_dir.glob("*.png")) + list(image_dir.glob("*.jpeg"))
        
        stats = {"processed": 0, "skipped": 0, "errors": 0, "new_images": []}
        
        for img_path in image_files:
            try:
                # 检查是否已处理
                if not self._is_new_image(img_path):
                    print(f"跳过已处理图片: {img_path.name}")
                    stats["skipped"] += 1
                    continue
                
                print(f"处理新图片: {img_path.name}")
                
                # 计算图片哈希
                img_hash = self._get_image_hash(img_path)
                
                # 提取文字（这里需要接入实际OCR）
                # 实际使用时，这里应该调用OCR API
                content = self._extract_text_from_image(img_path)
                
                # 智能分类
                category = self._classify_content(content)
                
                # 记录到配置
                self.config["processed_images"][img_hash] = {
                    "filename": img_path.name,
                    "category": category,
                    "date": datetime.now().isoformat(),
                    "hash": img_hash
                }
                
                # 添加到对应分类
                self.config["categories_content"][category].append({
                    "filename": img_path.name,
                    "content": content,
                    "date": datetime.now().isoformat()
                })
                
                stats["processed"] += 1
                stats["new_images"].append({
                    "filename": img_path.name,
                    "category": category
                })
                
            except Exception as e:
                print(f"处理图片出错 {img_path.name}: {e}")
                stats["errors"] += 1
        
        # 保存配置
        self._save_config()
        
        return stats
    
    def generate_documents(self):
        """
        生成所有分类文档（Markdown和Word）
        """
        print("\n=== 生成文档 ===")
        
        for category, info in self.CATEGORIES.items():
            contents = self.config["categories_content"].get(category, [])
            
            if not contents:
                print(f"分类 {category} 无内容，跳过")
                continue
            
            # 生成Markdown
            self._generate_markdown(category, info, contents)
            
            # 生成Word
            if DOCX_AVAILABLE:
                self._generate_word(category, info, contents)
            
            print(f"已生成: {category}")
    
    def _generate_markdown(self, category: str, info: Dict, contents: List[Dict]):
        """生成Markdown文档"""
        md_file = self.output_dir / f"{category}.md"
        
        md_content = f"""# {category.replace('01_', '').replace('02_', '').replace('03_', '').replace('04_', '')}

> {info['description']}
> 整理时间：{datetime.now().strftime('%Y年%m月%d日')}
> 内容来源：共{len(contents)}张图片整理

---

"""
        
        for i, item in enumerate(contents, 1):
            md_content += f"""## 内容 {i}（来源：{item['filename']}）

{item['content']}

---

"""
        
        md_content += """
*本文档由AI智能知识库系统自动生成*
*仅供参考学习使用*
"""
        
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(md_content)
    
    def _generate_word(self, category: str, info: Dict, contents: List[Dict]):
        """生成Word文档"""
        docx_file = self.output_dir / f"{category}.docx"
        
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)
        
        # 标题
        title = doc.add_heading(category.replace('01_', '').replace('02_', '').replace('03_', '').replace('04_', ''), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 说明
        doc.add_paragraph(f"{info['description']}")
        doc.add_paragraph(f"整理时间：{datetime.now().strftime('%Y年%m月%d日')}")
        doc.add_paragraph(f"内容来源：共{len(contents)}张图片整理")
        doc.add_paragraph('_' * 50)
        
        # 内容
        for i, item in enumerate(contents, 1):
            doc.add_heading(f'内容 {i}（来源：{item["filename"]}）', level=2)
            doc.add_paragraph(item['content'])
            doc.add_paragraph('_' * 50)
        
        # 页脚
        doc.add_paragraph()
        doc.add_paragraph('本文档由AI智能知识库系统自动生成')
        doc.add_paragraph('仅供参考学习使用')
        
        doc.save(str(docx_file))
    
    def get_statistics(self) -> Dict:
        """获取知识库统计信息"""
        stats = {
            "total_images": len(self.config["processed_images"]),
            "categories": {},
            "last_update": self.config.get("last_update")
        }
        
        for category, contents in self.config["categories_content"].items():
            stats["categories"][category] = len(contents)
        
        return stats
    
    def export_for_ima(self, export_dir: Optional[str] = None) -> str:
        """
        导出为ima友好的格式
        返回导出目录路径
        """
        if export_dir is None:
            export_dir = self.base_path / "ima_export"
        else:
            export_dir = Path(export_dir)
        
        export_dir.mkdir(parents=True, exist_ok=True)
        
        # 复制所有Markdown文件
        for md_file in self.output_dir.glob("*.md"):
            shutil.copy2(md_file, export_dir / md_file.name)
        
        # 生成导入说明
        readme = export_dir / "README.txt"
        with open(readme, 'w', encoding='utf-8') as f:
            f.write("""ima知识库导入说明
================

1. 打开ima.copilot应用
2. 进入"知识库"功能
3. 点击"导入"或"添加"按钮
4. 选择本文件夹中的 .md 文件
5. 等待ima解析完成

文件说明：
""")
            for category in self.CATEGORIES.keys():
                f.write(f"- {category}.md : {self.CATEGORIES[category]['description']}\n")
        
        print(f"\n已导出到: {export_dir}")
        return str(export_dir)


# ==================== ima导入模块 ====================

class ImaImporter:
    """
    ima知识库导入器
    提供多种导入方式
    """
    
    def __init__(self, ima_path: str = "C:/Users/LENOVO/AppData/Local/ima.copilot"):
        self.ima_path = Path(ima_path)
        self.user_data_path = self.ima_path / "User Data" / "Default"
    
    def get_import_methods(self) -> List[Dict]:
        """获取可用的导入方法"""
        methods = []
        
        # 方法1: 文件拖拽（最可靠）
        methods.append({
            "name": "文件拖拽导入",
            "description": "直接将文件拖入ima窗口",
            "reliability": "⭐⭐⭐⭐⭐",
            "automation": "需要模拟鼠标拖拽",
            "steps": [
                "启动ima.copilot",
                "打开知识库页面",
                "模拟拖拽文件到窗口"
            ]
        })
        
        # 方法2: 剪贴板导入
        methods.append({
            "name": "剪贴板导入",
            "description": "复制内容后粘贴到ima",
            "reliability": "⭐⭐⭐⭐",
            "automation": "可以自动化",
            "steps": [
                "读取文件内容到剪贴板",
                "模拟Ctrl+V粘贴",
                "触发ima的粘贴识别"
            ]
        })
        
        # 方法3: Chrome Extension通信
        methods.append({
            "name": "Chrome Extension API",
            "description": "通过Chrome扩展通信",
            "reliability": "⭐⭐⭐",
            "automation": "需要开发扩展",
            "steps": [
                "开发Chrome扩展",
                "与ima扩展通信",
                "调用内部API"
            ]
        })
        
        # 方法4: 数据库直接写入（不推荐）
        methods.append({
            "name": "数据库直接写入",
            "description": "直接操作ima的IndexedDB",
            "reliability": "⭐",
            "automation": "技术难度高",
            "steps": [
                "解析LevelDB格式",
                "理解ima数据结构",
                "直接写入数据库"
            ],
            "warning": "可能导致数据损坏，不推荐"
        })
        
        return methods
    
    def create_import_script(self, markdown_files: List[str]) -> str:
        """
        创建自动导入脚本（使用pyautogui模拟操作）
        """
        script_content = '''# -*- coding: utf-8 -*-
"""
ima知识库自动导入脚本
使用pyautogui模拟鼠标和键盘操作
"""

import time
import pyautogui
import subprocess
from pathlib import Path

# 配置
IMA_PATH = "C:/Users/LENOVO/AppData/Local/ima.copilot/Application/ima.exe"
MARKDOWN_FILES = [
'''
        
        for f in markdown_files:
            script_content += f'    r"{f}",\n'
        
        script_content += ''']

def launch_ima():
    """启动ima"""
    subprocess.Popen([IMA_PATH])
    time.sleep(5)  # 等待启动

def import_file(file_path):
    """导入单个文件"""
    print(f"导入: {file_path}")
    
    # 这里需要根据实际情况调整坐标
    # 1. 点击知识库按钮
    # 2. 点击导入按钮
    # 3. 选择文件
    # 4. 确认导入
    
    # 示例：模拟Ctrl+O打开文件对话框
    pyautogui.keyDown('ctrl')
    pyautogui.keyDown('o')
    pyautogui.keyUp('o')
    pyautogui.keyUp('ctrl')
    
    time.sleep(1)
    
    # 输入文件路径
    pyautogui.typewrite(str(file_path))
    time.sleep(0.5)
    
    # 按回车确认
    pyautogui.keyDown('return')
    pyautogui.keyUp('return')
    
    time.sleep(2)  # 等待导入完成

def main():
    """主函数"""
    print("=== ima知识库自动导入 ===")
    
    # 启动ima
    launch_ima()
    
    # 导入所有文件
    for file_path in MARKDOWN_FILES:
        if Path(file_path).exists():
            import_file(file_path)
        else:
            print(f"文件不存在: {file_path}")
    
    print("导入完成！")

if __name__ == "__main__":
    main()
'''
        
        return script_content


# ==================== 主程序 ====================

def main():
    """主程序入口"""
    print("=" * 60)
    print("智能图片知识库管理系统")
    print("=" * 60)
    
    # 初始化管理器
    manager = KnowledgeBaseManager()
    
    while True:
        print("\n请选择操作:")
        print("1. 处理新增图片")
        print("2. 生成/更新文档")
        print("3. 查看统计信息")
        print("4. 导出为ima格式")
        print("5. 查看ima导入方案")
        print("0. 退出")
        
        choice = input("\n输入选项: ").strip()
        
        if choice == "1":
            print("\n--- 处理新增图片 ---")
            stats = manager.process_new_images()
            print(f"\n处理完成:")
            print(f"  - 新处理: {stats['processed']} 张")
            print(f"  - 已存在跳过: {stats['skipped']} 张")
            print(f"  - 错误: {stats['errors']} 张")
            
            if stats['new_images']:
                print("\n新增图片分类:")
                for img in stats['new_images']:
                    print(f"  - {img['filename']} -> {img['category']}")
        
        elif choice == "2":
            print("\n--- 生成文档 ---")
            manager.generate_documents()
            print("文档生成完成！")
        
        elif choice == "3":
            print("\n--- 统计信息 ---")
            stats = manager.get_statistics()
            print(f"总图片数: {stats['total_images']}")
            print(f"最后更新: {stats['last_update']}")
            print("\n各类别数量:")
            for cat, count in stats['categories'].items():
                print(f"  - {cat}: {count} 张")
        
        elif choice == "4":
            print("\n--- 导出ima格式 ---")
            export_path = manager.export_for_ima()
            print(f"导出完成: {export_path}")
        
        elif choice == "5":
            print("\n--- ima导入方案 ---")
            importer = ImaImporter()
            methods = importer.get_import_methods()
            
            for i, method in enumerate(methods, 1):
                print(f"\n方法 {i}: {method['name']}")
                print(f"  可靠性: {method['reliability']}")
                print(f"  说明: {method['description']}")
                print(f"  自动化: {method['automation']}")
                if 'warning' in method:
                    print(f"  ⚠️ 警告: {method['warning']}")
        
        elif choice == "0":
            print("\n感谢使用，再见！")
            break
        
        else:
            print("无效选项，请重新输入")


if __name__ == "__main__":
    main()
