#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""修复SmartDocumentGenerator类"""

import re

# 读取原文件
with open('auto_process_all_v2.py', 'r', encoding='utf-8') as f:
    content = f.read()

# 新类内容
new_class = '''class SmartDocumentGenerator:
    """智能文档生成器 - 支持主题分组和文档合并"""

    def __init__(self, output_dir='处理结果'):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.merger = DocumentMerger(output_dir)
        self.existing_hashes = set()
        self._scan_hashes()
        # 追踪已生成的文档（避免重复生成Word）
        self.generated_docs = set()  # 格式: theme/safe_name

    def _scan_hashes(self):
        """扫描已有文档的hash"""
        for md_file in self.output_dir.rglob('*.md'):
            try:
                file_content = md_file.read_text(encoding='utf-8')
                hash_match = re.search(r'content_hash:\s*([a-f0-9]{8})', file_content)
                if hash_match:
                    self.existing_hashes.add(hash_match.group(1))
            except:
                pass

    def _get_doc_key(self, theme, doc_name):
        """获取文档唯一标识"""
        safe_name = re.sub(r'[<>:"/\\|?*]', '', doc_name)
        if len(safe_name) > 25:
            safe_name = safe_name[:25]
        return f"{theme}/{safe_name}"

    def is_duplicate(self, content_hash):
        """检测是否重复"""
        return content_hash in self.existing_hashes

    def is_doc_generated(self, theme, doc_name):
        """检查文档是否已生成"""
        return self._get_doc_key(theme, doc_name) in self.generated_docs

    def mark_doc_generated(self, theme, doc_name):
        """标记文档已生成"""
        self.generated_docs.add(self._get_doc_key(theme, doc_name))

    def generate_markdown(self, text, theme, doc_name, category, image_name, keywords=None, content_hash=None):
        """生成Markdown文档"""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')

        # 查找是否有相似文档可以合并
        similar_doc = self.merger.find_similar_doc(theme, doc_name, text)
        if similar_doc:
            print(f"  → 发现相似文档 {similar_doc.name}，合并内容...")
            if self.merger.merge_content(similar_doc, text, image_name, content_hash):
                # 标记文档已处理（避免重复生成Word）
                self.mark_doc_generated(theme, doc_name)
                return str(similar_doc)

        # 创建新文档
        keywords_str = ', '.join(keywords) if keywords else '无'
        md_content = f"""# {doc_name}

> 来源图片: {image_name}
> 识别时间: {timestamp}
> 主题分类: {theme}
> 原分类: {category}
> 关键词: {keywords_str}
> content_hash: {content_hash or 'N/A'}

---

## 内容

{text}

---

*本文档由图片知识库整理工具自动生成*
"""

        # 创建主题文件夹
        theme_folder = self.output_dir / theme
        theme_folder.mkdir(exist_ok=True)

        # 生成安全的文件名
        safe_name = re.sub(r'[<>:"/\\|?*]', '', doc_name)
        if len(safe_name) > 25:
            safe_name = safe_name[:25]

        # 检查是否已存在
        md_file = theme_folder / f"{safe_name}.md"
        counter = 1
        while md_file.exists():
            md_file = theme_folder / f"{safe_name}_{counter}.md"
            counter += 1

        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(md_content)

        # 更新hash记录
        if content_hash:
            self.existing_hashes.add(content_hash)
        # 更新合并器记录
        self.merger.theme_docs[theme].append({
            'file': md_file,
            'topic': doc_name
        })
        # 标记文档已生成
        self.mark_doc_generated(theme, doc_name)

        logger.info(f"[文档] Markdown已生成: {theme}/{md_file.name}")
        return str(md_file)

    def generate_word(self, text, theme, doc_name, image_name, keywords=None, content_hash=None, existing_md_file=None):
        """生成Word文档，如果已有同名文档则跳过"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.oxml.ns import qn
        except ImportError:
            logger.warning("[文档] python-docx未安装，跳过Word生成")
            return None

        # 检查是否已生成过Word文档
        if self.is_doc_generated(theme, doc_name):
            print(f"  → Word文档已存在，跳过")
            # 返回已有的Word文件路径
            safe_name = re.sub(r'[<>:"/\\|?*]', '', doc_name)
            if len(safe_name) > 25:
                safe_name = safe_name[:25]
            existing_docx = self.output_dir / theme / f"{safe_name}.docx"
            if existing_docx.exists():
                return str(existing_docx)
            # 检查带编号的版本
            for i in range(1, 100):
                existing_docx = self.output_dir / theme / f"{safe_name}_{i}.docx"
                if existing_docx.exists():
                    return str(existing_docx)
            return None

        # 保存到同一主题文件夹
        theme_folder = self.output_dir / theme
        theme_folder.mkdir(exist_ok=True)

        safe_name = re.sub(r'[<>:"/\\|?*]', '', doc_name)
        if len(safe_name) > 25:
            safe_name = safe_name[:25]

        docx_file = theme_folder / f"{safe_name}.docx"
        counter = 1
        while docx_file.exists():
            docx_file = theme_folder / f"{safe_name}_{counter}.docx"
            counter += 1

        # 如果有对应的Markdown文件已合并内容，则同步Word
        if existing_md_file and Path(existing_md_file).exists():
            try:
                md_content = Path(existing_md_file).read_text(encoding='utf-8')
                main_content = md_content.split('## 内容\\n\\n')[1] if '## 内容\\n\\n' in md_content else ''
                main_content = main_content.split('## 📌 补充内容')[0] if '## 📌 补充内容' in main_content else main_content
                main_content = main_content.replace('\\n\\n---\\n\\n*本文档由图片知识库整理工具自动生成*', '')
                text = main_content.strip()
            except Exception as e:
                logger.warning(f"[Word] 读取合并内容失败: {e}")

        # 创建Word文档
        doc = Document()
        doc.styles['Normal'].font.name = 'Microsoft YaHei'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        doc.add_heading(doc_name, 0)
        doc.add_paragraph(f"来源图片: {image_name}")
        doc.add_paragraph(f"识别时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"主题分类: {theme}")
        keywords_str = ', '.join(keywords) if keywords else '无'
        doc.add_paragraph(f"关键词: {keywords_str}")
        doc.add_paragraph("─" * 30)
        doc.add_paragraph(text)

        doc.save(str(docx_file))
        # 标记文档已生成
        self.mark_doc_generated(theme, doc_name)
        logger.info(f"[文档] Word已生成: {theme}/{docx_file.name}")
        return str(docx_file)

'''

# 替换类内容 - 使用不同的方法避免转义问题
pattern = r'class SmartDocumentGenerator:.*?(?=class IMASyncer:)'
match = re.search(pattern, content, flags=re.DOTALL)
if match:
    # 找到新类的起始和结束位置
    start_pos = match.start()
    end_pos = match.end()
    # 拼接新内容
    new_content = content[:start_pos] + new_class + content[end_pos:]
    print(f'已替换: {start_pos} - {end_pos}')
else:
    new_content = content
    print('未找到匹配')
    
# 写回文件
with open('auto_process_all_v2.py', 'w', encoding='utf-8') as f:
    f.write(new_content)

print('文件已更新')
