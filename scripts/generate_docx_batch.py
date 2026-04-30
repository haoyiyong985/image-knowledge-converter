#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
批量生成 Word 文档（01、02、03）
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx.oxml.ns
import re
import os

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = 'SimSun' if level > 1 else 'SimHei'
    run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'SimSun' if level > 1 else 'SimHei')
    run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    run.font.bold = True
    if level == 1:
        run.font.color.rgb = RGBColor(0, 0, 128)
    return heading

def add_paragraph_custom(doc, text, bold=False, italic=False, indent=False):
    """添加自定义段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    return p

def parse_markdown_to_docx(md_file, docx_file):
    """解析Markdown并生成Word文档"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    in_code_block = False
    code_content = []
    
    for line in lines:
        # 代码块处理
        if line.startswith('```'):
            if in_code_block:
                if code_content:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    run = p.add_run('\n'.join(code_content))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(50, 50, 50)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line)
            continue
        
        # 空行
        if not line.strip():
            continue
        
        # 标题
        if line.startswith('# '):
            add_heading_custom(doc, line[2:], level=1)
        elif line.startswith('## '):
            add_heading_custom(doc, line[3:], level=2)
        elif line.startswith('### '):
            add_heading_custom(doc, line[4:], level=3)
        elif line.startswith('#### '):
            add_heading_custom(doc, line[5:], level=4)
        
        # 引用块
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(line[2:])
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
        
        # 列表项
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line[2:])
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(11)
        
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\. ', '', line)
            run = p.add_run(text)
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(11)
        
        # 表格行（简单处理）
        elif line.startswith('|') and '|' in line[1:]:
            if not line.strip().replace('|', '').replace('-', '').replace(' ', ''):
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                p = doc.add_paragraph()
                run = p.add_run(' | '.join(cells))
                run.font.name = 'SimSun'
                run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'SimSun')
                run.font.size = Pt(10)
        
        # 普通段落
        else:
            parts = re.split(r'\*\*(.*?)\*\*', line)
            if len(parts) > 1:
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.name = 'SimSun'
                        run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'SimSun')
                        run.font.size = Pt(11)
                        run.font.bold = (i % 2 == 1)
            else:
                add_paragraph_custom(doc, line)
    
    doc.save(docx_file)
    print(f"[OK] {os.path.basename(docx_file)} 生成完成")

if __name__ == '__main__':
    base_dir = r'D:\新建文件夹\处理结果'
    
    files = [
        ('01_抗炎饮食与营养科普.md', '01_抗炎饮食与营养科普.docx'),
        ('02_肠道健康与饮食分类.md', '02_肠道健康与饮食分类.docx'),
        ('03_中医养生与食疗.md', '03_中医养生与食疗.docx'),
    ]
    
    for md_name, docx_name in files:
        md_file = os.path.join(base_dir, md_name)
        docx_file = os.path.join(base_dir, docx_name)
        
        if os.path.exists(md_file):
            parse_markdown_to_docx(md_file, docx_file)
        else:
            print(f"[ERROR] 找不到文件: {md_file}")
    
    print("\n" + "=" * 50)
    print("[DONE] 所有 Word 文档生成完成！")
    print("=" * 50)
