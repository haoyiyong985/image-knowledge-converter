# -*- coding: utf-8 -*-
"""
简单Word文档生成脚本
使用python-docx库
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    print("python-docx 已加载")
except ImportError:
    print("错误：未安装 python-docx")
    print("请运行: pip install python-docx")
    exit(1)

import os

OUTPUT_DIR = r"D:\新建文件夹\处理结果"

def create_doc_03():
    """生成 03_中医养生与食疗.docx"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('中医养生与食疗知识库', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('整理来源：微信读书、小红书、豆包等平台的中医养生内容')
    doc.add_paragraph('整理时间：2026年3月15日')
    doc.add_paragraph()
    
    # 一、三伏天养生指南
    doc.add_heading('一、三伏天养生指南', level=1)
    doc.add_paragraph('初伏 - 清内热：食用苦瓜、丝瓜、冬瓜、黄瓜降火')
    p = doc.add_paragraph()
    p.add_run('三伏茶配方：').bold = True
    p.add_run('金银花、夏枯草各10g，甘草、菊花、茯苓、麦冬、桑叶各5g，煮20分钟')
    doc.add_paragraph('中伏 - 健脾祛湿：三豆汤（黑豆、赤小豆、绿豆1:1:1）')
    doc.add_paragraph('末伏 - 排寒补阳：黄芪姜枣茶（生姜10g、红枣2个、黄芪5g、枸杞10g）')
    doc.add_paragraph('出伏 - 滋阴养血：枸杞、桑葚、麦冬煮水（1:1:1）')
    
    # 二、仲冬养生食疗方
    doc.add_heading('二、仲冬养生食疗方', level=1)
    p = doc.add_paragraph()
    p.add_run('仲冬，五脏要同补，重点是大补心和肾').bold = True
    p = doc.add_paragraph()
    p.add_run('补肾养心食疗方：').bold = True
    p.add_run('栗子6个、核桃6个、莲子6个、枸杞1把、葡萄干1把、陈皮1/4~1/2')
    
    # 三、经络穴位知识
    doc.add_heading('三、经络穴位知识', level=1)
    p = doc.add_paragraph()
    p.add_run('厥阴俞穴位置：').bold = True
    p.add_run('第4胸椎棘突下旁开1.5寸处')
    
    # 四、中药五味理论
    doc.add_heading('四、中药五味理论', level=1)
    p = doc.add_paragraph()
    p.add_run('来源：朱氏非遗话中医第十五课').bold = True
    doc.add_paragraph('酸（收敛）：山茱萸、五味子、五倍子、乌梅')
    doc.add_paragraph('苦（泄、燥、坚）：大黄、黄连')
    doc.add_paragraph('甘（补、缓、和）：人参、饴糖、甘草')
    doc.add_paragraph('辛（散、行）：麻黄、薄荷、木香、红花')
    doc.add_paragraph('咸（软、下）：海藻、昆布、鳖甲、芒硝')
    doc.add_paragraph('淡（渗、利）：猪苓、茯苓、薏苡仁、通草')
    
    # 五、传统中药方剂
    doc.add_heading('五、传统中药方剂', level=1)
    p = doc.add_paragraph()
    p.add_run('龙化丹：').bold = True
    p.add_run('治疗小儿耳后湿疹，配方含炉甘石')
    p = doc.add_paragraph()
    p.add_run('虎潜丸（健步强身丸）：').bold = True
    p.add_run('专调两腿酸软、腰腿痛、关节痛')
    
    # 六、柿子养生功效
    doc.add_heading('六、柿子养生功效', level=1)
    p = doc.add_paragraph()
    p.add_run('来源：一柿顶十药').bold = True
    p = doc.add_paragraph()
    p.add_run('青柿：').bold = True
    p.add_run('去火功效，应对秋燥上火')
    p = doc.add_paragraph()
    p.add_run('柿饼：').bold = True
    p.add_run('性平甘涩，归心、肺、大肠经')
    p = doc.add_paragraph()
    p.add_run('柿霜：').bold = True
    p.add_run('清热生津、润泽止咳，治口腔溃疡')
    p = doc.add_paragraph()
    p.add_run('柿核：').bold = True
    p.add_run('降逆下气，治呃逆（打嗝）')
    
    # 七、紫苏生姜红枣水
    doc.add_heading('七、紫苏生姜红枣水', level=1)
    p = doc.add_paragraph()
    p.add_run('来源：止水养生日记').bold = True
    p = doc.add_paragraph()
    p.add_run('功效：').bold = True
    p.add_run('驱寒暖胃、补益气血')
    p = doc.add_paragraph()
    p.add_run('适合人群：').bold = True
    p.add_run('容易感冒、手脚胃部不适、受凉后腹泻者')
    p = doc.add_paragraph()
    p.add_run('做法：').bold = True
    p.add_run('紫苏叶8-10片、生姜3-4片、红枣4-5颗，加水500ml煮沸后小火煮5-10分钟')
    
    # 八、醪糟养生
    doc.add_heading('八、醪糟养生', level=1)
    p = doc.add_paragraph()
    p.add_run('来源：酒窝姐爱养生').bold = True
    p = doc.add_paragraph()
    p.add_run('选购要点：').bold = True
    p.add_run('选配料表干净的醪糟（净化水、糯米、甜酒曲）')
    p = doc.add_paragraph()
    p.add_run('功效：').bold = True
    p.add_run('醪糟煮蛋，食养气血，巨养颜')
    
    # 九、紫苏陈皮水（新增）
    doc.add_heading('九、紫苏陈皮水', level=1)
    p = doc.add_paragraph()
    p.add_run('来源：止水养生日记').bold = True
    p = doc.add_paragraph()
    p.add_run('功效：').bold = True
    p.add_run('理气健脾、燥湿化痰')
    p = doc.add_paragraph()
    p.add_run('适合人群：').bold = True
    p.add_run('立秋后湿气仍重，感觉食欲不振、脘腹胀满、咳嗽痰多者')
    p = doc.add_paragraph()
    p.add_run('搭配原理：').bold = True
    p.add_run('陈皮能理气健脾，燥湿化痰；与紫苏搭配，增强行气宽中、化解脾胃湿浊的功效')
    p = doc.add_paragraph()
    p.add_run('做法：').bold = True
    p.add_run('紫苏叶8片、陈皮5克，开水冲泡焖5分钟，再放入紫苏叶焖3分钟')
    
    # 十、大雪后15天养生指南（新增）
    doc.add_heading('十、大雪后15天养生指南', level=1)
    p = doc.add_paragraph()
    p.add_run('来源：懒妈的养生账本').bold = True
    p = doc.add_paragraph()
    p.add_run('养生要点：').bold = True
    p.add_run('大雪后15天，就一个字：藏')
    p = doc.add_paragraph()
    p.add_run('食养原则：').bold = True
    p.add_run('藏精于黑咸，温润滋肾')
    p = doc.add_paragraph()
    p.add_run('推荐食材：').bold = True
    p.add_run('海带、紫菜、牡蛎、黑豆、黑芝麻、木耳')
    p = doc.add_paragraph()
    p.add_run('藏暖粥：').bold = True
    p.add_run('黑米、核桃、山药、枸杞同煮，早晚温食')
    p = doc.add_paragraph()
    p.add_run('地域饮食：').bold = True
    p.add_run('南方莲藕炖排骨，北方小米羊肉粥')
    
    # 十一、冬至养生羊肉汤（新增）
    doc.add_heading('十一、冬至养生羊肉汤', level=1)
    p = doc.add_paragraph()
    p.add_run('来源：微信读书《顺时生活》').bold = True
    p = doc.add_paragraph()
    p.add_run('原料：').bold = True
    p.add_run('羊肉1-2斤、黄芪100克、当归20克、甘蔗2-4节、带皮生姜2块、大枣8个')
    p = doc.add_paragraph()
    p.add_run('调料：').bold = True
    p.add_run('黄酒1两')
    p = doc.add_paragraph()
    p.add_run('善养生者，必养冬至').italic = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('本文档由AI整理生成，内容源自各平台中医养生知识').italic = True
    p.runs[0].font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run('仅供参考学习使用，如有不适请前往正规医院就诊').italic = True
    p.runs[0].font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run('更新记录：2026年3月18日 - 新增紫苏陈皮水、大雪养生、冬至羊肉汤').italic = True
    p.runs[0].font.size = Pt(10)
    
    output_path = os.path.join(OUTPUT_DIR, "03_中医养生与食疗.docx")
    doc.save(output_path)
    print("[OK] 03_中医养生与食疗.docx 生成完成")
    return True

def create_doc_04():
    """生成 04_日常饮食建议.docx"""
    doc = Document()
    
    title = doc.add_heading('日常饮食建议与食谱', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('整理来源：微信视频号、小红书等平台的健康饮食内容')
    doc.add_paragraph('整理时间：2026年3月15日')
    doc.add_paragraph()
    
    # 一、健康早餐建议
    doc.add_heading('一、健康早餐建议（以50岁人群为例）', level=1)
    p = doc.add_paragraph()
    p.add_run('推荐搭配：').bold = True
    doc.add_paragraph('• 全麦吐司1-2片 + 天然花生酱')
    doc.add_paragraph('• 煮鸡蛋1个')
    doc.add_paragraph('• 新鲜浆果（草莓/蓝莓）')
    doc.add_paragraph('• 酸奶一小杯')
    
    # 二、全日食谱推荐
    doc.add_heading('二、全日食谱推荐', level=1)
    p = doc.add_paragraph()
    p.add_run('标准成人一日饮食量：').bold = True
    doc.add_paragraph('• 粮食200g（粗粮50g）')
    doc.add_paragraph('• 鸡蛋1个、牛奶1袋')
    doc.add_paragraph('• 肉类100g、虾仁25g')
    doc.add_paragraph('• 蔬菜500g、水果200g')
    doc.add_paragraph('• 北豆腐50g、油20g、盐5g')
    p = doc.add_paragraph()
    p.add_run('营养搭配要点：').bold = True
    doc.add_paragraph('粗细搭配、荤素搭配、五色蔬菜、适量水果、少油少盐')
    
    # 三、逆转脂肪肝饮食指南（新增）
    doc.add_heading('三、逆转脂肪肝饮食指南', level=1)
    p = doc.add_paragraph()
    p.add_run('来源：医路向前巍子').bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('做好这几步，有可能会逆转脂肪肝：').bold = True
    doc.add_paragraph('1. 肥胖超重的人减肥')
    doc.add_paragraph('2. 爱喝酒的人戒酒')
    doc.add_paragraph('3. 荤素搭配，不一味吃素，补充优质蛋白')
    doc.add_paragraph('4. 适当吃杂粮')
    doc.add_paragraph('5. 适当摄入脂肪')
    doc.add_paragraph('6. 戒掉零食')
    doc.add_paragraph('7. 补充维生素')
    doc.add_paragraph('8. 适当的运动')
    p = doc.add_paragraph()
    p.add_run('做好这几点，对脂肪肝逆转有帮助').italic = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('本文档由AI整理生成，仅供参考学习使用').italic = True
    p.runs[0].font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run('更新记录：2026年3月18日 - 新增逆转脂肪肝饮食指南').italic = True
    p.runs[0].font.size = Pt(10)
    
    output_path = os.path.join(OUTPUT_DIR, "04_日常饮食建议.docx")
    doc.save(output_path)
    print("[OK] 04_日常饮食建议.docx 生成完成")
    return True

def create_doc_01():
    """生成 01_抗炎饮食与营养科普.docx"""
    doc = Document()
    
    title = doc.add_heading('抗炎饮食与营养科普知识库', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('整理来源：小红书、微信读书等平台的健康知识截图')
    doc.add_paragraph('整理时间：2026年3月15日')
    doc.add_paragraph()
    
    # 一、坚果的ω-6与ω-3比例指南
    doc.add_heading('一、坚果的ω-6与ω-3比例指南', level=1)
    p = doc.add_paragraph()
    p.add_run('来源：注册营养师盼盼').bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('坚持吃（比例优秀）：').bold = True
    doc.add_paragraph('亚麻籽 0.2:1 | 奇亚籽 1:3 | 核桃 4:1 | 夏威夷果 6.6:1 | 碧根果 21:1')
    p = doc.add_paragraph()
    p.add_run('经常吃（比例良好）：').bold = True
    doc.add_paragraph('榛子 90:1 | 腰果 125:1 | 开心果 62:1 | 巴旦木 410:1 | 南瓜子 178:1 | 花生 349:1')
    
    # 二、抗炎食物ORAC值排行榜
    doc.add_heading('二、抗炎食物ORAC值排行榜', level=1)
    p = doc.add_paragraph()
    p.add_run('来源：营养师黑皮').bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('重点推荐：肉桂和姜黄是日常最容易获取的高ORAC值香料！').italic = True
    
    # 三、21种抗炎最佳食物清单
    doc.add_heading('三、21种抗炎最佳食物清单', level=1)
    p = doc.add_paragraph()
    p.add_run('水果类：').bold = True
    p.add_run('余甘子、草莓、石榴、蓝莓、李子')
    p = doc.add_paragraph()
    p.add_run('谷物与种子：').bold = True
    p.add_run('荞麦、小米、燕麦、藜麦、核桃')
    p = doc.add_paragraph()
    p.add_run('其他抗炎食物：').bold = True
    p.add_run('咖啡、茶、开心果、黑巧、丁香、薄荷、肉桂、姜黄、紫甘蓝、甜菜、辣椒')
    p = doc.add_paragraph()
    p.add_run('动物源性：').bold = True
    p.add_run('酸奶、三文鱼、银鳕鱼、鸡蛋')
    
    # 四、抗炎膳食宝塔
    doc.add_heading('四、抗炎膳食宝塔', level=1)
    p = doc.add_paragraph()
    p.add_run('大量摄入（塔底）：').bold = True
    p.add_run('全谷物、豆类、蔬菜、水果')
    p = doc.add_paragraph()
    p.add_run('较多摄入：').bold = True
    p.add_run('大豆食品、鱼类贝类、健康脂肪')
    p = doc.add_paragraph()
    p.add_run('适量摄入：').bold = True
    p.add_run('补充剂、茶、香料、禽肉瘦肉')
    p = doc.add_paragraph()
    p.add_run('少量摄入（塔尖）：').bold = True
    p.add_run('红酒、健康甜点')
    
    # 五、饮食红黑榜
    doc.add_heading('五、饮食红黑榜', level=1)
    p = doc.add_paragraph()
    p.add_run('红榜（推荐）：').bold = True
    p.add_run('芹菜、油菜、茼蒿、春笋、银耳、红枣、山药、鲜蚕豆、豆芽、韭菜、紫甘蓝、香椿')
    p = doc.add_paragraph()
    p.add_run('黑榜（避免）：').bold = True
    p.add_run('鸡肉、螃蟹、羊肉、虾、生冷寒凉、海鱼、酸性食物、辛辣刺激、肥腻、发酵物')
    
    # 六、高膳食纤维食物清单（新增）
    doc.add_heading('六、高膳食纤维食物清单', level=1)
    p = doc.add_paragraph()
    p.add_run('来源：Rika黄老师').bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('谷物类：').bold = True
    p.add_run('燕麦、荞麦、糙米、玉米、小米')
    p = doc.add_paragraph()
    p.add_run('豆类：').bold = True
    p.add_run('黄豆、红豆、黑豆、鹰嘴豆、芸豆')
    p = doc.add_paragraph()
    p.add_run('蔬菜类：').bold = True
    p.add_run('菠菜、西兰花、芹菜、竹笋、南瓜')
    p = doc.add_paragraph()
    p.add_run('水果类：').bold = True
    p.add_run('苹果、石榴、猕猴桃、火龙果、蓝莓')
    p = doc.add_paragraph()
    p.add_run('菌藻类：').bold = True
    p.add_run('木耳、香菇、海带、裙带菜、银耳')
    p = doc.add_paragraph()
    p.add_run('坚果类：').bold = True
    p.add_run('黑芝麻、杏仁、核桃、巴旦木、腰果')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('本文档由AI整理生成，仅供参考学习使用').italic = True
    p.runs[0].font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run('更新记录：2026年3月18日 - 新增高膳食纤维食物清单').italic = True
    p.runs[0].font.size = Pt(10)
    
    output_path = os.path.join(OUTPUT_DIR, "01_抗炎饮食与营养科普.docx")
    doc.save(output_path)
    print("[OK] 01_抗炎饮食与营养科普.docx 生成完成")
    return True

if __name__ == "__main__":
    print("=" * 50)
    print("开始生成 Word 文档...")
    print("=" * 50)
    
    try:
        create_doc_01()
        create_doc_04()
        create_doc_03()
        print("\n" + "=" * 50)
        print("[DONE] 所有 Word 文档生成完成！")
        print("=" * 50)
    except Exception as e:
        print(f"\n[ERROR] 错误：{e}")
        import traceback
        traceback.print_exc()
