#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用 Word 文档生成脚本
=======================
自动将所有 Markdown 文档转换为 Word 文档。
支持多种模板，通过 --template 参数指定。

用法：
    python generate_word.py                         # 使用默认模板
    python generate_word.py --template minimalist   # 使用极简模板
    python generate_word.py --template detailed     # 使用详细模板
    python generate_word.py --list-templates        # 列出所有模板
    python generate_word.py --create-template 我的模板  # 创建自定义模板
    python generate_word.py --file 01_抗炎饮食.md   # 只处理指定文件
"""

import os
import sys
import re
import logging
from pathlib import Path
from typing import Optional, List, Tuple

# 确保 scripts 目录在导入路径中
_SCRIPTS_DIR = Path(__file__).parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("[错误] 未安装 python-docx")
    print("请运行: pip install python-docx")
    sys.exit(1)

try:
    from template_engine import TemplateEngine, TemplateConfig, DEFAULT_TEMPLATE_ID
    _TEMPLATE_SUPPORT = True
except ImportError:
    _TEMPLATE_SUPPORT = False

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(message)s")
logger = logging.getLogger(__name__)

# 配置路径
BASE_DIR = Path(__file__).parent.parent
OUTPUT_DIR = BASE_DIR / "处理结果"


# ============================================================
# Markdown 解析
# ============================================================

def parse_markdown(md_content: str) -> List[Tuple[str, object, int]]:
    """
    解析 Markdown 内容，返回结构化数据

    Returns:
        list of (type, content, level)
        type: 'heading' | 'paragraph' | 'list' | 'numbered_list' | 'table'
              | 'code' | 'hr' | 'blockquote' | 'blank'
    """
    lines = md_content.split("\n")
    result = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # 空行
        if not line.strip():
            result.append(("blank", "", 0))
            i += 1
            continue

        # 标题 #
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            content = line.lstrip("#").strip()
            result.append(("heading", content, level))
            i += 1
            continue

        # 水平线 --- 或 ***
        if re.match(r"^[\-\*\_]{3,}$", line.strip()):
            result.append(("hr", "", 0))
            i += 1
            continue

        # 代码块 ```
        if line.strip().startswith("```"):
            code_lines = []
            lang = line.strip()[3:].strip()
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束 ```
            result.append(("code", "\n".join(code_lines), 0))
            continue

        # 引用块 >
        if line.startswith("> "):
            content = line[2:].strip()
            result.append(("blockquote", content, 0))
            i += 1
            continue

        # 表格
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                rows = []
                for idx, row_line in enumerate(table_lines):
                    if idx == 1:  # 跳过分隔行 |---|---|
                        continue
                    cells = [c.strip() for c in row_line.split("|")[1:-1]]
                    if cells:
                        rows.append(cells)
                if rows:
                    result.append(("table", rows, 0))
            continue

        # 无序列表
        if line.strip().startswith(("- ", "* ", "+ ")):
            content = line.strip()[2:]
            result.append(("list", content, 0))
            i += 1
            continue

        # 有序列表
        numbered_match = re.match(r"^(\d+)\.\s+(.+)$", line.strip())
        if numbered_match:
            content = numbered_match.group(2)
            result.append(("numbered_list", content, 0))
            i += 1
            continue

        # 普通段落
        result.append(("paragraph", line.strip(), 0))
        i += 1

    return result


# ============================================================
# Word 生成（支持模板）
# ============================================================

def _apply_page_margins(doc: Document, tpl: Optional["TemplateConfig"]):
    """设置页面边距"""
    if tpl is None:
        return
    margins = tpl.get_page_margins_cm()
    for section in doc.sections:
        section.top_margin = Cm(margins["top"])
        section.bottom_margin = Cm(margins["bottom"])
        section.left_margin = Cm(margins["left"])
        section.right_margin = Cm(margins["right"])


def _add_paragraph_with_bold_italic(doc: Document, text: str) -> "docx.text.paragraph.Paragraph":
    """
    添加支持粗体/斜体的段落，解析 **bold** 和 *italic* 标记
    """
    p = doc.add_paragraph()
    # 先处理粗体 **...**
    parts = re.split(r"(\*\*[^*]+\*\*|__[^_]+__)", text)
    for part in parts:
        if (part.startswith("**") and part.endswith("**")) or \
           (part.startswith("__") and part.endswith("__")):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            # 再处理斜体 *...*
            sub_parts = re.split(r"(?<!\*)(\*(?!\*)[^*]+\*(?!\*)|_[^_]+_)(?!\*)", part)
            for sub in sub_parts:
                if (sub.startswith("*") and sub.endswith("*") and not sub.startswith("**")) or \
                   (sub.startswith("_") and sub.endswith("_")):
                    run = p.add_run(sub[1:-1])
                    run.italic = True
                else:
                    p.add_run(sub)
    return p


def create_word_from_markdown(
    md_file_path: Path,
    output_path: Path,
    template_id: str = DEFAULT_TEMPLATE_ID,
) -> bool:
    """
    将单个 Markdown 文件转换为 Word 文档

    Args:
        md_file_path: Markdown 文件路径
        output_path: 输出 Word 文件路径
        template_id: 使用的模板 ID

    Returns:
        bool: 是否成功
    """
    try:
        with open(md_file_path, "r", encoding="utf-8") as f:
            md_content = f.read()
    except Exception as e:
        logger.error(f"读取文件失败: {md_file_path} — {e}")
        return False

    # 加载模板
    tpl: Optional[TemplateConfig] = None
    if _TEMPLATE_SUPPORT:
        try:
            engine = TemplateEngine()
            tpl = engine.load(template_id)
        except Exception as e:
            logger.warning(f"模板加载失败（{template_id}），使用内置默认样式: {e}")

    try:
        parsed_content = parse_markdown(md_content)
        doc = Document()

        # 设置页面边距
        _apply_page_margins(doc, tpl)

        # 自动编号计数器
        numbering_counters = [0, 0, 0, 0]

        for item_type, content, level in parsed_content:
            if item_type == "blank":
                continue

            elif item_type == "heading":
                para = doc.add_heading(content, level=min(level, 4))
                # 主标题（level=1）特殊处理
                if level == 1:
                    if tpl:
                        tpl.apply_title(para)
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if tpl:
                    tpl.apply_heading(para, level, numbering_counters)

            elif item_type == "paragraph":
                if not content.strip():
                    continue
                para = _add_paragraph_with_bold_italic(doc, content)
                if tpl:
                    tpl.apply_body(para)

            elif item_type == "list":
                para = doc.add_paragraph(content, style="List Bullet")
                if tpl:
                    tpl.apply_list(para)

            elif item_type == "numbered_list":
                para = doc.add_paragraph(content, style="List Number")
                if tpl:
                    tpl.apply_list(para)

            elif item_type == "table":
                rows = content
                if rows:
                    num_cols = max(len(r) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    # 设置基础样式（再由模板覆盖）
                    table_style = "Table Grid"
                    if tpl:
                        table_style = tpl.table_cfg.get("style", "Table Grid")
                    try:
                        table.style = table_style
                    except Exception:
                        table.style = "Table Grid"

                    for row_idx, row_data in enumerate(rows):
                        row = table.rows[row_idx]
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < len(row.cells):
                                row.cells[col_idx].text = cell_text

                    if tpl:
                        tpl.apply_table(table)

                    doc.add_paragraph()

            elif item_type == "code":
                # 代码块：用等宽字体、灰色背景
                para = doc.add_paragraph(content)
                for run in para.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)

            elif item_type == "blockquote":
                para = doc.add_paragraph(content)
                para.style = "Quote" if "Quote" in [s.name for s in doc.styles] else "Normal"
                for run in para.runs:
                    run.italic = True

            elif item_type == "hr":
                # 水平线：加一个空段落+下划线
                para = doc.add_paragraph()
                pPr = para._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:color"), "AAAAAA")
                pBdr.append(bottom)
                pPr.append(pBdr)

            # 详细模板：章节间分页
            if tpl and tpl.should_page_break() and item_type == "heading" and level == 1:
                # 只在第一个 h1 之后分页，避免文档开头就多一页
                if numbering_counters[0] > 1:
                    para.runs[0].add_break() if para.runs else None

        # 底部版权声明
        if tpl is None or tpl.should_include_footer():
            footer_text = "本文档由AI整理生成"
            if tpl:
                footer_text = tpl.footer_text()
            if footer_text:
                doc.add_paragraph()
                footer_para = doc.add_paragraph()
                footer_run = footer_para.add_run(footer_text)
                footer_run.italic = True
                footer_run.font.size = Pt(9)
                footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(output_path)
        tpl_name = tpl.name if tpl else "默认"
        logger.info(f"[OK] {output_path.name} 生成完成（模板: {tpl_name}）")
        return True

    except Exception as e:
        logger.error(f"生成 {output_path.name} 失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def generate_all_word_docs(
    template_id: str = DEFAULT_TEMPLATE_ID,
    target_file: Optional[str] = None,
) -> Tuple[int, int]:
    """
    生成所有（或指定）Markdown 文档对应的 Word 文档

    Args:
        template_id: 模板 ID
        target_file: 仅处理指定文件名（不含路径），None 表示处理全部

    Returns:
        (成功数量, 失败数量)
    """
    success_count = 0
    fail_count = 0
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if target_file:
        md_files = [OUTPUT_DIR / target_file]
        md_files = [f for f in md_files if f.exists()]
        if not md_files:
            print(f"[错误] 文件不存在: {OUTPUT_DIR / target_file}")
            return 0, 1
    else:
        md_files = sorted(OUTPUT_DIR.glob("*.md"))

    if not md_files:
        print("[警告] 未找到 Markdown 文件")
        return 0, 0

    print(f"\n[发现 {len(md_files)} 个 Markdown 文件，模板: {template_id}]")
    print("=" * 50)

    for md_file in md_files:
        docx_name = md_file.stem + ".docx"
        output_path = OUTPUT_DIR / docx_name
        print(f"\n处理: {md_file.name}")

        # 跳过已是最新的 Word
        if output_path.exists() and not target_file:
            md_mtime = md_file.stat().st_mtime
            docx_mtime = output_path.stat().st_mtime
            if docx_mtime >= md_mtime:
                print("  [跳过] Word 文档已是最新")
                success_count += 1
                continue

        if create_word_from_markdown(md_file, output_path, template_id=template_id):
            success_count += 1
        else:
            fail_count += 1

    return success_count, fail_count


# ============================================================
# 命令行入口
# ============================================================

def main():
    import sys

    # 参数解析
    template_id = DEFAULT_TEMPLATE_ID
    list_templates = False
    create_template_name = None
    create_base = DEFAULT_TEMPLATE_ID
    target_file = None

    args = sys.argv[1:]
    i = 0
    while i < len(args):
        if args[i] == "--template" and i + 1 < len(args):
            template_id = args[i + 1]
            i += 2
        elif args[i] == "--list-templates":
            list_templates = True
            i += 1
        elif args[i] == "--create-template" and i + 1 < len(args):
            create_template_name = args[i + 1]
            i += 2
        elif args[i] == "--base" and i + 1 < len(args):
            create_base = args[i + 1]
            i += 2
        elif args[i] == "--file" and i + 1 < len(args):
            target_file = args[i + 1]
            i += 2
        else:
            i += 1

    # 列出模板
    if list_templates:
        if _TEMPLATE_SUPPORT:
            from template_engine import cmd_list
            cmd_list()
        else:
            print("[警告] template_engine 未找到，无法列出模板")
        return 0

    # 创建自定义模板
    if create_template_name:
        if _TEMPLATE_SUPPORT:
            from template_engine import cmd_create
            cmd_create(create_template_name, create_base)
        else:
            print("[警告] template_engine 未找到，无法创建模板")
        return 0

    # 生成 Word
    print("=" * 50)
    print("通用 Word 文档生成工具")
    print("=" * 50)
    print(f"输出目录: {OUTPUT_DIR}")
    print(f"使用模板: {template_id}")

    success, fail = generate_all_word_docs(
        template_id=template_id,
        target_file=target_file,
    )

    print("\n" + "=" * 50)
    print("生成完成!")
    print(f"  成功: {success} 个")
    print(f"  失败: {fail} 个")
    print("=" * 50)

    return 0 if fail == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
