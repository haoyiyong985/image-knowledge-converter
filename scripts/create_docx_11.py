from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 设置标题
title = doc.add_heading('植物养护指南', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加文档信息
doc.add_paragraph('整理来源：小红书、微信等平台的种植养护内容')
doc.add_paragraph('整理时间：2026年3月30日')
doc.add_paragraph()

# 月季病虫害识别与防治
doc.add_heading('月季病虫害识别与防治', 1)
doc.add_paragraph('来源：小红书（是柠檬吖）')
doc.add_paragraph()

# 创建表格
table = doc.add_table(rows=10, cols=4)
table.style = 'Light Grid Accent 1'

# 表头
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '害虫/病害'
hdr_cells[1].text = '危害特征'
hdr_cells[2].text = '高发时期'
hdr_cells[3].text = '防治方法'

# 表格数据
data = [
    ['蚜虫', '叶片、花和茎被啃食', '时高发', '甲维盐、高效氯氰菊酯、苏云金杆菌'],
    ['夜蛾幼虫', '花苞被啃食', '晚春初夏', '苏云金杆菌、甲维盐、高效氯氰菊酯'],
    ['刺蛾', '叶片背面啃食', '幼虫期', '苏云金杆菌、单维盐、高效氯氰菊酯'],
    ['叶蝉', '叶片表面变白', '全年', '并健、随奇易威、特富力'],
    ['茎蜂', '茎尖刺孔', '产卵期', '先正达森康、绿色威雷、杀虫双'],
    ['叶蜂', '枝条产卵', '春秋高发', '先正达森康、绿色威雷、杀虫双'],
    ['蓟马', '花瓣褐色斑点', '全年', '开健、特富力、艾绿士'],
    ['红蜘蛛', '叶片灰黄点', '高温干燥', '联苯肼酯、阿维菌素、乙唑螨腈'],
    ['根结线', '根部肿大', '定植期', '灭线膦、噻唑膦颗粒']
]

for i, row_data in enumerate(data, 1):
    row = table.rows[i].cells
    for j, text in enumerate(row_data):
        row[j].text = text

doc.add_paragraph()
doc.add_heading('防治要点', 2)
points = ['预防为主：早春清理杂草和垃圾', '物理防治：手工捕捉、修剪受害枝条', '药剂轮换：使用两种以上药剂轮换', '注意时机：蓟马防治需在傍晚进行']
for point in points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_page_break()

# 自制菜园杀虫水配方
doc.add_heading('自制菜园杀虫水配方', 1)
doc.add_paragraph('来源：小红书（种菜记小邱）')
doc.add_paragraph()

doc.add_heading('1. 白醋水', 2)
doc.add_paragraph('用途：防止白粉病、黑斑')

doc.add_heading('2. 花椒水', 2)
doc.add_paragraph('配方：花椒6克 + 水60克，煮沸过滤使用')
doc.add_paragraph('功效：防治蚜虫、介壳虫')

doc.add_heading('3. 大蒜水', 2)
doc.add_paragraph('配方：蒜泥10克 + 水160克，浸泡一天')
doc.add_paragraph('功效：抗菌、防治红蜘蛛、蚜虫、小黑飞')

doc.add_page_break()

# 盆栽番茄高产种植技巧
doc.add_heading('盆栽番茄高产种植技巧', 1)
doc.add_paragraph('来源：小红书（taco都市农妇）')
doc.add_paragraph()

doc.add_heading('底肥配方', 2)
doc.add_paragraph('煮熟的豆类（氮肥）+ 稻谷灰（钾肥）+ 虾蟹壳粉（磷钙肥）')

doc.add_heading('追肥方法', 2)
doc.add_paragraph('开花结果期浇过期酸奶加水100倍稀释的酸奶水')

doc.add_page_break()

# 蜡梅养护指南
doc.add_heading('蜡梅养护指南', 1)
doc.add_paragraph('来源：小红书')
doc.add_paragraph()

table3 = doc.add_table(rows=6, cols=2)
table3.style = 'Light Grid Accent 1'
hdr3 = table3.rows[0].cells
hdr3[0].text = '养护项目'
hdr3[1].text = '具体方法'

wax_plum_data = [
    ['光照', '喜阳植物，也耐半阴'],
    ['土壤', '盆栽需田园土加泥炭土改良'],
    ['浇水', '干透浇透，避免积水'],
    ['施肥', '早春用有机肥，4-6月用生长肥，6-9月用磷钾肥'],
    ['修剪', '发芽前剪短开花枝，10月停止修剪']
]

for i, row_data in enumerate(wax_plum_data, 1):
    row = table3.rows[i].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

doc.add_page_break()

# 金枝玉叶盆景造型技巧
doc.add_heading('金枝玉叶盆景造型技巧', 1)
doc.add_paragraph('来源：微信视频号（一分花田）')
doc.add_paragraph()

doc.add_heading('造型步骤', 2)
steps = ['选择主干：确定盆景的主干走向', '蟠扎定型：用铝线缠绕枝条弯曲造型', '修剪整形：剪除多余枝条', '养护定型：定期调整保持造型']
for step in steps:
    doc.add_paragraph(step, style='List Number')

# 保存文档
doc.save('处理结果/11_植物养护指南.docx')
print('Word文档已生成：处理结果/11_植物养护指南.docx')
