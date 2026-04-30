#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
全自动化图片处理工具
==============================
处理流程：
  1. OCR识别（优先级：腾讯云 → 百度 → 本地Tesseract）
  2. AI智能分类
  3. 生成Markdown文档
  4. 生成Word文档
  5. 同步到IMA笔记

使用方法：
  python auto_process_all.py
"""

import os
import sys
import io
import json
import time
import logging
from pathlib import Path
from datetime import datetime
from collections import Counter

# 修复 Windows 控制台编码
if sys.platform == 'win32':
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
    except Exception:
        pass

# 添加项目根目录到路径
sys.path.insert(0, '.')

# 加载 .env 环境变量配置
from dotenv import load_dotenv
env_path = Path('.') / '.env'
if env_path.exists():
    load_dotenv(dotenv_path=env_path)
else:
    print("[WARN] .env 文件不存在，使用系统环境变量")

# 导入各模块
from local_ocr import LocalOCR
from scripts.classifier_engine import ClassifierEngine

# 尝试导入云端OCR
try:
    from tencent_ocr import TencentOCR
    TENINCENT_AVAILABLE = True
except ImportError:
    TENINCENT_AVAILABLE = False

try:
    from baidu_ocr import BaiduOCR
    BAIDU_AVAILABLE = True
except ImportError:
    BAIDU_AVAILABLE = False

# 日志配置
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(message)s',
    handlers=[
        logging.FileHandler('处理结果/process.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


class MultiEngineOCR:
    """多引擎OCR管理器"""
    
    def __init__(self):
        self.current_engine = None
        
        # 按优先级检查引擎
        self.engine_status = []
        
        # 检查腾讯云
        secret_id = os.getenv('TENCENT_SECRET_ID', '')
        secret_key = os.getenv('TENCENT_SECRET_KEY', '')
        if secret_id and secret_key and '替换' not in secret_id and TENINCENT_AVAILABLE:
            self.engine_status.append(('腾讯云', True, '已配置'))
        else:
            self.engine_status.append(('腾讯云', False, '未配置'))
        
        # 检查百度云
        app_id = os.getenv('BAIDU_APP_ID', '')
        api_key = os.getenv('BAIDU_API_KEY', '')
        secret_key = os.getenv('BAIDU_SECRET_KEY', '')
        if app_id and api_key and secret_key and '替换' not in app_id and BAIDU_AVAILABLE:
            self.engine_status.append(('百度云', True, '已配置'))
        else:
            self.engine_status.append(('百度云', False, '未配置'))
        
        # 检查本地Tesseract
        local = LocalOCR()
        if local.tesseract_available:
            self.engine_status.append(('本地Tesseract', True, '可用'))
        else:
            self.engine_status.append(('本地Tesseract', False, getattr(local, 'error_message', '未知错误')))
        
        self._select_best_engine()
    
    def _select_best_engine(self):
        """选择最佳可用引擎"""
        for name, available, msg in self.engine_status:
            if available:
                self.current_engine = name
                logger.info(f"[OCR] 选用引擎: {name} ({msg})")
                return True
        logger.warning("[OCR] 没有可用的OCR引擎!")
        return False
    
    def recognize(self, image_path):
        """识别图片，按优先级尝试"""
        last_error = None
        
        for name, available, _ in self.engine_status:
            if not available:
                logger.info(f"[OCR] 跳过 {name}: 不可用")
                continue
            
            try:
                logger.info(f"[OCR] 尝试 {name}...")
                
                if name == '腾讯云' and TENINCENT_AVAILABLE:
                    engine = TencentOCR()
                elif name == '百度云' and BAIDU_AVAILABLE:
                    engine = BaiduOCR()
                else:
                    engine = LocalOCR()
                
                result = engine.recognize(image_path)
                
                if result and result.get('success'):
                    logger.info(f"[OCR] {name} 识别成功!")
                    return result
                else:
                    last_error = result.get('error', '未知错误') if result else '返回为空'
                    logger.warning(f"[OCR] {name} 识别失败: {last_error}")
            except Exception as e:
                last_error = str(e)
                logger.warning(f"[OCR] {name} 异常: {last_error}")
        
        # 所有引擎都失败
        return {'success': False, 'error': last_error or '无可用OCR引擎', 'text': ''}


class DocumentGenerator:
    """文档生成器"""
    
    def __init__(self, output_dir='处理结果'):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    def generate_markdown(self, text, category, image_name, keywords=None):
        """生成Markdown文档"""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
        
        md_content = f"""# {category}

> 来源图片: {image_name}  
> 识别时间: {timestamp}  
> 关键词: {', '.join(keywords) if keywords else '无'}

---

## 内容

{text}

---

*本文档由图片知识库整理工具自动生成*
"""
        
        # 文件名：分类_时间戳.md
        safe_name = image_name.replace('.jpg', '').replace('.png', '').replace('.jpeg', '')
        md_file = self.output_dir / f"{safe_name}.md"
        
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        logger.info(f"[文档] Markdown已生成: {md_file.name}")
        return str(md_file)
    
    def generate_word(self, text, category, image_name, keywords=None):
        """生成Word文档"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.oxml.ns import qn
        except ImportError:
            logger.warning("[文档] python-docx未安装，跳过Word生成")
            return None
        
        doc = Document()
        
        # 设置中文字体
        doc.styles['Normal'].font.name = 'Microsoft YaHei'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 标题
        doc.add_heading(category, 0)
        
        # 元信息
        doc.add_paragraph(f"来源图片: {image_name}")
        doc.add_paragraph(f"识别时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"关键词: {', '.join(keywords) if keywords else '无'}")
        doc.add_paragraph("─" * 30)
        
        # 正文
        doc.add_paragraph(text)
        
        # 保存
        safe_name = image_name.replace('.jpg', '').replace('.png', '').replace('.jpeg', '')
        docx_file = self.output_dir / f"{safe_name}.docx"
        doc.save(str(docx_file))
        
        logger.info(f"[文档] Word已生成: {docx_file.name}")
        return str(docx_file)


class IMASyncer:
    """IMA笔记同步器"""
    
    def __init__(self):
        self.client_id = os.getenv('IMA_OPENAPI_CLIENTID', '')
        self.api_key = os.getenv('IMA_OPENAPI_APIKEY', '')
        self.base_url = 'https://ima.qq.com/openapi/note/v1'
        self.enabled = bool(self.client_id and self.api_key and '填入' not in self.client_id)
        
        if self.enabled:
            logger.info("[IMA] 已配置，准备同步")
        else:
            logger.warning("[IMA] 未配置或凭证无效，跳过同步")
    
    def _api_call(self, endpoint, payload):
        """调用IMA API"""
        if not self.enabled:
            return None
        
        try:
            import requests
            
            headers = {
                'ima-openapi-clientid': self.client_id,
                'ima-openapi-apikey': self.api_key,
                'Content-Type': 'application/json'
            }
            
            url = f"{self.base_url}/{endpoint}"
            response = requests.post(url, json=payload, headers=headers, timeout=30)
            
            if response.status_code == 200:
                return response.json()
            else:
                logger.warning(f"[IMA] API响应错误: {response.status_code}")
                return None
        except requests.exceptions.Timeout:
            logger.warning("[IMA] API调用超时")
            return None
        except Exception as e:
            logger.warning(f"[IMA] API调用失败: {e}")
            return None
    
    def sync_note(self, title, content, category=None):
        """同步笔记到IMA"""
        if not self.enabled:
            return None
        
        # 格式化内容
        full_content = f"# {title}\n\n"
        if category:
            full_content += f"> 分类: {category}\n\n"
        full_content += f"{content}\n\n---\n*自动同步自图片知识库*\n"
        
        # 调用IMA API
        payload = {
            'content_format': 1,  # Markdown
            'content': full_content
        }
        
        result = self._api_call('import_doc', payload)
        
        # IMA API 返回格式: {"code":0,"msg":"success","data":{"note_id":"..."}}
        if result and result.get('code') == 0:
            doc_id = result.get('data', {}).get('note_id', '')
            if doc_id:
                logger.info(f"[IMA] 笔记已同步: {title} (ID: {doc_id})")
                return doc_id
            else:
                logger.warning(f"[IMA] 同步成功但无note_id: {result}")
                return "imported"
        else:
            error_msg = result.get('msg', '未知错误') if result else '无响应'
            logger.warning(f"[IMA] 同步失败: {error_msg}")
            return None


def process_single_image(ocr, classifier, doc_gen, ima_syncer, image_path, index, total):
    """处理单张图片"""
    image_name = Path(image_path).name
    print(f"\n{'='*50}")
    print(f"[{index}/{total}] 处理: {image_name}")
    print(f"{'='*50}")
    
    # 1. OCR识别
    print("\n📷 OCR识别中...")
    result = ocr.recognize(image_path)
    
    if not result.get('success'):
        print(f"  ⚠️ OCR失败: {result.get('error')}")
        return None
    
    text = result.get('text', '').strip()
    if not text:
        print("  ⚠️ 无文字内容")
        return None
    
    print(f"  ✓ 识别到 {len(text)} 个字符")
    
    # 2. AI分类
    print("\n🏷️ 内容分类中...")
    categories = classifier.classify(text, threshold=0.25)
    
    if categories:
        primary = categories[0]
        category_name = primary.category_name
        keywords = primary.matched_keywords if primary.matched_keywords else []
        confidence = primary.confidence
        print(f"  ✓ 分类: {category_name} (置信度 {confidence:.0%})")
        if keywords:
            print(f"     关键词: {', '.join(keywords[:5])}")
    else:
        category_name = "综合知识"
        keywords = []
        print("  ⚠️ 未匹配到分类，使用默认分类")
    
    # 3. 生成Markdown
    print("\n📝 生成Markdown...")
    md_file = doc_gen.generate_markdown(text, category_name, image_name, keywords)
    
    # 4. 生成Word
    print("\n📄 生成Word...")
    docx_file = doc_gen.generate_word(text, category_name, image_name, keywords)
    
    # 5. 同步到IMA
    print("\n☁️ 同步到IMA...")
    ima_id = ima_syncer.sync_note(image_name.replace('.jpg', '').replace('.png', ''), text, category_name)
    if ima_id:
        print(f"  ✓ IMA同步成功 (ID: {ima_id})")
    else:
        print("  ⚠️ IMA未配置或同步失败")
    
    return {
        'image': image_name,
        'text_length': len(text),
        'category': category_name,
        'keywords': keywords,
        'md_file': md_file,
        'docx_file': docx_file,
        'ima_id': ima_id
    }


def main():
    print("=" * 60)
    print("全自动化图片处理工具")
    print("OCR → 分类 → Markdown → Word → IMA同步")
    print("=" * 60)
    
    # 初始化组件
    print("\n[1/5] 初始化OCR引擎...")
    ocr = MultiEngineOCR()
    if not ocr.current_engine:
        print("❌ 没有可用的OCR引擎!")
        return
    
    print("\n[2/5] 初始化分类器...")
    classifier = ClassifierEngine()
    
    print("\n[3/5] 初始化文档生成器...")
    doc_gen = DocumentGenerator()
    
    print("\n[4/5] 初始化IMA同步器...")
    ima_syncer = IMASyncer()
    
    print("\n[5/5] 扫描待处理图片...")
    
    # 扫描待处理图片
    source_dir = Path('待处理图片')
    if not source_dir.exists():
        source_dir = Path('待处理图片/历史')
    
    images = list(source_dir.rglob('*.jpg')) + list(source_dir.rglob('*.png')) + list(source_dir.rglob('*.jpeg'))
    
    if not images:
        print("❌ 未找到待处理图片")
        return
    
    total = len(images)
    print(f"✓ 找到 {total} 张图片\n")
    
    # 批量处理
    results = []
    start_time = time.time()
    
    for i, img_path in enumerate(images, 1):
        try:
            result = process_single_image(ocr, classifier, doc_gen, ima_syncer, str(img_path), i, total)
            if result:
                results.append(result)
        except Exception as e:
            print(f"  ✗ 处理失败: {e}")
            logger.error(f"处理失败 {img_path}: {e}")
        
        # 每处理一张稍作停顿
        if i < total:
            time.sleep(0.3)
    
    # 统计
    elapsed = time.time() - start_time
    
    print("\n" + "=" * 60)
    print("处理完成!")
    print("=" * 60)
    print(f"总计处理: {len(results)}/{total} 张")
    print(f"总耗时: {elapsed:.1f} 秒")
    if results:
        print(f"平均速度: {elapsed/len(results):.1f} 秒/张")
    
    # 分类统计
    if results:
        print("\n分类统计:")
        cats = Counter(r['category'] for r in results)
        for cat, count in cats.most_common():
            print(f"  - {cat}: {count} 张")
    
    # 字符统计
    total_chars = sum(r['text_length'] for r in results)
    print(f"\n文字统计:")
    print(f"  - 总识别字符: {total_chars} 个")
    if results:
        print(f"  - 平均每张: {total_chars/len(results):.0f} 个字符")
    
    # 输出文件
    print(f"\n输出目录: 处理结果/")
    md_count = len([r for r in results if r.get('md_file')])
    docx_count = len([r for r in results if r.get('docx_file')])
    ima_count = len([r for r in results if r.get('ima_id')])
    print(f"  - Markdown: {md_count} 个")
    print(f"  - Word: {docx_count} 个")
    print(f"  - IMA同步: {ima_count} 个")
    
    print("\n" + "=" * 60)
    
    # 保存处理报告
    report_file = Path('处理结果/处理报告.json')
    report = {
        'timestamp': datetime.now().isoformat(),
        'total_images': total,
        'processed': len(results),
        'elapsed_seconds': elapsed,
        'results': results,
        'category_stats': dict(Counter(r['category'] for r in results))
    }
    with open(report_file, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    
    print(f"\n📊 详细报告已保存: {report_file}")


if __name__ == '__main__':
    main()
