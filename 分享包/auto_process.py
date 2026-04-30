#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
auto_process.py — 图片知识库自动处理脚本
=========================================
功能：
  1. 扫描 待处理图片/<主题>/ 下的所有新图片
  2. 【新增】自动读取现有 MD 文件结构，输出给 AI 用于精准插入判断
  3. 调用 AI 识别图片文字内容（需在 WorkBuddy 等支持图片读取的环境中使用）
  4. 分析内容，自动归类到已有文档或新建文档
  5. 智能合并到 处理结果/*.md 和 *.docx
  6. 将已处理图片移动到 已处理图片/<主题>/ 归档

使用方法：
  python auto_process.py         ← 扫描+读取现有文档结构（每次处理前运行）
  python auto_process.py --scan  ← 仅扫描待处理图片（不读取MD）

注意：
  此脚本是整套流程的"骨架"，图片识别部分由 AI（WorkBuddy）直接完成，
  脚本负责文件管理、目录操作和文档合并。
  实际使用时通过 WorkBuddy 的对话界面触发。
"""

import os
import re
import sys
import shutil
from pathlib import Path
from datetime import datetime

# ============================================================
# 路径配置
# ============================================================
BASE_DIR        = Path("D:/新建文件夹")
PENDING_DIR     = BASE_DIR / "待处理图片"       # 待处理图片根目录
PROCESSED_DIR   = BASE_DIR / "已处理图片"        # 已处理图片归档目录
RESULT_DIR      = BASE_DIR / "处理结果"          # 输出结果目录

# 已有文档（文件名 → 文档标题，用于判断归类）
EXISTING_DOCS = {
    "01_抗炎饮食与营养科普":   "抗炎饮食、坚果、ω-3、ω-6、营养素、维生素、矿物质、脂肪酸",
    "02_肠道健康与饮食分类":   "肠道、益生菌、益生元、绿灯食物、红灯食物、微生物、膳食纤维",
    "03_中医养生与食疗":       "中医、三伏、养生、食疗、茶饮、汤药、湿气、健脾、温补",
    "04_日常饮食建议":         "早餐、食谱、搭配、热量、蛋白质、减脂、饮食建议",
}

# 图片格式
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".bmp"}


def scan_pending_images():
    """扫描所有待处理图片，返回 {主题: [图片路径列表]} 的字典"""
    result = {}
    if not PENDING_DIR.exists():
        print(f"[错误] 待处理图片目录不存在：{PENDING_DIR}")
        return result

    for topic_dir in PENDING_DIR.iterdir():
        if not topic_dir.is_dir():
            continue
        images = [
            f for f in topic_dir.iterdir()
            if f.suffix.lower() in IMAGE_EXTENSIONS
        ]
        if images:
            result[topic_dir.name] = sorted(images)
            print(f"[扫描] 主题「{topic_dir.name}」：发现 {len(images)} 张图片")

    return result


def read_existing_docs_structure() -> dict:
    """
    读取 处理结果/ 目录下所有 MD 文件的结构（标题和章节），
    返回 {文件名: {title, sections: [章节名列表], keywords: str}} 的字典。
    每次处理新图片前调用，确保合并时能精准判断插入位置。
    """
    structure = {}
    md_files = sorted(RESULT_DIR.glob("[0-9][0-9]_*.md"))

    for md_path in md_files:
        doc_name = md_path.stem
        content  = md_path.read_text(encoding="utf-8")
        lines    = content.splitlines()

        title    = ""
        sections = []  # ## 级章节
        subsections = []  # ### 级子章节

        for line in lines:
            if line.startswith("# ") and not title:
                title = line[2:].strip()
            elif line.startswith("## "):
                sections.append(line[3:].strip())
            elif line.startswith("### "):
                subsections.append(line[4:].strip())

        structure[doc_name] = {
            "title":       title,
            "sections":    sections,
            "subsections": subsections,
            "char_count":  len(content),
        }

    return structure


def print_docs_structure(structure: dict):
    """将现有文档结构打印出来，供 AI 读取判断分类"""
    print("\n" + "=" * 60)
    print("现有文档结构（AI 合并参考）")
    print("=" * 60)
    if not structure:
        print("[提示] 尚无已有文档，所有内容将新建文档。")
    else:
        for doc_name, info in structure.items():
            print(f"\n[{doc_name}]")
            print(f"  标题：{info['title']}")
            print(f"  章节数：{len(info['sections'])}，字数：{info['char_count']}")
            if info["sections"]:
                print("  章节列表：")
                for s in info["sections"]:
                    print(f"    - {s}")
    print("=" * 60)


def archive_images(topic: str, image_paths: list):
    """将已处理图片移动到 已处理图片/<主题>/ 归档"""
    archive_dir = PROCESSED_DIR / topic
    archive_dir.mkdir(parents=True, exist_ok=True)

    moved = []
    for img_path in image_paths:
        dest = archive_dir / img_path.name
        # 如果目标已存在，加时间戳避免覆盖
        if dest.exists():
            stem = img_path.stem
            suffix = img_path.suffix
            ts = datetime.now().strftime("%Y%m%d%H%M%S")
            dest = archive_dir / f"{stem}_{ts}{suffix}"
        shutil.move(str(img_path), str(dest))
        moved.append(dest)
        print(f"  [归档] {img_path.name} → 已处理图片/{topic}/")

    return moved


def get_next_doc_number():
    """获取下一个文档编号（已有文档数 + 1）"""
    existing = list(RESULT_DIR.glob("[0-9][0-9]_*.md"))
    return len(existing) + 1


def create_new_md(doc_name: str, title: str, content: str):
    """新建 MD 文档"""
    md_path = RESULT_DIR / f"{doc_name}.md"
    today = datetime.now().strftime("%Y年%m月%d日")
    header = f"""# {title}

> 整理来源：微信、小红书等平台截图
> 整理时间：{today}

---

"""
    md_path.write_text(header + content, encoding="utf-8")
    print(f"  [新建] {md_path.name}")
    return md_path


def append_to_md(doc_name: str, content: str):
    """将内容追加到已有 MD 文档末尾"""
    md_path = RESULT_DIR / f"{doc_name}.md"
    if not md_path.exists():
        print(f"  [警告] 目标文档不存在：{md_path}")
        return

    existing = md_path.read_text(encoding="utf-8")
    separator = f"\n\n---\n\n> 补充内容（{datetime.now().strftime('%Y年%m月%d日')}）\n\n"
    md_path.write_text(existing + separator + content, encoding="utf-8")
    print(f"  [追加] 内容已追加到 {md_path.name}")


def md_to_docx(doc_name: str):
    """将 MD 文件转换为 Word 文档（调用 create_docx.py 的逻辑）"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        md_path  = RESULT_DIR / f"{doc_name}.md"
        docx_path = RESULT_DIR / f"{doc_name}.docx"

        if not md_path.exists():
            print(f"  [警告] MD 文件不存在：{md_path}")
            return

        doc = Document()
        # 设置默认字体
        style = doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style.font.size = Pt(11)

        lines = md_path.read_text(encoding="utf-8").splitlines()
        for line in lines:
            line = line.rstrip()
            if line.startswith("# "):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r"^\d+\. ", line):
                doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:])
                p.paragraph_format.left_indent = Pt(20)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            elif line.startswith("---"):
                doc.add_paragraph("─" * 30)
            elif line == "":
                doc.add_paragraph()
            else:
                # 处理加粗 **text**
                p = doc.add_paragraph()
                parts = re.split(r"\*\*(.*?)\*\*", line)
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True

        doc.save(str(docx_path))
        print(f"  [生成] Word 文档：{docx_path.name}")

    except ImportError:
        print("  [提示] 未安装 python-docx，跳过 Word 文档生成。")
        print("         请运行：pip install python-docx")
    except Exception as e:
        print(f"  [错误] 生成 Word 文档失败：{e}")


def generate_startup_prompt(structure: dict):
    """
    根据当前 处理结果/ 目录的实际文档状态，
    自动重新生成「启动提示词」文件。
    每次处理完图片后调用，确保文件始终是最新的。
    """
    prompt_path = BASE_DIR / "【启动提示词】新对话粘贴这段.txt"
    today = datetime.now().strftime("%Y年%m月%d日 %H:%M")

    # 动态生成已有文档列表部分
    doc_lines = []
    for doc_name, info in structure.items():
        sections_str = "、".join(info["sections"]) if info["sections"] else "（暂无章节）"
        doc_lines.append(
            f"编号{doc_name[:2]}：{doc_name}.docx / .md\n"
            f"  - 文档标题：{info['title']}\n"
            f"  - 现有章节：{sections_str}\n"
            f"  - 字数：约{info['char_count']}字"
        )
    docs_block = "\n\n".join(doc_lines) if doc_lines else "（尚无已有文档）"

    content = f"""你好，请帮我继续处理图片知识库项目。以下是项目完整背景，请读取后告诉我你已准备好。

========== 项目背景 ==========

【项目位置】D:\\新建文件夹\\

【任务说明】
将手机截图（小红书、微信读书、微信等来源）中的文字识别出来，
按内容分类整理，合并到对应的 Word 和 Markdown 文档中。

【目录结构】
- 待处理图片\\<主题名>\\   ← 新图片放在这里，按主题建子文件夹
- 已处理图片\\<主题名>\\   ← 处理完的图片自动归档到这里
- 处理结果\\              ← 所有输出的 Word (.docx) 和 Markdown (.md) 文件

【已有文档列表】（最后更新：{today}）
{docs_block}

【处理规则】
1. 识别新图片中的文字内容
2. 判断内容属于哪个分类：
   - 若属于上面已有文档之一 → 追加到对应的 .md 文件末尾（追加，不覆盖）
   - 若是全新主题 → 新建下一编号文档（如 {str(len(structure)+1).zfill(2)}_xxx.md 和 .docx）
3. 每次处理时先读取目标 .md 文件，了解已有章节，避免重复插入
4. 处理完毕后，将图片从 待处理图片\\<主题>\\ 移动到 已处理图片\\<主题>\\
5. 处理完成后调用 auto_process.py 自动刷新本启动提示词文件

【等待指令】
准备好后，我会告诉你接下来要做什么。
常用指令：
- "处理新图片"              → 扫描并处理 待处理图片\\ 下的所有新图片
- "处理待处理图片/<主题名>"  → 只处理指定主题文件夹的图片
- "查看现有文档"            → 列出处理结果目录中的所有文档和章节

==============================
"""

    prompt_path.write_text(content, encoding="utf-8")
    print(f"\n[更新] 启动提示词已自动刷新：{prompt_path.name}")
    print(f"       当前收录文档数：{len(structure)} 个")


def print_summary(pending_images: dict):
    """打印扫描汇总信息"""
    total = sum(len(v) for v in pending_images.values())
    print("\n" + "=" * 60)
    print("待处理图片扫描汇总")
    print("=" * 60)
    if total == 0:
        print("[OK] 没有发现新的待处理图片。")
    else:
        print(f"共发现 {total} 张待处理图片，分布如下：")
        for topic, images in pending_images.items():
            print(f"  [DIR] {topic}/  ({len(images)} 张)")
            for img in images:
                print(f"      - {img.name}")
    print("=" * 60)


def main():
    """
    主函数：
    1. 确保目录存在
    2. 【核心】读取现有 MD 文档结构 → 供 AI 合并时精准插入
    3. 扫描待处理图片，列出清单
    """
    scan_only = "--scan" in sys.argv

    print("=" * 60)
    print("图片知识库 — 处理前检查")
    print("=" * 60)

    # 确保目录存在
    RESULT_DIR.mkdir(parents=True, exist_ok=True)
    PROCESSED_DIR.mkdir(parents=True, exist_ok=True)

    # ---- 第一步：读取现有文档结构 ----
    if not scan_only:
        print("\n[Step 1] 读取现有文档结构...")
        structure = read_existing_docs_structure()
        print_docs_structure(structure)

        # ---- 自动刷新启动提示词 ----
        print("\n[Step 2] 自动刷新启动提示词文件...")
        generate_startup_prompt(structure)
    else:
        structure = {}
        print("\n[模式] 仅扫描，跳过文档结构读取")

    # ---- 第三步：扫描待处理图片 ----
    print("\n[Step 3] 扫描待处理图片...")
    pending = scan_pending_images()
    print_summary(pending)


    # ---- 输出操作指引 ----
    if not pending:
        print("\n[提示] 没有待处理图片。")
        print("       将新图片放入 待处理图片/<主题>/ 目录后，")
        print("       在 WorkBuddy 中发送「处理新图片」即可自动开始。")
    else:
        print("\n[提示] 准备就绪！请在 WorkBuddy 中发送「处理新图片」")
        print("       AI 将自动：识别 -> 分类 -> 合并 -> 生成文档 -> 归档图片")
        print()
        print("       合并规则：")
        print("         - 同类内容 → 追加到对应 .md 文件（不覆盖已有章节）")
        print("         - 新类别   → 自动新建下一编号文档")

    # ---- 第四步：自动同步到 ima（如果已配置凭证）----
    if not scan_only:
        _try_ima_sync()


def _try_ima_sync():
    """
    尝试调用 ima_sync.py 同步到 ima 个人笔记。
    如果凭证未配置或 ima_sync.py 不存在，静默跳过（不报错）。
    """
    ima_sync_path = BASE_DIR / "ima_sync.py"
    config_path   = BASE_DIR / "ima_config.txt"

    if not ima_sync_path.exists():
        return  # 脚本不存在，跳过

    if not config_path.exists():
        return  # 配置文件不存在，跳过

    # 检查配置文件里是否还是占位文字
    config_text = config_path.read_text(encoding="utf-8")
    if "填入你的" in config_text:
        print("\n[ima 同步] 跳过：ima_config.txt 尚未填写凭证")
        print("           填写后，每次运行此脚本将自动同步到 ima")
        return

    # 凭证已填写，调用 ima_sync
    print("\n[Step 4] 同步到 ima 个人笔记...")
    try:
        import importlib.util
        spec   = importlib.util.spec_from_file_location("ima_sync", ima_sync_path)
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        module.sync_all(force=False)
    except Exception as e:
        print(f"  [ima 同步] 执行出错：{e}")
        print("  可单独运行：python ima_sync.py")


if __name__ == "__main__":
    main()
