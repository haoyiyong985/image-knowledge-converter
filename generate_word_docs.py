#!/usr/bin/env python3
"""
生成 Word 文档
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_word_docs():
    """将所有 Markdown 文档转换为 Word"""
    
    result_dir = Path("D:/新建文件夹/处理结果")
    md_files = sorted(result_dir.glob("*.md"))
    
    print(f"[找到 {len(md_files)} 个 Markdown 文档]")
    
    for md_file in md_files:
        docx_file = md_file.with_suffix('.docx')
        
        # 读取 Markdown 内容
        content = md_file.read_text(encoding='utf-8')
        
        # 创建 Word 文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(md_file.stem, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 简单转换（按段落添加）
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # 标题处理
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('> '):
                # 引用
                p = doc.add_paragraph(line[2:])
                p.style = 'Quote'
            elif line.startswith('|') and line.endswith('|'):
                # 表格行（简化处理）
                continue
            elif line.startswith('---'):
                # 分隔线
                doc.add_paragraph('─' * 40)
            else:
                # 普通段落
                doc.add_paragraph(line)
        
        # 保存
        doc.save(str(docx_file))
        print(f"  [OK] {docx_file.name}")
    
    print(f"\n[OK] Word 文档生成完成！")

if __name__ == "__main__":
    generate_word_docs()
