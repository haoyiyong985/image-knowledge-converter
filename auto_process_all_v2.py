#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
全自动化图片处理工具 V3
==============================
优化功能：
  1. AI智能命名：根据内容自动生成有意义的文档名
  2. 主题分文件夹：同类内容自动归类到同一文件夹
  3. 重复检测：检测重复内容，避免重复存储
  4. 文档合并：同一主题的内容自动合并到同一文档
  5. 同时生成MD和Word文档
  6. IMA智能同步

处理流程：
  1. OCR识别（优先级：腾讯云 → 百度 → 本地Tesseract）
  2. AI智能分析内容主题 + 生成文档名
  3. 按主题分组到文件夹（动态创建新类别）
  4. 检测重复内容
  5. 同一主题文档合并
  6. 生成Markdown文档
  7. 生成Word文档
  8. 同步到IMA

使用方法：
  python auto_process_all_v2.py
"""

import os
import sys
import io
import json
import time
import re
import hashlib
import logging
from pathlib import Path
from datetime import datetime
from collections import Counter, defaultdict

# 修复 Windows 控制台编码
if sys.platform == 'win32':
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
    except Exception:
        pass

# 添加项目根目录到路径
sys.path.insert(0, '.')

# 加载 .env 环境变量配置
from dotenv import load_dotenv
env_path = Path('.') / '.env'
if env_path.exists():
    load_dotenv(dotenv_path=env_path)
else:
    print("[WARN] .env 文件不存在，使用系统环境变量")

# 导入各模块
from local_ocr import LocalOCR
from scripts.classifier_engine import ClassifierEngine

# 尝试导入云端OCR
try:
    from tencent_ocr import TencentOCR
    TENINCENT_AVAILABLE = True
except ImportError:
    TENINCENT_AVAILABLE = False

try:
    from baidu_ocr import BaiduOCR
    BAIDU_AVAILABLE = True
except ImportError:
    BAIDU_AVAILABLE = False

# 日志配置
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(message)s',
    handlers=[
        logging.FileHandler('处理结果/process.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# ============================================================
# 主题分类配置 - 更精准的关键词
# ============================================================
THEME_KEYWORDS = {
    "中医养生": ["中医", "养生", "药膳", "食疗", "中药", "经络", "穴位", "调理", "体质", "气血", "肝", "肾", "脾", "胃", "湿气", "虚", "寒", "热", "上火", "补", "滋阴", "温阳", "祛湿", "健脾", "润肺", "养心", "补肾"],
    "健康饮食": ["饮食", "营养", "食物", "蔬菜", "水果", "蛋白质", "碳水", "健康", "抗炎", "抗氧化", "免疫力", "肠道", "益生元", "膳食纤维", "维生素", "矿物质", "有机", "纯天然", "少油", "少盐", "清淡"],
    "疾病防治": ["疾病", "预防", "治疗", "症状", "指标", "血压", "血糖", "血脂", "胆固醇", "尿酸", "脂肪肝", "结节", "囊肿", "肿瘤", "癌症", "慢性病", "并发症", "吃药", "服药", "手术", "复查"],
    "生活方式": ["运动", "睡眠", "压力", "情绪", "心理健康", "作息", "习惯", "减肥", "增重", "美容", "护肤", "运动", "跑步", "走路", "瑜伽", "冥想", "放松"],
    "营养科普": ["科普", "知识", "研究", "发现", "实验", "数据", "结论", "专家", "建议", "指南", "推荐", "科学", "原理", "机制", "分析", "解读"],
    "旅游攻略": ["旅游", "旅行", "自驾", "景点", "美食", "打卡", "必去", "攻略", "路线", "酒店", "民宿", "拍照", "拍照"],
    "历史文化": ["历史", "文化", "典故", "人物", "朝代", "皇帝", "诗词", "文物", "博物馆", "遗址", "传承"]
}

# 主题特征词（更精准的识别）
TOPIC_SIGNATURES = {
    "养生方": ["方子", "秘方", "配方", "做法", "煮", "熬", "炖"],
    "食疗食谱": ["食谱", "菜谱", "做法", "食材", "配料", "烹饪"],
    "禁忌注意": ["禁忌", "注意", "不宜", "少吃", "避免", "千万", "不能", "不可"],
    "健康功效": ["功效", "作用", "好处", "益处", "价值", "可以", "有助于"],
    "疾病症状": ["症状", "表现", "特征", "体征", "不舒服", "疼痛", "难受"]
}


class ContentOrganizer:
    """内容整理器 - 对OCR识别的原始内容进行整理和优化"""

    def __init__(self):
        # 来源关键词
        self.source_keywords = [
            '来源', '摘录', '转载', '出处', '作者',
            '小红书', '微信', '抖音', '微博', 'B站',
            '公众号', '视频号', '知乎', '豆瓣', '百度'
        ]

        # 标题关键词
        self.heading_keywords = [
            '功效', '作用', '好处', '适合', '适宜', '适用',
            '原料', '食材', '配料', '材料',
            '做法', '制作', '步骤', '方法', '教程',
            '禁忌', '注意', '提醒', '警告', '不宜',
            '简介', '介绍', '概述', '说明',
            '推荐', '介绍', '建议', '指南'
        ]

    def extract_source(self, text):
        """从内容中提取来源信息"""
        lines = text.split('\n')
        source_info = {}

        for line in lines[:20]:  # 只在前20行查找
            line = line.strip()

            # 查找明确的来源标注
            for kw in ['来源', '出处', '作者', '摘录自']:
                if kw in line:
                    # 提取冒号后的内容
                    parts = line.split(kw)
                    if len(parts) > 1:
                        source = parts[1].strip().lstrip('：:').strip()
                        if source:
                            source_info['platform'] = source
                            break

            # 查找平台名称
            for platform in ['小红书', '微信读书', '微信视频号', '微信', '抖音', 'B站', 'BTV']:
                if platform in line:
                    source_info['platform'] = platform
                    break

            # 查找作者
            if '作者' in line:
                match = re.search(r'作者[：:]\s*([^\n]+)', line)
                if match:
                    source_info['author'] = match.group(1).strip()

        return source_info

    def clean_text(self, text):
        """清理OCR识别的原始文本"""
        lines = text.split('\n')
        cleaned_lines = []
        prev_line_empty = False

        for line in lines:
            line = line.strip()

            # 跳过极短的行（可能是OCR错误）
            if len(line) < 2 and line:
                continue

            # 跳过纯符号行
            if re.match(r'^[_\-=]{3,}$', line):
                continue

            # 跳过乱码行
            if re.match(r'^[\u4e00-\u9fa5]*[a-zA-Z0-9]{5,}[\u4e00-\u9fa5]*$', line):
                # 保留，但跳过纯英文行
                if not re.search(r'[\u4e00-\u9fa5]', line):
                    continue

            # 处理多余空行
            if not line:
                if not prev_line_empty:
                    cleaned_lines.append('')
                    prev_line_empty = True
                continue

            prev_line_empty = False

            # 修复常见的OCR错误
            line = self._fix_common_ocr_errors(line)
            cleaned_lines.append(line)

        # 合并结果
        result = '\n'.join(cleaned_lines)

        # 清理多余空行
        result = re.sub(r'\n{3,}', '\n\n', result)

        return result.strip()

    def _fix_common_ocr_errors(self, line):
        """修复常见的OCR错误"""
        # 修复特殊引号
        line = line.replace(chr(8220), '"')  # "
        line = line.replace(chr(8221), '"')  # "
        line = line.replace(chr(8216), "'")  # '
        line = line.replace(chr(8217), "'")  # '
        line = line.replace(chr(8218), "'")  # '
        line = line.replace(chr(8222), '"')  # "
        line = line.replace(chr(8223), '"')  # "
        line = line.replace(chr(8242), "'")  # '
        line = line.replace(chr(8243), "'")  # '
        line = line.replace(chr(180), "")    # '
        line = line.replace(chr(9032), "")   # ·
        line = line.replace(chr(8212), "-")  # —
        line = line.replace(chr(183), "-")   # ·
        line = line.replace("`", "")

        # 修复常见错误
        line = line.replace('V', '✓').replace('X', '✗')

        # 修复多余的空格
        line = re.sub(r' {2,}', ' ', line)

        return line

    def structure_content(self, text, source_info=None):
        """对内容进行结构化处理"""
        lines = text.split('\n')
        structured = []
        current_section = None
        section_content = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 检测是否为标题行
            is_heading = self._is_heading_line(line)

            if is_heading and current_section is None:
                # 第一级标题
                structured.append(f"### {line}")
                current_section = line
            elif is_heading:
                # 保存上一个章节内容
                if section_content:
                    structured.extend(section_content)
                    structured.append('')
                structured.append(f"#### {line}")
                current_section = line
                section_content = []
            elif current_section:
                # 章节内容
                section_content.append(line)
            else:
                # 没有标题的内容，直接添加
                structured.append(line)

        # 添加最后的章节内容
        if section_content:
            structured.extend(section_content)

        return '\n'.join(structured)

    def _is_heading_line(self, line):
        """判断是否为标题行"""
        # 太短的行不可能是标题
        if len(line) < 4:
            return False

        # 检查是否包含标题关键词
        for kw in self.heading_keywords:
            if kw in line and len(line) <= 30:
                return True

        # 检查格式特征
        # 1. 以特定符号开头
        if re.match(r'^[一二三四五六七八九十\d]+[、.．:：]', line):
            return True

        # 2. 特定格式
        if re.match(r'^[◆★▼▪▸●○]', line) and len(line) <= 25:
            return True

        return False

    def organize(self, raw_text, theme=None):
        """综合整理内容"""
        # 1. 提取来源信息
        source_info = self.extract_source(raw_text)

        # 2. 清理文本
        cleaned_text = self.clean_text(raw_text)

        # 3. 结构化内容
        structured_text = self.structure_content(cleaned_text, source_info)

        # 4. 组装最终结果
        result = {
            'content': structured_text,
            'source': source_info,
            'theme': theme
        }

        return result


class ContentAnalyzer:
    """AI内容分析器 - 自动识别主题并生成文档名"""

    def __init__(self):
        self.cache = {}
        self.organizer = ContentOrganizer()

    def extract_title(self, text):
        """从内容中提取标题"""
        lines = text.split('\n')
        for line in lines[:10]:
            line = line.strip()
            # 查找 # 标题
            if line.startswith('#'):
                title = line.lstrip('#').strip()
                if len(title) >= 2 and len(title) <= 30:
                    return title
            # 查找特殊符号标记的标题
            if re.match(r'^[◆★▼▪▸]+\s*.+', line):
                title = re.sub(r'^[◆★▼▪▸]+\s*', '', line)
                if len(title) >= 2:
                    return title[:20]
            # 查找居中或大标题模式
            if re.match(r'^[\u4e00-\u9fa5]{4,15}$', line) and len(line) >= 4:
                return line

        # 查找第一个有意义的短句
        for line in lines[:15]:
            line = line.strip()
            if len(line) >= 4 and len(line) <= 20 and re.match(r'^[\u4e00-\u9fa5]', line):
                if not line.startswith('#') and 'http' not in line.lower():
                    return line
        return None

    def detect_theme(self, text):
        """检测内容主题，返回主题文件夹名"""
        max_score = 0
        theme_name = "综合知识"

        for theme, keywords in THEME_KEYWORDS.items():
            score = sum(3 if kw in text else 0 for kw in keywords)
            # 计算关键词密度
            if score > 0:
                density = score / max(len(text), 1) * 1000
                if density > max_score:
                    max_score = density
                    theme_name = theme

        return theme_name

    def detect_topic_type(self, text):
        """检测内容类型（方子、食谱、禁忌等）"""
        for topic, signatures in TOPIC_SIGNATURES.items():
            score = sum(2 if sig in text else 0 for sig in signatures)
            if score >= 2:
                return topic
        return "知识整理"

    def generate_doc_name(self, text, category, index):
        """生成有意义的文档名"""
        # 1. 尝试从内容中提取标题
        title = self.extract_title(text)
        if title:
            title = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '', title)
            if len(title) >= 3:
                return title[:15]

        # 2. 根据内容类型命名
        topic_type = self.detect_topic_type(text)
        if topic_type != "知识整理":
            return f"{topic_type}"

        # 3. 根据分类和关键词生成
        if category and category != "综合知识":
            return f"{category}_整理"

        # 4. 默认命名
        timestamp = datetime.now().strftime('%m%d')
        return f"知识整理_{timestamp}_{index:02d}"

    def compute_content_hash(self, text):
        """计算内容哈希，用于检测重复"""
        cleaned = re.sub(r'\s+', '', text)
        cleaned = cleaned[:500]
        return hashlib.md5(cleaned.encode('utf-8')).hexdigest()[:8]

    def analyze(self, text, category, index):
        """综合分析，返回分析结果"""
        cache_key = self.compute_content_hash(text)

        if cache_key in self.cache:
            return self.cache[cache_key]

        theme = self.detect_theme(text)
        topic_type = self.detect_topic_type(text)
        doc_name = self.generate_doc_name(text, category, index)

        # 整理内容
        organized = self.organizer.organize(text, theme)
        organized_content = organized['content']
        source_info = organized['source']

        result = {
            'theme': theme,
            'topic_type': topic_type,
            'doc_name': doc_name,
            'content_hash': cache_key,
            'organized_content': organized_content,  # 整理后的内容
            'source': source_info  # 来源信息
        }

        self.cache[cache_key] = result
        return result


class MultiEngineOCR:
    """多引擎OCR管理器"""

    def __init__(self):
        self.current_engine = None
        self.engine_status = []

        secret_id = os.getenv('TENCENT_SECRET_ID', '')
        secret_key = os.getenv('TENCENT_SECRET_KEY', '')
        if secret_id and secret_key and '替换' not in secret_id and TENINCENT_AVAILABLE:
            self.engine_status.append(('腾讯云', True, '已配置'))
        else:
            self.engine_status.append(('腾讯云', False, '未配置'))

        app_id = os.getenv('BAIDU_APP_ID', '')
        api_key = os.getenv('BAIDU_API_KEY', '')
        secret_key = os.getenv('BAIDU_SECRET_KEY', '')
        if app_id and api_key and secret_key and '替换' not in app_id and BAIDU_AVAILABLE:
            self.engine_status.append(('百度云', True, '已配置'))
        else:
            self.engine_status.append(('百度云', False, '未配置'))

        local = LocalOCR()
        if local.tesseract_available:
            self.engine_status.append(('本地Tesseract', True, '可用'))
        else:
            self.engine_status.append(('本地Tesseract', False, getattr(local, 'error_message', '未知错误')))

        self._select_best_engine()

    def _select_best_engine(self):
        for name, available, msg in self.engine_status:
            if available:
                self.current_engine = name
                logger.info(f"[OCR] 选用引擎: {name} ({msg})")
                return True
        logger.warning("[OCR] 没有可用的OCR引擎!")
        return False

    def recognize(self, image_path):
        last_error = None
        for name, available, _ in self.engine_status:
            if not available:
                continue
            try:
                logger.info(f"[OCR] 尝试 {name}...")
                if name == '腾讯云' and TENINCENT_AVAILABLE:
                    engine = TencentOCR()
                elif name == '百度云' and BAIDU_AVAILABLE:
                    engine = BaiduOCR()
                else:
                    engine = LocalOCR()

                result = engine.recognize(image_path)
                if result and result.get('success'):
                    logger.info(f"[OCR] {name} 识别成功!")
                    return result
                else:
                    last_error = result.get('error', '未知错误') if result else '返回为空'
            except Exception as e:
                last_error = str(e)
        return {'success': False, 'error': last_error or '无可用OCR引擎', 'text': ''}


class DocumentMerger:
    """文档合并器 - 同一主题内容合并"""

    def __init__(self, output_dir='处理结果'):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        # 记录主题文档
        self.theme_docs = defaultdict(list)
        self._scan_existing_docs()

    def _scan_existing_docs(self):
        """扫描已存在的文档"""
        for theme_folder in self.output_dir.iterdir():
            if theme_folder.is_dir() and not theme_folder.name.startswith('.'):
                for md_file in theme_folder.glob('*.md'):
                    # 提取主题关键词
                    topic = self._extract_topic(md_file)
                    self.theme_docs[theme_folder.name].append({
                        'file': md_file,
                        'topic': topic
                    })

    def _extract_topic(self, md_file):
        """从文档名提取主题"""
        name = md_file.stem
        # 移除序号和后缀
        name = re.sub(r'^\d+_', '', name)
        name = re.sub(r'_\d+$', '', name)
        name = re.sub(r'_[a-f0-9]{8}$', '', name)
        return name

    def find_similar_doc(self, theme, doc_name, content):
        """查找相似的已有文档"""
        # 提取文档主题
        doc_topic = re.sub(r'^\d+_', '', doc_name)
        doc_topic = re.sub(r'_\d+$', '', doc_topic)

        # 在同主题文件夹中查找相似文档
        for existing_doc in self.theme_docs.get(theme, []):
            existing_topic = existing_doc['topic']
            # 主题相似度高
            if self._calc_similarity(doc_topic, existing_topic) > 0.6:
                return existing_doc['file']
        return None

    def _calc_similarity(self, s1, s2):
        """简单相似度计算"""
        s1_set = set(s1)
        s2_set = set(s2)
        if not s1_set or not s2_set:
            return 0
        return len(s1_set & s2_set) / len(s1_set | s2_set)

    def merge_content(self, existing_file, new_content, new_image_name, content_hash=None):
        """合并内容到已有文档"""
        try:
            with open(existing_file, 'r', encoding='utf-8') as f:
                existing = f.read()

            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
            append_section = f"""

---

## 📌 补充内容

> 来源图片: {new_image_name}
> 追加时间: {timestamp}
> 内容标识: {content_hash}

{new_content}

"""

            # 在自动生成标记前插入
            new_content_merged = existing.replace(
                '\n\n---\n\n*本文档由图片知识库整理工具自动生成*',
                append_section + '\n\n---\n\n*本文档由图片知识库整理工具自动生成*'
            )

            with open(existing_file, 'w', encoding='utf-8') as f:
                f.write(new_content_merged)

            return True
        except Exception as e:
            logger.warning(f"[合并] 失败: {e}")
            return False


class SmartDocumentGenerator:
    """智能文档生成器 - 支持主题分组和文档合并"""

    def __init__(self, output_dir='处理结果'):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.merger = DocumentMerger(output_dir)
        self.existing_hashes = set()
        self._scan_hashes()
        # 追踪已生成的文档（避免重复生成Word）
        self.generated_docs = set()  # 格式: theme/safe_name

    def _scan_hashes(self):
        """扫描已有文档的hash"""
        for md_file in self.output_dir.rglob('*.md'):
            try:
                file_content = md_file.read_text(encoding='utf-8')
                hash_match = re.search(r'content_hash:\s*([a-f0-9]{8})', file_content)
                if hash_match:
                    self.existing_hashes.add(hash_match.group(1))
            except:
                pass

    def _get_doc_key(self, theme, doc_name):
        """获取文档唯一标识"""
        safe_name = re.sub(r'[<>:"/\|?*]', '', doc_name)
        if len(safe_name) > 25:
            safe_name = safe_name[:25]
        return f"{theme}/{safe_name}"

    def is_duplicate(self, content_hash):
        """检测是否重复"""
        return content_hash in self.existing_hashes

    def is_doc_generated(self, theme, doc_name):
        """检查文档是否已生成"""
        return self._get_doc_key(theme, doc_name) in self.generated_docs

    def mark_doc_generated(self, theme, doc_name):
        """标记文档已生成"""
        self.generated_docs.add(self._get_doc_key(theme, doc_name))

    def generate_markdown(self, text, theme, doc_name, category, image_name, keywords=None, content_hash=None, source_info=None):
        """生成Markdown文档"""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')

        # 查找是否有相似文档可以合并
        similar_doc = self.merger.find_similar_doc(theme, doc_name, text)
        if similar_doc:
            print(f"  → 发现相似文档 {similar_doc.name}，合并内容...")
            if self.merger.merge_content(similar_doc, text, image_name, content_hash):
                # 标记文档已处理（避免重复生成Word）
                self.mark_doc_generated(theme, doc_name)
                return str(similar_doc)

        # 创建新文档
        keywords_str = ', '.join(keywords) if keywords else '无'

        # 来源信息
        source_section = ""
        if source_info:
            if source_info.get('platform'):
                source_section += f"**来源**：{source_info['platform']}\n\n"
            if source_info.get('author'):
                source_section += f"**作者**：{source_info['author']}\n\n"

        md_content = f"""# {doc_name}

> 来源图片: {image_name}
> 识别时间: {timestamp}
> 主题分类: {theme}
> 原分类: {category}
> 关键词: {keywords_str}
> content_hash: {content_hash or 'N/A'}

---



{source_section}{text}

---

*本文档由图片知识库整理工具自动生成*
"""

        # 创建主题文件夹
        theme_folder = self.output_dir / theme
        theme_folder.mkdir(exist_ok=True)

        # 生成安全的文件名
        safe_name = re.sub(r'[<>:"/\|?*]', '', doc_name)
        if len(safe_name) > 25:
            safe_name = safe_name[:25]

        # 检查是否已存在
        md_file = theme_folder / f"{safe_name}.md"
        counter = 1
        while md_file.exists():
            md_file = theme_folder / f"{safe_name}_{counter}.md"
            counter += 1

        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(md_content)

        # 更新hash记录
        if content_hash:
            self.existing_hashes.add(content_hash)
        # 更新合并器记录
        self.merger.theme_docs[theme].append({
            'file': md_file,
            'topic': doc_name
        })
        # 标记文档已生成
        self.mark_doc_generated(theme, doc_name)

        logger.info(f"[文档] Markdown已生成: {theme}/{md_file.name}")
        return str(md_file)

    def generate_word(self, text, theme, doc_name, image_name, keywords=None, content_hash=None, existing_md_file=None):
        """生成Word文档，如果已有同名文档则跳过"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.oxml.ns import qn
        except ImportError:
            logger.warning("[文档] python-docx未安装，跳过Word生成")
            return None

        # 检查是否已生成过Word文档（但文件不存在的情况仍需生成）
        if self.is_doc_generated(theme, doc_name):
            safe_name = re.sub(r'[<>:"/\|?*]', '', doc_name)
            if len(safe_name) > 25:
                safe_name = safe_name[:25]
            existing_docx = self.output_dir / theme / f"{safe_name}.docx"
            if existing_docx.exists():
                print(f"  → Word文档已存在，跳过")
                return str(existing_docx)
            # 检查带编号的版本
            for i in range(1, 100):
                existing_docx = self.output_dir / theme / f"{safe_name}_{i}.docx"
                if existing_docx.exists():
                    print(f"  → Word文档已存在，跳过")
                    return str(existing_docx)
            # 文件不存在，仍需生成（可能是同名文档被合并后重新生成）

        # 保存到同一主题文件夹
        theme_folder = self.output_dir / theme
        theme_folder.mkdir(exist_ok=True)

        safe_name = re.sub(r'[<>:"/\|?*]', '', doc_name)
        if len(safe_name) > 25:
            safe_name = safe_name[:25]

        docx_file = theme_folder / f"{safe_name}.docx"
        counter = 1
        while docx_file.exists():
            docx_file = theme_folder / f"{safe_name}_{counter}.docx"
            counter += 1

        # 如果有对应的Markdown文件已合并内容，则同步Word
        if existing_md_file and Path(existing_md_file).exists():
            try:
                md_content = Path(existing_md_file).read_text(encoding='utf-8')
                main_content = md_content.split('## 内容\n\n')[1] if '## 内容\n\n' in md_content else ''
                main_content = main_content.split('## 📌 补充内容')[0] if '## 📌 补充内容' in main_content else main_content
                main_content = main_content.replace('\n\n---\n\n*本文档由图片知识库整理工具自动生成*', '')
                text = main_content.strip()
            except Exception as e:
                logger.warning(f"[Word] 读取合并内容失败: {e}")

        # 创建Word文档
        doc = Document()
        doc.styles['Normal'].font.name = 'Microsoft YaHei'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        doc.add_heading(doc_name, 0)
        doc.add_paragraph(f"来源图片: {image_name}")
        doc.add_paragraph(f"识别时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"主题分类: {theme}")
        keywords_str = ', '.join(keywords) if keywords else '无'
        doc.add_paragraph(f"关键词: {keywords_str}")
        doc.add_paragraph("─" * 30)
        doc.add_paragraph(text)

        doc.save(str(docx_file))
        # 标记文档已生成
        self.mark_doc_generated(theme, doc_name)
        logger.info(f"[文档] Word已生成: {theme}/{docx_file.name}")
        return str(docx_file)

class IMASyncer:
    """IMA笔记同步器"""

    def __init__(self):
        self.client_id = os.getenv('IMA_OPENAPI_CLIENTID', '')
        self.api_key = os.getenv('IMA_OPENAPI_APIKEY', '')
        self.base_url = 'https://ima.qq.com/openapi/note/v1'
        self.enabled = bool(self.client_id and self.api_key and '填入' not in self.client_id)
        self.sync_log_file = Path('处理结果/ima_sync_log.json')
        self.sync_log = self._load_sync_log()
        self.rate_limited = False  # 标记是否被限流

        if self.enabled:
            logger.info("[IMA] 已配置")
        else:
            logger.warning("[IMA] 未配置")

    def _load_sync_log(self):
        if self.sync_log_file.exists():
            try:
                return json.loads(self.sync_log_file.read_text(encoding='utf-8'))
            except:
                pass
        return {}

    def _save_sync_log(self):
        self.sync_log_file.write_text(
            json.dumps(self.sync_log, ensure_ascii=False, indent=2),
            encoding='utf-8'
        )

    def _api_call(self, endpoint, payload, retries=3):
        """调用IMA API，支持重试"""
        if not self.enabled or self.rate_limited:
            return None

        for attempt in range(retries):
            try:
                import requests
                headers = {
                    'ima-openapi-clientid': self.client_id,
                    'ima-openapi-apikey': self.api_key,
                    'Content-Type': 'application/json'
                }
                url = f"{self.base_url}/{endpoint}"
                response = requests.post(url, json=payload, headers=headers, timeout=30)

                if response.status_code == 200:
                    return response.json()
                elif response.status_code == 403:
                    # 检查是否是限流
                    result = response.json()
                    if '请求超量' in result.get('msg', ''):
                        logger.warning("[IMA] API请求超量，已触发限流，明天再试")
                        self.rate_limited = True
                        return None
                return None
            except Exception as e:
                logger.warning(f"[IMA] API调用失败 ({attempt+1}/{retries}): {e}")
                if attempt < retries - 1:
                    time.sleep(2)
        return None

    def sync_note(self, title, content, theme=None, content_hash=None, doc_path=None):
        """同步笔记到IMA"""
        if not self.enabled:
            return None

        if self.rate_limited:
            return None

        # 检查是否已同步过
        doc_key = content_hash or title
        existing = self.sync_log.get(doc_key, {})

        # 格式化内容
        full_content = f"# {title}\n\n"
        if theme:
            full_content += f"> 主题: {theme}\n"
        if content_hash:
            full_content += f"> 内容标识: {content_hash}\n"
        full_content += f"\n{content}\n\n---\n*自动同步自图片知识库*\n"

        # 如果有本地文档文件，可以更新
        if doc_path and existing.get('doc_id'):
            # 追加模式
            result = self._api_call('append_doc', {
                'doc_id': existing['doc_id'],
                'content_format': 1,
                'content': f"\n\n---\n*更新于 {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n" + content
            })
            if result:
                self.sync_log[doc_key] = {
                    'doc_id': existing['doc_id'],
                    'last_sync': datetime.now().isoformat(),
                    'update_count': existing.get('update_count', 0) + 1
                }
                self._save_sync_log()
                return existing['doc_id']

        # 新建笔记
        result = self._api_call('import_doc', {
            'content_format': 1,
            'content': full_content
        })

        if result and result.get('code') == 0:
            doc_id = result.get('data', {}).get('note_id', '')
            if doc_id:
                self.sync_log[doc_key] = {
                    'doc_id': doc_id,
                    'first_sync': datetime.now().isoformat(),
                    'last_sync': datetime.now().isoformat(),
                    'theme': theme,
                    'update_count': 1
                }
                self._save_sync_log()
                logger.info(f"[IMA] 笔记已同步: {title[:20]}")
                return doc_id
            else:
                return "imported"

        return None


def process_single_image(ocr, classifier, content_analyzer, doc_gen, ima_syncer, image_path, index, total):
    """处理单张图片"""
    image_name = Path(image_path).name
    print(f"\n{'='*50}")
    print(f"[{index}/{total}] 处理: {image_name}")
    print(f"{'='*50}")

    # 1. OCR识别
    print("\n📷 OCR识别中...")
    result = ocr.recognize(image_path)

    if not result.get('success'):
        print(f"  ⚠️ OCR失败: {result.get('error')}")
        return None

    text = result.get('text', '').strip()
    if not text or len(text) < 10:
        print("  ⚠️ 无文字内容或内容过少")
        return None

    print(f"  ✓ 识别到 {len(text)} 个字符")

    # 2. AI分类
    print("\n🏷️ 内容分类中...")
    categories = classifier.classify(text, threshold=0.25)

    if categories:
        primary = categories[0]
        category_name = primary.category_name
        keywords = primary.matched_keywords if primary.matched_keywords else []
        confidence = primary.confidence
        print(f"  ✓ 分类: {category_name} (置信度 {confidence:.0%})")
        if keywords:
            print(f"     关键词: {', '.join(keywords[:5])}")
    else:
        category_name = "综合知识"
        keywords = []
        print("  ⚠️ 未匹配到分类，使用默认分类")

    # 3. AI内容分析（包括内容整理）
    print("\n🧠 AI内容分析中...")
    analysis = content_analyzer.analyze(text, category_name, index)
    theme = analysis['theme']
    topic_type = analysis['topic_type']
    doc_name = analysis['doc_name']
    content_hash = analysis['content_hash']
    organized_content = analysis.get('organized_content', text)  # 获取整理后的内容
    source_info = analysis.get('source', {})  # 获取来源信息
    print(f"  ✓ 主题分类: {theme}")
    print(f"  ✓ 内容类型: {topic_type}")
    print(f"  ✓ 文档命名: {doc_name}")
    if source_info.get('platform'):
        print(f"  ✓ 来源: {source_info['platform']}")

    # 4. 检测重复
    print("\n🔍 检测重复内容...")
    if doc_gen.is_duplicate(content_hash):
        print("  ⚠️ 检测到重复内容，跳过存储")
        return {
            'image': image_name,
            'text_length': len(text),
            'category': category_name,
            'theme': theme,
            'doc_name': doc_name,
            'is_duplicate': True,
            'content_hash': content_hash
        }
    print("  ✓ 新内容")

    # 5. 生成Markdown（使用整理后的内容）
    print("\n📝 生成Markdown...")
    md_file = doc_gen.generate_markdown(
        organized_content, theme, doc_name, category_name,
        image_name, keywords, content_hash, source_info
    )

    # 6. 生成Word（同步Markdown的合并内容）
    print("\n📄 生成Word...")
    docx_file = doc_gen.generate_word(
        organized_content, theme, doc_name, image_name, keywords, content_hash, existing_md_file=md_file
    )
    if docx_file:
        rel_path = Path(docx_file).relative_to(doc_gen.output_dir)
        print(f"  ✓ 已保存到: {rel_path}")
    else:
        print("  ⚠️ Word文档生成失败或未安装python-docx")

    # 7. 同步到IMA
    print("\n☁️ 同步到IMA...")
    if ima_syncer.rate_limited:
        print("  ⚠️ IMA被限流，请明天再试")
    else:
        ima_id = ima_syncer.sync_note(doc_name, organized_content, theme, content_hash, md_file)
        if ima_id:
            print(f"  ✓ IMA同步成功")
        else:
            print("  ⚠️ IMA同步失败")

    return {
        'image': image_name,
        'text_length': len(text),
        'category': category_name,
        'theme': theme,
        'topic_type': topic_type,
        'doc_name': doc_name,
        'keywords': keywords,
        'is_duplicate': False,
        'md_file': md_file,
        'docx_file': docx_file,
        'content_hash': content_hash,
        'source': source_info
    }


def main():
    print("=" * 60)
    print("全自动化图片处理工具 V3")
    print("OCR → AI分类 → 主题分组 → 文档合并 → MD+Word → IMA")
    print("=" * 60)

    # 初始化组件
    print("\n[1/6] 初始化OCR引擎...")
    ocr = MultiEngineOCR()
    if not ocr.current_engine:
        print("❌ 没有可用的OCR引擎!")
        return

    print("\n[2/6] 初始化分类器...")
    classifier = ClassifierEngine()

    print("\n[3/6] 初始化内容分析器...")
    content_analyzer = ContentAnalyzer()

    print("\n[4/6] 初始化文档生成器...")
    doc_gen = SmartDocumentGenerator()

    print("\n[5/6] 初始化IMA同步器...")
    ima_syncer = IMASyncer()
    if ima_syncer.rate_limited:
        print("  ⚠️ IMA被限流，本次将跳过IMA同步")

    print("\n[6/6] 扫描待处理图片...")

    # 扫描待处理图片
    source_dir = Path('待处理图片')
    if not source_dir.exists():
        source_dir = Path('待处理图片/历史')

    images = list(source_dir.rglob('*.jpg')) + list(source_dir.rglob('*.png')) + list(source_dir.rglob('*.jpeg'))

    if not images:
        print("❌ 未找到待处理图片")
        return

    total = len(images)
    print(f"✓ 找到 {total} 张图片\n")

    # 批量处理
    results = []
    start_time = time.time()

    for i, img_path in enumerate(images, 1):
        try:
            result = process_single_image(
                ocr, classifier, content_analyzer,
                doc_gen, ima_syncer,
                str(img_path), i, total
            )
            if result:
                results.append(result)
        except Exception as e:
            print(f"  ✗ 处理失败: {e}")
            logger.error(f"处理失败 {img_path}: {e}")

        if i < total:
            time.sleep(0.5)

    # 统计
    elapsed = time.time() - start_time

    print("\n" + "=" * 60)
    print("处理完成!")
    print("=" * 60)
    print(f"总计处理: {len(results)}/{total} 张")
    print(f"总耗时: {elapsed:.1f} 秒")

    # 主题分布
    if results:
        print("\n📊 主题分布:")
        themes = Counter(r['theme'] for r in results if not r.get('is_duplicate'))
        for theme, count in themes.most_common():
            print(f"  - {theme}: {count} 张")

    # 重复检测
    dup_count = sum(1 for r in results if r.get('is_duplicate'))
    if dup_count > 0:
        print(f"\n🔍 重复检测: 跳过 {dup_count} 条重复内容")

    # 输出统计
    md_count = len([r for r in results if r.get('md_file')])
    docx_count = len([r for r in results if r.get('docx_file')])
    print(f"\n📁 输出统计:")
    print(f"  - Markdown: {md_count} 个")
    print(f"  - Word: {docx_count} 个")

    if ima_syncer.rate_limited:
        print(f"\n☁️ IMA: 被限流，请明天再试同步")

    print("\n" + "=" * 60)

    # 保存报告
    report_file = Path('处理结果/处理报告.json')
    report = {
        'timestamp': datetime.now().isoformat(),
        'total_images': total,
        'processed': len(results),
        'new_documents': md_count,
        'word_documents': docx_count,
        'duplicates': dup_count,
        'elapsed_seconds': elapsed,
        'results': [{k: v for k, v in r.items() if k not in ['md_file', 'docx_file']} for r in results],
        'theme_stats': dict(Counter(r['theme'] for r in results if not r.get('is_duplicate')))
    }
    with open(report_file, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    print(f"\n📊 报告已保存: {report_file}")


if __name__ == '__main__':
    main()
