#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
全自动化图片处理工具 V6.0
==============================
整合优化：
  ✅ V2功能（恢复）：
    1. 智能分类器 ClassifierEngine（支持置信度）
    2. AI内容分析器 ContentAnalyzer（提取标题、来源、主题类型）
    3. Word文档生成（MD+Word双格式）
    4. 文档合并（同一主题合并）
  ✅ V4优化（保留）：
    1. 实时进度条（带ETA）
    2. 图片归档（按主题分类）
    3. 丰富的主题分类体系
  ✅ V6新增（集成）：
    1. 智能分批处理（按文件大小自动分批）
    2. 断点续传（状态保存到JSON）
    3. 失败重试机制（自动重试3次）
    4. 批次进度追踪

处理流程：
  1. 分批初始化（扫描并智能分批）
  2. OCR识别（腾讯云 → 百度 → 本地Tesseract）
  3. 智能分类（ClassifierEngine + 置信度）
  4. AI内容分析（标题提取、来源识别、主题分类）
  5. 重复检测
  6. 文档生成（Markdown + Word）
  7. IMA同步
  8. 图片归档（按主题分类）

使用方法：
  python auto_process_all_v6.py           # 全自动处理所有图片
  python auto_process_all_v6.py --batch    # 分批模式（可中断续传）
  python auto_process_all_v6.py --init     # 仅初始化分批
  python auto_process_all_v6.py --progress # 查看进度
"""

import os
import sys
import io
import json
import time
import re
import hashlib
import shutil
import logging
from pathlib import Path
from datetime import datetime
from collections import Counter, defaultdict
from typing import Optional, List, Dict, Tuple
from dataclasses import dataclass, asdict

# 修复 Windows 控制台编码
if sys.platform == 'win32':
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
    except Exception:
        pass

# 添加项目根目录到路径
sys.path.insert(0, '.')

# 加载 .env 环境变量配置
from dotenv import load_dotenv
env_path = Path('.') / '.env'
if env_path.exists():
    load_dotenv(dotenv_path=env_path)
else:
    print("[WARN] .env 文件不存在，使用系统环境变量")

# 导入各模块
from local_ocr import LocalOCR
from scripts.classifier_engine import ClassifierEngine

# 尝试导入云端OCR
try:
    from tencent_ocr import TencentOCR
    TENINCENT_AVAILABLE = True
except ImportError:
    TENINCENT_AVAILABLE = False

try:
    from baidu_ocr import BaiduOCR
    BAIDU_AVAILABLE = True
except ImportError:
    BAIDU_AVAILABLE = False

# 日志配置
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(message)s',
    handlers=[
        logging.FileHandler('处理结果/process.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# P2-17: .env配置格式校验
def validate_env_config() -> Dict[str, bool]:
    """
    启动时校验API Key格式（长度、前缀）
    返回各配置的校验结果
    """
    results = {}
    
    # 腾讯云OCR校验
    tencent_secret_id = os.getenv('TENCENT_SECRET_ID', '')
    tencent_secret_key = os.getenv('TENCENT_SECRET_KEY', '')
    if tencent_secret_id and '填入' not in tencent_secret_id:
        # 腾讯云SecretId通常是16-32位字母数字
        results['tencent'] = 16 <= len(tencent_secret_id) <= 64
    else:
        results['tencent'] = False
    
    # 百度OCR校验
    baidu_api_key = os.getenv('BAIDU_API_KEY', '')
    if baidu_api_key and '填入' not in baidu_api_key:
        # 百度API Key通常是24位
        results['baidu'] = 20 <= len(baidu_api_key) <= 40
    else:
        results['baidu'] = False
    
    # IMA配置校验
    ima_client_id = os.getenv('IMA_OPENAPI_CLIENTID', '')
    ima_api_key = os.getenv('IMA_OPENAPI_APIKEY', '')
    if ima_client_id and '填入' not in ima_client_id:
        results['ima'] = len(ima_client_id) >= 16
    else:
        results['ima'] = False
    
    return results

# P2-12: 处理前预检
def preflight_check() -> Tuple[bool, List[str]]:
    """
    启动时检查OCR可用性、输出目录可写、python-docx已安装
    返回: (是否通过, 问题列表)
    """
    issues = []
    
    # 1. 检查python-docx
    try:
        import docx
    except ImportError:
        issues.append("python-docx未安装，Word文档生成将不可用")
    
    # 2. 检查输出目录可写
    output_dirs = ['处理结果', '已处理图片', 'logs', 'progress']
    for dir_name in output_dirs:
        dir_path = Path(dir_name)
        if dir_path.exists():
            if not os.access(dir_path, os.W_OK):
                issues.append(f"目录'{dir_name}'不可写")
        else:
            try:
                dir_path.mkdir(parents=True, exist_ok=True)
            except Exception as e:
                issues.append(f"无法创建目录'{dir_name}': {e}")
    
    # 3. 检查OCR引擎可用性
    ocr_engines = []
    if TENINCENT_AVAILABLE:
        secret_id = os.getenv('TENCENT_SECRET_ID', '')
        secret_key = os.getenv('TENCENT_SECRET_KEY', '')
        if secret_id and secret_key and '填入' not in secret_id:
            ocr_engines.append('腾讯云OCR')
    
    if BAIDU_AVAILABLE:
        api_key = os.getenv('BAIDU_API_KEY', '')
        secret_key = os.getenv('BAIDU_SECRET_KEY', '')
        if api_key and secret_key and '填入' not in api_key:
            ocr_engines.append('百度OCR')
    
    # 检查本地Tesseract
    try:
        local_ocr = LocalOCR()
        if local_ocr.is_available():
            ocr_engines.append('本地Tesseract')
    except:
        pass
    
    if not ocr_engines:
        issues.append("没有可用的OCR引擎，请配置腾讯云或百度OCR")
    
    # 4. P2-17: 校验API Key格式
    env_results = validate_env_config()
    if not env_results['tencent'] and not env_results['baidu']:
        issues.append("云端OCR的API Key格式可能不正确")
    
    passed = len(issues) == 0
    return passed, issues


# ============================================================
# 分批处理配置（从 batch_processor.py 集成）
# ============================================================
# 使用脚本所在目录的绝对路径，避免从其他目录运行脚本时定位错误
_BATCH_DIR = Path(__file__).parent / '.workbuddy' / 'memory'
BATCH_STATE_FILE = _BATCH_DIR / 'batch_state.json'

BATCH_CONFIG = {
    "small": {"max_size": 500 * 1024, "batch_size": 10},      # <500KB: 10张/批
    "medium": {"max_size": 2 * 1024 * 1024, "batch_size": 6},  # 500KB-2MB: 6张/批
    "large": {"max_size": 5 * 1024 * 1024, "batch_size": 4},   # 2MB-5MB: 4张/批
    "xlarge": {"max_size": float('inf'), "batch_size": 2}      # >5MB: 2张/批
}


# ============================================================
# 分批处理数据类（从 batch_processor.py 集成）
# ============================================================
@dataclass
class ImageInfo:
    """图片信息"""
    path: str
    name: str
    size: int
    folder: str
    hash: str = ""
    
    def to_dict(self):
        return asdict(self)
    
    @classmethod
    def from_dict(cls, data):
        return cls(**data)

@dataclass
class BatchInfo:
    """批次信息"""
    batch_id: str
    folder: str
    images: List[str]
    status: str  # pending/processing/completed/failed
    created_at: str
    started_at: Optional[str] = None
    completed_at: Optional[str] = None
    result: Optional[str] = None
    retry_count: int = 0
    
    def to_dict(self):
        return asdict(self)
    
    @classmethod
    def from_dict(cls, data):
        return cls(**data)

@dataclass
class ProcessingState:
    """处理状态"""
    session_id: str
    started_at: str
    total_images: int
    total_batches: int
    completed_batches: int = 0
    failed_batches: int = 0
    batches: Optional[Dict] = None
    
    def __post_init__(self):
        if self.batches is None:
            self.batches = {}
    
    def to_dict(self):
        data = asdict(self)
        if self.batches:
            data['batches'] = {k: v.to_dict() if isinstance(v, BatchInfo) else v 
                              for k, v in self.batches.items()}
        return data
    
    @classmethod
    def from_dict(cls, data):
        if 'batches' in data and data['batches']:
            data['batches'] = {k: BatchInfo.from_dict(v) for k, v in data['batches'].items()}
        return cls(**data)


# ============================================================
# 分批处理管理器（从 batch_processor.py 集成并优化）
# ============================================================
class BatchManager:
    """分批处理管理器"""
    
    def __init__(self, source_dir: str = '待处理图片'):
        self.source_dir = Path(source_dir)
        self.state: Optional[ProcessingState] = None
        
    def _load_state(self) -> Optional[ProcessingState]:
        """加载处理状态"""
        if BATCH_STATE_FILE.exists():
            try:
                with open(BATCH_STATE_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    return ProcessingState.from_dict(data)
            except Exception as e:
                print(f"[WARN] 加载状态失败: {e}")
        return None
    
    def _save_state(self):
        """保存处理状态"""
        try:
            BATCH_STATE_FILE.parent.mkdir(parents=True, exist_ok=True)
            with open(BATCH_STATE_FILE, 'w', encoding='utf-8') as f:
                json.dump(self.state.to_dict(), f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"[ERROR] 保存状态失败: {e}")
    
    def _calculate_hash(self, file_path: str) -> str:
        """计算文件MD5哈希"""
        try:
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()[:16]
        except Exception:
            return ""
    
    def _get_batch_size_limit(self, image_path: str) -> Tuple[int, str]:
        """根据图片大小获取批次大小"""
        try:
            size = os.path.getsize(image_path)
            for category, config in BATCH_CONFIG.items():
                if size < config['max_size']:
                    return config['batch_size'], category
        except Exception:
            pass
        return 5, "medium"
    
    def scan_images(self) -> List[ImageInfo]:
        """扫描所有待处理图片（含根目录和所有子文件夹）"""
        images = []
        seen_paths = set()
        
        if not self.source_dir.exists():
            return images
        
        # 1. 先扫根目录下直接放置的图片
        for ext in ['*.jpg', '*.jpeg', '*.png', '*.webp', '*.bmp']:
            for img_path in self.source_dir.glob(ext):
                path_str = str(img_path)
                if path_str not in seen_paths:
                    seen_paths.add(path_str)
                    images.append(ImageInfo(
                        path=path_str,
                        name=img_path.name,
                        size=img_path.stat().st_size,
                        folder='.',   # 根目录用 '.' 标识
                        hash=self._calculate_hash(path_str)
                    ))
        
        # 2. 再扫各子文件夹
        for folder in self.source_dir.iterdir():
            if folder.is_dir():
                for ext in ['*.jpg', '*.jpeg', '*.png', '*.webp', '*.bmp']:
                    for img_path in folder.glob(ext):
                        path_str = str(img_path)
                        if path_str not in seen_paths:
                            seen_paths.add(path_str)
                            images.append(ImageInfo(
                                path=path_str,
                                name=img_path.name,
                                size=img_path.stat().st_size,
                                folder=folder.name,
                                hash=self._calculate_hash(path_str)
                            ))
        
        root_count = sum(1 for img in images if img.folder == '.')
        if root_count > 0:
            logger.info(f"[扫描] 根目录发现 {root_count} 张图片")
        
        return images
    
    def create_batches(self, images: List[ImageInfo]) -> List[BatchInfo]:
        """创建批次（按文件夹和大小智能分批）"""
        folder_groups: Dict[str, List[ImageInfo]] = {}
        for img in images:
            if img.folder not in folder_groups:
                folder_groups[img.folder] = []
            folder_groups[img.folder].append(img)
        
        batches = []
        batch_counter = 0
        
        for folder_name, folder_images in folder_groups.items():
            folder_images.sort(key=lambda x: x.size)
            current_batch = []
            
            for img in folder_images:
                batch_size_limit, category = self._get_batch_size_limit(img.path)
                
                if len(current_batch) >= batch_size_limit:
                    batch_counter += 1
                    batch_id = f"{folder_name}_{batch_counter:03d}"
                    batches.append(BatchInfo(
                        batch_id=batch_id,
                        folder=folder_name,
                        images=[i.path for i in current_batch],
                        status="pending",
                        created_at=datetime.now().isoformat()
                    ))
                    current_batch = [img]
                else:
                    current_batch.append(img)
            
            if current_batch:
                batch_counter += 1
                batch_id = f"{folder_name}_{batch_counter:03d}"
                batches.append(BatchInfo(
                    batch_id=batch_id,
                    folder=folder_name,
                    images=[i.path for i in current_batch],
                    status="pending",
                    created_at=datetime.now().isoformat()
                ))
        
        return batches
    
    def initialize(self) -> Optional[ProcessingState]:
        """初始化处理会话"""
        # 检查是否有未完成的会话
        existing_state = self._load_state()
        if existing_state:
            pending = sum(1 for b in existing_state.batches.values() if b.status == "pending")
            processing = sum(1 for b in existing_state.batches.values() if b.status == "processing")
            
            if pending > 0 or processing > 0:
                print(f"[INFO] 发现未完成的会话: {existing_state.session_id}")
                print(f"       已完成: {existing_state.completed_batches}/{existing_state.total_batches}")
                self.state = existing_state
                return existing_state
        
        # 创建新会话
        images = self.scan_images()
        if not images:
            print("[INFO] 没有待处理的图片")
            return None
        
        batches = self.create_batches(images)
        
        self.state = ProcessingState(
            session_id=datetime.now().strftime("%Y%m%d_%H%M%S"),
            started_at=datetime.now().isoformat(),
            total_images=len(images),
            total_batches=len(batches),
            batches={b.batch_id: b for b in batches}
        )
        
        self._save_state()
        
        print(f"[INFO] 新会话已创建: {self.state.session_id}")
        print(f"       共 {self.state.total_images} 张图片，分成 {self.state.total_batches} 批次")
        
        return self.state
    
    def get_next_batch(self) -> Optional[BatchInfo]:
        """获取下一个待处理的批次"""
        if not self.state:
            self.state = self._load_state()
        
        if not self.state:
            return None
        
        for batch_id, batch in self.state.batches.items():
            if batch.status == "pending":
                batch.status = "processing"
                batch.started_at = datetime.now().isoformat()
                self._save_state()
                return batch
        
        return None
    
    def mark_batch_completed(self, batch_id: str, result: str = ""):
        """标记批次为已完成"""
        if not self.state:
            self.state = self._load_state()
        
        if self.state and batch_id in self.state.batches:
            batch = self.state.batches[batch_id]
            if batch.status != "completed":
                batch.status = "completed"
                batch.completed_at = datetime.now().isoformat()
                batch.result = result
                self.state.completed_batches += 1
                self._save_state()
                print(f"[INFO] 批次 {batch_id} 已完成")
    
    def mark_batch_failed(self, batch_id: str, error: str = ""):
        """标记批次为失败"""
        if self.state and batch_id in self.state.batches:
            batch = self.state.batches[batch_id]
            batch.status = "failed"
            batch.result = error
            batch.retry_count += 1
            
            if batch.retry_count < 3:
                batch.status = "pending"
                print(f"[WARN] 批次 {batch_id} 失败，将重试 ({batch.retry_count}/3)")
            else:
                print(f"[ERROR] 批次 {batch_id} 失败超过3次，已放弃")
            
            self._save_state()
    
    def get_progress(self) -> Dict:
        """获取处理进度"""
        if not self.state:
            self.state = self._load_state()
        
        if not self.state:
            return {}
        
        pending = sum(1 for b in self.state.batches.values() if b.status == "pending")
        processing = sum(1 for b in self.state.batches.values() if b.status == "processing")
        completed = sum(1 for b in self.state.batches.values() if b.status == "completed")
        failed = sum(1 for b in self.state.batches.values() if b.status == "failed")
        
        total = self.state.total_batches
        progress_pct = (completed / total * 100) if total > 0 else 0
        
        return {
            "session_id": self.state.session_id,
            "total": total,
            "pending": pending,
            "processing": processing,
            "completed": completed,
            "failed": failed,
            "progress": progress_pct,
            "is_complete": pending == 0 and processing == 0
        }
    
    def print_progress(self):
        """打印进度"""
        progress = self.get_progress()
        if not progress:
            print("[INFO] 没有正在进行的处理任务")
            return
        
        print("\n" + "=" * 60)
        print(f"处理进度 - 会话 {progress['session_id']}")
        print("=" * 60)
        print(f"总批次: {progress['total']}")
        print(f"已完成: {progress['completed']} ({progress['progress']:.1f}%)")
        print(f"处理中: {progress['processing']}")
        print(f"待处理: {progress['pending']}")
        print(f"失败: {progress['failed']}")
        print("=" * 60)
        
        if progress['is_complete']:
            print("[INFO] 所有批次处理完成！")
    
    def clear_state(self):
        """清除处理状态"""
        if BATCH_STATE_FILE.exists():
            BATCH_STATE_FILE.unlink()
            print("[INFO] 处理状态已清除")


# ============================================================
# 主题分类配置
# ============================================================
THEME_KEYWORDS = {
    "中医养生": ["中医", "养生", "药膳", "食疗", "中药", "经络", "穴位", "调理", "体质", "气血", "肝", "肾", "脾", "胃", "湿气", "虚", "寒", "热", "上火", "补", "滋阴", "温阳", "祛湿", "健脾", "润肺", "养心", "补肾", "中医师", "中医说", "中医认为"],
    "健康饮食": ["饮食", "营养", "食物", "蔬菜", "水果", "蛋白质", "碳水", "健康", "抗炎", "抗氧化", "免疫力", "肠道", "益生元", "膳食纤维", "维生素", "矿物质", "有机", "纯天然", "少油", "少盐", "清淡", "养生餐", "健康餐"],
    "疾病防治": ["疾病", "预防", "治疗", "症状", "指标", "血压", "血糖", "血脂", "胆固醇", "尿酸", "脂肪肝", "结节", "囊肿", "肿瘤", "癌症", "慢性病", "并发症", "吃药", "服药", "手术", "复查", "就医", "医院", "医生", "确诊"],
    "生活方式": ["运动", "睡眠", "压力", "情绪", "心理健康", "作息", "习惯", "减肥", "增重", "美容", "护肤", "跑步", "走路", "瑜伽", "冥想", "放松", "健身", "锻炼"],
    "营养科普": ["科普", "知识", "研究", "发现", "实验", "数据", "结论", "专家", "建议", "指南", "推荐", "科学", "原理", "机制", "分析", "解读", "揭秘", "真相", "为什么", "是什么"],
    "育儿教育": ["育儿", "宝宝", "孩子", "教育", "辅食", "喂养", "早教", "亲子", "成长", "发育", "妈妈", "孕妇", "备孕", "新生儿", "婴儿", "幼儿", "儿童"],
    "美食烹饪": ["美食", "烹饪", "做法", "菜谱", "食谱", "食材", "配料", "做饭", "煮", "炒", "炖", "蒸", "烤", "炸", "凉拌", "汤", "粥", "面食", "烘焙", "蛋糕", "面包"]
}

TOPIC_SIGNATURES = {
    "养生方": ["方子", "秘方", "配方", "做法", "煮", "熬", "炖"],
    "食疗食谱": ["食谱", "菜谱", "做法", "食材", "配料", "烹饪"],
    "禁忌注意": ["禁忌", "注意", "不宜", "少吃", "避免", "千万", "不能", "不可"],
    "健康功效": ["功效", "作用", "好处", "益处", "价值", "可以", "有助于"],
    "疾病症状": ["症状", "表现", "特征", "体征", "不舒服", "疼痛", "难受"]
}


# ============================================================
# 进度条显示（V4优化）
# ============================================================
class ProgressBar:
    """终端进度条"""
    
    def __init__(self, total: int, width: int = 40, prefix: str = "进度"):
        self.total = total
        self.width = width
        self.prefix = prefix
        self.current = 0
        self.start_time = time.time()
        
    def update(self, current: int, info: str = ""):
        self.current = current
        percent = current / self.total if self.total > 0 else 0
        filled = int(self.width * percent)
        bar = '█' * filled + '░' * (self.width - filled)
        
        if current > 0:
            elapsed = time.time() - self.start_time
            eta = elapsed / current * (self.total - current)
            if eta > 60:
                eta_str = f"{int(eta/60)}分{int(eta%60)}秒"
            else:
                eta_str = f"{int(eta)}秒"
        else:
            eta_str = "--"
            
        print(f"\r{self.prefix}: |{bar}| {current}/{self.total} ({percent:.0%}) ETA:{eta_str} {info[:15]}", 
              end='', flush=True)
        
    def finish(self):
        self.current = self.total
        self.update(self.total, "完成!")
        print()


# ============================================================
# 图片归档器（V4优化）
# ============================================================
class ImageArchiver:
    """图片归档器 - 按主题分类归档"""
    
    def __init__(self, processed_dir: str = '已处理图片'):
        self.processed_dir = Path(processed_dir)
        self.processed_dir.mkdir(parents=True, exist_ok=True)
        self.archived_count = 0
        
    def archive_image(self, image_path: str, theme: str = None) -> Optional[str]:
        """移动图片到按主题分类的已处理文件夹"""
        src = Path(image_path)
        if not src.exists():
            return None
            
        if theme and theme in THEME_KEYWORDS:
            archive_subdir = self.processed_dir / theme
        else:
            source_folder = src.parent.name
            if source_folder and source_folder not in ['待处理图片', 'images', 'source']:
                archive_subdir = self.processed_dir / source_folder
            else:
                archive_subdir = self.processed_dir / "其他"
        
        archive_subdir.mkdir(parents=True, exist_ok=True)
        
        dest_name = src.name
        dest = archive_subdir / dest_name
        
        if dest.exists():
            timestamp = datetime.now().strftime('%H%M%S')
            name_parts = src.stem, src.suffix
            dest = archive_subdir / f"{name_parts[0]}_{timestamp}{name_parts[1]}"
        
        try:
            shutil.move(str(src), str(dest))
            self.archived_count += 1
            return str(dest)
        except Exception as e:
            logger.warning(f"[归档] 移动失败 {src.name}: {e}")
            try:
                shutil.copy2(str(src), str(dest))
                src.unlink()
                self.archived_count += 1
                return str(dest)
            except:
                return None
                
    def get_archive_stats(self) -> Dict:
        """获取归档统计"""
        total = 0
        subdirs = {}
        for item in self.processed_dir.iterdir():
            if item.is_dir():
                count = len(list(item.glob('*.*')))
                subdirs[item.name] = count
                total += count
        return {
            'archived_count': self.archived_count,
            'total_archived': total,
            'subdirs': subdirs
        }


# ============================================================
# 内容整理器（V2功能）
# ============================================================
class ContentOrganizer:
    """内容整理器"""

    def __init__(self):
        self.source_keywords = [
            '来源', '摘录', '转载', '出处', '作者',
            '小红书', '微信', '抖音', '微博', 'B站',
            '公众号', '视频号', '知乎', '豆瓣', '百度'
        ]
        self.heading_keywords = [
            '功效', '作用', '好处', '适合', '适宜', '适用',
            '原料', '食材', '配料', '材料',
            '做法', '制作', '步骤', '方法', '教程',
            '禁忌', '注意', '提醒', '警告', '不宜',
            '简介', '介绍', '概述', '说明',
            '推荐', '建议', '指南'
        ]
        self.section_keywords = [
            '一、', '二、', '三、', '四、', '五、', '六、',
            '1.', '2.', '3.', '4.', '5.', '6.',
            '（1）', '（2）', '（3）', '（4）',
            '首先', '其次', '然后', '最后', '另外', '此外'
        ]

    def extract_source(self, text: str) -> Dict:
        """从内容中提取来源信息"""
        lines = text.split('\n')
        source_info = {}

        for line in lines[:20]:
            line = line.strip()
            for kw in ['来源', '出处', '作者', '摘录自']:
                if kw in line:
                    parts = line.split(kw)
                    if len(parts) > 1:
                        source = parts[1].strip().lstrip('：:').strip()
                        if source:
                            source_info['platform'] = source
                            break

            for platform in ['小红书', '微信读书', '微信视频号', '微信', '抖音', 'B站', 'BTV']:
                if platform in line:
                    source_info['platform'] = platform
                    break

            if '作者' in line:
                match = re.search(r'作者[：:]\s*([^\n]+)', line)
                if match:
                    source_info['author'] = match.group(1).strip()

        return source_info

    def clean_text(self, text: str) -> str:
        """清理OCR识别的原始文本"""
        lines = text.split('\n')
        cleaned_lines = []
        prev_line_empty = False

        for line in lines:
            line = line.strip()

            if len(line) < 2 and line:
                continue

            if re.match(r'^[_\-=]{3,}$', line):
                continue

            if re.match(r'^[a-zA-Z0-9\s]{20,}$', line) and not re.search(r'[\u4e00-\u9fa5]', line):
                continue

            ui_patterns = ['开始', '插入', '绘图', '设计', '切换', '动画', '审阅', '视图', '帮助', '文件', '编辑', '格式', '工具', '表格']
            if any(ui in line for ui in ui_patterns) and len(line) < 10:
                continue

            if not line:
                if not prev_line_empty:
                    cleaned_lines.append('')
                    prev_line_empty = True
                continue

            prev_line_empty = False
            line = self._fix_common_ocr_errors(line)
            cleaned_lines.append(line)

        result = '\n'.join(cleaned_lines)
        result = re.sub(r'\n{3,}', '\n\n', result)
        return result.strip()

    def _fix_common_ocr_errors(self, line: str) -> str:
        """修复常见的OCR错误"""
        replacements = [
            (chr(8220), '"'), (chr(8221), '"'),
            (chr(8216), "'"), (chr(8217), "'"), (chr(8218), "'"),
            (chr(8222), '"'), (chr(8223), '"'),
            (chr(8242), "'"), (chr(8243), "'"),
            (chr(180), ""), (chr(9032), ""),
            (chr(8212), "-"), (chr(183), "-"),
            ('`', ''), ('√', '✓'), ('×', '✗')
        ]
        for old, new in replacements:
            line = line.replace(old, new)
        line = re.sub(r' {2,}', ' ', line)
        return line

    def _is_heading_line(self, line: str) -> Tuple[bool, str]:
        """
        判断一行是否构成标题，并返回标题级别和格式化后的文本。
        返回 (is_heading, formatted_line)
        """
        line = line.strip()
        if not line:
            return False, line

        # ---- H3 级标题（带序号/前缀的短标题）----

        # ① ② ③ 等圆圈数字：
        #   - 短内容（<15字）只带物品名的是纯列表项，不当标题
        #   - 有实质内容（≥15字或含冒号）才是小标题
        if re.match(r'^[\u2460-\u24ff\u2776-\u277f\u2780-\u2789\u24ea\u24B6-\u24CF①-⑳]\s*.+', line):
            if len(line) >= 15 or '：' in line or '：' in line:
                return True, f"### {line}"
            else:
                return False, line   # 纯列表项（如"① 苹果"），不处理

        # 1) 2) 编号格式：短内容（<15字）是列表项，有实质内容才是标题
        if re.match(r'^\d+[\.\)]\s+.+', line):
            if len(line) >= 15:
                return True, f"### {line}"
            else:
                return False, line

        # 中文序号 + 标点：一、二、三、| 第一章 | 1. 2. 3.
        if re.match(r'^[一二三四五六七八九十百千零\d]+[、.．:：]\s*.+', line):
            if len(line) <= 40:
                return True, f"### {line}"
        if re.match(r'^(第[一二三四五六七八九十百千零\d]+[章节篇部])\s*[：:、]?\s*.+', line):
            if len(line) <= 50:
                return True, f"### {line}"

        # **加粗**包裹的短行
        if line.startswith('**') and line.endswith('**') and len(line) <= 40:
            return True, f"### {line.strip('*')}"

        # 以 section_keywords 开头的短行（如"【功效】"、"【做法】"）
        if any(line.startswith(kw) or line.startswith(kw.replace('【', '').replace('】', ''))
               for kw in self.section_keywords) and len(line) <= 50:
            return True, f"### {line}"

        # ---- H2 级标题（纯中文短语，无标点，长度适中）----
        if re.match(r'^[\u4e00-\u9fa5]{4,25}$', line):
            if not re.search(r'[a-zA-Z0-9]', line):
                return True, f"## {line}"

        # ---- 非标题 ----
        return False, line

    def structure_content(self, text: str) -> str:
        """对内容进行智能结构化处理"""
        lines = text.split('\n')
        structured = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            is_heading, formatted = self._is_heading_line(line)
            structured.append(formatted)

        return '\n'.join(structured)
    
    def organize(self, text: str, theme: str) -> Dict:
        """整理内容"""
        cleaned = self.clean_text(text)
        structured = self.structure_content(cleaned)
        source = self.extract_source(text)
        return {
            'content': structured,
            'source': source
        }


# ============================================================
# AI内容分析器（V2功能）
# ============================================================
class ContentAnalyzer:
    """AI内容分析器"""

    def __init__(self):
        self.cache = {}
        self.organizer = ContentOrganizer()

    def extract_title(self, text: str) -> Optional[str]:
        """从内容中提取标题"""
        lines = text.split('\n')
        for line in lines[:10]:
            line = line.strip()
            if line.startswith('#'):
                title = line.lstrip('#').strip()
                if 2 <= len(title) <= 30:
                    return title
            if re.match(r'^[◆★▼▪▸]+\s*.+', line):
                title = re.sub(r'^[◆★▼▪▸]+\s*', '', line)
                if len(title) >= 2:
                    return title[:20]
            if re.match(r'^[\u4e00-\u9fa5]{4,15}$', line) and len(line) >= 4:
                return line

        for line in lines[:15]:
            line = line.strip()
            if 4 <= len(line) <= 20 and re.match(r'^[\u4e00-\u9fa5]', line):
                if not line.startswith('#') and 'http' not in line.lower():
                    return line
        return None

    def detect_theme(self, text: str) -> str:
        """检测内容主题"""
        max_score = 0
        theme_name = "综合知识"

        for theme, keywords in THEME_KEYWORDS.items():
            score = sum(3 if kw in text else 0 for kw in keywords)
            if score > 0:
                density = score / max(len(text), 1) * 1000
                if density > max_score:
                    max_score = density
                    theme_name = theme

        return theme_name

    def detect_topic_type(self, text: str) -> str:
        """检测内容类型"""
        for topic, signatures in TOPIC_SIGNATURES.items():
            score = sum(2 if sig in text else 0 for sig in signatures)
            if score >= 2:
                return topic
        return "知识整理"

    def generate_doc_name(self, text: str, category: str, index: int) -> str:
        """生成有意义的文档名"""
        title = self.extract_title(text)
        if title:
            title = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '', title)
            if len(title) >= 3:
                return title[:15]

        topic_type = self.detect_topic_type(text)
        if topic_type != "知识整理":
            return f"{topic_type}"

        if category and category != "综合知识":
            return f"{category}_整理"

        timestamp = datetime.now().strftime('%m%d')
        return f"知识整理_{timestamp}_{index:02d}"

    def compute_content_hash(self, text: str) -> str:
        """计算内容哈希（全量文本，不截断）"""
        cleaned = re.sub(r'\s+', '', text)
        return hashlib.md5(cleaned.encode('utf-8')).hexdigest()[:8]

    def analyze(self, text: str, category: str, index: int) -> Dict:
        """综合分析"""
        cache_key = self.compute_content_hash(text)

        if cache_key in self.cache:
            return self.cache[cache_key]

        theme = self.detect_theme(text)
        topic_type = self.detect_topic_type(text)
        doc_name = self.generate_doc_name(text, category, index)

        organized = self.organizer.organize(text, theme)
        organized_content = organized['content']
        source_info = organized['source']

        result = {
            'theme': theme,
            'topic_type': topic_type,
            'doc_name': doc_name,
            'content_hash': cache_key,
            'organized_content': organized_content,
            'source': source_info
        }

        self.cache[cache_key] = result
        return result


# ============================================================
# 文档合并器（V2功能）
# ============================================================
class DocumentMerger:
    """文档合并器"""

    def __init__(self, output_dir: str = '处理结果'):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.theme_docs = defaultdict(list)
        self._scan_existing_docs()

    def _scan_existing_docs(self):
        """扫描已存在的文档"""
        for theme_folder in self.output_dir.iterdir():
            if theme_folder.is_dir() and not theme_folder.name.startswith('.'):
                for md_file in theme_folder.glob('*.md'):
                    topic = self._extract_topic(md_file)
                    self.theme_docs[theme_folder.name].append({
                        'file': md_file,
                        'topic': topic
                    })

    def _extract_topic(self, md_file: Path) -> str:
        """从文档名提取主题"""
        name = md_file.stem
        name = re.sub(r'^\d+_', '', name)
        name = re.sub(r'_\d+$', '', name)
        name = re.sub(r'_[a-f0-9]{8}$', '', name)
        return name

    def find_similar_doc(self, theme: str, doc_name: str, content: str) -> Optional[Path]:
        """
        查找相似的已有文档（P1-10 修复）：
        1. 先要求主题名高度相似（≥0.7）
        2. 再要求内容也有一定重叠（≥30字符相同片段）
        两个条件都满足才合并，避免"健康饮食"和"健康饮食_1"内容完全不同却合并
        """
        doc_topic = re.sub(r'^\d+_', '', doc_name)
        doc_topic = re.sub(r'_\d+$', '', doc_topic)

        for existing_doc in self.theme_docs.get(theme, []):
            existing_topic = existing_doc['topic']

            # 条件1：主题名相似度足够高
            if self._calc_similarity(doc_topic, existing_topic) < 0.7:
                continue

            # 条件2：内容也有重叠（防止主题名碰巧相似但内容完全不同）
            existing_file = existing_doc['file']
            if existing_file.exists():
                try:
                    existing_content = existing_file.read_text(encoding='utf-8')
                    # 取两段内容的前300字比较
                    c1 = content[:300]
                    c2 = existing_content[:300]
                    overlap_len = len(set(c1) & set(c2))
                    if overlap_len < 30:   # 至少30个相同字符才合并
                        continue
                except Exception:
                    pass

            return existing_doc['file']
        return None

    def _calc_similarity(self, s1: str, s2: str) -> float:
        """
        改进的相似度计算：
        1. 包含关系：完全包含得高分
        2. 字符 bigram Jaccard（对中文更有意义）
        3. 关键词重叠（营养/食疗/做法等核心词精确匹配）
        """
        if not s1 or not s2:
            return 0

        # 归一化：去除序号前缀
        s1_norm = re.sub(r'^[一二三四五六七八九十百千零\d]+[、.．:：]+', '', s1).strip()
        s2_norm = re.sub(r'^[一二三四五六七八九十百千零\d]+[、.．:：]+', '', s2).strip()
        s1_norm = re.sub(r'^\d+[\.\)]\s*', '', s1_norm)
        s2_norm = re.sub(r'^\d+[\.\)]\s*', '', s2_norm)

        # 1. 完全相同
        if s1_norm == s2_norm:
            return 1.0

        # 2. 包含关系（短串完全包含在长串中）
        shorter, longer = (s1_norm, s2_norm) if len(s1_norm) <= len(s2_norm) else (s2_norm, s1_norm)
        if len(shorter) >= 4 and shorter in longer:
            # 部分包含：按覆盖比例给分
            return 0.85 + 0.1 * (len(shorter) / len(longer))

        # 3. 字符 bigram Jaccard（适用于中文）
        def bigrams(s):
            return set(s[i:i+2] for i in range(len(s) - 1))
        b1, b2 = bigrams(s1_norm), bigrams(s2_norm)
        if b1 and b2:
            jaccard = len(b1 & b2) / len(b1 | b2)
        else:
            jaccard = 0

        # 4. 关键词精确重叠（中医/养生/食疗/做法等）
        KEYWORDS = {'中医', '养生', '食疗', '药膳', '中药', '功效', '做法',
                     '营养', '健康', '疾病', '预防', '饮食', '减肥', '睡眠'}
        words1 = set(s1_norm) & KEYWORDS
        words2 = set(s2_norm) & KEYWORDS
        if words1 and words2:
            keyword_score = len(words1 & words2) / len(words1 | words2)
        else:
            keyword_score = 0

        return max(jaccard * 0.6 + keyword_score * 0.4, jaccard)

    def merge_content(self, existing_file: Path, new_content: str, 
                      new_image_name: str, content_hash: str = None) -> bool:
        """合并内容到已有文档"""
        try:
            with open(existing_file, 'r', encoding='utf-8') as f:
                existing = f.read()

            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
            append_section = f"""

---

## 📌 补充内容

> 来源图片: {new_image_name}
> 追加时间: {timestamp}
> 内容标识: {content_hash or 'N/A'}

{new_content}

"""

            new_content_merged = existing.replace(
                '\n\n---\n\n*本文档由图片知识库整理工具自动生成*',
                append_section + '\n\n---\n\n*本文档由图片知识库整理工具自动生成*'
            )

            with open(existing_file, 'w', encoding='utf-8') as f:
                f.write(new_content_merged)

            return True
        except Exception as e:
            logger.warning(f"[合并] 失败: {e}")
            return False


# ============================================================
# 智能文档生成器（V2功能）
# ============================================================
class SmartDocumentGenerator:
    """智能文档生成器"""

    def __init__(self, output_dir: str = '处理结果'):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.merger = DocumentMerger(output_dir)
        self.existing_hashes = set()
        self._scan_hashes()
        self.generated_docs = set()

    def _scan_hashes(self):
        """扫描已有文档的hash"""
        for md_file in self.output_dir.rglob('*.md'):
            try:
                file_content = md_file.read_text(encoding='utf-8')
                hash_match = re.search(r'content_hash:\s*([a-f0-9]{8})', file_content)
                if hash_match:
                    self.existing_hashes.add(hash_match.group(1))
            except:
                pass

    def _get_doc_key(self, theme: str, doc_name: str) -> str:
        """获取文档唯一标识"""
        safe_name = re.sub(r'[<>:"/\|?*]', '', doc_name)
        if len(safe_name) > 25:
            safe_name = safe_name[:25]
        return f"{theme}/{safe_name}"

    def is_duplicate(self, content_hash: str) -> bool:
        """检测是否重复"""
        return content_hash in self.existing_hashes

    def is_doc_generated(self, theme: str, doc_name: str) -> bool:
        """检查文档是否已生成"""
        return self._get_doc_key(theme, doc_name) in self.generated_docs

    def mark_doc_generated(self, theme: str, doc_name: str):
        """标记文档已生成"""
        self.generated_docs.add(self._get_doc_key(theme, doc_name))

    def generate_markdown(self, text: str, theme: str, doc_name: str, 
                          category: str, image_name: str, keywords: List = None,
                          content_hash: str = None, source_info: Dict = None) -> str:
        """生成Markdown文档"""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')

        similar_doc = self.merger.find_similar_doc(theme, doc_name, text)
        if similar_doc:
            print(f"  → 发现相似文档 {similar_doc.name}，合并内容...")
            if self.merger.merge_content(similar_doc, text, image_name, content_hash):
                self.mark_doc_generated(theme, doc_name)
                return str(similar_doc)

        keywords_str = ', '.join(keywords) if keywords else '无'

        source_section = ""
        if source_info:
            if source_info.get('platform'):
                source_section += f"**来源**：{source_info['platform']}\n\n"
            if source_info.get('author'):
                source_section += f"**作者**：{source_info['author']}\n\n"

        md_content = f"""# {doc_name}

> 来源图片: {image_name}
> 识别时间: {timestamp}
> 主题分类: {theme}
> 原分类: {category}
> 关键词: {keywords_str}
> content_hash: {content_hash or 'N/A'}

---

<!-- CONTENT_START -->
{source_section}{text}
<!-- CONTENT_END -->

---

*本文档由图片知识库整理工具自动生成*
"""

        theme_folder = self.output_dir / theme
        theme_folder.mkdir(exist_ok=True)

        safe_name = re.sub(r'[<>:"/\|?*]', '', doc_name)
        if len(safe_name) > 25:
            safe_name = safe_name[:25]

        md_file = theme_folder / f"{safe_name}.md"
        counter = 1
        while md_file.exists():
            md_file = theme_folder / f"{safe_name}_{counter}.md"
            counter += 1

        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(md_content)

        if content_hash:
            self.existing_hashes.add(content_hash)
        self.merger.theme_docs[theme].append({
            'file': md_file,
            'topic': doc_name
        })
        self.mark_doc_generated(theme, doc_name)

        logger.info(f"[文档] Markdown已生成: {theme}/{md_file.name}")
        return str(md_file)

    def generate_word(self, text: str, theme: str, doc_name: str, image_name: str,
                      keywords: List = None, content_hash: str = None, 
                      existing_md_file: str = None) -> Optional[str]:
        """生成Word文档"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.oxml.ns import qn
        except ImportError:
            logger.warning("[文档] python-docx未安装，跳过Word生成")
            return None

        if self.is_doc_generated(theme, doc_name):
            safe_name = re.sub(r'[<>:"/\|?*]', '', doc_name)
            if len(safe_name) > 25:
                safe_name = safe_name[:25]
            existing_docx = self.output_dir / theme / f"{safe_name}.docx"
            if existing_docx.exists():
                print(f"  → Word文档已存在，跳过")
                return str(existing_docx)
            for i in range(1, 100):
                existing_docx = self.output_dir / theme / f"{safe_name}_{i}.docx"
                if existing_docx.exists():
                    print(f"  → Word文档已存在，跳过")
                    return str(existing_docx)

        theme_folder = self.output_dir / theme
        theme_folder.mkdir(exist_ok=True)

        safe_name = re.sub(r'[<>:"/\|?*]', '', doc_name)
        if len(safe_name) > 25:
            safe_name = safe_name[:25]

        docx_file = theme_folder / f"{safe_name}.docx"
        counter = 1
        while docx_file.exists():
            docx_file = theme_folder / f"{safe_name}_{counter}.docx"
            counter += 1

        if existing_md_file and Path(existing_md_file).exists():
            try:
                md_content = Path(existing_md_file).read_text(encoding='utf-8')
                # 优先用明确标记提取（新格式）
                if '<!-- CONTENT_START -->' in md_content and '<!-- CONTENT_END -->' in md_content:
                    start = md_content.index('<!-- CONTENT_START -->') + len('<!-- CONTENT_START -->')
                    end = md_content.index('<!-- CONTENT_END -->')
                    main_content = md_content[start:end]
                else:
                    # 兼容旧格式：取第1个 --- 之后、最后一个 --- 之前的内容
                    parts = md_content.split('\n---\n')
                    if len(parts) >= 2:
                        # parts[0]=元数据头, parts[1]=正文, parts[-1]=尾注
                        # 如果只有2段说明无尾注，取parts[1]；3段以上则去掉最后一段
                        main_content = '\n---\n'.join(parts[1:-1]) if len(parts) > 2 else parts[1]
                        main_content = main_content.replace('*本文档由图片知识库整理工具自动生成*', '')
                    else:
                        main_content = md_content
                text = main_content.strip()
            except Exception as e:
                logger.warning(f"[Word] 读取合并内容失败: {e}")

        doc = Document()
        doc.styles['Normal'].font.name = 'Microsoft YaHei'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        doc.add_heading(doc_name, 0)
        doc.add_paragraph(f"来源图片: {image_name}")
        doc.add_paragraph(f"识别时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"主题分类: {theme}")
        keywords_str = ', '.join(keywords) if keywords else '无'
        doc.add_paragraph(f"关键词: {keywords_str}")
        doc.add_paragraph("─" * 30)
        doc.add_paragraph(text)

        doc.save(str(docx_file))
        self.mark_doc_generated(theme, doc_name)
        logger.info(f"[文档] Word已生成: {theme}/{docx_file.name}")
        return str(docx_file)


# ============================================================
# 多引擎OCR
# ============================================================
class MultiEngineOCR:
    """多引擎OCR管理器"""

    def __init__(self):
        self.current_engine = None
        self.engine_status = []

        secret_id = os.getenv('TENCENT_SECRET_ID', '')
        secret_key = os.getenv('TENCENT_SECRET_KEY', '')
        if secret_id and secret_key and '替换' not in secret_id and TENINCENT_AVAILABLE:
            self.engine_status.append(('腾讯云', True, '已配置'))
        else:
            self.engine_status.append(('腾讯云', False, '未配置'))

        app_id = os.getenv('BAIDU_APP_ID', '')
        api_key = os.getenv('BAIDU_API_KEY', '')
        secret_key = os.getenv('BAIDU_SECRET_KEY', '')
        if app_id and api_key and secret_key and '替换' not in app_id and BAIDU_AVAILABLE:
            self.engine_status.append(('百度云', True, '已配置'))
        else:
            self.engine_status.append(('百度云', False, '未配置'))

        local = LocalOCR()
        if local.tesseract_available:
            self.engine_status.append(('本地Tesseract', True, '可用'))
        else:
            self.engine_status.append(('本地Tesseract', False, getattr(local, 'error_message', '未知错误')))

        self._select_best_engine()

    def _select_best_engine(self):
        for name, available, msg in self.engine_status:
            if available:
                self.current_engine = name
                logger.info(f"[OCR] 选用引擎: {name} ({msg})")
                return True
        logger.warning("[OCR] 没有可用的OCR引擎!")
        return False

    def recognize(self, image_path: str, timeout: int = 60) -> Dict:
        """识别图片文字"""
        last_error = None
        tried_engines = []
        for name, available, _ in self.engine_status:
            if not available:
                continue
            tried_engines.append(name)
            try:
                if name == '腾讯云' and TENINCENT_AVAILABLE:
                    engine = TencentOCR()
                elif name == '百度云' and BAIDU_AVAILABLE:
                    engine = BaiduOCR()
                else:
                    engine = LocalOCR()

                result = engine.recognize(image_path)
                if result and result.get('success'):
                    if len(tried_engines) > 1:
                        # 降级成功，打印提示
                        print(f"  ⚡ 引擎降级：{tried_engines[0]} → {name}（{name}识别成功）")
                    logger.info(f"[OCR] {name} 识别成功!")
                    return result
                else:
                    last_error = result.get('error', '未知错误') if result else '返回为空'
                    if len(tried_engines) > 1:
                        print(f"  ⚠️ {name} 识别失败，尝试下一引擎...")
            except Exception as e:
                last_error = str(e)
                if len(tried_engines) > 1:
                    print(f"  ⚠️ {name} 异常（{e}），尝试下一引擎...")
        return {'success': False, 'error': last_error or '无可用OCR引擎', 'text': ''}


# ============================================================
# IMA同步器
# ============================================================
class IMASyncer:
    """IMA笔记同步器"""

    def __init__(self):
        self.client_id = os.getenv('IMA_OPENAPI_CLIENTID', '')
        self.api_key = os.getenv('IMA_OPENAPI_APIKEY', '')
        self.base_url = 'https://ima.qq.com/openapi/note/v1'
        self.enabled = bool(self.client_id and self.api_key and '填入' not in self.client_id)
        self.sync_log_file = Path('处理结果/ima_sync_log.json')
        self.sync_log = self._load_sync_log()
        self.rate_limited = False
        
        # P2-16: 支持同步到指定笔记本
        self.default_notebook_id = os.getenv('IMA_NOTEBOOK_ID', '')
        self._notebook_cache = {}  # 缓存笔记本ID映射

        if self.enabled:
            logger.info("[IMA] 已配置")
            if self.default_notebook_id:
                logger.info(f"[IMA] P2-16: 将同步到指定笔记本: {self.default_notebook_id}")
        else:
            logger.warning("[IMA] 未配置")

    def _load_sync_log(self) -> Dict:
        if self.sync_log_file.exists():
            try:
                return json.loads(self.sync_log_file.read_text(encoding='utf-8'))
            except:
                pass
        return {}

    def _save_sync_log(self):
        self.sync_log_file.write_text(
            json.dumps(self.sync_log, ensure_ascii=False, indent=2),
            encoding='utf-8'
        )

    def _api_call(self, endpoint: str, payload: Dict, retries: int = 3) -> Optional[Dict]:
        if not self.enabled or self.rate_limited:
            return None

        for attempt in range(retries):
            try:
                import requests
                headers = {
                    'ima-openapi-clientid': self.client_id,
                    'ima-openapi-apikey': self.api_key,
                    'Content-Type': 'application/json'
                }
                url = f"{self.base_url}/{endpoint}"
                response = requests.post(url, json=payload, headers=headers, timeout=30)

                if response.status_code == 200:
                    return response.json()
                elif response.status_code == 403:
                    result = response.json()
                    if '请求超量' in result.get('msg', ''):
                        logger.warning(f"[IMA] API请求超量（{attempt+1}/{retries}），等待60秒后重试...")
                        if attempt < retries - 1:
                            time.sleep(60)   # 等待60秒后重试（不设全局跳过）
                        continue
                return None
            except Exception as e:
                logger.warning(f"[IMA] API调用失败 ({attempt+1}/{retries}): {e}")
                if attempt < retries - 1:
                    time.sleep(2)
        return None

    def _validate_doc_id(self, doc_id: str) -> bool:
        """验证 doc_id 是否在 IMA 中仍然有效（通过 get_doc API）"""
        result = self._api_call('get_doc', {'doc_id': doc_id})
        if result and result.get('code') == 0:
            return True
        return False

    def sync_note(self, title: str, content: str, theme: str = None,
                  content_hash: str = None, doc_path: str = None,
                  notebook_id: str = None) -> Optional[str]:
        """
        同步笔记到IMA
        
        Args:
            notebook_id: 指定笔记本ID（可选），优先使用；否则使用默认笔记本
        """
        if not self.enabled or self.rate_limited:
            return None

        # P2-16: 确定目标笔记本
        target_notebook = notebook_id or self.default_notebook_id
        if target_notebook:
            logger.info(f"[IMA] P2-16: 将同步到笔记本: {target_notebook[:16]}...")

        doc_key = content_hash or title
        existing = self.sync_log.get(doc_key, {})

        full_content = f"# {title}\n\n"
        if theme:
            full_content += f"> 主题: {theme}\n"
        if content_hash:
            full_content += f"> 内容标识: {content_hash}\n"
        full_content += f"\n{content}\n\n---\n*自动同步自图片知识库*\n"

        if doc_path and existing.get('doc_id'):
            # P1-6 修复：追加前先验证 doc_id 是否有效
            if not self._validate_doc_id(existing['doc_id']):
                logger.warning(f"[IMA] doc_id {existing['doc_id']} 已失效，清理记录并创建新笔记")
                del self.sync_log[doc_key]
                self._save_sync_log()
                existing = {}

            if existing.get('doc_id'):
                result = self._api_call('append_doc', {
                    'doc_id': existing['doc_id'],
                    'content_format': 1,
                    'content': f"\n\n---\n*更新于 {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n" + content
                })
                if result:
                    self.sync_log[doc_key] = {
                        'doc_id': existing['doc_id'],
                        'last_sync': datetime.now().isoformat(),
                        'update_count': existing.get('update_count', 0) + 1
                    }
                    self._save_sync_log()
                    return existing['doc_id']

        # P2-16: 构建导入请求，添加笔记本ID
        import_payload = {
            'content_format': 1,
            'content': full_content
        }
        if target_notebook:
            import_payload['notebook_id'] = target_notebook

        result = self._api_call('import_doc', import_payload)

        if result and result.get('code') == 0:
            doc_id = result.get('data', {}).get('note_id', '')
            if doc_id:
                self.sync_log[doc_key] = {
                    'doc_id': doc_id,
                    'first_sync': datetime.now().isoformat(),
                    'last_sync': datetime.now().isoformat(),
                    'theme': theme,
                    'notebook_id': target_notebook,
                    'update_count': 1
                }
                self._save_sync_log()
                nb_info = f"[笔记本: {target_notebook[:8]}...]" if target_notebook else ""
                logger.info(f"[IMA] 笔记已同步 {nb_info}: {title[:20]}")
                return doc_id
            else:
                return "imported"

        return None
    
    def get_notebooks(self) -> List[Dict]:
        """获取笔记本列表"""
        result = self._api_call('list_notebooks', {})
        if result and result.get('code') == 0:
            return result.get('data', {}).get('notebooks', [])
        return []
    
    def create_notebook(self, name: str) -> Optional[str]:
        """创建笔记本并返回ID"""
        result = self._api_call('create_notebook', {'name': name})
        if result and result.get('code') == 0:
            return result.get('data', {}).get('notebook_id')
        # 尝试旧API
        result = self._api_call('add_notebook', {'name': name})
        if result and result.get('code') == 0:
            return result.get('data', {}).get('notebook_id')
        return None


# ============================================================
# 处理单张图片
# ============================================================
def process_single_image(ocr, classifier, content_analyzer, doc_gen, ima_syncer, archiver,
                         image_path: str, index: int, total: int, 
                         progress: ProgressBar = None):
    """处理单张图片"""
    image_name = Path(image_path).name
    source_folder = Path(image_path).parent.name
    
    if progress:
        progress.update(index, f"[{index}/{total}]")
    
    print(f"\n{'='*50}")
    print(f"[{index}/{total}] 处理: {image_name}")
    print(f"  📁 源文件夹: {source_folder}")
    print(f"{'='*50}")

    # 1. OCR识别
    print("\n📷 OCR识别中...")
    start_time = time.time()
    result = ocr.recognize(image_path)
    ocr_time = time.time() - start_time

    if not result.get('success'):
        print(f"  ⚠️ OCR失败: {result.get('error')}")
        archiver.archive_image(image_path)
        return None

    text = result.get('text', '').strip()
    if not text or len(text) < 10:
        print("  ⚠️ 无文字内容或内容过少")
        archiver.archive_image(image_path)
        return None

    print(f"  ✓ 识别到 {len(text)} 个字符 (耗时 {ocr_time:.1f}秒)")

    # 2. 智能分类
    print("\n🏷️ 内容分类中...")
    categories = classifier.classify(text, threshold=0.25)

    if categories:
        primary = categories[0]
        category_name = primary.category_name
        keywords = primary.matched_keywords if primary.matched_keywords else []
        confidence = primary.confidence
        
        # P2-15: 低置信度分类告警（置信度<30%记录到日志）
        if confidence < 0.30:
            logger.warning(f"[P2-15] 低置信度分类告警 | 图片: {image_name} | 分类: {category_name} | 置信度: {confidence:.0%} | 内容摘要: {text[:50]}...")
            print(f"  ⚠️ 分类: {category_name} (置信度 {confidence:.0%} - 低置信度)")
        else:
            print(f"  ✓ 分类: {category_name} (置信度 {confidence:.0%})")
        
        if keywords:
            print(f"     关键词: {', '.join(keywords[:5])}")
    else:
        category_name = "综合知识"
        keywords = []
        print("  ⚠️ 未匹配到分类，使用默认分类")

    # 3. AI内容分析
    print("\n🧠 AI内容分析中...")
    analysis = content_analyzer.analyze(text, category_name, index)
    theme = analysis['theme']
    topic_type = analysis['topic_type']
    doc_name = analysis['doc_name']
    content_hash = analysis['content_hash']
    organized_content = analysis.get('organized_content', text)
    source_info = analysis.get('source', {})
    print(f"  ✓ 主题分类: {theme}")
    print(f"  ✓ 内容类型: {topic_type}")
    print(f"  ✓ 文档命名: {doc_name}")
    if source_info.get('platform'):
        print(f"  ✓ 来源: {source_info['platform']}")

    # 4. 检测重复
    print("\n🔍 检测重复内容...")
    if doc_gen.is_duplicate(content_hash):
        print("  ⚠️ 检测到重复内容，跳过存储")
        archiver.archive_image(image_path, theme)
        return {
            'image': image_name,
            'text_length': len(text),
            'category': category_name,
            'theme': theme,
            'doc_name': doc_name,
            'is_duplicate': True,
            'content_hash': content_hash
        }
    print("  ✓ 新内容")

    # 5. 生成Markdown
    print("\n📝 生成Markdown...")
    md_file = doc_gen.generate_markdown(
        organized_content, theme, doc_name, category_name,
        image_name, keywords, content_hash, source_info
    )

    # 6. 生成Word
    print("\n📄 生成Word...")
    docx_file = doc_gen.generate_word(
        organized_content, theme, doc_name, image_name, keywords, content_hash, existing_md_file=md_file
    )
    if docx_file:
        rel_path = Path(docx_file).relative_to(doc_gen.output_dir)
        print(f"  ✓ 已保存: {rel_path}")
    else:
        print("  ⚠️ Word文档生成失败或未安装python-docx")

    # 7. 同步到IMA
    print("\n☁️ 同步到IMA...")
    if ima_syncer.rate_limited:
        print("  ⚠️ IMA被限流，跳过")
    else:
        ima_id = ima_syncer.sync_note(doc_name, organized_content, theme, content_hash, md_file)
        if ima_id:
            print(f"  ✓ IMA同步成功")
        else:
            print("  ⚠️ IMA同步失败")

    # 8. 图片归档
    print("\n📦 归档图片...")
    archived_path = archiver.archive_image(image_path, theme)
    if archived_path:
        archived_name = Path(archived_path).name
        archived_dir = Path(archived_path).parent.name
        print(f"  ✓ 已归档: 已处理图片/{archived_dir}/{archived_name}")

    return {
        'image': image_name,
        'source_folder': source_folder,
        'text_length': len(text),
        'category': category_name,
        'theme': theme,
        'topic_type': topic_type,
        'doc_name': doc_name,
        'keywords': keywords,
        'is_duplicate': False,
        'md_file': md_file,
        'docx_file': docx_file,
        'content_hash': content_hash,
        'source': source_info
    }



def _print_global_summary(batch_num: int, total_batches: int,
                          processed: int, total: int,
                          elapsed_seconds: float):
    """在批次间打印全局进度汇总行（单行覆盖）"""
    if total == 0:
        return
    pct = processed / total * 100
    if elapsed_seconds > 0 and processed < total:
        rate = processed / elapsed_seconds
        remaining = total - processed
        eta = remaining / rate if rate > 0 else 0
        eta_str = f"{int(eta // 60)}分{int(eta % 60)}秒"
    else:
        eta_str = "已完成"
    bar_width = 20
    filled = int(bar_width * processed / total)
    bar = '█' * filled + '░' * (bar_width - filled)
    line = (f"\r  全局进度 | {bar} {pct:5.1f}% "
            f"| 批次 {batch_num}/{total_batches} "
            f"| 已处理 {processed}/{total} 张 "
            f"| ETA {eta_str}   ")
    print(line, end='', flush=True)


# ============================================================
# 分批处理流程
# ============================================================
def run_batch_mode(batch_manager: BatchManager):
    """分批处理模式"""
    print("\n" + "=" * 60)
    print("分批处理模式")
    print("=" * 60)
    
    # P2-12: 处理前预检
    print("\n🔍 [P2-12] 执行预检...")
    check_passed, issues = preflight_check()
    if check_passed:
        print("  ✓ 预检通过")
    else:
        print("  ⚠️ 预检发现问题：")
        for issue in issues:
            print(f"    - {issue}")
        print("\n  继续处理（部分功能可能受限）...")
    
    # P2-17: 显示API Key配置状态
    env_results = validate_env_config()
    print("\n📋 [P2-17] API配置状态：")
    print(f"  腾讯云OCR: {'✓ 格式正确' if env_results.get('tencent') else '✗ 未配置/格式错误'}")
    print(f"  百度OCR: {'✓ 格式正确' if env_results.get('baidu') else '✗ 未配置/格式错误'}")
    print(f"  IMA同步: {'✓ 已配置' if env_results.get('ima') else '✗ 未配置'}")
    
    # 初始化
    state = batch_manager.initialize()
    if not state:
        return
    
    batch_manager.print_progress()
    
    # 初始化组件
    print("\n[1/7] 初始化OCR引擎...")
    ocr = MultiEngineOCR()
    if not ocr.current_engine:
        print("❌ 没有可用的OCR引擎!")
        return

    print("\n[2/7] 初始化智能分类器...")
    classifier = ClassifierEngine()

    print("\n[3/7] 初始化AI内容分析器...")
    content_analyzer = ContentAnalyzer()

    print("\n[4/7] 初始化文档生成器...")
    doc_gen = SmartDocumentGenerator()

    print("\n[5/7] 初始化IMA同步器...")
    ima_syncer = IMASyncer()

    print("\n[6/7] 初始化图片归档器...")
    archiver = ImageArchiver()

    print("\n[7/7] 准备处理...")
    
    # 逐批处理
    all_results = []
    batch_num = 0
    total_images = state.total_images   # 全局总图片数
    total_batches = state.total_batches
    start_time = time.time()

    # 打印全局进度表头（仅首次）
    print(f"\n  ┌{'─'*60}┐")
    print(f"  │ 全局进度汇总                                            │")
    print(f"  └{'─'*60}┘")
    print()

    while True:
        batch = batch_manager.get_next_batch()
        if not batch:
            print("\n\n[INFO] 所有批次处理完成！")
            break

        batch_num += 1
        print(f"\n{'#'*60}")
        print(f"批次 {batch_num}/{total_batches}: {batch.batch_id}  |  文件夹: {batch.folder}  |  图片: {len(batch.images)} 张")
        print(f"{'#'*60}")

        batch_results = []
        batch_success = True

        # 初始化批次进度条
        batch_progress = ProgressBar(len(batch.images), prefix=f"批次{batch_num}")

        for i, img_path in enumerate(batch.images, 1):
            try:
                result = process_single_image(
                    ocr, classifier, content_analyzer, doc_gen, ima_syncer, archiver,
                    img_path, i, len(batch.images), batch_progress
                )
                if result:
                    batch_results.append(result)
                else:
                    batch_results.append({'image': Path(img_path).name, 'is_duplicate': False, 'failed': True})
            except Exception as e:
                print(f"\n  ✗ 处理失败: {e}")
                logger.error(f"处理失败 {img_path}: {e}")
                batch_results.append({'image': Path(img_path).name, 'is_duplicate': False, 'failed': True, 'error': str(e)})
                batch_success = False

            # 每处理完一张就更新全局进度
            elapsed = time.time() - start_time
            _print_global_summary(
                batch_num, total_batches,
                len(all_results) + len(batch_results), total_images,
                elapsed
            )

        batch_progress.finish()
        print()   # 换行，避免与后续内容混在一起

        # 标记批次完成
        if batch_success:
            batch_manager.mark_batch_completed(batch.batch_id, f"成功处理 {len(batch_results)} 张")
        else:
            failed_count = sum(1 for r in batch_results if r.get('failed'))
            batch_manager.mark_batch_failed(batch.batch_id, f"失败 {failed_count} 张")

        all_results.extend(batch_results)
        
        # 批次间休息
        if batch_num < total_batches:
            print("\n[INFO] 批次间隔休息 1 秒...")
            time.sleep(1)
    
    # -------------------------------------------------------
    # 全部批次完成后：汇总统计 + 生成报告
    # -------------------------------------------------------
    total_elapsed = time.time() - start_time
    valid_results = [r for r in all_results if not r.get('failed')]
    dup_count     = sum(1 for r in valid_results if r.get('is_duplicate'))
    md_count      = sum(1 for r in valid_results if r.get('md_file'))
    docx_count    = sum(1 for r in valid_results if r.get('docx_file'))
    failed_count  = sum(1 for r in all_results if r.get('failed'))

    print("\n" + "=" * 60)
    print("📊 批次处理完成 — 汇总统计")
    print("=" * 60)
    print(f"  - 总图片数:   {total_images} 张")
    print(f"  - 成功处理:   {len(valid_results)} 张")
    print(f"  - 失败:       {failed_count} 张")
    print(f"  - 重复跳过:   {dup_count} 张")
    print(f"  - 新建MD:     {md_count} 个")
    print(f"  - 生成Word:   {docx_count} 个")
    print(f"  - 总耗时:     {int(total_elapsed // 60)}分{int(total_elapsed % 60)}秒")
    print("=" * 60)

    # 保存报告（与 full 模式格式一致）
    Path('处理结果').mkdir(exist_ok=True)
    report_file = Path('处理结果/处理报告.json')
    report = {
        'timestamp': datetime.now().isoformat(),
        'version': 'V6.0',
        'mode': 'batch',
        'total_batches': batch_num,
        'total_images': total_images,
        'processed': len(valid_results),
        'new_documents': md_count,
        'word_documents': docx_count,
        'duplicates': dup_count,
        'failed': failed_count,
        'elapsed_seconds': round(total_elapsed, 1),
        'theme_stats': dict(Counter(r['theme'] for r in valid_results if r.get('theme') and not r.get('is_duplicate'))),
        'category_stats': dict(Counter(r['category'] for r in valid_results if r.get('category') and not r.get('is_duplicate'))),
        'results': [{k: v for k, v in r.items() if k not in ['md_file', 'docx_file']} for r in all_results]
    }
    with open(report_file, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    print(f"\n📋 报告已保存: {report_file}")

    return all_results


# ============================================================
# 全自动处理流程（不分批）
# ============================================================
def run_full_mode():
    """全自动处理模式"""
    print("=" * 60)
    print("全自动化图片处理工具 V6.0")
    print("OCR → 智能分类 → AI分析 → 文档合并 → MD+Word → IMA → 归档")
    print("=" * 60)

    # P2-12: 处理前预检
    print("\n🔍 [P2-12] 执行预检...")
    check_passed, issues = preflight_check()
    if check_passed:
        print("  ✓ 预检通过")
    else:
        print("  ⚠️ 预检发现问题：")
        for issue in issues:
            print(f"    - {issue}")
        print("\n  继续处理（部分功能可能受限）...")
    
    # P2-17: 显示API Key配置状态
    env_results = validate_env_config()
    print("\n📋 [P2-17] API配置状态：")
    print(f"  腾讯云OCR: {'✓ 格式正确' if env_results.get('tencent') else '✗ 未配置/格式错误'}")
    print(f"  百度OCR: {'✓ 格式正确' if env_results.get('baidu') else '✗ 未配置/格式错误'}")
    print(f"  IMA同步: {'✓ 已配置' if env_results.get('ima') else '✗ 未配置'}")

    # 初始化组件
    print("\n[1/7] 初始化OCR引擎...")
    ocr = MultiEngineOCR()
    if not ocr.current_engine:
        print("❌ 没有可用的OCR引擎!")
        return

    print("\n[2/7] 初始化智能分类器...")
    classifier = ClassifierEngine()

    print("\n[3/7] 初始化AI内容分析器...")
    content_analyzer = ContentAnalyzer()

    print("\n[4/7] 初始化文档生成器...")
    doc_gen = SmartDocumentGenerator()

    print("\n[5/7] 初始化IMA同步器...")
    ima_syncer = IMASyncer()

    print("\n[6/7] 初始化图片归档器...")
    archiver = ImageArchiver()

    print("\n[7/7] 扫描待处理图片...")
    source_dir = Path('待处理图片')
    if not source_dir.exists():
        print("❌ 待处理图片目录不存在!")
        return

    images = []
    for ext in ['*.jpg', '*.png', '*.jpeg', '*.webp', '*.bmp']:
        images.extend(source_dir.rglob(ext))

    if not images:
        print("❌ 未找到待处理图片")
        return

    total = len(images)
    print(f"✓ 找到 {total} 张图片")
    subfolders = [d.name for d in source_dir.iterdir() if d.is_dir() and not d.name.startswith('.')]
    if subfolders:
        print(f"  子文件夹: {subfolders}")

    # 初始化进度条
    progress = ProgressBar(total, prefix="📊 处理进度")

    # 批量处理
    results = []
    start_time = time.time()

    for i, img_path in enumerate(images, 1):
        try:
            result = process_single_image(
                ocr, classifier, content_analyzer, doc_gen, ima_syncer, archiver,
                str(img_path), i, total, progress
            )
            if result:
                results.append(result)
        except Exception as e:
            print(f"\n  ✗ 处理失败: {e}")
            logger.error(f"处理失败 {img_path}: {e}")

    progress.finish()
    
    # 统计
    elapsed = time.time() - start_time

    print("\n" + "=" * 60)
    print("处理完成!")
    print("=" * 60)
    print(f"总计处理: {len(results)}/{total} 张")
    print(f"总耗时: {elapsed:.1f} 秒")

    # 主题分布
    if results:
        print("\n📊 主题分布:")
        themes = Counter(r['theme'] for r in results if not r.get('is_duplicate'))
        for theme, count in themes.most_common():
            print(f"  - {theme}: {count} 张")

    # 分类分布
    if results:
        print("\n🏷️ 分类分布:")
        categories = Counter(r['category'] for r in results if not r.get('is_duplicate'))
        for cat, count in categories.most_common():
            print(f"  - {cat}: {count} 张")

    # 重复检测
    dup_count = sum(1 for r in results if r.get('is_duplicate'))
    if dup_count > 0:
        print(f"\n🔍 重复检测: 跳过 {dup_count} 条重复内容")

    # 输出统计
    md_count = len([r for r in results if r.get('md_file')])
    docx_count = len([r for r in results if r.get('docx_file')])
    print(f"\n📁 输出统计:")
    print(f"  - Markdown: {md_count} 个")
    print(f"  - Word: {docx_count} 个")

    # 归档统计
    archive_stats = archiver.get_archive_stats()
    print(f"\n📦 归档统计:")
    print(f"  - 本次归档: {archive_stats['archived_count']} 张")
    print(f"  - 已归档总数: {archive_stats['total_archived']} 张")

    if ima_syncer.rate_limited:
        print(f"\n☁️ IMA: 被限流，请明天再试")

    print("\n" + "=" * 60)

    # 保存报告
    report_file = Path('处理结果/处理报告.json')
    report = {
        'timestamp': datetime.now().isoformat(),
        'version': 'V6.0',
        'mode': 'full',
        'total_images': total,
        'processed': len(results),
        'new_documents': md_count,
        'word_documents': docx_count,
        'duplicates': dup_count,
        'archived': archive_stats['archived_count'],
        'elapsed_seconds': elapsed,
        'theme_stats': dict(Counter(r['theme'] for r in results if not r.get('is_duplicate'))),
        'category_stats': dict(Counter(r['category'] for r in results if not r.get('is_duplicate'))),
        'results': [{k: v for k, v in r.items() if k not in ['md_file', 'docx_file']} for r in results]
    }
    with open(report_file, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    print(f"\n📊 报告已保存: {report_file}")


# ============================================================
# 主入口
# ============================================================
def main():
    batch_manager = BatchManager()
    
    # 命令行参数处理
    if len(sys.argv) > 1:
        cmd = sys.argv[1]
        
        if cmd == '--init':
            # 仅初始化分批
            state = batch_manager.initialize()
            if state:
                batch_manager.print_progress()
        
        elif cmd == '--progress':
            # 查看进度
            batch_manager.print_progress()
        
        elif cmd == '--clear':
            # 清除状态
            batch_manager.clear_state()
        
        elif cmd == '--batch':
            # 分批处理模式
            run_batch_mode(batch_manager)
        
        else:
            print(f"[ERROR] 未知参数: {cmd}")
            print_usage()
    
    else:
        # 默认：全自动处理
        run_full_mode()


def print_usage():
    """打印使用说明"""
    print("\n用法:")
    print("  python auto_process_all_v6.py           # 全自动处理所有图片")
    print("  python auto_process_all_v6.py --batch    # 分批模式（可中断续传）")
    print("  python auto_process_all_v6.py --init     # 仅初始化分批")
    print("  python auto_process_all_v6.py --progress # 查看进度")
    print("  python auto_process_all_v6.py --clear    # 清除分批状态")


if __name__ == '__main__':
    main()
